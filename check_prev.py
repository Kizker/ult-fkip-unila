import docx

doc = docx.Document(r"c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text.startswith("Gambar 37") or text.startswith("Gambar 38") or text.startswith("Gambar 39"):
        print(f"Index {i}: {text}")
        prev_p = doc.paragraphs[i-1]
        print(f"  Prev P text: '{prev_p.text}'")
        has_drawing = 'drawing' in prev_p._p.xml
        print(f"  Prev P has drawing: {has_drawing}")
