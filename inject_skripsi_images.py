import os
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def inject_images(doc_path, img_dir, save_path):
    print(f"Memproses dokumen: {doc_path}")
    doc = Document(doc_path)
    
    # Track injections
    injected_count = 0
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        # Look for [Placeholder Gambar XX]
        match = re.match(r'^\[Placeholder Gambar (4[3-9]|5[0-2])\]', text, re.IGNORECASE)
        if match:
            img_num = match.group(1)
            img_filename = f"Gambar {img_num}.png"
            img_path = os.path.join(img_dir, img_filename)
            
            if os.path.exists(img_path):
                print(f"  -> Menginjeksi {img_filename} di tempat placeholder.")
                # Clear the text of the placeholder paragraph
                paragraph.text = ""
                # Add picture
                run = paragraph.add_run()
                run.add_picture(img_path, width=Inches(5.5))
                # Set alignment to Center
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                injected_count += 1
            else:
                print(f"  -> PERINGATAN: Gambar {img_filename} tidak ditemukan!")
                
    if injected_count > 0:
        doc.save(save_path)
        print(f"Selesai! Disimpan ke: {save_path}\n")
    else:
        print("Tidak ada gambar yang diinjeksi.\n")

if __name__ == "__main__":
    base_dir = r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update"
    img_dir = os.path.join(base_dir, "screenshots_baru")
    
    doc_clean = os.path.join(base_dir, "001_Skripsi_Andricha Dea Mitra_Clean.docx")
    doc_highlighted = os.path.join(base_dir, "001_Skripsi_Andricha Dea Mitra_Highlighted.docx")
    
    inject_images(doc_clean, img_dir, doc_clean)
    inject_images(doc_highlighted, img_dir, doc_highlighted)
