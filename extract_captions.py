import docx
import re

def extract_captions(file_path):
    try:
        doc = docx.Document(file_path)
        print(f"Berhasil membuka dokumen: {file_path}")
        
        captions = []
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            # Mencari teks yang dimulai dengan "Gambar " diikuti angka 43 sampai 52
            match = re.match(r'^Gambar (4[3-9]|5[0-2])\..*', text)
            if match:
                prev_text = doc.paragraphs[i-1].text.strip() if i > 0 else ""
                captions.append({
                    "caption": text,
                    "prev_text": prev_text
                })
        
        if not captions:
            print("Tidak ditemukan caption Gambar 43 - 52.")
        else:
            for c in captions:
                print(f"Caption: {c['caption']}")
                print(f"Prev   : {c['prev_text']}")
                print("-" * 50)
                
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    file_path = r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx"
    extract_captions(file_path)
