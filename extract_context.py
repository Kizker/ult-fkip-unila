import docx
import re

def extract_context(file_path):
    try:
        doc = docx.Document(file_path)
        
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            match = re.match(r'^Gambar (4[3-9]|5[0-2])\..*', text)
            if match:
                print(f"=== {text} ===")
                # Print 3 paragraphs before and 3 after
                for j in range(max(0, i-5), min(len(doc.paragraphs), i+5)):
                    if j == i:
                        continue
                    print(f"[{j}]: {doc.paragraphs[j].text.strip()}")
                print("\n")
                
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    file_path = r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx"
    extract_context(file_path)
