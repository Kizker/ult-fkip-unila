import os
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def revert_images(doc_path, save_path):
    print(f"Memulihkan dokumen: {doc_path}")
    doc = Document(doc_path)
    
    reverted_count = 0
    
    # We will iterate through paragraphs.
    for i in range(len(doc.paragraphs) - 1):
        p = doc.paragraphs[i]
        
        # Check if this paragraph contains a drawing/picture
        has_picture = '<w:drawing>' in p._element.xml
        
        if has_picture:
            # Look ahead for the caption
            caption_p = None
            for j in range(i + 1, min(i + 4, len(doc.paragraphs))):
                next_p = doc.paragraphs[j].text.strip()
                if next_p:
                    caption_p = next_p
                    break
            
            if caption_p:
                match = re.search(r'Gambar\s*(4[3-9]|5[0-2])\.', caption_p, re.IGNORECASE)
                if match:
                    img_num = match.group(1)
                    print(f"  -> Menemukan gambar untuk {caption_p[:30]}... Mengembalikan ke [Placeholder Gambar {img_num}]")
                    # Clear the paragraph (removes the picture)
                    p.clear()
                    # Add back the placeholder
                    p.text = f"[Placeholder Gambar {img_num}]"
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    reverted_count += 1
                    
    if reverted_count > 0:
        doc.save(save_path)
        print(f"Selesai! Disimpan ke: {save_path}\n")
    else:
        print("Tidak ada gambar yang dipulihkan.\n")

if __name__ == "__main__":
    base_dir = r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update"
    
    doc_clean = os.path.join(base_dir, "001_Skripsi_Andricha Dea Mitra_Clean.docx")
    doc_highlighted = os.path.join(base_dir, "001_Skripsi_Andricha Dea Mitra_Highlighted.docx")
    
    revert_images(doc_clean, doc_clean)
    revert_images(doc_highlighted, doc_highlighted)
