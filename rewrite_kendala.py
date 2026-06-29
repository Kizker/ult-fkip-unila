import docx
import os
import re

FILES_TO_PROCESS = [
    r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx',
    r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
]

# Ordered list of operations based on prefix matching
REPLACEMENTS = {
    "Dalam proses pengembangan dan implementa": "Dalam proses pengembangan dan implementasi sistem pada lingkungan Universitas Lampung, terdapat beberapa kendala teknis yang memerlukan penyesuaian fungsional dan arsitektural. Resolusi atas kendala-kendala tersebut dijabarkan sebagai berikut:",
    
    "Evaluasi tahapan purwarupa awal sistem w": "Evaluasi pada purwarupa awal menunjukkan adanya ketidaksesuaian format teks ketika dokumen hasil pengisian form pengguna (berbasis WYSIWYG editor) dikonversi menjadi berkas biner Word (.docx).",
    
    "Ketidakhadiran fungsi penyelarasan peren": "Sistem bawaan Tiptap Editor pada antarmuka web memproduksi keluaran dalam bentuk tag HTML mentah. Akibatnya, elemen pemformatan dasar seperti paragraf baru, teks tebal, maupun penomoran daftar urut (list) tidak dapat dirender secara visual saat diekstraksi ke dalam berkas Word, melainkan hanya tercetak sebagai deretan tag HTML (seperti `<p>` dan `<b>`).",
    
    "Kasus inkompatibilitas konversi struktur": "Ketidakmampuan sistem merender tag HTML secara orisinal ini dapat memicu disrupsi visual yang mengurangi nilai formalitas dan legalitas dari draf surat akademik yang dihasilkan oleh portal layanan.",
    
    "Efektivitas hasil penyuntikan komponen p": "Untuk mengatasi isu degradasi pemformatan tersebut, diimplementasikan algoritma parser HTML-to-OpenXML secara dinamis yang terpusat pada file XML mentah (`document.xml`). Modul ini secara efektif menelusuri tag-tag HTML visual (seperti `<b>`, `<i>`, dan `<ul>`) dan mengonversinya langsung menjadi elemen struktur OpenXML (`<w:b/>`, `<w:i/>`) saat proses perakitan dokumen berlangsung. Hasil resolusi ini sukses menduplikasi gaya teks orisinal tanpa merusak tata letak margin, ukuran fon, maupun spasi dokumen (Gambar 44).",
    
    "Proses pemindaian kelayakan infrastruktur": "Fase pengujian sistem juga memprioritaskan validasi keamanan pada setiap antarmuka masukan data, mengingat potensi ancaman siber berbasis Cross-Site Scripting (XSS).",
    
    "Formulir isian penyedia antarmuka masuka": "Pada uji coba keamanan awal (Gambar 45), simulasi injeksi skrip XSS yang dilakukan pada kolom masukan berhasil lolos dan dieksekusi oleh peramban klien dalam bentuk kotak peringatan (alert box). Hal ini membuktikan adanya celah kerentanan yang memungkinkan pihak tidak bertanggung jawab menyisipkan kode berbahaya.",
    
    "Satuan kerja perbaikan teknisi arsitektu": "Sebagai tindakan perbaikan, diterapkan kebijakan Content Security Policy (CSP) pada header HTTP, serta penggunaan pustaka sanitasi HTML pihak ketiga (seperti `mews/purifier`) pada level pengendali (controller) aplikasi.",
    
    "Laporan log pemeriksaan uji penetrasi se": "Hasil uji coba ulang (Gambar 46) menunjukkan bahwa seluruh skrip injeksi XSS telah dieliminasi secara preventif. Setiap input berbahaya dinetralisasi menjadi teks biasa, memastikan keamanan integritas pangkalan data serta melindungi para pengambil keputusan (pejabat fakultas) dari eksploitasi peretasan.",
    
    # Deletions (Poin C)
    "Penyesuaian Komposisi Warna dan Tata Let": None,
    "Tim evaluator antarmuka interaktif membe": None,
    "[Placeholder Gambar 47]": None,
    "Gambar 47. Komposisi Warna Antarmuka Web": None,
    "Integrasi pengaturan atribut pewarnaan l": None,
    "Tahapan restorasi arsitektur penyusunan ": None,
    "[Placeholder Gambar 48]": None,
    "Gambar 48. Komposisi Warna Antarmuka Web": None,
    "Standardisasi palet desain perancangan a": None,
    
    # Deletions (Poin D)
    "Otomatisasi Penomoran Surat dan Penjadwa": None,
    "Tahapan asesmen identifikasi observasi o": None,
    "[Placeholder Gambar 49]": None,
    "Gambar 49. Penomoran Surat dan Penjadwal": None,
    "Operasional pengelolaan kalender meja ma": None,
    "Implementasi penyelesaian sistem layanan": None,
    "Penggabungan infrastruktur perangkat fun": None,
    "[Placeholder Gambar 50]": None,
    "Gambar 50. Penomoran Surat dan Penjadwal": None,
    "Penyelarasan instruksi algoritma penomor": None,
    
    # Poin E (becomes C)
    "Tanggapan rujukan pengawasan instrumen p": "Selain masalah fungsionalitas, evaluasi visual turut mencatat keluhan mengenai kepadatan antarmuka pada tampilan dasbor tabel manajemen data.",
    
    "[Placeholder Gambar 51]": "[Placeholder Gambar 47]",
    "Gambar 51. Tata Letak Ruang Kosong Tabel": "Gambar 47. Tata Letak Ruang Kosong Tabel - Kondisi Sebelum.",
    
    "Konfigurasi rancangan matriks pangkalan ": "Pada rancangan awal (Gambar 47), susunan sel tabel memiliki jarak antarbaris (padding) yang terlampau sempit. Hal ini membuat rincian data pemohon terlihat sesak dan membebani daya visual staf layanan yang harus membaca puluhan baris data setiap harinya.",
    
    "Eksekusi revisi perbaikan fungsional imp": "Revisi arsitektur antarmuka dilakukan dengan memperluas rasio padding sel tabel dan menyesuaikan ukuran ketebalan fon pada bagian tajuk.",
    
    "Kesesuaian penerapan parameter penyusuna": "Penerapan prinsip white space (ruang kosong) yang optimal ini bertujuan untuk meningkatkan harmoni visual, hierarki informasi, dan keterbacaan data, sesuai dengan pedoman desain antarmuka pengguna yang terstruktur.",
    
    "[Placeholder Gambar 52]": "[Placeholder Gambar 48]",
    "Gambar 52. Tata Letak Ruang Kosong Tabel": "Gambar 48. Tata Letak Ruang Kosong Tabel - Kondisi Sesudah.",
    
    "Implementasi pelepasan batas perluasan r": "Setelah perbaikan (Gambar 48), kelancaran pembacaan data administrasi oleh staf peninjau meningkat drastis. Penyesuaian metrik ruang pada tabel telah mengakomodasi kenyamanan ergonomi visual dan mengurangi kelelahan kognitif selama masa operasional harian.",
    
    "Dinamika kendala teknis yang sempat terj": "Penyelesaian atas berbagai disrupsi pemformatan, perisai keamanan siber, hingga ergonomi tabel antarmuka membuktikan pentingnya pengujian yang komprehensif. Konvergensi perbaikan strategis tersebut mentransformasi purwarupa awal menjadi ekosistem digital yang andal dan aman bagi kebutuhan operasional.",
    
    "Untuk mengatasi stagnasi pengujian terse": "Selama fase pengujian kepraktisan sistem, salah satu tantangan lapangan utama adalah mobilitas responden yang sangat padat. Tingginya ritme operasional harian pimpinan, staf ULT, maupun mahasiswa menyebabkan keterlambatan yang berisiko pada pengumpulan instrumen kuesioner.",
    
    "Dari dimensi manajemen lapangan, tantang": "Oleh karena itu, diimplementasikan strategi \"jemput bola\" melalui pendampingan teknis secara langsung (on-the-spot) dan ringkas. Responden diberikan skenario uji yang dikondensasi sedemikian rupa, sehingga durasi simulasi dapat diminimalkan tanpa mengaburkan penilaian terhadap fitur-fitur utama sistem aplikasi.",
    
    "Dalam merealisasikan produk pengembangan": "Lebih lanjut, tantangan migrasi sistem dari peladen lokal menuju peladen awan (cloud) memerlukan harmonisasi spesifikasi komputasi, seperti pembaruan lingkungan PHP versi 8.4+ agar sesuai dengan prasyarat framework Laravel 12. Sistem keamanan penyimpanan berkas privat juga dikonfigurasi menggunakan tautan simbolik (symlink) khusus untuk memastikan isolasi perlindungan dari ancaman Insecure Direct Object Reference (IDOR).",
    
    "Transisi sistem birokrasi dari model kon": "Secara keseluruhan, rangkaian resolusi dari kendala teknis maupun manajemen lapangan yang timbul ini merupakan bagian integral dari proses penyempurnaan adaptasi arsitektur peranti lunak, yang bermuara pada kesiapan fungsional portal administratif fakultas yang optimal."
}

def clean_paragraph_text(p):
    return p.text.strip().replace('\u200b', '')

for filepath in FILES_TO_PROCESS:
    print(f"Processing: {filepath}")
    doc = docx.Document(filepath)
    paragraphs_to_remove = []
    
    pengujian_keamanan_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        text = clean_paragraph_text(p)
        if not text:
            continue
            
        for prefix, replacement in REPLACEMENTS.items():
            if text.startswith(prefix):
                if replacement is None:
                    paragraphs_to_remove.append(p)
                else:
                    p.text = replacement
                    if 'Highlighted' in filepath:
                        for run in p.runs:
                            run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
                break
                
        if text.startswith("Pengujian Keamanan Tingkat Lanjut:"):
            pengujian_keamanan_idx = i

    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)
        
    if pengujian_keamanan_idx != -1:
        saran_text = "Integrasi Penomoran Surat dan Penjadwalan Pimpinan: Disarankan untuk mengembangkan fitur kalender digital terintegrasi pada dasbor pimpinan fakultas yang tersinkronisasi langsung dengan mesin penomoran surat. Hal ini bertujuan untuk mencegah terjadinya tumpang tindih alokasi waktu pimpinan dan meminimalisir kesalahan cetak pada nomor dokumen kelulusan yang saat ini masih rentan terjadi."
        
        target_idx = pengujian_keamanan_idx + 1
        if target_idx < len(doc.paragraphs):
            new_p = doc.paragraphs[target_idx].insert_paragraph_before(saran_text)
            new_p.style = doc.paragraphs[pengujian_keamanan_idx].style
        else:
            new_p = doc.add_paragraph(saran_text)
            new_p.style = doc.paragraphs[pengujian_keamanan_idx].style
            
        if 'Highlighted' in filepath:
            for run in new_p.runs:
                run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW

    doc.save(filepath)
    print("Saved successfully.")
