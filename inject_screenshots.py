import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

base_dir = r"c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update"
img_dir = os.path.join(base_dir, "screenshots_baru")

clean = f"{base_dir}\\001_Skripsi_Andricha Dea Mitra_Clean.docx"
highlight = f"{base_dir}\\001_Skripsi_Andricha Dea Mitra_Highlighted.docx"

def inject_images(docx_path):
    doc = docx.Document(docx_path)
    
    img_map = {
        "Gambar 37": "Gambar 37.png",
        "Gambar 38": "Gambar 38.png",
        "Gambar 39": "Gambar 39.png",
        "Gambar 40": "Gambar 40.png",
        "Gambar 41": "Gambar 41.png",
        "Gambar 42": "Gambar 42.png",
    }
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        for g_num, img_file in img_map.items():
            if text.startswith(g_num):
                # Found the caption. The image goes in the paragraph before it.
                prev_p = doc.paragraphs[i-1]
                prev_p.clear() # Remove old placeholder text and old image
                run = prev_p.add_run()
                img_path = os.path.join(img_dir, img_file)
                if os.path.exists(img_path):
                    run.add_picture(img_path, width=Inches(5.5))
                    prev_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    print(f"Injected {img_file} at {g_num}")
                else:
                    print(f"MISSING: {img_path}")
    
    doc.save(docx_path)
    print(f"Saved {docx_path}")

inject_images(clean)
inject_images(highlight)
