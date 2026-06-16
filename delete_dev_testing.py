import docx

def delete_dev_testing(doc_path):
    print(f"Processing {doc_path}")
    doc = docx.Document(doc_path)
    
    start_idx = -1
    end_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text == "c. Developmental Testing":
            start_idx = i
        elif text == "Tahap Implementasi (Implementation)":
            end_idx = i
            break
            
    if start_idx != -1 and end_idx != -1:
        print(f"Deleting paragraphs from {start_idx} to {end_idx - 1}")
        for i in range(start_idx, end_idx):
            p_elem = doc.paragraphs[i]._element
            p_elem.getparent().remove(p_elem)
        doc.save(doc_path)
        print("Success")
    else:
        print(f"Could not find markers: start_idx={start_idx}, end_idx={end_idx}")

delete_dev_testing(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx")
delete_dev_testing(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx")
