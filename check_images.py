import docx

doc = docx.Document(r"c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text.startswith("Gambar 37") or text.startswith("Gambar 38") or text.startswith("Gambar 39") or text.startswith("Gambar 40") or text.startswith("Gambar 41") or text.startswith("Gambar 42"):
        print(f"Index {i}: {text}")
        # check preceding and succeeding paragraphs for images
        for offset in [-2, -1, 1, 2]:
            idx = i + offset
            if 0 <= idx < len(doc.paragraphs):
                p_offset = doc.paragraphs[idx]
                has_img = 'Graphic' in p_offset._p.xml or 'pic:pic' in p_offset._p.xml or 'drawing' in p_offset._p.xml
                if has_img:
                    print(f"  Paragraph {offset} away contains an image!")
