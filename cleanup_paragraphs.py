import docx
import re

def remove_respondent_lines(doc_path):
    print(f"Cleaning up {doc_path}")
    doc = docx.Document(doc_path)
    
    # We will search for any paragraph starting with "Responden"
    # and "Admin Program Studi", "Mahasiswa", "Staf Unit" etc.
    # since we already put them in the table.
    
    # We only want to remove the ones outside the table!
    # doc.paragraphs only contains paragraphs in the body (not inside tables)
    
    removed_count = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if re.match(r'^Responden \d+ \(.*?\):', text):
            p._element.getparent().remove(p._element)
            removed_count += 1
            
    # Also remove empty paragraphs around the area if possible, 
    # but let's just focus on the Respondent lines first.
    doc.save(doc_path)
    print(f"Removed {removed_count} lines")

remove_respondent_lines(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx")
remove_respondent_lines(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx")
