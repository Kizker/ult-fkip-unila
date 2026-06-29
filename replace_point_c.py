import docx
import os

FILES = [
    r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx',
    r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
]

NEW_CONTENT = [
    ("c. Integrasi Dasbor Analitik dan Visualisasi Data", False),
    ("Pengembangan antarmuka peladen komputasi juga berfokus pada integrasi dasbor analitik berbasis visualisasi grafik. Kebutuhan ini muncul dari keluhan staf administrasi yang kesulitan membaca laporan rekapitulasi permohonan dalam bentuk tumpukan data tabel statis. Penyediaan modul grafik statistik menjadi solusi krusial untuk mempercepat proses pengambilan keputusan oleh pimpinan fakultas.", True),
    ("[Placeholder Gambar 47]", False),
    ("Gambar 47. Dasbor Statistik - Kondisi Sebelum (Tanpa Grafik).", False),
    ("Rancangan antarmuka versi purwarupa (Gambar 47) memperlihatkan kegagalan sistem dalam memuat pustaka grafik secara dinamis. Modul visualisasi pada dasbor awal hanya menampilkan ruang kosong atau pesan galat akibat bentrokan kompatibilitas antara kerangka kerja antarmuka dan pustaka eksternal. Kegagalan rendering visualisasi ini menyebabkan pemantauan laju dokumen permohonan akademik berjalan sangat tidak efisien.", True),
    ("Perombakan arsitektural dilakukan dengan menginjeksi pustaka Chart.js secara asinkron menggunakan modul bundler Vite pada ekosistem Laravel. Modul penerjemah ini secara efektif memproses data mentah agregat dari pangkalan data dan mengubahnya menjadi visualisasi interaktif berbentuk grafik tren maupun diagram. Sinkronisasi data grafik ini dijamin mutakhir secara real-time karena langsung terhubung dengan antarmuka pemrograman aplikasi internal sistem.", True),
    ("[Placeholder Gambar 48]", False),
    ("Gambar 48. Dasbor Analitik Statistik - Kondisi Sesudah.", False),
    ("Kondisi pasca perbaikan (Gambar 48) membuktikan bahwa seluruh modul dasbor analitik telah berfungsi secara optimal. Visualisasi grafik interaktif ini sukses memberikan gambaran komprehensif terkait rasio status permohonan dan tren pendaftaran harian mahasiswa. Pencapaian teknis ini secara drastis meningkatkan efisiensi waktu staf pengelola dalam menyusun laporan rekapitulasi operasional layanan persuratan akademik.", True),
    ("Penyelesaian holistik terhadap berbagai kendala degradasi pemformatan, perisai celah keamanan siber, hingga integrasi dasbor analitik membuktikan urgensi dari uji kelayakan sistem. Konvergensi perbaikan strategis ini sukses mentransformasi kelemahan purwarupa awal menjadi fondasi kekuatan ekosistem digital layanan pendidikan yang modern. Produk akhir peranti lunak kini hadir sebagai sebuah sistem persuratan terpadu yang andal, aman, dan berdaya tahan tinggi menghadapi lonjakan kebutuhan administrasi perguruan tinggi.", True),
]

def apply_highlight(p):
    for run in p.runs:
        run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW

def replace_point_c():
    for filepath in FILES:
        print(f"Processing: {filepath}")
        doc = docx.Document(filepath)
        
        # Find the starting index
        start_idx = -1
        for i, p in enumerate(doc.paragraphs):
            if 'Optimalisasi Tata Letak dan Ruang Kosong Tabel' in p.text or 'c. Optimalisasi' in p.text:
                start_idx = i
                if p.text.strip() == 'Optimalisasi Tata Letak dan Ruang Kosong Tabel':
                    break
        
        if start_idx == -1:
            print(f"Start index not found in {filepath}")
            continue
            
        print(f"Found starting paragraph at {start_idx}")
        
        # In current document, lines 1209 to 1219 are exactly 11 paragraphs.
        # We will replace these 11 paragraphs with the 10 paragraphs of NEW_CONTENT.
        # To do this safely without breaking XML structure (like images/styles),
        # we can modify the text of the first 10, and clear the 11th.
        
        for i, (new_text, highlight) in enumerate(NEW_CONTENT):
            p = doc.paragraphs[start_idx + i]
            # preserve paragraph style if needed, but just changing text is safer
            style = p.style
            p.text = new_text
            p.style = style
            
            if highlight and 'Highlighted' in filepath:
                apply_highlight(p)
                
        # Clear the 11th paragraph (which was the old conclusion, now moved to index 1218)
        # Wait, NEW_CONTENT has 10 items. The old block had 11 items (1209 to 1219).
        # Let's count old items: 
        # 1209: Heading (1)
        # 1210: Par 1 (2)
        # 1211: [Placeholder 47] (3)
        # 1212: Gambar 47 title (4)
        # 1213: Par 2 (5)
        # 1214: Par 3 (6)
        # 1215: Par 4 (7)
        # 1216: [Placeholder 48] (8)
        # 1217: Gambar 48 title (9)
        # 1218: Par 5 (10)
        # 1219: Conclusion (11)
        
        # If I only have 10 items in NEW_CONTENT, I should add a dummy empty text for the 11th item to effectively remove it, or just combine two paragraphs.
        # Let's clear the 11th paragraph.
        doc.paragraphs[start_idx + 10].text = ""

        doc.save(filepath)
        print("Saved.")

if __name__ == '__main__':
    replace_point_c()
