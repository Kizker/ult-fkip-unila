from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image


ROOT = Path(r"c:/laragon/www/ult-fkip-unila")
OUTPUT = ROOT / "docs/buku-panduan/buku-panduan-pemohon-layanan/buku-panduan-pemohon-layanan.docx"
ASSETS = ROOT / "docs/buku-panduan/assets/screenshots"
FONT_NAME = "Poppins"
COLOR_PRIMARY = RGBColor(91, 33, 182)
COLOR_TEXT = RGBColor(31, 41, 55)
PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
TOP_MARGIN_CM = 2.5
BOTTOM_MARGIN_CM = 2.5
LEFT_MARGIN_CM = 3.0
RIGHT_MARGIN_CM = 2.5
LIST_TEXT_INDENT = Cm(1.2)
LIST_HANGING_INDENT = Cm(-0.6)

CHAPTERS = [
    {
        "title": "BAB I PENDAHULUAN",
        "intro": [
            "Bab ini menjelaskan alasan penyusunan buku panduan dan batas pembahasan yang dipakai pada dokumen. Fokus utama panduan diarahkan pada pengalaman pemohon saat mengakses layanan, mengajukan permohonan, memantau status, menindaklanjuti catatan, hingga memperoleh hasil layanan melalui Website ULT FKIP Universitas Lampung.",
        ],
        "sections": [
            {
                "title": "1.1 Latar Belakang",
                "body": [
                    "Website Unit Layanan Terpadu (ULT) FKIP Universitas Lampung dikembangkan untuk membuat proses layanan lebih terpusat, terdokumentasi, dan mudah ditelusuri. Namun, keberadaan sistem saja belum cukup jika pengguna belum memiliki acuan yang jelas saat pertama kali berhadapan dengan alur digital yang berbeda dari layanan manual.",
                    "Dari sudut pandang pemohon, proses pada website tidak berhenti pada tahap membuka layanan lalu mengirim form. Pemohon juga harus memahami cara memilih layanan yang tepat, membaca persyaratan dan SOP, menyiapkan dokumen, mengunggah file sesuai kebutuhan, memantau status, menanggapi catatan petugas, melakukan perbaikan, sampai mengambil hasil layanan. Tanpa panduan yang tersusun rapi, pengguna berisiko salah langkah, salah unggah dokumen, atau keliru menafsirkan status proses.",
                    "Buku panduan ini diperlukan sebagai pedoman operasional yang menjembatani antarmuka sistem dengan kebutuhan pengguna nyata. Melalui panduan yang runtut dan berbasis tampilan aktual, pemohon dapat memahami alasan setiap langkah, mengetahui tindakan yang harus dilakukan pada setiap tahap, dan menggunakan sistem secara lebih tertib. Dengan demikian, buku ini bukan hanya pelengkap dokumentasi, tetapi juga alat bantu penting agar penggunaan website berjalan konsisten, efisien, dan mengurangi kesalahan selama proses layanan.",
                ],
            },
            {
                "title": "1.2 Tujuan Buku Panduan",
                "body": [
                    "Buku panduan ini bertujuan memberi acuan operasional yang lebih faktual bagi pemohon layanan. Dokumen ini juga menjadi pelengkap dokumentasi produk sehingga alur penggunaan sistem dapat dijelaskan secara konsisten kepada pengguna baru, penguji, dan pihak pengelola layanan.",
                    "Secara lebih rinci, tujuan penyusunan buku panduan ini dapat dijabarkan ke dalam beberapa poin berikut.",
                ],
                "bullets": [
                    "menjelaskan langkah penggunaan portal pemohon dari login sampai unduh output;",
                    "membantu pemohon membaca status, riwayat, dan catatan tindak lanjut secara benar;",
                    "menunjukkan area antarmuka yang penting agar pengguna lebih cepat beradaptasi;",
                    "mencatat fitur tambahan yang hanya muncul pada layanan tertentu, seperti tanda tangan pemohon, signer lain, lampiran umum, dan preview dokumen.",
                ],
            },
            {
                "title": "1.3 Capaian Buku Panduan",
                "body": [
                    "Setelah mempelajari buku panduan ini, pemohon diharapkan tidak hanya mengenali tampilan sistem, tetapi juga mampu menjalankan alur penggunaan portal secara mandiri, tertib, dan sesuai kebutuhan layanan yang dipilih, sehingga buku panduan benar-benar berfungsi sebagai acuan praktik penggunaan portal pemohon.",
                    "Secara lebih operasional, capaian buku panduan ini dapat dijabarkan ke dalam beberapa poin berikut.",
                ],
                "bullets": [
                    "mampu membuat akun, melakukan verifikasi, dan login ke portal pemohon dengan benar;",
                    "mampu menelusuri layanan, membaca persyaratan, dan memahami SOP sebelum mengajukan permohonan;",
                    "mampu mengisi form layanan, mengunggah file, dan mengirim permohonan sesuai kebutuhan layanan;",
                    "mampu memantau status, membaca catatan, melakukan perbaikan, dan mengambil hasil layanan bila sudah tersedia.",
                ],
            },
            {
                "title": "1.4 Sasaran Pengguna Buku Panduan",
                "body": [
                    "Sasaran utama dokumen ini adalah pengguna yang memiliki role sebagai pemohon layanan. Pada implementasi saat ini, role tersebut paling sering dipakai oleh mahasiswa, tetapi panduan tetap relevan untuk pengguna lain yang diberi hak mengajukan layanan melalui portal pemohon.",
                ],
            },
            {
                "title": "1.5 Ruang Lingkup Pembahasan",
                "body": [
                    "Pembahasan difokuskan pada portal pemohon dan tahap penggunaan yang langsung berhubungan dengan aktivitas pengajuan layanan. Area pengelola layanan, penandatangan, dan superadmin hanya disebut seperlunya untuk menjelaskan konteks alur status dari sudut pandang pemohon.",
                ],
                "bullets": [
                    "akses masuk ke sistem dan orientasi dashboard pemohon;",
                    "penelusuran layanan, pembacaan detail, persyaratan, dan SOP;",
                    "pengisian formulir, pengunggahan file, serta submit permohonan;",
                    "pemantauan status, riwayat, catatan, lampiran, dan revisi data;",
                    "preview dokumen dan unduh hasil layanan bila output sudah tersedia.",
                ],
            },
        ],
    },
    {
        "title": "BAB II MENGENAL PORTAL PEMOHON",
        "intro": [
            "Portal pemohon adalah area kerja utama setelah proses autentikasi berhasil. Dari portal ini, pengguna dapat melihat ringkasan pengajuan, mengakses daftar permohonan, membuka form layanan, membaca detail permohonan, dan menindaklanjuti proses yang sedang berjalan.",
            "Portal pemohon bukan hanya tempat mengirim formulir. Portal ini juga berfungsi sebagai pusat pelacakan proses karena pemohon dapat melihat perkembangan status, membaca catatan dari petugas, mengunggah lampiran tambahan, melakukan perbaikan, membuka preview dokumen, dan mengunduh output akhir bila layanan telah selesai.",
        ],
        "sections": [
            {
                "title": "2.1 Dashboard Pemohon",
                "body": [
                    "Setelah login, sistem mengarahkan pemohon ke dashboard pengajuan layanan. Dashboard menampilkan ringkasan cepat seperti total permohonan, permohonan terbaru, dan status pengajuan yang masih aktif. Halaman ini memudahkan pengguna memahami posisi awal sebelum membuka detail pengajuan tertentu.",
                ],
                "figure": ("03-dashboard-pemohon.png", "Gambar 1. Dashboard pengajuan layanan pada portal pemohon"),
                "bullets": [
                    "bagian hero menampilkan identitas pengguna dan tombol ajukan layanan;",
                    "kartu permohonan terbaru memudahkan akses cepat ke detail pengajuan terakhir;",
                    "ringkasan status memperlihatkan distribusi permohonan berdasarkan status saat ini.",
                ],
            },
            {
                "title": "2.2 Halaman Permohonan Saya",
                "body": [
                    "Menu permohonan membuka daftar seluruh pengajuan yang pernah dibuat oleh pemohon. Halaman ini menjadi tempat utama untuk melacak status banyak permohonan sekaligus, lalu memilih satu entri untuk dibuka lebih lanjut.",
                ],
                "figure": ("05-daftar-permohonan-pemohon.png", "Gambar 2. Halaman Permohonan Saya"),
                "steps": [
                    "gunakan daftar ini saat ingin melihat seluruh pengajuan dalam satu tampilan;",
                    "perhatikan kode permohonan, nama layanan, waktu pengajuan, dan badge status;",
                    "buka tombol Detail pada kartu yang relevan untuk masuk ke halaman pemantauan yang lebih lengkap.",
                ],
            },
        ],
    },
    {
        "title": "BAB III AKSES MASUK DAN ORIENTASI AWAL",
        "intro": [
            "Bab ini membahas tahapan awal sebelum pengajuan dilakukan, yaitu membuat akun pemohon, melakukan verifikasi yang diperlukan, login ke sistem, dan membaca halaman awal portal pemohon. Pemahaman pada tahap ini penting agar pengguna masuk ke area kerja yang benar dan tidak mengalami kendala dasar saat akan mengajukan layanan.",
        ],
        "sections": [
            {
                "title": "3.1 Membuat Akun Pemohon",
                "body": [
                    "Sebelum menggunakan portal pemohon, pengguna yang belum memiliki akun perlu melakukan registrasi terlebih dahulu. Pada implementasi saat ini, halaman pendaftaran menyediakan form yang memuat identitas dasar, pilihan jenis akun, data akademik, dan unggah foto profil.",
                    "Untuk pendaftaran manual, pengguna perlu mengisi nama, email, jenis akun, jurusan, program studi, password, konfirmasi password, lalu mengunggah foto profil. Jenis akun yang tersedia pada form registrasi adalah Mahasiswa dan Dosen. Program studi harus dipilih sesuai jurusan yang dipakai pada form.",
                ],
                "figure": ("01-register-akun-pemohon.png", "Gambar 3. Form pendaftaran akun pemohon"),
                "steps": [
                    "buka halaman daftar dari tautan registrasi yang tersedia pada website atau dari halaman login;",
                    "isi nama, email, jenis akun, jurusan, program studi, password, dan konfirmasi password dengan benar;",
                    "unggah foto profil sesuai ketentuan sistem, lalu periksa kembali seluruh data akademik yang dipilih;",
                    "kirim form pendaftaran lalu cek instruksi verifikasi email yang diberikan sistem.",
                ],
                "bullets": [
                    "gunakan email aktif karena proses verifikasi akun dikirim ke alamat tersebut;",
                    "pilih jenis akun, jurusan, dan program studi sesuai data akademik yang benar;",
                    "unggah foto profil yang jelas dengan rasio yang sesuai agar tidak ditolak saat validasi form;",
                    "buat password yang kuat dan mudah Anda simpan dengan aman.",
                ],
            },
            {
                "title": "3.2 Verifikasi Akun",
                "body": [
                    "Jika email verifikasi belum diterima, pengguna sebaiknya memeriksa folder spam atau junk, lalu memastikan kembali bahwa alamat email yang dimasukkan saat registrasi sudah benar. Tahap ini penting karena akun yang belum terverifikasi dapat membatasi akses ke portal.",
                ],
                "steps": [
                    "buka inbox email yang digunakan saat registrasi;",
                    "cari pesan verifikasi dari sistem lalu buka tautannya;",
                    "setelah verifikasi berhasil, kembali ke website dan lanjutkan ke halaman login;",
                    "jika belum menerima email, periksa folder spam atau ulangi pemeriksaan alamat email yang digunakan.",
                ],
            },
            {
                "title": "3.3 Halaman Login",
                "body": [
                    "Halaman login adalah gerbang autentikasi untuk masuk ke sistem. Pengguna perlu memasukkan kredensial akun yang terdaftar agar sistem dapat mengenali hak akses dan mengarahkan ke portal yang sesuai.",
                ],
                "figure": ("02-halaman-login-pemohon.png", "Gambar 4. Halaman login Website ULT"),
            },
            {
                "title": "3.4 Langkah Login ke Portal Pemohon",
                "body": [
                    "Proses login dirancang ringkas. Namun, pengguna tetap perlu memastikan bahwa akun yang dipakai memang memiliki akses sebagai pemohon layanan.",
                ],
                "steps": [
                    "buka halaman login Website ULT melalui menu masuk yang tersedia pada website;",
                    "isi email atau identitas akun dan kata sandi dengan benar;",
                    "tekan tombol masuk lalu tunggu pengalihan halaman;",
                    "pastikan sistem mengarahkan Anda ke dashboard pengajuan layanan, bukan ke portal role lain.",
                ],
            },
            {
                "title": "3.5 Ringkasan Dashboard",
                "body": [
                    "Setelah login, dashboard menjadi titik orientasi pertama. Pemohon dapat melihat apakah sudah ada pengajuan yang berjalan, status mana yang paling dominan, dan jalur tercepat untuk membuat layanan baru atau membuka permohonan yang sudah ada.",
                    "Jika belum ada pengajuan, dashboard dan halaman Permohonan Saya akan menampilkan keadaan kosong yang tetap memberi arah tindakan, yaitu mengajukan layanan baru melalui katalog layanan.",
                ],
            },
            {
                "title": "3.6 Pencarian dan Filter Permohonan",
                "body": [
                    "Pada saat jumlah pengajuan mulai bertambah, fitur pencarian realtime dan filter status atau layanan menjadi penting. Fitur ini membantu pemohon menemukan permohonan tertentu tanpa harus membaca seluruh daftar secara manual.",
                ],
                "figure": ("06-filter-dan-pencarian-permohonan.png", "Gambar 5. Pencarian dan filter pada daftar permohonan"),
                "bullets": [
                    "kolom pencarian dapat dipakai untuk mencari kode, nama layanan, status, atau waktu pengajuan;",
                    "filter status membantu menyaring permohonan yang masih perlu ditindaklanjuti;",
                    "filter layanan berguna saat pemohon pernah mengajukan beberapa jenis layanan yang berbeda.",
                ],
            },
        ],
    },
    {
        "title": "BAB IV MENELUSURI INFORMASI LAYANAN",
        "intro": [
            "Sebelum mengirim permohonan, pemohon perlu memastikan layanan yang dipilih sudah benar. Tahap ini berfokus pada pembacaan informasi publik yang menjadi dasar sebelum form pengajuan dibuka.",
        ],
        "sections": [
            {
                "title": "4.1 Membuka Detail Layanan",
                "body": [
                    "Setelah memilih layanan pada katalog, pengguna diarahkan ke halaman detail layanan. Halaman ini memperlihatkan judul, ringkasan, dokumen contoh atau preview, serta tombol ajukan layanan.",
                ],
                "figure": ("07-detail-layanan-publik.png", "Gambar 6. Halaman detail layanan"),
            },
            {
                "title": "4.2 Membaca Persyaratan dan SOP",
                "body": [
                    "Persyaratan dan SOP adalah dasar sebelum pengisian form dimulai. Pemohon perlu membaca kedua bagian ini agar data dan dokumen yang dikirim sesuai dengan kebutuhan layanan yang dipilih.",
                ],
                "figure": ("08-persyaratan-dan-sop-layanan.png", "Gambar 7. Bagian persyaratan dan SOP layanan"),
                "steps": [
                    "baca seluruh persyaratan untuk memastikan data dan file yang dibutuhkan sudah tersedia;",
                    "pelajari SOP agar Anda memahami alur status yang mungkin muncul setelah pengajuan dikirim;",
                    "lanjutkan ke tombol Ajukan Layanan hanya jika layanan yang dipilih memang sesuai kebutuhan Anda.",
                ],
            },
        ],
    },
    {
        "title": "BAB V MENGAJUKAN LAYANAN",
        "intro": [
            "Bab ini menyusun proses pengajuan layanan berdasarkan sumber dokumen awal yang dipakai oleh pemohon. Pengelompokan ini dibuat agar pembaca tidak perlu menelusuri terlalu banyak subbab kecil yang tercampur antara layanan biasa, layanan dengan upload awal, dan layanan kategori sertifikat atau piagam.",
        ],
        "sections": [
            {
                "title": "5.1 Jenis Pengajuan Layanan",
                "body": [
                    "Sebelum mengajukan layanan, pemohon perlu mengenali lebih dahulu jenis sumber dokumen awal yang dipakai oleh layanan tersebut. Tidak semua layanan memakai pola yang sama. Ada layanan yang cukup diawali dengan isian form biasa, ada layanan yang meminta upload file atau tanda tangan sejak awal, dan ada layanan sertifikat atau piagam yang memakai dokumen sumber .pptx dari pemohon.",
                    "Dengan memahami pengelompokan ini, pemohon dapat langsung menyesuaikan persiapan sebelum membuka form. Pendekatan ini juga membantu pembaca memilih subbab yang benar tanpa harus membaca seluruh variasi alur pengajuan dari awal hingga akhir.",
                    "Secara umum, jenis pengajuan layanan pada portal pemohon dapat dibedakan ke dalam beberapa kelompok berikut.",
                ],
                "bullets": [
                    "kelompok pertama: layanan dengan isian form biasa sebagai sumber awal pengajuan;",
                    "kelompok kedua: layanan yang meminta upload file, tanda tangan, atau signer tertentu sejak awal;",
                    "kelompok ketiga: layanan kategori sertifikat atau piagam dengan dokumen sumber .pptx dari pemohon.",
                ],
            },
            {
                "title": "5.2 Pengajuan Layanan Biasa",
                "body": [
                    "Pada kelompok ini, sumber awal pengajuan berasal dari field yang diisi langsung oleh pemohon pada form layanan. Pola ini paling umum ditemui pada layanan yang lebih menekankan isian data daripada dokumen sumber khusus. Karena seluruh isi form akan menjadi dasar verifikasi petugas, pemohon tetap perlu membaca struktur field secara teliti sebelum mulai mengisi.",
                    "Pada panduan ini, pembahasan kelompok layanan berbasis isian dibagi ke dalam beberapa bagian agar pembaca dapat mengikuti urutan kerja dari form awal sampai halaman detail setelah submit tanpa perlu berpindah ke subbab lain.",
                ],
                "figures": [
                    ("09-form-pengajuan-layanan-biasa.png", "Gambar 8. Form pengajuan layanan"),
                    ("15-detail-permohonan-setelah-submit.png", "Gambar 9. Detail permohonan setelah submit"),
                ],
                "steps": [
                    "baca kembali ringkasan layanan, persyaratan, dan SOP sebelum membuka form;",
                    "amati struktur field pada form lalu isi data dari bagian paling dasar ke bagian berikutnya secara berurutan;",
                    "setelah semua field terisi, lakukan pemeriksaan ulang terhadap nama, nomor, tanggal, dan uraian yang dimasukkan;",
                    "tekan submit lalu lanjutkan pemeriksaan pada halaman detail permohonan sebagai tanda bahwa pengajuan telah tersimpan.",
                ],
                "bullets": [
                    "jangan mengisi field hanya agar form cepat selesai; pastikan setiap jawaban sesuai dengan konteks layanan;",
                    "pastikan nama, nomor, tanggal, dan uraian penting sudah benar sebelum submit dijalankan;",
                    "bila layanan ini tidak meminta upload awal, lampiran tambahan tetap dapat ditambahkan dari halaman detail permohonan.",
                ],
            },
            {
                "title": "5.3 Pengajuan dengan Upload Awal",
                "body": [
                    "Pada kelompok ini, sumber awal pengajuan tidak hanya berasal dari isian field, tetapi juga dari file yang harus diunggah sejak tahap form dibuka. File tersebut bisa berupa dokumen persyaratan tertentu, tanda tangan pemohon, tanda tangan signer custom, atau pilihan signer yang harus ditentukan sebelum submit.",
                    "Karena ada komponen upload awal, pemohon perlu menempatkan ketelitian pada dua area sekaligus, yaitu isian data dan file yang dipilih. Kesalahan pada salah satu area ini dapat menyebabkan hasil verifikasi tertunda, signer tidak sesuai, atau dokumen yang diproses menjadi tidak lengkap.",
                    "Agar lebih mudah dipahami, pembahasan pada subbab ini dibagi menjadi bagian form utama dan bagian area unggah. Dengan demikian, pemohon dapat membedakan mana komponen isian dasar dan mana komponen file yang harus diperiksa sebelum submit.",
                ],
                "figures": [
                    ("10-form-pengajuan-dengan-signer.png", "Gambar 10. Bagian tanda tangan dan penandatangan tambahan pada form layanan"),
                ],
                "steps": [
                    "buka form layanan lalu identifikasi bagian mana yang meminta upload file, tanda tangan, atau signer tambahan;",
                    "isi field dasar terlebih dahulu, kemudian unggah file pada komponen yang sesuai dengan labelnya;",
                    "jika ada signer tambahan, tentukan data signer dengan benar sebelum submit;",
                    "periksa kembali bahwa file yang dipilih, tanda tangan yang diunggah, dan data signer sudah cocok dengan kebutuhan layanan.",
                ],
                "bullets": [
                    "untuk file tanda tangan, gunakan file yang jelas dan mudah terbaca saat ditampilkan pada dokumen;",
                    "untuk signer custom, periksa kembali identitas dan jabatan agar tidak salah tampil di dokumen;",
                    "meskipun ada upload awal, lampiran umum tambahan tetap dapat dikelola lagi dari halaman detail permohonan.",
                ],
            },
            {
                "title": "5.4 Pengajuan Sertifikat atau Piagam",
                "body": [
                    "Layanan kategori sertifikat atau piagam memiliki pola pengajuan yang berbeda dari layanan biasa. Pada kategori ini, pemohon tidak hanya mengisi field umum, tetapi juga mengunggah dokumen sumber berformat .pptx, menentukan daftar signer sejak awal, dan menyiapkan tanda tangan gambar untuk signer tertentu sesuai kebutuhan layanan.",
                    "Karena output akhirnya disusun dari dokumen sumber pemohon, ketelitian pada kategori layanan ini harus lebih tinggi. Kesalahan kecil pada template, token, signer, atau file tanda tangan dapat menyebabkan hasil dokumen bergeser, nama penandatangan tidak sesuai, atau preview dokumen sulit dibaca dengan baik.",
                    "Supaya alurnya lebih mudah diikuti, pembahasan layanan sertifikat atau piagam pada subbab ini dipecah menjadi beberapa bagian, mulai dari form utama, upload dokumen sumber dan signer, pedoman template, halaman detail setelah submit, sampai area preview dan output hasil layanan.",
                ],
                "figures": [
                    ("12-form-sertifikat-piagam.png", "Gambar 11. Form layanan sertifikat atau piagam"),
                    ("13-upload-pptx-dan-signer-sertifikat.png", "Gambar 12. Upload file .pptx dan pengaturan signer sertifikat"),
                    ("14-panel-pedoman-sertifikat-piagam.png", "Gambar 13. Panel pedoman dan contoh layout sertifikat atau piagam"),
                ],
                "steps": [
                    "siapkan file .pptx final sebelum membuka form pengajuan;",
                    "tentukan urutan signer sesuai kebutuhan dokumen sertifikat atau piagam dan isi data signer dengan tepat;",
                    "unggah file tanda tangan untuk signer yang memerlukannya dengan format yang sesuai;",
                    "gunakan preview dokumen setelah submit untuk memeriksa nama signer, posisi tanda tangan, isi utama dokumen, dan keterbacaan tata letak;",
                    "jika ada masalah, catat bagian yang tidak sesuai agar lebih mudah saat melakukan perbaikan atau saat petugas meminta revisi.",
                ],
                "bullets": [
                    "pastikan file .pptx menjadi sumber utama hasil dokumen, bukan lampiran tambahan biasa;",
                    "jika ada lebih dari satu versi file, beri nama yang jelas agar tidak salah pilih saat upload;",
                    "jangan menukar urutan signer tanpa alasan yang jelas karena hasil dokumen mengikuti urutan tersebut;",
                    "lebih aman memakai file tanda tangan dengan latar yang bersih dan mudah terbaca;",
                    "jangan lewatkan tahap preview setelah submit karena di situlah masalah layout biasanya mulai terlihat;",
                    "hindari memakai file draft, signer yang salah, atau tanda tangan yang buram karena kesalahan itu paling sering memicu perbaikan.",
                ],
            },
        ],
    },
    {
        "title": "BAB VI MEMANTAU DAN MENINDAKLANJUTI PERMOHONAN",
        "intro": [
            "Setelah permohonan tersimpan, perhatian pemohon berpindah ke tahap pemantauan. Pada implementasi saat ini, halaman detail permohonan memuat data pengajuan, lampiran, catatan, riwayat status, tombol aksi, preview dokumen, serta akses unduh output bila hasil layanan sudah tersedia.",
        ],
        "sections": [
            {
                "title": "6.1 Detail Permohonan",
                "body": [
                    "Halaman detail permohonan adalah pusat pemantauan utama. Pemohon dapat melihat status terkini, tanggal pengajuan, jumlah riwayat status, jumlah lampiran, nomor dokumen, tombol preview, dan tombol unduh output bila hasil layanan sudah tersedia.",
                ],
                "figure": ("15-detail-permohonan-setelah-submit.png", "Gambar 14. Halaman detail permohonan"),
            },
            {
                "title": "6.2 Status Permohonan",
                "body": [
                    "Status menunjukkan posisi permohonan dalam alur layanan. Badge status membantu pemohon memahami apakah pengajuan sedang diverifikasi, menunggu tindakan lanjutan, perlu diperbaiki, sedang dalam proses penandatanganan, atau telah selesai.",
                ],
                "figure": ("16-status-permohonan.png", "Gambar 15. Status permohonan"),
                "bullets": [
                    "status tampil pada kartu ringkasan dan daftar permohonan;",
                    "status perlu dibaca bersama catatan dan riwayat agar konteks perubahan lebih jelas;",
                    "saat status menjadi PERLU_PERBAIKAN, sistem membuka akses untuk memperbarui data dan mengirim revisi.",
                ],
            },
            {
                "title": "6.3 Mengelola Lampiran Tambahan",
                "body": [
                    "Pada halaman detail, tersedia section Lampiran untuk mengunggah file pendukung dan mengunduh lampiran yang sudah tersimpan. Fitur ini penting karena tidak semua lampiran harus diunggah sejak awal submit.",
                ],
                "steps": [
                    "buka halaman detail permohonan yang relevan;",
                    "unggah file pada form Lampiran bila ada dokumen pendukung tambahan yang perlu dilampirkan;",
                    "gunakan tombol download pada daftar lampiran untuk memeriksa file yang sudah tersimpan di sistem.",
                ],
            },
            {
                "title": "6.4 Catatan Permohonan",
                "body": [
                    "Section Catatan menampilkan percakapan yang terlihat oleh pemohon dan petugas, kecuali catatan internal. Pemohon dapat menulis catatan opsional untuk menyampaikan klarifikasi atau informasi tambahan terkait permohonan yang sedang diproses.",
                ],
                "bullets": [
                    "bacalah catatan petugas sebelum mengambil tindakan lain;",
                    "gunakan catatan untuk komunikasi yang tetap terdokumentasi di dalam permohonan;",
                    "hindari menulis pesan yang tidak berkaitan dengan proses layanan agar riwayat tetap rapi dan relevan.",
                ],
            },
            {
                "title": "6.5 Perbaikan Permohonan",
                "body": [
                    "Jika status berubah menjadi PERLU_PERBAIKAN, halaman detail akan menampilkan form revisi data dan tombol Kirim perbaikan. Pada fase ini, pemohon dapat memperbarui field pengajuan, mengganti tanda tangan tertentu, mengubah signer dosen bila layanan memerlukannya, lalu mengirim perbaikan kembali ke sistem.",
                ],
                "figures": [
                    ("17-form-perbaikan-permohonan.png", "Gambar 16. Form perbaikan permohonan"),
                    ("18-revisi-signer-dan-tandatangan.png", "Gambar 17. Perubahan signer dan tanda tangan saat perbaikan"),
                ],
                "steps": [
                    "baca catatan perbaikan dari petugas dan identifikasi bagian yang harus diubah;",
                    "perbarui data pada section Data permohonan atau ganti file yang memang diminta untuk direvisi;",
                    "simpan perubahan data terlebih dahulu bila diperlukan;",
                    "kirim revisi melalui form Kirim perbaikan agar petugas dapat memproses ulang permohonan.",
                ],
            },
            {
                "title": "6.6 Riwayat Status",
                "body": [
                    "Riwayat status memperlihatkan perpindahan status dari waktu ke waktu beserta aktor dan catatan yang terkait. Informasi ini membantu pemohon memahami perjalanan proses secara utuh, bukan hanya melihat status terkini.",
                ],
                "figure": ("20-riwayat-status-permohonan.png", "Gambar 18. Riwayat status permohonan"),
            },
        ],
    },
    {
        "title": "BAB VII MENGAKSES HASIL LAYANAN",
        "intro": [
            "Tahap akhir dari sudut pandang pemohon adalah ketika dokumen hasil sudah tersedia atau layanan dinyatakan selesai. Pada sistem saat ini, akses hasil layanan terintegrasi langsung di halaman detail permohonan.",
        ],
        "sections": [
            {
                "title": "7.1 Preview Dokumen",
                "body": [
                    "Untuk layanan yang mendukung preview dokumen, halaman detail menyediakan tombol Buka preview. Fitur ini berguna untuk meninjau tampilan hasil dokumen tanpa harus langsung mengunduh file output, terutama pada layanan yang memiliki template dokumen dinamis.",
                ],
            },
            {
                "title": "7.2 Unduh Hasil Layanan",
                "body": [
                    "Jika output sudah tersedia, tombol Unduh berkas akan tampil pada area ringkasan permohonan. Melalui tombol ini, pemohon dapat memperoleh hasil layanan secara langsung dari sistem tanpa meminta file secara manual kepada petugas.",
                ],
                "figure": ("20-output-layanan-unduh-berkas.png", "Gambar 19. Area output layanan dan tombol unduh berkas"),
                "steps": [
                    "pastikan status layanan sudah mendukung akses hasil atau output telah muncul pada halaman detail;",
                    "gunakan tombol Unduh berkas untuk mengunduh dokumen hasil;",
                    "bila tersedia preview dokumen, buka terlebih dahulu untuk memeriksa kesesuaian isi sebelum menyimpan file.",
                ],
            },
        ],
    },
    {
        "title": "BAB VIII KETENTUAN DAN PRAKTIK BAIK",
        "intro": [
            "Selain memahami tombol dan alur layar, pemohon perlu mengikuti beberapa ketentuan umum agar penggunaan sistem tetap tertib, aman, dan tidak menimbulkan hambatan verifikasi di tahap berikutnya.",
        ],
        "sections": [
            {
                "title": "8.1 Penggunaan Akun",
                "body": [
                    "Akun portal pemohon harus dijaga kerahasiaannya dan digunakan sesuai kewenangan. Hindari berbagi kredensial dengan pihak lain karena semua aktivitas pada permohonan tercatat berdasarkan akun yang digunakan.",
                ],
            },
            {
                "title": "8.2 Ketelitian Pengisian Data",
                "body": [
                    "Seluruh data pada form harus sesuai dengan dokumen resmi dan konteks layanan. Kesalahan kecil pada nama, nomor identitas, pilihan signer, tanggal, atau file yang diunggah dapat menyebabkan proses verifikasi tertunda atau permohonan dikembalikan untuk perbaikan.",
                ],
            },
            {
                "title": "8.3 Kualitas File dan Lampiran",
                "body": [
                    "File yang diunggah harus terbaca dengan jelas dan sesuai ketentuan format atau ukuran dari sistem. Untuk file tanda tangan, dokumen sumber .pptx, atau lampiran lain yang digunakan dalam perakitan dokumen, kualitas file berpengaruh langsung terhadap hasil akhir layanan.",
                ],
                "bullets": [
                    "gunakan file yang benar-benar final sebelum submit;",
                    "pastikan nama dan isi file tidak tertukar antar layanan atau antar tahap revisi;",
                    "periksa kembali hasil preview atau output bila layanan menyediakan fitur tersebut.",
                ],
            },
            {
                "title": "8.4 Disiplin Memantau Status",
                "body": [
                    "Pemohon tidak cukup hanya mengirim permohonan lalu menunggu tanpa pemantauan. Status, catatan, dan riwayat perlu diperiksa secara berkala agar tindakan perbaikan atau pengunduhan hasil dapat dilakukan tepat waktu.",
                ],
            },
        ],
    },
    {
        "title": "BAB IX PENUTUP",
        "intro": [
            "Buku panduan pemohon layanan ini disusun untuk menyesuaikan dokumentasi dengan antarmuka dan alur Website ULT versi saat ini. Melalui panduan yang lebih faktual, pemohon diharapkan dapat menggunakan portal secara lebih terarah sejak tahap login, pengajuan, pemantauan, revisi, hingga pengambilan hasil layanan.",
            "Pembaruan dokumen ini juga memperkuat fungsi panduan sebagai referensi operasional dan pelengkap dokumentasi produk. Dengan demikian, proses pendampingan pengguna baru maupun evaluasi penggunaan sistem dapat dilakukan dengan acuan yang lebih konsisten.",
        ],
        "sections": [],
    },
]


def set_run_font(run, size=10, bold=False, italic=False, color=COLOR_TEXT):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_doc_defaults(doc):
    for section in doc.sections:
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.top_margin = Cm(TOP_MARGIN_CM)
        section.bottom_margin = Cm(BOTTOM_MARGIN_CM)
        section.left_margin = Cm(LEFT_MARGIN_CM)
        section.right_margin = Cm(RIGHT_MARGIN_CM)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(4)

    for name, size in [("Heading 1", 12), ("Heading 2", 11)]:
        style = doc.styles[name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = COLOR_PRIMARY
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=12 if level == 1 else 11, bold=True, color=COLOR_PRIMARY)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = LIST_TEXT_INDENT
        p.paragraph_format.first_line_indent = LIST_HANGING_INDENT
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.tab_stops.add_tab_stop(LIST_TEXT_INDENT)
        bullet = p.add_run("•\t")
        set_run_font(bullet, size=10, bold=True, color=COLOR_PRIMARY)
        body = p.add_run(item)
        set_run_font(body, size=10)


def build_section_steps_intro(title):
    lowered = title.lower()
    if "mengunggah" in lowered or "upload" in lowered or "dokumen" in lowered:
        return "Agar proses pada subbab ini berjalan tertib, langkah-langkah yang perlu dilakukan adalah sebagai berikut."
    if "meninjau" in lowered or "submit" in lowered or "mengirim" in lowered:
        return "Sebelum berpindah ke tahap berikutnya, urutan kerja yang disarankan pada bagian ini adalah sebagai berikut."
    return "Untuk memudahkan penerapan saat praktik penggunaan, langkah-langkah pada subbab ini dapat diikuti sebagai berikut."


def build_section_bullets_intro(title):
    lowered = title.lower()
    if "hasil" in lowered or "status" in lowered:
        return "Selain langkah penggunaan di atas, beberapa hal berikut penting diperhatikan pada bagian ini."
    if "mengunggah" in lowered or "dokumen" in lowered or "file" in lowered:
        return "Di samping langkah operasional tersebut, beberapa catatan berikut perlu menjadi perhatian."
    return "Sebagai pelengkap dari uraian sebelumnya, beberapa poin penting berikut perlu diperhatikan."


def add_steps_intro(doc, title):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(build_section_steps_intro(title))
    set_run_font(run, size=10)


def add_bullets_intro(doc, title):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(build_section_bullets_intro(title))
    set_run_font(run, size=10)


def add_steps(doc, items, variant="section"):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if variant == "figure_sub":
            p.paragraph_format.left_indent = Cm(1.55)
            p.paragraph_format.first_line_indent = LIST_HANGING_INDENT
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.tab_stops.add_tab_stop(p.paragraph_format.left_indent)
            number = p.add_run(f"{index}.\t")
        elif variant == "figure":
            p.paragraph_format.left_indent = LIST_TEXT_INDENT
            p.paragraph_format.first_line_indent = LIST_HANGING_INDENT
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.tab_stops.add_tab_stop(p.paragraph_format.left_indent)
            number = p.add_run(f"{index}.\t")
        else:
            p.paragraph_format.left_indent = LIST_TEXT_INDENT
            p.paragraph_format.first_line_indent = LIST_HANGING_INDENT
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.tab_stops.add_tab_stop(p.paragraph_format.left_indent)
            number = p.add_run(f"{index}.\t")
        set_run_font(number, size=10, bold=True, color=COLOR_PRIMARY)
        body = p.add_run(item)
        set_run_font(body, size=10)


def add_field_note(doc, title, instruction):
    add_heading(doc, title, 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(instruction)
    set_run_font(run, size=9)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(85)
    run = p.add_run("BUKU PANDUAN PEMOHON LAYANAN")
    set_run_font(run, size=16, bold=True, color=COLOR_PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "WEBSITE UNIT LAYANAN TERPADU (ULT)\n"
        "FAKULTAS KEGURUAN DAN ILMU PENDIDIKAN\n"
        "UNIVERSITAS LAMPUNG"
    )
    set_run_font(run, size=13, bold=True, color=COLOR_PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    run = p.add_run("Panduan Operasional Penggunaan Portal Pemohon")
    set_run_font(run, size=11, bold=True, color=COLOR_PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Dokumen ini menjelaskan alur penggunaan portal pemohon mulai dari login,\n"
        "pengajuan layanan, pemantauan status, revisi, hingga pengunduhan hasil layanan."
    )
    set_run_font(run, size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(150)
    run = p.add_run("Bandar Lampung\n2026")
    set_run_font(run, size=10)


def add_front_matter(doc):
    add_heading(doc, "KATA PENGANTAR", 1)
    add_body(
        doc,
        "Buku panduan ini disusun sebagai acuan operasional penggunaan Website Unit Layanan Terpadu (ULT) FKIP Universitas Lampung dari sudut pandang pemohon layanan. Pembaruan dokumen dilakukan agar isi panduan selaras dengan antarmuka dan alur sistem yang telah berkembang, khususnya pada area dashboard pemohon, daftar permohonan, halaman detail, catatan, revisi, preview dokumen, dan akses hasil layanan.",
    )
    add_body(
        doc,
        "Melalui buku ini, pengguna diharapkan dapat memahami langkah penggunaan sistem secara lebih terarah sejak tahap awal sampai akhir. Tangkapan layar yang digunakan pada setiap bab dipilih dari implementasi aktual agar uraian yang disajikan lebih mudah diikuti ketika praktik penggunaan dilakukan.",
    )
    add_body(
        doc,
        "Dokumen ini juga berfungsi sebagai pelengkap dokumentasi produk hasil pengembangan, sehingga dapat dipakai sebagai referensi saat orientasi pengguna baru, peninjauan tugas akhir, maupun evaluasi pemanfaatan sistem dalam konteks layanan nyata.",
    )

    doc.add_page_break()
    add_field_note(
        doc,
        "DAFTAR ISI",
        'Daftar isi dapat diperbarui otomatis melalui fitur "Update Field" pada Microsoft Word setelah peninjauan akhir dokumen.',
    )

    doc.add_page_break()
    add_field_note(
        doc,
        "DAFTAR GAMBAR",
        'Daftar gambar dapat diperbarui otomatis melalui fitur "Update Field" pada Microsoft Word setelah caption final diperiksa.',
    )


def resolve_image_dimensions(path: Path, subfigure=False):
    with Image.open(path) as image:
        width_px, height_px = image.size

    aspect_ratio = width_px / max(height_px, 1)
    content_width = PAGE_WIDTH_CM - LEFT_MARGIN_CM - RIGHT_MARGIN_CM
    content_height = PAGE_HEIGHT_CM - TOP_MARGIN_CM - BOTTOM_MARGIN_CM
    max_width = min(12.2 if subfigure else 13.4, content_width - (1.8 if subfigure else 0.2))
    max_height = min(12.4 if subfigure else 15.2, content_height - (1.8 if subfigure else 1.2))

    if aspect_ratio >= 4.5:
        target_width = 10.4
    elif aspect_ratio >= 2.8:
        target_width = 11.2
    elif aspect_ratio >= 1.6:
        target_width = 12.4
    elif aspect_ratio >= 1.0:
        target_width = 11.6
    elif aspect_ratio >= 0.72:
        target_width = 9.8
    else:
        target_width = 8.4

    if subfigure:
        target_width = min(target_width, 11.2)

    target_width = min(target_width, max_width)
    target_height = target_width / max(aspect_ratio, 0.01)

    if target_height > max_height:
        scale = max_height / target_height
        target_width *= scale
        target_height = max_height

    return Cm(round(target_width, 2)), Cm(round(target_height, 2))


def build_figure_explanation(caption):
    base = caption.split(". ", 1)[1] if ". " in caption else caption
    desc = base[:1].lower() + base[1:] if base else "tampilan terkait"
    lowered = desc.lower()
    subject = desc[:1].upper() + desc[1:] if desc else "Tampilan terkait"

    if lowered.startswith("halaman detail layanan"):
        return f"{subject} menampilkan ringkasan informasi layanan sebelum pemohon membuka form pengajuan. Pada bagian ini, pengguna dapat memastikan judul layanan, persyaratan awal, dan konteks layanan yang dipilih sudah sesuai."
    if lowered.startswith("halaman detail"):
        return f"{subject} menampilkan informasi yang lebih lengkap daripada daftar ringkas. Pada bagian ini, pemohon dapat membaca konteks permohonan secara menyeluruh sebelum mengambil tindakan lanjutan."
    if lowered.startswith("halaman login"):
        return f"{subject} menjadi gerbang autentikasi awal sebelum pengguna masuk ke portal sesuai role. Karena itu, elemen yang tampil pada halaman ini perlu dipahami sebagai titik awal akses sistem."
    if lowered.startswith("dashboard"):
        return f"{subject} berfungsi sebagai titik orientasi awal setelah login. Melalui tampilan ini, pemohon dapat membaca ringkasan cepat sebelum berpindah ke daftar permohonan atau mengajukan layanan baru."
    if lowered.startswith("form"):
        return f"{subject} menunjukkan area pengisian data yang menjadi inti proses pengajuan. Setiap bagian yang tampil pada form perlu dibaca dengan cermat agar data yang dikirim tetap sesuai dengan kebutuhan layanan."
    if lowered.startswith("bagian form untuk menentukan signer"):
        return f"{subject} memperlihatkan komponen yang dipakai pemohon untuk menyusun urutan penandatangan. Pada bagian ini, ketepatan memilih tipe signer, identitas, dan urutan penandatangan akan memengaruhi hasil dokumen akhir."
    if lowered.startswith("bagian upload dokumen sumber .pptx dan penyusunan signer"):
        return f"{subject} memperlihatkan dua area paling penting pada layanan sertifikat atau piagam, yaitu pemilihan file sumber .pptx dan penyusunan signer. Pada tahap ini, pemohon harus memastikan file sumber yang dipilih benar serta urutan signer telah disusun sesuai kebutuhan dokumen."
    if lowered.startswith("area upload dokumen sumber .pptx dan penyusunan signer"):
        return f"{subject} memperlihatkan dua area paling penting pada layanan sertifikat atau piagam, yaitu pemilihan file sumber .pptx dan penyusunan signer. Pada tahap ini, pemohon harus memastikan file sumber yang dipilih benar serta urutan signer telah disusun sesuai kebutuhan dokumen."
    if lowered.startswith("upload file .pptx dan signer"):
        return f"{subject} memperlihatkan dua komponen penting pada layanan sertifikat atau piagam, yaitu upload file sumber .pptx dan penyusunan signer. Pada tahap ini, pemohon harus memastikan file yang dipilih benar serta urutan signer telah disusun sesuai kebutuhan dokumen."
    if lowered.startswith("area unggah"):
        return f"{subject} memperlihatkan komponen unggah file yang berperan langsung terhadap kelengkapan permohonan. Bagian ini perlu diperhatikan secara khusus agar file yang dikirim sesuai dengan format dan tujuan penggunaannya."
    if lowered.startswith("area upload dokumen sumber .pptx"):
        return f"{subject} menunjukkan titik paling penting pada layanan sertifikat atau piagam karena file inilah yang akan dijadikan sumber perakitan dokumen. Pemohon perlu memastikan file yang dipilih benar-benar final sebelum melanjutkan proses."
    if lowered.startswith("komponen upload tanda tangan"):
        return f"{subject} menampilkan area unggah tanda tangan untuk signer yang memerlukannya. Kualitas file pada bagian ini akan berpengaruh langsung terhadap keterbacaan tanda tangan pada hasil dokumen."
    if lowered.startswith("bagian status"):
        return f"{subject} membantu pemohon membaca posisi permohonan pada alur layanan yang sedang berjalan. Melalui bagian ini, pengguna dapat mengenali apakah permohonan cukup dipantau, perlu diperbaiki, atau sudah selesai."
    if lowered.startswith("riwayat"):
        return f"{subject} menyajikan jejak proses secara kronologis. Dengan membaca bagian ini, pemohon dapat memahami perpindahan status dan tindakan yang pernah terjadi pada permohonannya."
    if lowered.startswith("contoh layout sertifikat") or lowered.startswith("contoh tampilan sertifikat"):
        return f"{subject} berfungsi sebagai acuan visual terhadap susunan elemen pada dokumen sertifikat atau piagam. Gambar ini membantu pemohon memeriksa apakah isi, komposisi, dan area tanda tangan pada file sumber sudah masuk akal sebelum diunggah."
    if lowered.startswith("panel pedoman sertifikat"):
        return f"{subject} memperlihatkan bahwa layanan sertifikat atau piagam dilengkapi petunjuk teknis langsung di dalam form pengajuan. Dari panel ini, pemohon dapat membaca aturan dokumen sumber, token penting, dan contoh layout sebelum submit dilakukan."
    if lowered.startswith("panel pedoman sertifikat atau piagam"):
        return f"{subject} memperlihatkan petunjuk teknis yang tersedia langsung di dalam form pengajuan. Dari panel ini, pemohon dapat membaca aturan file sumber, token penting, dan contoh layout sebelum submit dilakukan."
    if lowered.startswith("hasil layanan"):
        return f"{subject} menunjukkan bahwa output telah tersedia dan dapat diakses oleh pemohon. Bagian ini menandai tahap akhir penggunaan sistem dari sisi pengguna pada satu permohonan tertentu."
    if lowered.startswith("output layanan"):
        return f"{subject} menunjukkan bahwa file hasil sudah tersedia dan dapat diakses oleh pemohon. Bagian ini menandai tahap akhir penggunaan sistem dari sisi pengguna pada satu permohonan tertentu."
    if lowered.startswith("area akses preview") or lowered.startswith("area akses preview atau output"):
        return f"{subject} memperlihatkan titik akses yang dipakai pemohon untuk memeriksa tampilan hasil dokumen atau mengunduh dokumen akhir. Bagian ini penting sebagai tahap validasi akhir setelah proses sertifikat atau piagam diproses sistem."
    if lowered.startswith("area preview dan output"):
        return f"{subject} memperlihatkan titik akses yang dipakai pemohon untuk membuka preview dokumen dan mengambil file hasil layanan. Bagian ini penting sebagai tahap validasi akhir setelah proses sertifikat atau piagam dinyatakan selesai."
    if lowered.startswith("contoh hasil preview dokumen sertifikat"):
        return f"{subject} menunjukkan tampilan awal hasil dokumen setelah file sumber diproses oleh sistem. Dari gambar ini, pemohon dapat menilai apakah isi utama, komposisi layout, dan penempatan elemen penting sudah terbaca dengan baik."
    if lowered.startswith("contoh hasil akhir sertifikat"):
        return f"{subject} memperlihatkan bentuk hasil akhir yang sudah lebih dekat dengan dokumen keluaran layanan. Gambar ini berguna sebagai acuan akhir untuk memeriksa apakah hasil sertifikat atau piagam sudah layak sebelum diunduh atau ditindaklanjuti."
    return f"{subject} disajikan sebagai acuan visual untuk membantu pembaca mengenali susunan antarmuka, letak informasi utama, dan hubungan gambar dengan langkah penggunaan pada subbab ini."


def build_figure_points(caption):
    base = caption.split(". ", 1)[1] if ". " in caption else caption
    desc = base[:1].lower() + base[1:] if base else "tampilan terkait"
    lowered = desc.lower()

    if lowered.startswith("dashboard"):
        return [
            f"Perhatikan ringkasan utama yang tampil pada {desc} untuk memahami posisi awal penggunaan setelah login.",
            "Gunakan tombol aksi dan kartu ringkasan pada halaman ini sebagai jalur tercepat menuju pengajuan baru atau pemantauan permohonan yang sudah ada.",
        ]
    if lowered.startswith("halaman permohonan"):
        return [
            f"Amati susunan kartu dan badge status pada {desc} agar lebih mudah mencari permohonan yang ingin dibuka.",
            "Gunakan tombol detail pada setiap entri sebagai pintu masuk ke pusat pemantauan permohonan.",
        ]
    if lowered.startswith("halaman login"):
        return [
            f"Fokuskan perhatian pada kolom kredensial dan tombol masuk yang tampil pada {desc}.",
            "Pastikan informasi yang dimasukkan benar agar sistem mengarahkan Anda ke portal pemohon tanpa kendala autentikasi.",
        ]
    if lowered.startswith("form") or lowered.startswith("struktur umum form"):
        return [
            f"Gunakan {desc} untuk mengenali susunan field sebelum mulai mengisi data.",
            "Baca urutan field dari atas ke bawah agar tidak ada bagian yang terlewat saat pengajuan dilakukan.",
        ]
    if lowered.startswith("bagian form untuk menentukan signer"):
        return [
            f"Perhatikan susunan komponen pada {desc} untuk memahami bagaimana signer ditentukan oleh pemohon.",
            "Pastikan urutan signer, tipe signer, dan data identitas yang diisi sudah sesuai dengan kebutuhan dokumen.",
        ]
    if lowered.startswith("bagian upload dokumen sumber .pptx dan penyusunan signer"):
        return [
            f"Gunakan {desc} untuk memeriksa bahwa file sumber dan daftar signer diisi pada area yang benar.",
            "Sebelum submit, pastikan nama file .pptx dan susunan signer yang tampil sudah sesuai dengan rancangan dokumen.",
        ]
    if lowered.startswith("area upload dokumen sumber .pptx dan penyusunan signer"):
        return [
            f"Gunakan {desc} untuk memeriksa bahwa file sumber dan daftar signer diisi pada area yang benar.",
            "Sebelum submit, pastikan nama file .pptx dan susunan signer yang tampil sudah sesuai dengan rancangan dokumen.",
        ]
    if lowered.startswith("upload file .pptx dan signer"):
        return [
            f"Gunakan {desc} untuk memastikan file sumber dan daftar signer diisi pada area yang benar.",
            "Sebelum submit, pastikan nama file .pptx dan susunan signer yang tampil sudah sesuai dengan rancangan dokumen.",
        ]
    if lowered.startswith("area unggah"):
        return [
            f"Perhatikan komponen yang tampil pada {desc} untuk memastikan jenis file yang diminta sudah sesuai.",
            "Sebelum melanjutkan, cek kembali file yang dipilih agar tidak tertukar dengan dokumen lain.",
        ]
    if lowered.startswith("area upload dokumen sumber .pptx"):
        return [
            f"Gunakan {desc} untuk memastikan bahwa file .pptx diunggah pada komponen yang benar.",
            "Sebelum submit, periksa lagi nama file dan pastikan file yang terpilih adalah versi final dokumen sumber.",
        ]
    if lowered.startswith("komponen upload tanda tangan"):
        return [
            f"Gunakan {desc} untuk memeriksa bahwa file tanda tangan diunggah pada signer yang tepat.",
            "Perhatikan kembali file yang terpilih agar tanda tangan tidak tertukar antar signer.",
        ]
    if lowered.startswith("halaman detail layanan"):
        return [
            f"Gunakan {desc} untuk memastikan layanan yang dibuka memang sesuai dengan kebutuhan pengajuan Anda.",
            "Perhatikan judul layanan, ringkasan, panel persyaratan, dan area SOP sebelum melanjutkan ke tombol ajukan layanan.",
        ]
    if lowered.startswith("halaman detail"):
        return [
            f"Gunakan {desc} sebagai titik review utama setelah submit berhasil.",
            "Perhatikan status, kode permohonan, tombol aksi, dan informasi ringkasan yang akan dipakai pada tahap pemantauan berikutnya.",
        ]
    if lowered.startswith("bagian status"):
        return [
            f"Baca {desc} bersama catatan dan riwayat agar posisi permohonan lebih mudah dipahami.",
            "Gunakan informasi status untuk menentukan apakah Anda cukup menunggu, perlu mengunggah lampiran, atau harus mengirim perbaikan.",
        ]
    if lowered.startswith("riwayat"):
        return [
            f"Perhatikan urutan waktu dan aktor pada {desc} untuk memahami perjalanan proses.",
            "Gunakan bagian ini saat ingin menelusuri kapan perubahan status tertentu pernah terjadi.",
        ]
    if lowered.startswith("hasil layanan"):
        return [
            f"Gunakan {desc} untuk memastikan bahwa dokumen hasil benar-benar sudah tersedia.",
            "Jika sistem menyediakan preview dan download, manfaatkan keduanya untuk memeriksa hasil sebelum menyimpan file.",
        ]
    if lowered.startswith("output layanan"):
        return [
            f"Gunakan {desc} untuk memastikan bahwa file hasil benar-benar sudah tersedia.",
            "Jika sistem menyediakan preview dan download, manfaatkan keduanya untuk memeriksa hasil sebelum menyimpan file.",
        ]
    if lowered.startswith("contoh layout sertifikat") or lowered.startswith("contoh tampilan sertifikat"):
        return [
            f"Amati {desc} untuk menilai apakah komposisi judul, isi utama, dan area penandatanganan pada dokumen sudah proporsional.",
            "Gunakan gambar contoh ini sebagai pembanding sebelum file .pptx diunggah ke sistem.",
        ]
    if lowered.startswith("panel pedoman sertifikat"):
        return [
            f"Baca {desc} untuk memastikan aturan file sumber, signer, dan token sudah dipahami sebelum submit.",
            "Gunakan contoh layout pada panel ini sebagai acuan cepat saat memeriksa kesiapan dokumen .pptx Anda.",
        ]
    if lowered.startswith("panel pedoman sertifikat atau piagam"):
        return [
            f"Baca {desc} untuk memastikan aturan file sumber, signer, dan token sudah dipahami sebelum submit.",
            "Gunakan petunjuk pada panel ini sebagai acuan cepat saat memeriksa kesiapan dokumen .pptx Anda.",
        ]
    if lowered.startswith("area akses preview") or lowered.startswith("area akses preview atau output"):
        return [
            f"Gunakan {desc} untuk memeriksa apakah tampilan hasil dokumen atau file output sudah bisa diakses.",
            "Setelah membuka preview atau output, telusuri kembali isi dokumen untuk memastikan hasil akhir sudah sesuai.",
        ]
    if lowered.startswith("area preview dan output"):
        return [
            f"Gunakan {desc} untuk memastikan tombol preview dokumen dan download output sudah tersedia bagi pemohon.",
            "Sebelum menyimpan hasil, buka preview terlebih dahulu agar isi dan tata letak akhir dapat diperiksa kembali.",
        ]
    if lowered.startswith("preview dan output sertifikat atau piagam"):
        return [
            f"Gunakan {desc} untuk memastikan preview dokumen dan file output sudah tersedia bagi pemohon.",
            "Sebelum menyimpan hasil, buka preview terlebih dahulu agar isi dan tata letak akhir dapat diperiksa kembali.",
        ]
    if lowered.startswith("contoh hasil preview dokumen sertifikat"):
        return [
            f"Amati {desc} untuk menilai apakah hasil awal dokumen sudah terbaca dengan baik setelah sistem memproses file sumber.",
            "Fokuskan perhatian pada isi utama, posisi elemen penting, dan keterbacaan layout sebelum melanjutkan ke tahap berikutnya.",
        ]
    if lowered.startswith("contoh hasil akhir sertifikat"):
        return [
            f"Gunakan {desc} untuk menilai bentuk dokumen keluaran yang paling mendekati hasil final.",
            "Bandingkan hasil dokumen ini dengan rancangan awal agar lebih mudah menemukan pergeseran atau isi yang tidak sesuai.",
        ]
    return [
        f"Gunakan {desc} sebagai acuan visual untuk memahami tampilan yang sedang dibahas pada subbab ini.",
        "Perhatikan letak informasi utama dan tombol aksi yang tersedia sebelum melanjutkan ke tahap berikutnya.",
    ]


def add_figure_explanation(doc, caption, subfigure=False):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if subfigure:
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(0)
    else:
        p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(build_figure_explanation(caption))
    set_run_font(run, size=10)
    add_steps(doc, build_figure_points(caption), variant="figure_sub" if subfigure else "figure")


def build_figure_subheading(caption):
    return caption.split(". ", 1)[1] if ". " in caption else caption


def add_figure_subheading(doc, caption, index):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(5 if index > 1 else 3)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"Bagian {index}. {build_figure_subheading(caption)}")
    set_run_font(run, size=10, bold=True, color=COLOR_PRIMARY)


def add_figure(doc, filename, caption, subfigure=False):
    candidates = [
        ASSETS / filename,
        ROOT / "public/example" / filename,
    ]
    path = next((candidate for candidate in candidates if candidate.exists()), None)
    if path is None:
        add_body(doc, f"[Gambar tidak ditemukan: {filename}]")
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subfigure:
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.right_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    width, height = resolve_image_dimensions(path, subfigure=subfigure)
    p.add_run().add_picture(str(path), width=width, height=height)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subfigure:
        cap.paragraph_format.left_indent = Cm(0.8)
        cap.paragraph_format.right_indent = Cm(0.8)
    cap.paragraph_format.space_after = Pt(3)
    run = cap.add_run(caption)
    set_run_font(run, size=8, italic=True)
    add_figure_explanation(doc, caption, subfigure=subfigure)


def add_figures(doc, figures):
    multiple = len(figures) > 1
    for index, (filename, caption) in enumerate(figures, start=1):
        if multiple:
            add_figure_subheading(doc, caption, index)
        add_figure(doc, filename, caption, subfigure=multiple)


def add_chapters(doc):
    for chapter in CHAPTERS:
        doc.add_page_break()
        add_heading(doc, chapter["title"], 1)
        for paragraph in chapter.get("intro", []):
            add_body(doc, paragraph)

        for section in chapter.get("sections", []):
            add_heading(doc, section["title"], 2)
            for paragraph in section.get("body", []):
                add_body(doc, paragraph)
            if section.get("figure"):
                add_figure(doc, section["figure"][0], section["figure"][1])
            if section.get("figures"):
                add_figures(doc, section["figures"])
            if section.get("steps"):
                add_steps_intro(doc, section["title"])
                add_steps(doc, section["steps"])
            if section.get("bullets"):
                if section.get("steps") or section.get("figure") or section.get("figures"):
                    add_bullets_intro(doc, section["title"])
                add_bullets(doc, section["bullets"])


def main():
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc)
    doc.add_page_break()
    add_front_matter(doc)
    add_chapters(doc)
    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    main()
