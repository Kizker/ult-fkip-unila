from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image


ROOT = Path(r"c:/laragon/www/ult-fkip-unila")
OUTPUT = ROOT / "docs/buku-panduan/buku-panduan-penandatangan/buku-panduan-penandatangan.docx"
ASSETS = ROOT / "docs/buku-panduan/assets/screenshots"
FONT_NAME = "Poppins"
COLOR_PRIMARY = RGBColor(91, 33, 182)
COLOR_TEXT = RGBColor(31, 41, 55)
PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
TOP_MARGIN_CM = 2.5
BOTTOM_MARGIN_CM = 2.5
LEFT_MARGIN_CM = 3.0
RIGHT_MARGIN_CM = 2.5
LIST_TEXT_INDENT = Cm(1.2)
LIST_HANGING_INDENT = Cm(-0.6)

CHAPTERS = [
    {
        "title": "BAB I PENDAHULUAN",
        "intro": [
            "Bab ini menjelaskan alasan penyusunan buku panduan, tujuan, capaian, sasaran pengguna, dan ruang lingkup pembahasan portal penandatangan. Fokus utama dokumen diarahkan pada pengalaman penandatangan saat menerima antrian dokumen, membuka detail permohonan, meninjau preview dan data snapshot, lalu menyimpan keputusan secara tertib melalui Website ULT FKIP Universitas Lampung.",
        ],
        "sections": [
            {
                "title": "1.1 Latar Belakang",
                "body": [
                    "Portal penandatangan merupakan titik kontrol formal dalam alur layanan digital Website ULT. Pada tahap ini, keputusan penandatangan dapat melanjutkan dokumen ke proses berikutnya, mengembalikan dokumen untuk revisi, atau menghentikan proses bila permohonan dinilai tidak layak dilanjutkan.",
                    "Walaupun area kerja penandatangan lebih ringkas dibanding portal pemohon atau pengelola, tahap ini tetap memerlukan panduan yang jelas. Penandatangan perlu memahami bahwa antrian yang tampil pada inbox hanya berisi permohonan yang memang sedang menunggu tindakannya, sehingga proses peninjauan harus dilakukan dengan cermat sebelum keputusan disimpan.",
                    "Buku panduan ini disusun agar penandatangan memiliki acuan operasional yang selaras dengan tampilan sistem terbaru. Dengan panduan yang runtut, pengguna dapat lebih cepat mengenali susunan inbox, membaca ringkasan permohonan, memeriksa preview dokumen dan data snapshot, serta memberikan keputusan yang sesuai dengan konteks layanan.",
                ],
            },
            {
                "title": "1.2 Tujuan Buku Panduan",
                "body": [
                    "Buku panduan ini bertujuan memberikan acuan operasional yang lebih faktual bagi pengguna portal penandatangan. Dokumen ini juga berfungsi sebagai pelengkap dokumentasi produk agar alur peninjauan dan pengambilan keputusan dapat dijelaskan secara konsisten kepada pengguna baru, penguji, maupun pihak pengelola sistem.",
                    "Secara lebih rinci, tujuan penyusunan buku panduan ini dapat dijabarkan ke dalam beberapa poin berikut.",
                ],
                "bullets": [
                    "menjelaskan langkah penggunaan portal penandatangan dari login sampai keputusan disimpan;",
                    "membantu penandatangan memahami informasi yang tampil pada inbox dan halaman detail permohonan;",
                    "menunjukkan area antarmuka yang penting agar peninjauan dokumen dapat dilakukan secara lebih tertib;",
                    "menerangkan dampak keputusan menyetujui, meminta revisi, dan menolak terhadap alur layanan.",
                ],
            },
            {
                "title": "1.3 Capaian Buku Panduan",
                "body": [
                    "Setelah mempelajari buku panduan ini, penandatangan diharapkan tidak hanya mengenali tampilan portal, tetapi juga mampu menjalankan alur penggunaan sistem secara mandiri, tertib, dan sesuai peran penandatangan pada tahap proses yang sedang aktif.",
                    "Secara lebih operasional, capaian buku panduan ini dapat dijabarkan ke dalam beberapa poin berikut.",
                ],
                "bullets": [
                    "mampu login dan masuk ke Signer Inbox dengan benar;",
                    "mampu membaca daftar antrian dan memilih permohonan yang relevan untuk ditinjau;",
                    "mampu meninjau ringkasan permohonan, preview dokumen, dan data snapshot sebelum mengambil keputusan;",
                    "mampu memberikan keputusan menyetujui, meminta revisi, atau menolak beserta catatan yang diperlukan;",
                    "mampu memahami perubahan kondisi antrian setelah keputusan penandatangan disimpan.",
                ],
            },
            {
                "title": "1.4 Sasaran Pengguna Buku Panduan",
                "body": [
                    "Sasaran utama dokumen ini adalah pengguna yang memiliki kewenangan sebagai penandatangan pada alur layanan. Pada implementasi saat ini, peran tersebut dapat melekat pada jabatan struktural tertentu, penandatangan berbasis pengguna tertentu, maupun penandatangan khusus yang diatur pada rantai penandatangan layanan.",
                ],
            },
            {
                "title": "1.5 Ruang Lingkup Buku Panduan",
                "body": [
                    "Pembahasan difokuskan pada portal penandatangan dan tahap penggunaan yang langsung berhubungan dengan antrian penandatangan. Area pemohon, pengelola layanan, dan superadmin hanya disebut seperlunya untuk menjelaskan konteks alur layanan dari sudut pandang penandatangan.",
                ],
                "bullets": [
                    "akses masuk ke sistem dan pengenalan Signer Inbox;",
                    "pembacaan daftar permohonan dan informasi ringkas pada antrian penandatangan;",
                    "peninjauan detail permohonan, preview dokumen, dan data snapshot;",
                    "pemberian keputusan, catatan, dan unggah tanda tangan bila diwajibkan sistem;",
                    "pemahaman dampak keputusan terhadap alur layanan dan kondisi inbox setelah keputusan disimpan.",
                ],
            },
        ],
    },
    {
        "title": "BAB II MENGENAL PORTAL PENANDATANGAN",
        "intro": [
            "Portal penandatangan adalah area kerja yang menampung tahap persetujuan formal pada alur layanan. Dari portal ini, penandatangan dapat melihat antrian yang menunggu tindakan, membuka detail permohonan, membaca dokumen atau data pendukung, lalu menyimpan keputusan pada tahap yang sedang aktif.",
        ],
        "sections": [
            {
                "title": "2.1 Portal Penandatangan",
                "body": [
                    "Pada implementasi saat ini, penandatangan tidak melihat seluruh permohonan dalam sistem. Portal hanya menampilkan antrian yang memang sedang menunggu keputusan penandatangan terkait. Karena itu, halaman inbox berfungsi sebagai daftar kerja aktif yang perlu ditinjau secara bertahap.",
                    "Dari sisi operasional, portal penandatangan harus dipahami sebagai ruang keputusan, bukan ruang pengelolaan data umum. Informasi yang ditampilkan sengaja dibuat ringkas agar penandatangan dapat fokus pada dokumen yang sedang aktif pada tahap dirinya.",
                ],
            },
            {
                "title": "2.2 Fitur Utama Portal Penandatangan",
                "body": [
                    "Fitur utama portal penandatangan dirancang untuk mendukung peninjauan singkat namun tetap akurat. Komponen inti yang tersedia pada area ini adalah Signer Inbox, daftar permohonan, halaman detail permohonan, preview dokumen, data snapshot, dan form keputusan penandatangan.",
                ],
                "bullets": [
                    "Signer Inbox menampilkan ringkasan antrian yang menunggu keputusan penandatangan;",
                    "daftar permohonan menampilkan identitas layanan, pemohon, unit akademik, tahap proses, dan waktu pembaruan terakhir;",
                    "halaman detail memperlihatkan ringkasan permohonan sebelum keputusan diambil;",
                    "preview dokumen dan data snapshot membantu penandatangan meninjau isi dokumen dan konteks pengajuan;",
                    "form keputusan dipakai untuk mengirim pilihan menyetujui, meminta revisi, atau menolak beserta catatan yang diperlukan.",
                ],
            },
        ],
    },
    {
        "title": "BAB III AKSES MASUK KE SISTEM",
        "intro": [
            "Bab ini membahas tahap awal sebelum penandatangan mulai meninjau permohonan, yaitu login ke Website ULT, masuk ke portal penandatangan, dan mengenali halaman awal Signer Inbox. Pemahaman pada tahap ini penting agar pengguna masuk ke area kerja yang benar dan dapat segera memulai peninjauan dokumen pada antrian aktif.",
        ],
        "sections": [
            {
                "title": "3.1 Halaman Login",
                "body": [
                    "Halaman login menjadi gerbang autentikasi awal sebelum penandatangan masuk ke portal sesuai hak aksesnya. Pengguna perlu memasukkan kredensial akun yang benar agar sistem dapat mengenali kewenangan yang dimiliki dan mengarahkan ke area kerja yang sesuai.",
                ],
                "figure": ("01-halaman-login-penandatangan.png", "Gambar 1. Halaman login Website ULT"),
            },
            {
                "title": "3.2 Langkah Masuk ke Portal Penandatangan",
                "body": [
                    "Proses login penandatangan dirancang ringkas, tetapi pengguna tetap perlu memastikan bahwa akun yang dipakai memang memiliki kewenangan sebagai penandatangan pada layanan tertentu.",
                ],
                "steps": [
                    "buka halaman login Website ULT melalui menu masuk yang tersedia pada website;",
                    "isi email atau identitas akun dan kata sandi dengan benar;",
                    "tekan tombol masuk lalu tunggu proses autentikasi selesai;",
                    "pastikan sistem mengarahkan Anda ke portal penandatangan atau halaman Signer Inbox yang sesuai.",
                ],
            },
            {
                "title": "3.3 Halaman Awal Signer Inbox",
                "body": [
                    "Setelah login, penandatangan diarahkan ke halaman awal Signer Inbox. Halaman ini menampilkan ringkasan jumlah antrian dan menjadi titik orientasi pertama sebelum pengguna membuka daftar permohonan yang sedang menunggu tindakannya.",
                ],
                "figure": ("02-signer-inbox-hero.png", "Gambar 2. Halaman awal Signer Inbox"),
            },
        ],
    },
    {
        "title": "BAB IV MELIHAT DAFTAR PERMOHONAN",
        "intro": [
            "Bab ini menjelaskan cara membaca daftar antrian pada portal penandatangan. Pemahaman terhadap informasi ringkas pada daftar permohonan penting agar penandatangan dapat memilih dokumen yang perlu ditinjau terlebih dahulu dan tidak salah membaca konteks layanan yang sedang aktif.",
        ],
        "sections": [
            {
                "title": "4.1 Halaman Daftar Permohonan",
                "body": [
                    "Daftar permohonan pada portal penandatangan memuat kartu-kartu antrian yang sedang menunggu keputusan penandatangan. Halaman ini menjadi area kerja utama untuk memilih dokumen yang akan dibuka lebih lanjut.",
                    "Pada setiap kartu antrian, sistem menampilkan informasi ringkas seperti nama layanan, nomor permohonan, identitas pemohon, unit akademik, tahap proses yang sedang aktif, dan waktu pembaruan terakhir. Susunan informasi ini membantu penandatangan membaca konteks permohonan tanpa harus langsung membuka halaman detail.",
                ],
                "figure": ("03-daftar-permohonan-signer.png", "Gambar 3. Daftar permohonan pada portal penandatangan"),
            },
            {
                "title": "4.2 Informasi yang Ditampilkan pada Daftar Permohonan",
                "body": [
                    "Setiap kartu pada daftar permohonan menampilkan ringkasan informasi yang cukup untuk membantu penandatangan menyusun prioritas peninjauan. Informasi ini perlu dibaca dengan benar sebelum tombol detail digunakan.",
                ],
                "bullets": [
                    "nama layanan menunjukkan konteks dokumen yang sedang diproses;",
                    "nomor permohonan memudahkan identifikasi dokumen pada antrian penandatangan;",
                    "nama pemohon dan unit akademik membantu penandatangan memastikan konteks pihak yang mengajukan;",
                    "tahap proses saat ini menunjukkan tahap aktif yang sedang menunggu keputusan penandatangan;",
                    "update terakhir membantu menilai kapan permohonan terakhir kali berubah.",
                ],
            },
            {
                "title": "4.3 Memilih Permohonan untuk Ditinjau",
                "body": [
                    "Setelah membaca ringkasan pada daftar, penandatangan dapat membuka permohonan yang relevan melalui tombol detail. Langkah ini membawa pengguna ke halaman detail penandatangan yang menampilkan informasi lebih lengkap sebelum keputusan diambil.",
                ],
                "steps": [
                    "baca nama layanan, pemohon, unit akademik, dan tahap proses pada kartu antrian;",
                    "pastikan permohonan yang dipilih memang sedang menunggu keputusan Anda pada tahap aktif;",
                    "tekan tombol Buka Detail untuk masuk ke halaman peninjauan penandatangan;",
                    "lanjutkan pembacaan ringkasan, preview dokumen, dan data snapshot pada halaman detail.",
                ],
            },
        ],
    },
    {
        "title": "BAB V MENINJAU DETAIL PERMOHONAN",
        "intro": [
            "Bab ini membahas susunan halaman detail penandatangan yang dipakai untuk menilai dokumen sebelum keputusan disimpan. Pada tahap ini, penandatangan tidak cukup hanya membaca judul layanan, tetapi juga perlu meninjau ringkasan permohonan, preview dokumen, data snapshot, dan status permohonan yang sedang aktif.",
        ],
        "sections": [
            {
                "title": "5.1 Membuka Detail Permohonan",
                "body": [
                    "Halaman detail penandatangan menampilkan ringkasan permohonan sebagai konteks utama sebelum keputusan diambil. Pada bagian ini, penandatangan dapat membaca identitas layanan, mahasiswa, nomor surat, tahap aktif, dan status permohonan.",
                ],
                "figure": ("05-ringkasan-permohonan-signer.png", "Gambar 4. Ringkasan permohonan pada halaman detail penandatangan"),
            },
            {
                "title": "5.2 Membaca Ringkasan Informasi Dokumen",
                "body": [
                    "Ringkasan informasi dokumen membantu penandatangan memahami dokumen yang sedang ditinjau tanpa harus langsung berpindah ke seluruh data pendukung. Bagian ini perlu dibaca lebih dahulu agar peninjauan preview dan snapshot dilakukan dengan konteks yang benar.",
                ],
                "steps": [
                    "periksa nama layanan dan identitas permohonan yang tampil pada halaman detail;",
                    "baca nomor surat, tahap aktif, dan status yang sedang berjalan;",
                    "pastikan dokumen yang dibuka memang sesuai dengan antrian yang dipilih dari inbox;",
                    "lanjutkan ke preview dokumen atau data snapshot jika konteks dasar sudah dipahami.",
                ],
            },
            {
                "title": "5.3 Meninjau Preview Dokumen dan Data Snapshot",
                "body": [
                    "Preview dokumen dan data snapshot membantu penandatangan membaca isi dokumen sambil tetap memahami data input yang mendasari pengajuan. Kombinasi kedua area ini penting agar keputusan tidak diambil hanya berdasarkan ringkasan singkat atau judul layanan saja.",
                ],
                "figure": ("06-preview-dokumen-dan-snapshot-signer.png", "Gambar 5. Preview dokumen dan data snapshot pada portal penandatangan"),
                "bullets": [
                    "preview dokumen dipakai untuk melihat hasil dokumen yang akan disetujui atau dikembalikan;",
                    "data snapshot menampilkan ringkasan input pemohon saat pengajuan dibuat;",
                    "lampiran atau file yang muncul pada snapshot dapat membantu penandatangan memeriksa konteks pengajuan secara lebih utuh;",
                    "baca preview dan snapshot secara berpasangan agar keputusan lebih akurat.",
                ],
            },
            {
                "title": "5.4 Memahami Status Permohonan",
                "body": [
                    "Status pada halaman detail membantu penandatangan membaca posisi dokumen dalam alur layanan. Pada tahap penandatanganan, status ini penting untuk memastikan bahwa dokumen memang sedang berada pada fase penandatanganan atau keputusan formal lainnya.",
                ],
                "figure": ("07-status-permohonan-signer.png", "Gambar 6. Status permohonan pada halaman detail penandatangan"),
            },
        ],
    },
    {
        "title": "BAB VI MEMBERIKAN KEPUTUSAN",
        "intro": [
            "Bab ini menjelaskan penggunaan form keputusan penandatangan. Pada implementasi saat ini, keputusan yang tersedia di sistem adalah APPROVE, REVISION, dan REJECT. Setiap pilihan memiliki dampak yang berbeda terhadap kelanjutan alur layanan, sehingga penandatangan perlu memahami fungsi masing-masing sebelum tombol kirim keputusan ditekan.",
        ],
        "sections": [
            {
                "title": "6.1 Menyetujui Dokumen",
                "body": [
                    "Keputusan approve dipakai ketika penandatangan menilai dokumen dan data pendukung telah sesuai untuk dilanjutkan ke tahap berikutnya. Sebelum memilih approve, penandatangan perlu memastikan bahwa preview dokumen, data snapshot, dan konteks layanan telah diperiksa dengan cukup.",
                ],
            },
            {
                "title": "6.2 Meminta Revisi",
                "body": [
                    "Keputusan revision dipakai ketika dokumen belum dapat dilanjutkan, tetapi masih dapat diperbaiki oleh pihak terkait. Opsi ini berbeda dari reject karena tujuan utamanya adalah mengembalikan dokumen untuk perbaikan, bukan menghentikan proses sepenuhnya.",
                ],
            },
            {
                "title": "6.3 Menolak Dokumen",
                "body": [
                    "Keputusan reject dipakai ketika dokumen tidak layak dilanjutkan pada tahap penandatanganan. Karena dampaknya lebih tegas daripada revision, penggunaan opsi ini sebaiknya dilakukan setelah penandatangan benar-benar yakin terhadap hasil peninjauan dokumen dan data pendukung.",
                ],
            },
            {
                "title": "6.4 Memberikan Catatan atau Pertimbangan",
                "body": [
                    "Field catatan dapat digunakan untuk menjelaskan pertimbangan keputusan atau memberi instruksi tindak lanjut. Walaupun bersifat opsional pada sistem, catatan tetap penting ketika keputusan revision atau reject memerlukan arahan yang lebih jelas bagi pihak berikutnya.",
                ],
                "figure": ("08-form-keputusan-signer.png", "Gambar 7. Form keputusan penandatangan"),
                "bullets": [
                    "pilih keputusan sesuai hasil peninjauan dokumen dan data snapshot;",
                    "tambahkan catatan jika keputusan memerlukan penjelasan atau instruksi lanjutan;",
                    "baca kembali pilihan keputusan sebelum form dikirim;",
                    "gunakan bahasa catatan yang singkat, jelas, dan operasional.",
                ],
            },
            {
                "title": "6.5 Menyimpan Hasil Keputusan",
                "body": [
                    "Setelah keputusan dan catatan diisi, penandatangan dapat mengirim form melalui tombol kirim keputusan. Pada kondisi tertentu, sistem juga dapat menampilkan field unggah file tanda tangan sebagai syarat tambahan, terutama ketika langkah penandatangan memang mewajibkan unggah tanda tangan untuk keputusan approve.",
                ],
                "steps": [
                    "pastikan pilihan keputusan sudah sesuai dengan hasil peninjauan;",
                    "isi catatan bila diperlukan untuk menjelaskan alasan keputusan;",
                    "unggah file tanda tangan jika field tersebut diwajibkan sistem pada tahap penandatangan aktif;",
                    "tekan tombol kirim keputusan lalu tunggu sistem menyimpan hasil tindakan penandatangan.",
                ],
            },
        ],
    },
    {
        "title": "BAB VII DAMPAK KEPUTUSAN DAN TINDAK LANJUT",
        "page_break_before": False,
        "intro": [
            "Setelah keputusan penandatangan disimpan, sistem akan memperbarui kondisi antrian dan melanjutkan alur layanan sesuai keputusan yang dipilih. Bab ini menjelaskan cara memahami perubahan tersebut dari sudut pandang penandatangan, tanpa mengasumsikan adanya halaman hasil keputusan yang berdiri sendiri.",
        ],
        "sections": [
            {
                "title": "7.1 Keputusan yang Sudah Disimpan",
                "body": [
                    "Sesudah form keputusan dikirim, sistem menyimpan hasil tindakan penandatangan dan mengarahkan pengguna kembali ke inbox atau kondisi antrian yang relevan. Dari titik ini, penandatangan dapat mengetahui bahwa tindakan sudah tercatat dan permohonan akan mengikuti alur berikutnya sesuai keputusan yang diberikan.",
                ],
                "figure": ("10-inbox-setelah-keputusan-signer.png", "Gambar 8. Perubahan kondisi inbox setelah keputusan penandatangan disimpan"),
            },
            {
                "title": "7.2 Dampak Keputusan terhadap Alur Layanan",
                "body": [
                    "Keputusan approve, revision, dan reject memiliki konsekuensi yang berbeda terhadap perjalanan dokumen. Karena itu, penandatangan perlu memahami bahwa keputusan yang dipilih tidak hanya mengubah catatan pada form, tetapi juga memengaruhi status permohonan dan perpindahan proses ke tahap berikutnya.",
                ],
                "bullets": [
                    "approve dapat melanjutkan dokumen ke tahap berikutnya pada rantai proses;",
                    "revision mengembalikan dokumen agar pihak terkait melakukan perbaikan sesuai catatan;",
                    "reject menghentikan atau menutup kelanjutan dokumen pada jalur yang sedang berjalan;",
                    "perubahan kondisi antrian setelah keputusan disimpan menjadi indikator awal bahwa sistem telah mencatat tindakan penandatangan.",
                ],
            },
            {
                "title": "7.3 Tindak Lanjut Setelah Keputusan",
                "body": [
                    "Setelah keputusan tersimpan, penandatangan umumnya tidak lagi melakukan edit pada permohonan yang sama pada tahap tersebut. Tindak lanjut berikutnya adalah memantau apakah inbox masih memuat antrian lain yang menunggu tindakan, lalu melanjutkan peninjauan pada dokumen berikutnya yang relevan.",
                ],
                "figure": ("11-status-atau-dampak-keputusan-signer.png", "Gambar 9. Dampak keputusan penandatangan terhadap alur layanan"),
            },
        ],
    },
    {
        "title": "BAB VIII KETENTUAN UMUM PENGGUNAAN",
        "intro": [
            "Bab ini memuat ketentuan umum yang perlu diperhatikan oleh pengguna portal penandatangan. Tujuannya adalah menjaga konsistensi penggunaan akun, ketelitian peninjauan dokumen, dan kehati-hatian saat mengirim keputusan yang berdampak langsung pada alur layanan.",
        ],
        "sections": [
            {
                "title": "8.1 Ketentuan Penggunaan Akun",
                "body": [
                    "Akun penandatangan harus digunakan sesuai kewenangan yang diberikan pada sistem. Pengguna tidak disarankan meminjamkan akun atau membiarkan sesi login aktif tanpa pengawasan, karena setiap keputusan yang dikirim melalui portal akan tercatat sebagai tindakan pengguna yang sedang login.",
                ],
            },
            {
                "title": "8.2 Ketentuan Peninjauan Dokumen",
                "body": [
                    "Peninjauan dokumen sebaiknya dilakukan dengan membaca ringkasan permohonan, preview dokumen, dan data snapshot secara utuh. Mengambil keputusan hanya berdasarkan judul layanan atau informasi yang terbatas berisiko menimbulkan kesalahan tindak lanjut pada alur layanan.",
                ],
            },
            {
                "title": "8.3 Ketentuan Pemberian Keputusan",
                "body": [
                    "Keputusan approve, revision, dan reject harus dipilih sesuai hasil peninjauan yang nyata. Jika diperlukan, penandatangan perlu menambahkan catatan agar alasan keputusan dapat dipahami oleh pihak berikutnya dan tidak menimbulkan ambiguitas saat tindak lanjut dilakukan.",
                ],
            },
            {
                "title": "8.4 Ketentuan Kehati-hatian dan Keamanan Penggunaan",
                "body": [
                    "Pengguna perlu berhati-hati ketika mengunggah file tanda tangan pada tahap yang mewajibkannya. Pastikan file yang digunakan benar, sesuai ketentuan sistem, dan tidak tertukar dengan file milik pengguna lain. Setelah selesai menggunakan portal, penandatangan juga disarankan mengakhiri sesi dengan logout dari sistem.",
                ],
            },
        ],
    },
    {
        "title": "BAB IX PENUTUP",
        "page_break_before": False,
        "intro": [
            "Buku panduan penandatangan ini disusun sebagai acuan operasional penggunaan portal penandatangan pada Website ULT FKIP Universitas Lampung. Dengan memahami susunan inbox, detail permohonan, preview dokumen, data snapshot, dan form keputusan, penandatangan diharapkan dapat menjalankan perannya secara lebih tertib dan konsisten sesuai alur layanan yang berlaku.",
            "Pembaruan isi dan gambar pada dokumen ini perlu terus dilakukan ketika antarmuka atau alur kerja sistem mengalami perubahan. Karena itu, buku panduan ini sebaiknya diperlakukan sebagai dokumen kerja yang hidup dan selalu diselaraskan dengan implementasi portal penandatangan terbaru.",
        ],
        "sections": [],
    },
]


def set_run_font(run, size=10, bold=False, italic=False, color=COLOR_TEXT):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def set_doc_defaults(doc):
    for section in doc.sections:
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.top_margin = Cm(TOP_MARGIN_CM)
        section.bottom_margin = Cm(BOTTOM_MARGIN_CM)
        section.left_margin = Cm(LEFT_MARGIN_CM)
        section.right_margin = Cm(RIGHT_MARGIN_CM)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(4)

    for name, size in [("Heading 1", 12), ("Heading 2", 11)]:
        style = doc.styles[name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = COLOR_PRIMARY
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = name == "Heading 2"


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=12 if level == 1 else 11, bold=True, color=COLOR_PRIMARY)


def add_body(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = LIST_TEXT_INDENT
        p.paragraph_format.first_line_indent = LIST_HANGING_INDENT
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.tab_stops.add_tab_stop(p.paragraph_format.left_indent)
        bullet = p.add_run("\u2022\t")
        set_run_font(bullet, size=10, bold=True, color=COLOR_PRIMARY)
        body = p.add_run(item)
        set_run_font(body, size=10)


def add_steps(doc, items):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = LIST_TEXT_INDENT
        p.paragraph_format.first_line_indent = LIST_HANGING_INDENT
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.tab_stops.add_tab_stop(p.paragraph_format.left_indent)
        number = p.add_run(f"{index}.\t")
        set_run_font(number, size=10, bold=True, color=COLOR_PRIMARY)
        body = p.add_run(item)
        set_run_font(body, size=10)


def add_steps_intro(doc, title):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"Untuk mempraktikkan {title.lower()}, ikuti langkah-langkah berikut.")
    set_run_font(run, size=10)


def add_bullets_intro(doc, title):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(
        f"Beberapa hal penting yang perlu diperhatikan pada bagian {title.lower()} adalah sebagai berikut."
    )
    set_run_font(run, size=10)


def add_field_note(doc, title, instruction):
    add_heading(doc, title, 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(instruction)
    set_run_font(run, size=9)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(85)
    run = p.add_run("BUKU PANDUAN PENANDATANGAN")
    set_run_font(run, size=16, bold=True, color=COLOR_PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "WEBSITE UNIT LAYANAN TERPADU (ULT)\n"
        "FAKULTAS KEGURUAN DAN ILMU PENDIDIKAN\n"
        "UNIVERSITAS LAMPUNG"
    )
    set_run_font(run, size=13, bold=True, color=COLOR_PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    run = p.add_run("Panduan Operasional Penggunaan Portal Penandatangan")
    set_run_font(run, size=11, bold=True, color=COLOR_PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Dokumen ini menjelaskan alur penggunaan portal penandatangan mulai dari login,\n"
        "membaca inbox, meninjau detail permohonan, hingga menyimpan keputusan penandatangan."
    )
    set_run_font(run, size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(150)
    run = p.add_run("Bandar Lampung\n2026")
    set_run_font(run, size=10)


def add_front_matter(doc):
    add_heading(doc, "KATA PENGANTAR", 1)
    add_body(
        doc,
        "Buku panduan ini disusun sebagai acuan operasional penggunaan Website Unit Layanan Terpadu (ULT) FKIP Universitas Lampung dari sudut pandang penandatangan. Pembaruan dokumen dilakukan agar isi panduan selaras dengan antarmuka dan alur sistem terbaru, khususnya pada area Signer Inbox, halaman detail permohonan, preview dokumen, data snapshot, dan form keputusan penandatangan.",
    )
    add_body(
        doc,
        "Melalui buku ini, pengguna diharapkan dapat memahami langkah penggunaan sistem secara lebih terarah sejak tahap login sampai keputusan disimpan. Tangkapan layar yang digunakan pada setiap bab dipilih dari implementasi aktual atau aset kerja yang diselaraskan dengan struktur portal penandatangan agar uraian yang disajikan lebih mudah diikuti ketika praktik penggunaan dilakukan.",
    )
    add_body(
        doc,
        "Dokumen ini juga berfungsi sebagai pelengkap dokumentasi produk hasil pengembangan, sehingga dapat dipakai sebagai referensi saat orientasi pengguna baru, peninjauan tugas akhir, maupun evaluasi alur persetujuan layanan pada sistem.",
    )

    doc.add_page_break()
    add_field_note(
        doc,
        "DAFTAR ISI",
        'Daftar isi dapat diperbarui otomatis melalui fitur "Update Field" pada Microsoft Word setelah peninjauan akhir dokumen.',
    )

    doc.add_page_break()
    add_field_note(
        doc,
        "DAFTAR GAMBAR",
        'Daftar gambar dapat diperbarui otomatis melalui fitur "Update Field" pada Microsoft Word setelah caption final diperiksa.',
    )


def resolve_image_dimensions(path: Path):
    with Image.open(path) as image:
        width_px, height_px = image.size

    aspect_ratio = width_px / max(height_px, 1)
    content_width = PAGE_WIDTH_CM - LEFT_MARGIN_CM - RIGHT_MARGIN_CM
    content_height = PAGE_HEIGHT_CM - TOP_MARGIN_CM - BOTTOM_MARGIN_CM
    max_width = min(13.4, content_width - 0.2)
    max_height = min(15.2, content_height - 1.2)

    if aspect_ratio >= 4.5:
        target_width = 10.4
    elif aspect_ratio >= 2.8:
        target_width = 11.2
    elif aspect_ratio >= 1.6:
        target_width = 12.4
    elif aspect_ratio >= 1.0:
        target_width = 11.6
    elif aspect_ratio >= 0.72:
        target_width = 9.8
    else:
        target_width = 8.4

    target_width = min(target_width, max_width)
    target_height = target_width / max(aspect_ratio, 0.01)

    if target_height > max_height:
        scale = max_height / target_height
        target_width *= scale
        target_height = max_height

    return Cm(round(target_width, 2)), Cm(round(target_height, 2))


def add_figure(doc, filename, caption):
    candidates = [
        ASSETS / filename,
        ROOT / "public/example" / filename,
    ]
    path = next((candidate for candidate in candidates if candidate.exists()), None)
    if path is None:
        add_body(doc, f"[Gambar tidak ditemukan: {filename}]")
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    width, height = resolve_image_dimensions(path)
    p.add_run().add_picture(str(path), width=width, height=height)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(4)
    run = cap.add_run(caption)
    set_run_font(run, size=8, italic=True)


def add_chapters(doc):
    for chapter in CHAPTERS:
        if chapter.get("page_break_before", True):
            doc.add_page_break()
        add_heading(doc, chapter["title"], 1)
        for paragraph in chapter.get("intro", []):
            add_body(doc, paragraph)

        for section in chapter.get("sections", []):
            add_heading(doc, section["title"], 2)
            for paragraph in section.get("body", []):
                add_body(doc, paragraph)
            if section.get("figure"):
                add_figure(doc, section["figure"][0], section["figure"][1])
            if section.get("steps"):
                add_steps_intro(doc, section["title"])
                add_steps(doc, section["steps"])
            if section.get("bullets"):
                if section.get("steps") or section.get("figure"):
                    add_bullets_intro(doc, section["title"])
                add_bullets(doc, section["bullets"])


def main():
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc)
    doc.add_page_break()
    add_front_matter(doc)
    add_chapters(doc)
    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    main()
