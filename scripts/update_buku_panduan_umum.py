from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image


ROOT = Path(r"c:/laragon/www/ult-fkip-unila")
OUTPUT = ROOT / "docs/buku-panduan/buku-panduan-umum/buku-panduan-umum-website-ult.docx"
COVER_OUTPUT = ROOT / "docs/buku-panduan/buku-panduan-umum/buku-panduan-umum-website-ult - Cover.docx"
ASSETS = ROOT / "docs/buku-panduan/assets/screenshots"
FONT_NAME = "Poppins"
COLOR_PRIMARY = RGBColor(91, 33, 182)
COLOR_TEXT = RGBColor(31, 41, 55)
PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
TOP_MARGIN_CM = 2.5
BOTTOM_MARGIN_CM = 2.5
LEFT_MARGIN_CM = 3.0
RIGHT_MARGIN_CM = 2.5
LIST_TEXT_INDENT = Cm(1.2)
LIST_HANGING_INDENT = Cm(-0.6)

ROLE_ROWS = [
    ("Pengunjung publik", "Membaca informasi, katalog layanan, detail layanan, dan panduan publik.", "Pengguna yang belum login."),
    ("Pemohon layanan", "Mengajukan layanan, memantau status, memperbaiki data, dan mengunduh hasil.", "Mahasiswa, alumni, dosen, atau tenaga kependidikan."),
    ("Pengelola layanan", "Memeriksa permohonan, menjalankan workflow, dan mengelola layanan dokumen.", "Staf admin, operator, atau reviewer ULT."),
    ("Penandatangan", "Meninjau dokumen, membaca snapshot, dan memberi keputusan.", "Pimpinan atau signer dokumen."),
    ("Superadmin", "Mengelola konfigurasi global, layanan, pengguna, konten, dan audit.", "Administrator utama sistem."),
]

CHAPTERS = [
    {
        "title": "BAB I PENDAHULUAN",
        "intro": [
            "Buku panduan umum ini disusun sebagai panduan orientasi sistem untuk Website Unit Layanan Terpadu (ULT) FKIP Universitas Lampung. Fokusnya adalah membantu pembaca memahami peta besar modul, hubungan antarperan, dan arah penggunaan sistem secara menyeluruh.",
            "Berbeda dari buku panduan berbasis peran yang fokus pada pemohon, pengelola, atau penandatangan, buku ini merangkum ketiga perspektif tersebut dalam satu alur besar agar hubungan antarportal lebih mudah dipahami.",
        ],
        "sections": [
            {"title": "1.1 Latar Belakang", "body": ["Website ULT dibangun untuk memusatkan layanan administrasi berbasis web dalam satu sistem yang menghubungkan area publik, portal pemohon, portal admin, portal penandatangan, dan modul administratif pendukung. Karena itulah diperlukan buku panduan umum yang menjelaskan konteks besar sistem, bukan hanya langkah teknis tiap peran."]},
            {"title": "1.2 Tujuan Buku Panduan", "body": ["Buku panduan ini bertujuan menyediakan acuan orientatif bagi pembaca yang perlu memahami Website ULT secara menyeluruh sebelum masuk ke panduan yang lebih khusus. Secara lebih rinci, tujuan buku panduan ini dapat dijabarkan ke dalam beberapa poin berikut."], "bullets": ["menjelaskan hubungan antara area publik, portal berbasis peran, dan modul administratif;", "membantu pembaca memahami alur layanan dari informasi awal sampai hasil layanan tersedia;", "menjadi panduan orientasi sebelum menggunakan buku panduan khusus;", "mendukung dokumentasi sistem agar struktur penggunaan Website ULT dapat dipahami secara konsisten;"]},
            {"title": "1.3 Capaian Buku Panduan", "body": ["Setelah mempelajari buku panduan ini, pembaca diharapkan mampu memahami struktur besar Website ULT, mengenali alur layanan lintas peran, dan menentukan buku panduan khusus yang perlu dibuka untuk pendalaman operasional. Secara lebih operasional, capaian buku panduan ini dapat dijabarkan ke dalam beberapa poin berikut."], "bullets": ["mampu mengenali area publik dan akses awal;", "mampu memahami ringkasan alur kerja portal pemohon, pengelola, dan penandatangan;", "mampu membaca hubungan antara layanan dokumen, nomor surat, monitoring status, dan modul pendukung;", "mampu menggunakan buku umum sebagai peta orientasi sebelum membuka buku khusus;"]},
            {"title": "1.4 Kedudukan Buku Panduan Umum", "body": ["Dokumen ini berkedudukan sebagai panduan umum atau panduan induk. Artinya, isi buku merangkum area penting dari buku panduan pemohon layanan, pengelola layanan, dan penandatangan, lalu menambahkan konteks lintas peran yang diperlukan untuk membaca sistem secara utuh."]},
            {"title": "1.5 Ruang Lingkup", "body": ["Ruang lingkup pembahasan mencakup area publik website, akses awal, ringkasan portal pemohon, ringkasan portal pengelola, ringkasan portal penandatangan, alur umum layanan, serta modul administratif yang penting dipahami oleh superadmin."], "bullets": ["halaman publik dan navigasi umum;", "detail layanan, persyaratan, SOP, dan panduan publik;", "ringkasan alur portal pemohon, pengelola, dan penandatangan;", "monitoring status dan alur layanan lintas peran;", "modul administratif strategis pada area superadmin;"]},
            {"title": "1.6 Cara Menggunakan Buku Ini", "body": ["Agar buku ini efektif digunakan, pembaca disarankan membaca bab secara berurutan ketika pertama kali mempelajari sistem. Setelah memahami struktur globalnya, pembaca dapat memakai bab tertentu sebagai referensi cepat sesuai konteks yang sedang dipelajari."], "steps": ["mulailah dari gambaran umum sistem untuk memahami peran dan modul;", "lanjutkan ke area publik dan akses awal untuk melihat pengalaman pengguna sebelum login;", "gunakan bab portal pemohon, pengelola, dan penandatangan untuk membaca gambaran operasional tiap peran;", "gunakan bab alur umum dan administrasi superadmin saat melakukan monitoring atau evaluasi;"]},
        ],
    },
    {
        "title": "BAB II GAMBARAN UMUM WEBSITE ULT",
        "intro": ["Bab ini menempatkan Website ULT sebagai sistem layanan berbasis web yang menghubungkan pengunjung publik, pemohon, pengelola, penandatangan, dan superadmin dalam satu alur kerja digital."],
        "table_title": "Tabel 1. Ringkasan peran pengguna pada Website ULT",
        "table_rows": ROLE_ROWS,
        "sections": [
            {"title": "2.1 Pengertian Website ULT", "body": ["Website Unit Layanan Terpadu (ULT) FKIP Universitas Lampung merupakan platform layanan administrasi berbasis web yang memusatkan penyampaian informasi, pengajuan layanan, verifikasi, persetujuan, dan distribusi hasil layanan secara digital."]},
            {"title": "2.2 Tujuan Penggunaan Website ULT", "body": ["Penerapan Website ULT bertujuan mempercepat akses informasi, mengefisienkan penyampaian permohonan, menjaga dokumentasi proses, dan meningkatkan transparansi status layanan."], "bullets": ["mengurangi ketergantungan pada proses manual;", "menjaga jejak status dan keputusan;", "memudahkan pemantauan performa layanan;", "mendukung distribusi hasil layanan dalam bentuk digital;"]},
            {"title": "2.3 Peta Modul Utama", "body": ["Website ULT perlu dipahami sebagai kumpulan modul yang saling terhubung. Tiap modul melayani kebutuhan pengguna yang berbeda, tetapi seluruhnya berkontribusi pada satu siklus layanan yang utuh."], "bullets": ["modul publik dan autentikasi;", "modul pemohon;", "modul pengelola;", "modul penandatangan;", "modul administratif master;"]},
            {"title": "2.4 Relasi Buku Panduan Umum dengan Buku Panduan Khusus", "body": ["Buku panduan umum ini menghimpun isi penting dari tiga buku panduan berbasis peran. Karena itu, buku umum berfungsi sebagai jembatan sebelum pembaca membuka buku khusus sesuai kebutuhan detail perannya."]},
        ],
    },
    {
        "title": "BAB III AREA PUBLIK DAN AKSES AWAL",
        "intro": ["Bagian ini menjelaskan pengalaman awal pengguna sebelum masuk ke portal berbasis peran. Pemahaman terhadap area publik penting karena area inilah yang pertama kali membentuk persepsi pengguna tentang layanan yang tersedia."],
        "sections": [
            {"title": "3.1 Beranda Website ULT", "body": ["Beranda berfungsi sebagai pintu masuk utama Website ULT dan menjadi wajah awal sistem sebelum pengguna membuka katalog layanan atau login."], "figure": ("01-beranda-website-ult.png", "Gambar 1. Beranda Website ULT")},
            {"title": "3.2 Menu Navigasi Utama", "body": ["Navigasi utama menjadi alat orientasi pertama bagi pengguna publik dan memudahkan perpindahan antarhalaman informasi."], "figure": ("02-menu-navigasi-utama.png", "Gambar 2. Menu navigasi utama Website ULT")},
            {"title": "3.3 Daftar Layanan", "body": ["Halaman daftar layanan menyajikan katalog layanan yang dapat dijelajahi pengguna untuk menilai kecocokan layanan sebelum membuka detail."], "figure": ("03-daftar-layanan.png", "Gambar 3. Daftar layanan pada halaman publik")},
            {"title": "3.4 Pencarian dan Filter Layanan", "body": ["Ketika jumlah layanan bertambah, fitur pencarian dan filter menjadi instrumen penting untuk mempercepat penemuan layanan."], "figure": ("04-filter-layanan.png", "Gambar 4. Pencarian dan filter layanan")},
            {"title": "3.5 Halaman Detail Layanan", "body": ["Detail layanan adalah halaman kunci sebelum pengajuan dilakukan karena di sinilah pengguna memeriksa kecocokan layanan dan arah tindakan berikutnya."], "figure": ("07-detail-layanan-publik.png", "Gambar 5. Halaman detail layanan publik")},
            {"title": "3.6 Persyaratan dan SOP Layanan", "body": ["Bagian persyaratan dan SOP merupakan inti informasi sebelum pengguna mengisi formulir dan menyiapkan dokumen."], "figure": ("08-persyaratan-dan-sop-layanan.png", "Gambar 6. Persyaratan dan SOP layanan")},
            {"title": "3.7 Halaman Login", "body": ["Halaman login berfungsi sebagai gerbang autentikasi menuju portal berbasis peran dan mengarahkan pengguna ke dashboard sesuai role yang dimiliki."], "figure": ("07-halaman-login.png", "Gambar 7. Halaman login Website ULT")},
            {"title": "3.8 Panduan Pengguna Publik", "body": ["Website ULT menyediakan halaman panduan pengguna publik sebagai sarana pembelajaran mandiri dan distribusi dokumentasi ke pengguna akhir."], "figure": ("08-panduan-pengguna-publik.png", "Gambar 8. Halaman panduan pengguna publik")},
        ],
    },
    {
        "title": "BAB IV PORTAL PEMOHON LAYANAN",
        "intro": ["Bab ini merangkum pengalaman utama pemohon layanan. Tujuannya bukan mengulang seluruh langkah teknis, melainkan membantu pembaca memahami bagaimana pengguna mengajukan layanan, memantau status, melakukan perbaikan, dan menerima hasil layanan dari portal pemohon."],
        "sections": [
            {"title": "4.1 Dashboard Pemohon", "body": ["Setelah login, pemohon diarahkan ke dashboard sebagai pusat orientasi awal dan akses cepat menuju daftar permohonan maupun layanan yang tersedia."], "figure": ("09-dashboard-pemohon.png", "Gambar 9. Dashboard portal pemohon")},
            {"title": "4.2 Daftar Permohonan", "body": ["Daftar permohonan menjadi tempat pemohon meninjau seluruh layanan yang pernah diajukan dan membaca status ringkasnya."], "figure": ("10-daftar-permohonan-pemohon.png", "Gambar 10. Daftar permohonan pada portal pemohon")},
            {"title": "4.3 Pengajuan Layanan", "body": ["Formulir pengajuan adalah sarana utama untuk mengirim data permohonan ke dalam sistem. Pengalaman pengajuan pemohon dapat berbeda antar layanan meskipun dimulai dari portal yang sama."], "figure": ("11-form-pengajuan-layanan.png", "Gambar 11. Form pengajuan layanan")},
            {"title": "4.4 Jenis Pengajuan Layanan", "body": ["Pada portal pemohon, ada layanan yang cukup memakai form biasa, ada layanan yang meminta upload file atau signer sejak awal, dan ada layanan sertifikat atau piagam yang memakai dokumen sumber .pptx."], "figure": ("12-form-sertifikat-piagam.png", "Gambar 12. Variasi pengajuan layanan pada portal pemohon")},
            {"title": "4.5 Detail, Status, dan Riwayat Permohonan", "body": ["Setelah pengajuan dikirim, pemohon memantau detail permohonan sebagai pusat informasi utama yang memuat ringkasan data, status, dan riwayat proses."], "figure": ("16-status-dan-riwayat-permohonan.png", "Gambar 13. Detail, status, dan riwayat permohonan")},
            {"title": "4.6 Perbaikan dan Output Layanan", "body": ["Pada kondisi tertentu, pemohon diminta memperbaiki permohonan sebelum proses dapat dilanjutkan. Setelah layanan selesai, sistem juga menyediakan output dokumen atau hasil layanan yang bisa diperiksa dan diunduh pengguna."], "figure": ("17-perbaikan-permohonan.png", "Gambar 14. Perbaikan permohonan dan output layanan")},
        ],
    },
    {
        "title": "BAB V PORTAL PENGELOLA LAYANAN",
        "intro": ["Bab ini merangkum area kerja utama pada portal pengelola layanan. Dari perspektif buku umum, portal admin dibaca sebagai pusat kendali proses karena di sinilah permohonan diperiksa, workflow dijalankan, layanan dokumen disiapkan, dan hasil layanan ditindaklanjuti."],
        "sections": [
            {"title": "5.1 Halaman Awal Portal Admin", "body": ["Setelah login sebagai admin, pengguna diarahkan ke halaman awal portal admin untuk membaca ringkasan kerja dan menentukan modul yang perlu dibuka lebih dahulu."], "figure": ("18-dashboard-pengelola.png", "Gambar 15. Halaman awal portal admin")},
            {"title": "5.2 Daftar Permohonan dan Detail Permohonan", "body": ["Daftar permohonan dan halaman detail merupakan inti kerja portal pengelola. Dari dua area ini, pengelola membaca layanan yang masuk, lalu membuka konteks permohonan secara utuh pada halaman detail.", "Pada buku umum, gambar yang ditampilkan cukup mewakili daftar permohonan sebagai titik masuk kerja pengelola. Langkah rinci pembacaan detail dijelaskan lebih lanjut pada buku panduan pengelola layanan."], "figure": ("20-daftar-permohonan-pengelola.png", "Gambar 16. Daftar permohonan pada portal admin")},
            {"title": "5.3 Verifikasi, Review, dan Workflow", "body": ["Area workflow memuat tindakan inti pengelola, seperti verifikasi awal, pemberian catatan, pengembalian untuk perbaikan, dan tindak lanjut ke tahap berikutnya."], "figure": ("24-aksi-verifikasi-dan-workflow.png", "Gambar 17. Verifikasi, review, dan workflow pada portal admin")},
            {"title": "5.4 Pengelolaan Layanan Dokumen", "body": ["Selain memproses permohonan, pengelola atau superadmin tertentu juga mengelola layanan dokumen melalui setup template, placeholder, gate, signer, dan publish readiness."], "figure": ("68-ringkasan-setup-layanan-dokumen-admin.png", "Gambar 18. Pengelolaan layanan dokumen")},
            {"title": "5.5 Output, Nomor Surat, dan Modul Pendukung", "body": ["Pada tahap lanjut, pengelola juga berhadapan dengan output layanan, template nomor surat, serta modul pendukung seperti kritik dan saran atau pedoman placeholder. Dalam buku umum, gambar pada bagian ini dipakai untuk mewakili modul nomor surat sebagai salah satu area pendukung yang paling penting pada portal admin."], "figure": ("45-daftar-format-nomor-surat-admin.png", "Gambar 19. Template nomor surat pada portal admin")},
        ],
    },
    {
        "title": "BAB VI PORTAL PENANDATANGAN",
        "intro": ["Portal penandatangan menampung tahap persetujuan formal pada alur layanan. Walaupun lebih ringkas dibanding portal pengelola, area ini sangat penting karena keputusan yang diambil di sini dapat melanjutkan, menahan, atau mengembalikan dokumen dalam proses layanan."],
        "sections": [
            {"title": "6.1 Signer Inbox", "body": ["Signer Inbox menampilkan permohonan yang sedang menunggu keputusan penandatangan dan menjadi titik peralihan dari proses pengelola ke tahap penandatanganan."], "figure": ("02-signer-inbox-hero.png", "Gambar 20. Signer Inbox")},
            {"title": "6.2 Detail Permohonan", "body": ["Setelah satu permohonan dibuka, penandatangan melihat halaman detail yang memuat ringkasan permohonan dan konteks dokumen yang sedang dimintakan keputusan."], "figure": ("05-ringkasan-permohonan-signer.png", "Gambar 21. Detail permohonan pada portal penandatangan")},
            {"title": "6.3 Preview Dokumen dan Data Snapshot", "body": ["Preview dokumen dan data snapshot membantu penandatangan membaca dokumen yang akan disetujui sambil tetap memahami konteks layanan."], "figure": ("06-preview-dokumen-dan-snapshot-signer.png", "Gambar 22. Preview dokumen dan data snapshot")},
            {"title": "6.4 Keputusan Penandatangan", "body": ["Keputusan penandatangan dicatat melalui form tindakan yang tersedia pada halaman permohonan dan akan memengaruhi status layanan secara langsung."], "figure": ("08-form-keputusan-signer.png", "Gambar 23. Form keputusan penandatangan")},
            {"title": "6.5 Dampak Keputusan terhadap Alur Layanan", "body": ["Setelah keputusan disimpan, status layanan dan antrian berikutnya akan berubah sesuai alur yang berlaku."], "figure": ("11-status-atau-dampak-keputusan-signer.png", "Gambar 24. Dampak keputusan penandatangan terhadap alur layanan")},
        ],
    },
    {
        "title": "BAB VII ALUR UMUM LAYANAN DAN MONITORING STATUS",
        "intro": ["Bab ini menjelaskan hubungan antarportal dan cara membaca progress layanan dari titik awal hingga penyelesaian. Pada level buku umum, bagian ini penting agar pembaca dapat melihat satu permohonan sebagai proses lintas peran."],
        "sections": [
            {"title": "7.1 Alur Layanan dari Pemohon ke Pengelola", "body": ["Alur umum layanan dimulai dari pemohon yang membaca informasi layanan, mengajukan permohonan, lalu menunggu verifikasi dan tindak lanjut dari pengelola."], "figure": ("35-diagram-alur-umum-layanan.png", "Gambar 25. Alur layanan dari pemohon ke pengelola")},
            {"title": "7.2 Alur Layanan ke Penandatangan", "body": ["Pada layanan yang membutuhkan persetujuan formal, alur proses berlanjut dari pengelola ke portal penandatangan sampai hasil keputusan memengaruhi langkah berikutnya."], "figure": ("11-status-atau-dampak-keputusan-signer.png", "Gambar 26. Alur layanan ke penandatangan")},
            {"title": "7.3 Monitoring Status dan Titik Kontrol", "body": ["Monitoring status dan titik kontrol membantu pembaca memahami di mana sebuah permohonan sedang berada, siapa yang memegang proses saat ini, dan bagian mana yang paling sering menjadi sumber ketidaksesuaian alur."], "figure": ("48-audit-aktivitas-admin.png", "Gambar 27. Monitoring status dan titik kontrol layanan")},
        ],
    },
    {
        "title": "BAB VIII ADMINISTRASI MASTER OLEH SUPERADMIN",
        "intro": ["Bab ini difokuskan pada area administrasi master yang dikelola oleh superadmin. Dalam buku umum, pembahasannya dibuat ringkas agar pembaca cukup memahami modul strategis yang menjaga konsistensi sistem."],
        "sections": [
            {"title": "8.1 Pengguna dan Peran", "body": ["Modul pengguna dan peran menjadi fondasi pengelolaan hak akses pada Website ULT. Melalui modul ini, superadmin memastikan bahwa setiap akun melihat menu, portal, dan aksi yang sesuai dengan kewenangannya."], "figure": ("41-daftar-pengguna-admin.png", "Gambar 28. Pengguna dan peran pada portal admin")},
            {"title": "8.2 Layanan dan Kategori Layanan", "body": ["Modul layanan dan kategori layanan memengaruhi bagaimana informasi tampil pada area publik dan bagaimana pengajuan dibentuk pada portal pemohon."], "figure": ("43-daftar-layanan-admin.png", "Gambar 29. Layanan dan kategori layanan")},
            {"title": "8.3 Setup Dokumen dan Placeholder", "body": ["Setup dokumen dan placeholder menjadi area penting ketika layanan menghasilkan dokumen secara otomatis dan perlu dipastikan siap sebelum layanan digunakan."], "figure": ("68-ringkasan-setup-layanan-dokumen-admin.png", "Gambar 30. Setup dokumen dan placeholder")},
            {"title": "8.4 Template Nomor Surat", "body": ["Template nomor surat menjadi modul penting untuk menjaga konsistensi keluaran resmi pada layanan yang membutuhkan penomoran. Melalui modul ini, superadmin dapat meninjau format, sequence, dan riwayat penggunaan nomor surat."], "figure": ("45-daftar-format-nomor-surat-admin.png", "Gambar 31. Template nomor surat pada portal admin")},
            {"title": "8.5 Panduan Pengguna", "body": ["Modul panduan pengguna membantu distribusi dokumentasi ke pengguna internal maupun publik. Pada level superadmin, modul ini penting untuk memastikan file panduan, visibilitas akses, dan status publikasi selalu selaras dengan versi sistem terbaru."], "figure": ("44-daftar-panduan-pengguna-admin.png", "Gambar 32. Panduan pengguna pada portal admin")},
            {"title": "8.6 CMS, Publikasi, dan Audit", "body": ["CMS, publikasi, dan audit menjadi instrumen kontrol yang membantu superadmin menjaga konsistensi website serta memantau aktivitas penting pada sistem."], "figure": ("48-audit-aktivitas-admin.png", "Gambar 33. CMS, publikasi, dan audit")},
        ],
    },
    {
        "title": "BAB IX KETENTUAN UMUM PENGGUNAAN",
        "intro": ["Bab ini memuat ketentuan umum yang perlu diperhatikan oleh pengguna Website ULT dari berbagai peran agar penggunaan akun, data, dan proses layanan tetap tertib dan aman."],
        "sections": [
            {"title": "9.1 Ketentuan Penggunaan Akun", "body": ["Setiap akun harus digunakan sesuai kewenangan yang diberikan di dalam sistem. Pengguna tidak disarankan meminjamkan akun atau membiarkan sesi login aktif tanpa pengawasan."]},
            {"title": "9.2 Ketentuan Pengisian Data dan Unggah Dokumen", "body": ["Data yang diisi pada form layanan dan file yang diunggah harus sesuai dengan konteks layanan yang dipilih karena ketidaksesuaian akan memengaruhi kecepatan verifikasi dan kualitas hasil layanan."]},
            {"title": "9.3 Ketentuan Verifikasi, Persetujuan, dan Finalisasi", "body": ["Setiap tindakan verifikasi, keputusan penandatangan, dan finalisasi hasil perlu dilakukan secara cermat sesuai kewenangan yang dimiliki."]},
            {"title": "9.4 Ketentuan Keamanan dan Kehati-hatian", "body": ["Pengguna perlu berhati-hati ketika mengunggah file, mengubah konfigurasi layanan, memproses nomor surat, atau mempublikasikan perubahan yang berdampak ke pengguna lain."]},
        ],
    },
    {
        "title": "BAB X PENUTUP",
        "intro": ["Buku panduan umum ini disusun sebagai peta orientasi penggunaan Website ULT FKIP Universitas Lampung. Dengan memahami area publik, portal berbasis peran, alur layanan, dan modul administratif penting, pembaca diharapkan dapat membaca sistem secara lebih utuh sebelum bekerja lebih rinci pada satu peran tertentu."],
        "sections": [
            {"title": "10.1 Penegasan Akhir", "body": ["Pada akhirnya, buku panduan umum tidak dimaksudkan untuk menggantikan buku khusus, melainkan untuk menjelaskan hubungan antarmodul dan memberi gambaran menyeluruh tentang Website ULT. Dengan memahami peta besarnya terlebih dahulu, pengguna akan lebih mudah menempatkan setiap langkah teknis ke dalam konteks layanan yang benar."]},
        ],
    },
]


def set_run_font(run, size=10, bold=False, italic=False, color=COLOR_TEXT):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def set_doc_defaults(doc):
    for section in doc.sections:
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.top_margin = Cm(TOP_MARGIN_CM)
        section.bottom_margin = Cm(BOTTOM_MARGIN_CM)
        section.left_margin = Cm(LEFT_MARGIN_CM)
        section.right_margin = Cm(RIGHT_MARGIN_CM)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(4)

    for name, size in [("Heading 1", 12), ("Heading 2", 11)]:
        style = doc.styles[name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = COLOR_PRIMARY
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = name == "Heading 2"


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    if level == 1 and text.startswith("BAB "):
        p.paragraph_format.page_break_before = True
    run = p.add_run(text)
    set_run_font(run, size=12 if level == 1 else 11, bold=True, color=COLOR_PRIMARY)


def add_body(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10)


def add_list_intro(doc, title):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f"Beberapa hal penting yang perlu diperhatikan pada bagian {title.lower()} adalah sebagai berikut.")
    set_run_font(run, size=10)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = LIST_TEXT_INDENT
        p.paragraph_format.first_line_indent = LIST_HANGING_INDENT
        p.paragraph_format.tab_stops.add_tab_stop(LIST_TEXT_INDENT)
        bullet = p.add_run("\u2022\t")
        set_run_font(bullet, size=10, bold=True, color=COLOR_PRIMARY)
        run = p.add_run(item)
        set_run_font(run, size=10)


def add_steps(doc, items):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = LIST_TEXT_INDENT
        p.paragraph_format.first_line_indent = LIST_HANGING_INDENT
        p.paragraph_format.tab_stops.add_tab_stop(LIST_TEXT_INDENT)
        number = p.add_run(f"{index}.\t")
        set_run_font(number, size=10, bold=True, color=COLOR_PRIMARY)
        run = p.add_run(item)
        set_run_font(run, size=10)


def add_field_note(doc, title, instruction):
    add_heading(doc, title, 1)
    p = doc.add_paragraph()
    run = p.add_run(instruction)
    set_run_font(run, size=9)


def set_keep_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(85)
    run = p.add_run("BUKU PANDUAN UMUM WEBSITE ULT")
    set_run_font(run, size=16, bold=True, color=COLOR_PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("WEBSITE UNIT LAYANAN TERPADU (ULT)\nFAKULTAS KEGURUAN DAN ILMU PENDIDIKAN\nUNIVERSITAS LAMPUNG")
    set_run_font(run, size=13, bold=True, color=COLOR_PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    run = p.add_run("Panduan Orientasi Umum Penggunaan Sistem")
    set_run_font(run, size=11, bold=True, color=COLOR_PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Dokumen ini menjelaskan gambaran umum Website ULT,\nhubungan antarportal, dan modul utama yang perlu dipahami sebelum menggunakan buku panduan khusus.")
    set_run_font(run, size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(150)
    run = p.add_run("Bandar Lampung\n2026")
    set_run_font(run, size=10)


def add_front_matter(doc):
    add_heading(doc, "KATA PENGANTAR", 1)
    add_body(doc, "Buku panduan umum ini disusun sebagai acuan orientasi penggunaan Website Unit Layanan Terpadu (ULT) FKIP Universitas Lampung. Pembaruan dokumen dilakukan agar isi panduan selaras dengan antarmuka dan alur sistem terbaru, sekaligus menjaga hubungan yang jelas antara area publik, portal berbasis peran, dan modul administratif sistem.")
    add_body(doc, "Melalui buku ini, pembaca diharapkan dapat memahami struktur besar Website ULT secara lebih terarah sebelum mempelajari buku panduan pemohon, pengelola layanan, atau penandatangan secara lebih rinci. Tangkapan layar yang digunakan pada setiap bab dipilih dari aset kerja terbaru agar uraian yang disajikan lebih mudah diikuti.")
    add_body(doc, "Dokumen ini juga berfungsi sebagai pelengkap dokumentasi sistem sehingga dapat dipakai sebagai referensi saat orientasi pengguna internal, peninjauan tugas akhir, maupun evaluasi implementasi layanan digital pada lingkungan ULT FKIP Universitas Lampung.")

    doc.add_page_break()
    add_field_note(doc, "DAFTAR ISI", 'Daftar isi dapat diperbarui otomatis melalui fitur "Update Field" pada Microsoft Word setelah peninjauan akhir dokumen.')
    doc.add_page_break()
    add_field_note(doc, "DAFTAR GAMBAR", 'Daftar gambar dapat diperbarui otomatis melalui fitur "Update Field" pada Microsoft Word setelah caption final diperiksa.')


def add_role_table(doc, title, rows):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(title)
    set_run_font(run, size=9, italic=True)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Peran", "Fungsi Utama", "Contoh Pengguna"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=9, bold=True)
    for role, function, example in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, [role, function, example]):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            run = p.add_run(text)
            set_run_font(run, size=9)
    doc.add_paragraph()


def resolve_image_dimensions(path: Path):
    with Image.open(path) as image:
        width_px, height_px = image.size
    aspect_ratio = width_px / max(height_px, 1)
    content_width = PAGE_WIDTH_CM - LEFT_MARGIN_CM - RIGHT_MARGIN_CM
    max_width = min(content_width - 0.2, 12.0)
    max_height = min(PAGE_HEIGHT_CM - TOP_MARGIN_CM - BOTTOM_MARGIN_CM - 1.2, 10.8)
    if aspect_ratio >= 4.5:
        target_width = 9.4
    elif aspect_ratio >= 2.8:
        target_width = 10.0
    elif aspect_ratio >= 1.6:
        target_width = 10.8
    elif aspect_ratio >= 1.0:
        target_width = 10.4
    elif aspect_ratio >= 0.72:
        target_width = 9.2
    else:
        target_width = 7.8
    target_width = min(target_width, max_width)
    target_height = target_width / max(aspect_ratio, 0.01)
    if target_height > max_height:
        scale = max_height / target_height
        target_width *= scale
        target_height = max_height
    return Cm(round(target_width, 2)), Cm(round(target_height, 2))


def add_figure(doc, filename, caption):
    path = ASSETS / filename
    if not path.exists():
        add_body(doc, f"[Gambar tidak ditemukan: {filename}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_keep_next(p)
    width, height = resolve_image_dimensions(path)
    p.add_run().add_picture(str(path), width=width, height=height)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, size=8, italic=True)


def add_chapters(doc):
    for chapter in CHAPTERS:
        add_heading(doc, chapter["title"], 1)
        for paragraph in chapter.get("intro", []):
            add_body(doc, paragraph)
        if chapter.get("table_title") and chapter.get("table_rows"):
            add_role_table(doc, chapter["table_title"], chapter["table_rows"])
        for section in chapter.get("sections", []):
            add_heading(doc, section["title"], 2)
            for paragraph in section.get("body", []):
                add_body(doc, paragraph)
            if section.get("figure"):
                add_figure(doc, section["figure"][0], section["figure"][1])
            if section.get("bullets"):
                add_list_intro(doc, section["title"])
                add_bullets(doc, section["bullets"])
            if section.get("steps"):
                add_list_intro(doc, section["title"])
                add_steps(doc, section["steps"])


def main():
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc)
    doc.add_page_break()
    add_front_matter(doc)
    add_chapters(doc)
    doc.save(OUTPUT)
    doc.save(COVER_OUTPUT)
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    main()
