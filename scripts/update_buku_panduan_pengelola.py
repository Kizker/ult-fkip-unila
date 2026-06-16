from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image


ROOT = Path(r"c:/laragon/www/ult-fkip-unila")
OUTPUT = ROOT / "docs/buku-panduan/buku-panduan-pengelola-layanan/buku-panduan-pengelola-layanan.docx"
ASSETS = ROOT / "docs/buku-panduan/assets/screenshots"
FONT_NAME = "Poppins"
COLOR_PRIMARY = RGBColor(91, 33, 182)
COLOR_TEXT = RGBColor(31, 41, 55)
PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
TOP_MARGIN_CM = 2.5
BOTTOM_MARGIN_CM = 2.5
LEFT_MARGIN_CM = 3.0
RIGHT_MARGIN_CM = 2.5
LIST_TEXT_INDENT = Cm(1.2)
LIST_HANGING_INDENT = Cm(-0.6)


CHAPTERS = [
    {
        "title": "BAB I PENDAHULUAN",
        "intro": [
            "Bab ini menjelaskan alasan penyusunan buku panduan, tujuan, capaian, sasaran pengguna, dan ruang lingkup pembahasan portal pengelola layanan. Fokus utama dokumen diarahkan pada pengalaman pengguna internal saat memantau permohonan, memverifikasi data, menindaklanjuti workflow, mengelola layanan dokumen, serta menjaga modul administratif yang mendukung operasional Website ULT FKIP Universitas Lampung.",
        ],
        "sections": [
            {
                "title": "1.1 Latar Belakang",
                "body": [
                    "Website ULT tidak hanya mempermudah pemohon dalam mengajukan layanan, tetapi juga mengubah pola kerja pengelola internal yang menangani proses layanan di balik sistem. Pengelola tidak lagi bekerja dengan alur manual yang terpisah, melainkan melalui portal yang menghubungkan daftar permohonan, detail pengajuan, workflow, dokumen hasil, layanan dokumen, dan modul pendukung lain dalam satu area kerja.",
                    "Perubahan ini membuat kebutuhan panduan operasional menjadi semakin penting. Tanpa panduan yang jelas, pengelola berisiko memahami fitur secara parsial, padahal tindakan pada portal admin berpengaruh langsung terhadap status permohonan, kelengkapan dokumen, penomoran surat, hingga kesiapan layanan dokumen yang dipublikasikan kepada pemohon.",
                    "Buku panduan ini disusun agar pengelola memiliki acuan kerja yang lebih runtut dan selaras dengan implementasi sistem terbaru. Dengan demikian, portal admin dapat dipahami bukan sekadar sebagai kumpulan menu, tetapi sebagai pusat kendali proses layanan yang perlu dijalankan secara tertib, tepat, dan sesuai kewenangan.",
                ],
            },
            {
                "title": "1.2 Tujuan Buku Panduan",
                "body": [
                    "Buku panduan ini bertujuan memberikan acuan operasional yang lebih faktual bagi pengguna portal pengelola layanan. Dokumen ini juga berfungsi sebagai pelengkap dokumentasi produk agar alur pengelolaan permohonan, layanan dokumen, dan modul admin dapat dijelaskan secara konsisten kepada pengguna baru, penguji, maupun pengelola sistem.",
                    "Secara lebih rinci, tujuan penyusunan buku panduan ini dapat dijabarkan ke dalam beberapa poin berikut.",
                ],
                "bullets": [
                    "menjelaskan langkah penggunaan portal admin dari login sampai pengelolaan modul layanan;",
                    "membantu pengelola memahami daftar permohonan, detail permohonan, dan aksi workflow sesuai kewenangan;",
                    "menunjukkan area antarmuka penting pada setup layanan dokumen, nomor surat, dan modul pendukung;",
                    "membantu pengguna internal memahami hubungan antarfitur agar pemrosesan layanan berjalan lebih tertib.",
                ],
            },
            {
                "title": "1.3 Capaian Buku Panduan",
                "body": [
                    "Setelah mempelajari buku panduan ini, pengelola diharapkan tidak hanya mengenali tampilan portal admin, tetapi juga mampu menjalankan alur penggunaan sistem secara mandiri, tertib, dan sesuai kewenangan yang dimiliki pada proses layanan yang sedang aktif.",
                    "Secara lebih operasional, capaian buku panduan ini dapat dijabarkan ke dalam beberapa poin berikut.",
                ],
                "bullets": [
                    "mampu login dan masuk ke portal admin dengan benar;",
                    "mampu membaca daftar permohonan, membuka detail, dan memahami data pemohon, lampiran, serta status layanan;",
                    "mampu melakukan verifikasi, memberikan catatan, dan menindaklanjuti workflow sesuai kewenangan;",
                    "mampu mengelola setup layanan dokumen, nomor surat, dan modul pendukung yang tersedia pada portal admin.",
                ],
            },
            {
                "title": "1.4 Sasaran Pengguna Buku Panduan",
                "body": [
                    "Sasaran utama dokumen ini adalah pengguna internal yang memiliki akses ke portal admin dan terlibat dalam pengelolaan layanan. Pada implementasi saat ini, pengguna tersebut dapat mencakup pengelola unit, reviewer ULT, admin yang menangani layanan dokumen, serta pihak lain yang diberi kewenangan administratif pada modul tertentu.",
                ],
            },
            {
                "title": "1.5 Ruang Lingkup Buku Panduan",
                "body": [
                    "Pembahasan difokuskan pada portal pengelola layanan dan urutan penggunaan yang langsung berhubungan dengan akses masuk, pengelolaan permohonan, workflow, layanan dokumen, nomor surat, dan modul administratif pendukung. Area pemohon dan penandatangan hanya disebut seperlunya untuk menjelaskan konteks alur layanan dari sudut pandang pengelola.",
                ],
                "bullets": [
                    "akses masuk ke sistem dan navigasi menu admin;",
                    "pembacaan daftar permohonan, filter, detail, data, dan lampiran;",
                    "verifikasi, review, catatan, dan tindak lanjut workflow layanan;",
                    "setup layanan dokumen, placeholder, gate, signer, dan publish readiness;",
                    "preview output, nomor surat, kritik dan saran, serta pedoman placeholder.",
                ],
            },
        ],
    },
    {
        "title": "BAB II MENGENAL PORTAL PENGELOLA LAYANAN",
        "intro": [
            "Portal pengelola layanan merupakan area kerja internal yang dipakai untuk memantau, menindaklanjuti, dan mengendalikan proses layanan setelah permohonan diajukan oleh pemohon. Dari portal ini, pengelola dapat berpindah antara daftar permohonan, detail layanan, setup dokumen, nomor surat, dan modul administratif lain yang mendukung operasional sistem.",
        ],
        "sections": [
            {
                "title": "2.1 Portal Pengelola Layanan",
                "body": [
                    "Pada implementasi saat ini, portal pengelola berfungsi sebagai pusat kendali proses layanan. Halaman-halaman di dalamnya tidak hanya menampilkan daftar data, tetapi juga menyediakan konteks pengambilan keputusan, tindak lanjut workflow, pengelolaan output, serta pengaturan modul dokumen yang memengaruhi hasil layanan secara langsung.",
                    "Pengelola juga tidak selalu melihat fungsi yang sama pada setiap akun. Tampilan dan aksi yang muncul dipengaruhi oleh hak akses, sehingga portal admin perlu dipahami sebagai ruang kerja berbasis kewenangan, bukan sebagai kumpulan menu yang identik untuk semua pengguna internal.",
                ],
            },
            {
                "title": "2.2 Fitur Utama Portal Pengelola",
                "body": [
                    "Fitur utama portal pengelola dirancang untuk mendukung pemrosesan layanan dari tahap awal sampai modul administratif lanjut. Komponen inti pada area ini mencakup dashboard admin, daftar permohonan, detail permohonan, aksi workflow, setup layanan dokumen, template nomor surat, serta modul pendukung seperti kritik dan saran dan pedoman placeholder.",
                ],
                "bullets": [
                    "dashboard admin menampilkan orientasi awal dan jalur menuju area kerja yang paling sering dipakai;",
                    "daftar permohonan dan detail permohonan menjadi pusat pembacaan layanan yang sedang berjalan;",
                    "aksi workflow dipakai untuk verifikasi, review, catatan, dan tindak lanjut permohonan;",
                    "setup layanan dokumen memuat template, placeholder, gate, signer, dan publish readiness;",
                    "modul pendukung membantu pengelola menjaga konsistensi layanan dan data administratif.",
                ],
            },
            {
                "title": "2.3 Menu Kerja Utama pada Sidebar Admin",
                "body": [
                    "Sidebar admin menjadi penghubung utama antarhalaman. Dari menu ini, pengelola dapat berpindah dari daftar permohonan ke modul layanan, setup dokumen, nomor surat, kritik dan saran, atau area kerja lain tanpa harus kembali ke halaman awal setiap kali berpindah konteks kerja.",
                ],
                "figure": ("18-dashboard-pengelola.png", "Gambar 1. Halaman awal portal admin dan area menu utama"),
                "bullets": [
                    "gunakan sidebar sebagai peta kerja utama untuk memahami hubungan antar modul;",
                    "perhatikan bahwa menu yang tampil dapat berbeda sesuai hak akses masing-masing akun;",
                    "biasakan berpindah modul melalui sidebar agar alur kerja tetap konsisten dan mudah ditelusuri.",
                ],
            },
        ],
    },
    {
        "title": "BAB III AKSES MASUK DAN NAVIGASI PORTAL",
        "intro": [
            "Bab ini membahas tahap awal sebelum pengelola mulai bekerja pada portal admin, yaitu login ke Website ULT, masuk ke area kerja internal, dan memahami navigasi awal yang tersedia pada sidebar dan halaman admin.",
        ],
        "sections": [
            {
                "title": "3.1 Halaman Login",
                "body": [
                    "Halaman login menjadi gerbang autentikasi sebelum pengelola masuk ke portal sesuai hak aksesnya. Pengguna perlu memasukkan kredensial akun yang benar agar sistem dapat mengenali kewenangan yang dimiliki dan menampilkan area kerja yang sesuai.",
                ],
                "figure": ("02-halaman-login-pemohon.png", "Gambar 2. Halaman login Website ULT"),
            },
            {
                "title": "3.2 Langkah Masuk ke Portal Pengelola",
                "body": [
                    "Proses login pengelola dirancang ringkas, tetapi pengguna tetap perlu memastikan bahwa akun yang dipakai memang memiliki hak akses administratif pada modul yang akan dibuka.",
                ],
                "steps": [
                    "buka halaman login Website ULT melalui menu masuk yang tersedia pada website;",
                    "isi email atau identitas akun dan kata sandi dengan benar;",
                    "tekan tombol masuk lalu tunggu proses autentikasi selesai;",
                    "pastikan sistem mengarahkan Anda ke area admin atau portal pengelola yang sesuai.",
                ],
            },
            {
                "title": "3.3 Halaman Awal Portal Admin",
                "body": [
                    "Setelah login, pengelola diarahkan ke halaman awal portal admin. Halaman ini berfungsi sebagai titik orientasi pertama untuk membaca ringkasan kerja, mengenali konteks layanan yang sedang aktif, dan menentukan area mana yang perlu dibuka lebih dahulu.",
                ],
                "figure": ("19-ringkasan-kpi-dan-antrian.png", "Gambar 3. Ringkasan awal pada dashboard pengelola"),
            },
            {
                "title": "3.4 Navigasi Menu pada Sidebar",
                "body": [
                    "Setelah orientasi awal dipahami, perpindahan kerja berikutnya dilakukan melalui sidebar. Navigasi ini memudahkan pengguna berpindah antar kelompok kerja tanpa mengulang proses pencarian menu, misalnya dari pemrosesan permohonan ke setup dokumen atau modul pendukung lain.",
                ],
                "figure": ("43-daftar-layanan-admin.png", "Gambar 4. Sidebar admin dan menu kerja utama"),
            },
        ],
    },
    {
        "title": "BAB IV PENGELOLAAN PERMOHONAN LAYANAN",
        "intro": [
            "Setelah masuk ke portal admin, aktivitas utama pengelola berpusat pada permohonan yang diajukan pemohon. Bab ini memuat tahapan membaca daftar permohonan, melakukan pencarian, membuka detail, serta memeriksa data dan lampiran yang menjadi dasar verifikasi layanan.",
        ],
        "sections": [
            {
                "title": "4.1 Melihat Daftar Permohonan",
                "body": [
                    "Daftar permohonan menyajikan layanan yang masuk ke sistem dalam satu tampilan kerja. Dari halaman ini, pengelola dapat memantau permohonan yang sedang berjalan dan menentukan entri mana yang perlu diperiksa lebih dahulu.",
                ],
                "figure": ("20-daftar-permohonan-pengelola.png", "Gambar 5. Daftar permohonan pada portal admin"),
            },
            {
                "title": "4.2 Pencarian dan Filter Permohonan",
                "body": [
                    "Ketika jumlah permohonan mulai bertambah, pengelola memerlukan sarana pencarian dan filter untuk menemukan data yang relevan secara cepat. Fitur ini membantu mempercepat penelusuran berdasarkan kata kunci atau kondisi tertentu pada layanan yang sedang diproses.",
                    "Pada praktiknya, area pencarian dan filter berada dalam konteks yang sama dengan halaman daftar permohonan. Karena itu, pengelola sebaiknya membaca kedua fungsi ini sebagai satu kesatuan kerja: daftar dipakai untuk melihat antrian, sedangkan pencarian dan filter dipakai untuk mempersempit fokus pada permohonan yang benar-benar ingin ditindaklanjuti.",
                ],
                "bullets": [
                    "gunakan kolom pencarian untuk menemukan permohonan berdasarkan kata kunci yang relevan;",
                    "manfaatkan filter untuk menyaring layanan atau kondisi proses tertentu;",
                    "selalu baca hasil filter bersama daftar permohonan agar konteks antrian tetap utuh.",
                ],
            },
            {
                "title": "4.3 Membuka Detail Permohonan",
                "body": [
                    "Setelah satu permohonan dipilih, pengelola masuk ke halaman detail yang menjadi pusat keputusan operasional. Seluruh konteks layanan, identitas pemohon, status, dan area aksi dipusatkan pada halaman ini.",
                ],
                "figure": ("22-detail-permohonan-pengelola.png", "Gambar 6. Halaman detail permohonan pada portal admin"),
            },
            {
                "title": "4.4 Membaca Data Pemohon, Lampiran, dan Status",
                "body": [
                    "Pemeriksaan data dan lampiran dilakukan sebelum keputusan workflow diambil. Pada tahap ini, pengelola membandingkan data isian dengan dokumen pendukung untuk memastikan permohonan layak diteruskan, dikembalikan, atau ditindaklanjuti sesuai prosedur.",
                ],
                "figure": ("23-data-dan-lampiran-permohonan.png", "Gambar 7. Data pemohon, lampiran, dan status permohonan"),
                "bullets": [
                    "baca identitas pemohon dan konteks layanan sebelum masuk ke aksi workflow;",
                    "periksa lampiran yang tersedia untuk memastikan dokumen pendukung sudah sesuai;",
                    "gunakan status yang tampil sebagai penanda posisi permohonan pada alur layanan.",
                ],
            },
        ],
    },
    {
        "title": "BAB V VERIFIKASI, REVIEW, DAN TINDAK LANJUT",
        "intro": [
            "Tahap verifikasi dan review merupakan inti kerja pengelola karena pada bagian inilah permohonan dinilai, diteruskan, dikembalikan, atau ditandai untuk tindak lanjut tertentu. Setiap tindakan dilakukan melalui area workflow yang terhubung langsung dengan alur layanan di dalam sistem.",
        ],
        "sections": [
            {
                "title": "5.1 Melakukan Verifikasi Awal",
                "body": [
                    "Verifikasi awal dilakukan dengan membaca detail permohonan, memeriksa data pemohon, dan membandingkan lampiran yang tersedia dengan kebutuhan layanan. Pada tahap ini, pengelola perlu memastikan bahwa konteks permohonan sudah dipahami dengan benar sebelum menekan aksi workflow apa pun pada halaman detail.",
                    "Pembacaan awal ini penting karena tindakan pada portal admin tidak dapat dipisahkan dari status dan kewenangan yang sedang aktif. Jika pengelola salah membaca konteks layanan, permohonan dapat diteruskan ke tahap yang tidak tepat atau dikembalikan tanpa alasan yang cukup jelas.",
                ],
            },
            {
                "title": "5.2 Memberikan Catatan atau Umpan Balik",
                "body": [
                    "Dalam kondisi tertentu, pengelola perlu memberikan catatan kepada pemohon apabila terdapat kekurangan, ketidaksesuaian data, atau lampiran yang belum memenuhi kebutuhan layanan. Catatan ini sebaiknya ditulis dengan bahasa yang ringkas, jelas, dan langsung menunjuk bagian yang harus ditindaklanjuti.",
                    "Bagi pemohon, catatan dari pengelola akan menjadi dasar saat melakukan perbaikan. Karena itu, catatan yang terlalu umum justru dapat memperlambat proses, sedangkan catatan yang spesifik akan membantu perbaikan dilakukan dengan lebih cepat dan tepat.",
                ],
            },
            {
                "title": "5.3 Mengembalikan Permohonan untuk Perbaikan",
                "body": [
                    "Jika permohonan belum dapat dilanjutkan, pengelola dapat mengembalikannya untuk diperbaiki melalui alur yang tersedia pada workflow layanan. Tindakan ini perlu dilakukan secara jelas agar pemohon memahami bagian mana yang harus dilengkapi, diperbarui, atau diunggah ulang sebelum proses dilanjutkan kembali.",
                    "Pada praktiknya, pengembalian untuk perbaikan sebaiknya dipakai ketika permohonan masih layak diproses setelah koreksi dilakukan. Dengan demikian, tindakan ini berbeda dari penolakan final karena fokusnya adalah memperbaiki kelengkapan atau ketepatan data agar layanan dapat masuk kembali ke alur proses yang benar.",
                ],
            },
            {
                "title": "5.4 Menindaklanjuti Permohonan Sesuai Workflow",
                "body": [
                    "Apabila permohonan dinilai sesuai, pengelola dapat menindaklanjutinya ke tahap berikutnya sesuai alur yang berlaku pada sistem. Tindakan ini harus dibaca sebagai bagian dari workflow layanan, bukan sekadar persetujuan umum, karena setiap aksi admin akan memindahkan permohonan ke status dan aktor berikutnya yang berbeda.",
                    "Sebelum aksi dijalankan, pengelola perlu memastikan bahwa status saat ini sudah benar, kewenangan yang dimiliki memang mengizinkan tindakan tersebut, dan seluruh data yang dibutuhkan pada tahap sekarang telah diperiksa dengan cukup.",
                ],
                "figure": ("24-aksi-verifikasi-dan-workflow.png", "Gambar 8. Aksi verifikasi dan tindak lanjut pada permohonan"),
            },
            {
                "title": "5.5 Memahami Dampak Tindakan terhadap Workflow",
                "body": [
                    "Setiap tindakan pada workflow akan memengaruhi status permohonan, pihak yang menerima giliran proses berikutnya, dan kesiapan layanan untuk masuk ke tahap lanjutan seperti layanan dokumen atau penandatanganan. Karena itu, pengelola perlu memahami bahwa aksi pada portal admin bukan sekadar perubahan tampilan, tetapi perubahan nyata pada alur sistem.",
                ],
                "bullets": [
                    "verifikasi dan tindak lanjut workflow dapat meneruskan permohonan ke tahap berikutnya sesuai peran yang sedang aktif;",
                    "catatan atau pengembalian untuk perbaikan mengembalikan tanggung jawab ke pemohon atau pihak sebelumnya;",
                    "perubahan status sesudah aksi dijalankan perlu dibaca sebagai indikator siapa yang memegang proses berikutnya;",
                    "pemahaman alur ini membantu pengelola menilai apakah permohonan siap masuk ke tahap dokumen, nomor surat, atau penandatanganan.",
                ],
            },
        ],
    },
    {
        "title": "BAB VI PENGELOLAAN LAYANAN DOKUMEN",
        "intro": [
            "Selain memproses permohonan, pengelola atau admin tertentu juga berhadapan dengan area setup layanan dokumen. Bab ini menjelaskan konfigurasi template, placeholder, gate, signer, dan publish readiness yang menjadi fondasi generator dokumen pada layanan tertentu.",
        ],
        "sections": [
            {
                "title": "6.1 Membuka Setup Layanan Dokumen",
                "body": [
                    "Halaman setup layanan dokumen menjadi titik audit utama untuk melihat kesiapan konfigurasi sebuah layanan. Dari halaman ini, pengelola dapat meninjau ringkasan readiness sebelum masuk ke panel yang lebih spesifik.",
                ],
                "figure": ("68-ringkasan-setup-layanan-dokumen-admin.png", "Gambar 9. Halaman setup layanan dokumen"),
            },
            {
                "title": "6.2 Mengunggah Template Dokumen",
                "body": [
                    "Template utama dipakai sebagai sumber pembentukan dokumen layanan. Pada tahap ini, pengelola perlu memastikan file template yang diunggah benar, final, dan sesuai dengan mode dokumen yang dipakai layanan terkait.",
                ],
                "figure": ("69-upload-template-layanan-dokumen-admin.png", "Gambar 10. Area unggah template dokumen"),
            },
            {
                "title": "6.3 Mengekstrak dan Memetakan Placeholder",
                "body": [
                    "Sesudah template tersedia, pengelola masuk ke tahap ekstraksi dan mapping placeholder. Pada bagian ini dipastikan bahwa token pada dokumen berhasil dibaca sistem dan dipetakan ke sumber data yang sesuai.",
                ],
                "figure": ("70-mapping-placeholder-layanan-dokumen-admin.png", "Gambar 11. Ekstraksi dan mapping placeholder dokumen"),
            },
            {
                "title": "6.4 Mengatur Gate Proses",
                "body": [
                    "Gate proses menentukan siapa yang memegang verifikasi awal dan langkah penting lain seperti input nomor surat pada layanan dokumen. Karena itu, panel gate perlu dipahami sebagai pengaturan kendali awal yang menghubungkan konfigurasi layanan dengan pelaksana proses di lapangan.",
                    "Pada tahap ini, pengelola tidak hanya mengisi nama peran atau langkah gate, tetapi juga memastikan bahwa alur awal layanan memang dibuka oleh unit atau petugas yang tepat. Kesalahan pada gate dapat membuat permohonan berhenti pada tahap yang salah atau berpindah ke aktor yang tidak sesuai dengan alur yang direncanakan.",
                ],
            },
            {
                "title": "6.5 Mengatur Signer Dokumen",
                "body": [
                    "Panel signer dipakai untuk menyusun rantai penandatangan pada dokumen. Urutan, peran, dan label signer yang disusun di sini akan berpengaruh langsung terhadap alur persetujuan dokumen pada tahap berikutnya.",
                    "Pengelola perlu memastikan bahwa urutan signer sesuai dengan kebutuhan layanan dan struktur persetujuan yang berlaku. Jika urutan, peran, atau label signer tidak tepat, dokumen dapat masuk ke tahap penandatanganan yang keliru, menyebabkan kebingungan pada penandatangan maupun keterlambatan pada proses akhir.",
                ],
            },
            {
                "title": "6.6 Mempublikasikan Layanan Dokumen",
                "body": [
                    "Publish readiness membantu pengelola memastikan bahwa layanan dokumen telah cukup siap sebelum dipublikasikan. Panel ini sebaiknya dibaca sebagai ringkasan kontrol mutu akhir sebelum layanan digunakan secara operasional.",
                    "Sebelum layanan dipublikasikan, pengelola sebaiknya membaca kembali kesiapan template, placeholder, gate, signer, dan hasil pengujian dasar yang sudah dilakukan. Dengan cara ini, proses publish tidak sekadar menjadi tindakan administratif, tetapi juga penegasan bahwa layanan telah siap dipakai tanpa menimbulkan gangguan pada pemohon maupun pihak internal berikutnya.",
                ],
            },
        ],
    },
    {
        "title": "BAB VII PENGELOLAAN OUTPUT DAN NOMOR SURAT",
        "intro": [
            "Pada tahap tertentu, pengelola tidak hanya berhenti pada verifikasi permohonan, tetapi juga berhadapan dengan output layanan dan penomoran surat. Bab ini menjelaskan area yang berkaitan dengan hasil dokumen dan pengelolaan nomor surat pada portal admin.",
        ],
        "sections": [
            {
                "title": "7.1 Meninjau Output atau Preview Dokumen",
                "body": [
                    "Preview atau output dokumen membantu pengelola memeriksa hasil layanan sebelum didistribusikan lebih lanjut. Bagian ini penting karena menjadi titik pemeriksaan akhir terhadap isi dokumen yang telah diproses sistem.",
                ],
                "figure": ("27-halaman-finalisasi-dokumen.png", "Gambar 12. Preview atau output dokumen layanan"),
            },
            {
                "title": "7.2 Mengelola Hasil Layanan",
                "body": [
                    "Pada beberapa layanan, pengelola tidak hanya meninjau preview, tetapi juga perlu mengelola hasil layanan yang sudah siap dipakai atau diakses pemohon. Tahap ini dapat mencakup unggah hasil, pengecekan ulang file keluaran, dan memastikan bahwa dokumen yang tersedia benar-benar sesuai dengan konteks layanan yang diproses.",
                    "Dengan demikian, subbab ini berbeda dari tahap preview. Jika preview berfungsi untuk memeriksa tampilan hasil sebelum final, maka pengelolaan hasil layanan berfokus pada kesiapan dokumen yang benar-benar akan diterima, diunduh, atau ditindaklanjuti oleh pengguna berikutnya.",
                ],
            },
            {
                "title": "7.3 Mengelola Template Nomor Surat",
                "body": [
                    "Template nomor surat menjadi bagian penting dari standardisasi keluaran resmi. Melalui modul ini, pengelola dapat meninjau dan mengatur pola penomoran sesuai unit dan konteks dokumen yang digunakan.",
                ],
                "figure": ("45-daftar-format-nomor-surat-admin.png", "Gambar 13. Halaman template nomor surat"),
            },
            {
                "title": "7.4 Meninjau History Nomor Surat",
                "body": [
                    "History nomor surat membantu pengelola membaca jejak penerbitan nomor dan memastikan konsistensi penggunaan format pada dokumen yang sudah diterbitkan.",
                    "Dalam praktik penggunaan, history ini sebaiknya dibaca sebagai kelanjutan dari halaman template nomor surat. Setelah pola format dipahami, pengelola dapat meninjau jejak nomor yang sudah terbit untuk memastikan bahwa penggunaan format berjalan konsisten antar dokumen dan antar unit yang terlibat.",
                    "Dari sudut pandang audit, history nomor surat juga membantu menelusuri apakah sebuah nomor pernah diterbitkan, dipakai pada jenis dokumen yang benar, dan mengikuti pola yang seharusnya berlaku pada unit terkait. Karena itu, fungsi history tidak hanya informatif, tetapi juga penting untuk kontrol operasional.",
                ],
            },
        ],
    },
    {
        "title": "BAB VIII PENGELOLAAN MODUL PENDUKUNG",
        "intro": [
            "Portal admin juga memuat modul pendukung yang tidak selalu berada pada alur verifikasi permohonan, tetapi tetap penting untuk operasional layanan. Bab ini merangkum modul yang membantu pengelola menjaga kualitas layanan dan memahami hubungan antararea kerja.",
        ],
        "sections": [
            {
                "title": "8.1 Mengelola Kritik dan Saran",
                "body": [
                    "Modul kritik dan saran berguna untuk membaca masukan yang datang dari pengguna. Walaupun tidak selalu dipakai setiap saat, modul ini penting untuk memahami masalah layanan dari sudut pandang publik dan menindaklanjutinya secara administratif.",
                ],
                "figure": ("61-daftar-kritik-dan-saran-admin.png", "Gambar 14. Halaman kritik dan saran pada portal admin"),
            },
            {
                "title": "8.2 Membaca Pedoman Placeholder",
                "body": [
                    "Pedoman placeholder membantu pengelola memahami aturan teknis penulisan token pada template dokumen. Modul ini penting terutama ketika terjadi perubahan template atau saat layanan dokumen baru sedang disiapkan.",
                ],
                "figure": ("52-pedoman-placeholder-layanan-dokumen.png", "Gambar 15. Pedoman placeholder layanan dokumen"),
            },
            {
                "title": "8.3 Memahami Hubungan Portal Pengelola dengan Signer Inbox",
                "body": [
                    "Walaupun Signer Inbox merupakan area kerja penandatangan, pengelola tetap perlu memahami hubungannya dengan portal admin. Tindakan yang dilakukan pengelola pada workflow layanan dapat memengaruhi kapan dokumen masuk ke tahap penandatanganan, siapa penandatangan aktif berikutnya, dan kapan permohonan berpindah ke inbox signer.",
                ],
                "bullets": [
                    "portal admin dan signer inbox berada pada rantai proses yang saling terhubung;",
                    "keputusan pengelola pada tahap verifikasi dapat menentukan kapan dokumen berpindah ke tahap penandatanganan;",
                    "pemahaman hubungan ini membantu pengelola membaca workflow secara lebih utuh.",
                ],
            },
        ],
    },
    {
        "title": "BAB IX KETENTUAN UMUM PENGGUNAAN",
        "intro": [
            "Bab ini memuat ketentuan umum yang perlu diperhatikan oleh pengguna portal pengelola layanan. Tujuannya adalah menjaga konsistensi penggunaan akun, ketelitian verifikasi, kehati-hatian pada dokumen dan nomor surat, serta keamanan penggunaan sistem secara umum.",
        ],
        "sections": [
            {
                "title": "9.1 Ketentuan Penggunaan Akun",
                "body": [
                    "Akun pengelola harus digunakan sesuai kewenangan yang diberikan dalam sistem. Pengguna tidak disarankan meminjamkan akun atau membiarkan sesi login aktif tanpa pengawasan karena setiap tindakan administratif akan tercatat sebagai aktivitas pengguna yang sedang login.",
                ],
            },
            {
                "title": "9.2 Ketentuan Verifikasi dan Pemrosesan",
                "body": [
                    "Setiap verifikasi harus dilakukan secara cermat dengan membandingkan data permohonan dan dokumen pendukung yang tersedia. Pengelola juga perlu menghindari tindakan di luar kewenangan atau melompati tahapan yang telah ditentukan workflow sistem.",
                ],
            },
            {
                "title": "9.3 Ketentuan Pengelolaan Dokumen dan Nomor Surat",
                "body": [
                    "Dokumen hasil, template dokumen, placeholder, dan nomor surat harus diperlakukan sebagai bagian sensitif operasional. Setiap perubahan perlu diperiksa kembali sebelum dinyatakan siap dipakai atau dipublikasikan.",
                ],
            },
            {
                "title": "9.4 Ketentuan Kehati-hatian dan Keamanan Penggunaan",
                "body": [
                    "Pengelola perlu berhati-hati ketika mengunggah file, mengubah konfigurasi layanan dokumen, atau memproses nomor surat. Setelah selesai menggunakan portal, sesi login sebaiknya diakhiri dengan logout agar risiko penggunaan akun tanpa izin dapat dikurangi.",
                ],
            },
        ],
    },
    {
        "title": "BAB X PENUTUP",
        "intro": [
            "Buku panduan pengelola layanan ini disusun sebagai acuan operasional penggunaan portal admin pada Website ULT FKIP Universitas Lampung. Dengan memahami daftar permohonan, detail layanan, workflow, layanan dokumen, nomor surat, dan modul pendukung, pengelola diharapkan dapat menjalankan perannya secara lebih tertib dan konsisten sesuai kebutuhan operasional sistem.",
            "Pembaruan isi dan gambar pada dokumen ini perlu terus dilakukan ketika antarmuka atau alur kerja sistem mengalami perubahan. Karena itu, buku panduan ini sebaiknya diperlakukan sebagai dokumen kerja yang hidup dan selalu diselaraskan dengan implementasi portal admin terbaru.",
        ],
    },
]


def set_run_font(run, size=10, bold=False, italic=False, color=COLOR_TEXT):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def set_doc_defaults(doc):
    for section in doc.sections:
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.top_margin = Cm(TOP_MARGIN_CM)
        section.bottom_margin = Cm(BOTTOM_MARGIN_CM)
        section.left_margin = Cm(LEFT_MARGIN_CM)
        section.right_margin = Cm(RIGHT_MARGIN_CM)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(4)

    for name, size in [("Heading 1", 12), ("Heading 2", 11)]:
        style = doc.styles[name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = COLOR_PRIMARY
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = name == "Heading 2"


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    if level == 1 and text.startswith("BAB "):
        p.paragraph_format.page_break_before = True
    run = p.add_run(text)
    set_run_font(run, size=12 if level == 1 else 11, bold=True, color=COLOR_PRIMARY)


def add_body(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = LIST_TEXT_INDENT
        p.paragraph_format.first_line_indent = LIST_HANGING_INDENT
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.tab_stops.add_tab_stop(LIST_TEXT_INDENT)
        bullet = p.add_run("•\t")
        set_run_font(bullet, size=10, bold=True, color=COLOR_PRIMARY)
        body = p.add_run(item)
        set_run_font(body, size=10)


def add_steps(doc, items):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = LIST_TEXT_INDENT
        p.paragraph_format.first_line_indent = LIST_HANGING_INDENT
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.tab_stops.add_tab_stop(p.paragraph_format.left_indent)
        number = p.add_run(f"{index}.\t")
        set_run_font(number, size=10, bold=True, color=COLOR_PRIMARY)
        body = p.add_run(item)
        set_run_font(body, size=10)


def add_steps_intro(doc, title):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"Untuk mempraktikkan {title.lower()}, ikuti langkah-langkah berikut.")
    set_run_font(run, size=10)


def add_bullets_intro(doc, title):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(
        f"Beberapa hal penting yang perlu diperhatikan pada bagian {title.lower()} adalah sebagai berikut."
    )
    set_run_font(run, size=10)


def add_field_note(doc, title, instruction):
    add_heading(doc, title, 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(instruction)
    set_run_font(run, size=9)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(85)
    run = p.add_run("BUKU PANDUAN PENGELOLA LAYANAN")
    set_run_font(run, size=16, bold=True, color=COLOR_PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "WEBSITE UNIT LAYANAN TERPADU (ULT)\n"
        "FAKULTAS KEGURUAN DAN ILMU PENDIDIKAN\n"
        "UNIVERSITAS LAMPUNG"
    )
    set_run_font(run, size=13, bold=True, color=COLOR_PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    run = p.add_run("Panduan Operasional Penggunaan Portal Pengelola Layanan")
    set_run_font(run, size=11, bold=True, color=COLOR_PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Dokumen ini menjelaskan alur penggunaan portal admin mulai dari login,\n"
        "pengelolaan permohonan, layanan dokumen, nomor surat, hingga modul pendukung administrasi."
    )
    set_run_font(run, size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(150)
    run = p.add_run("Bandar Lampung\n2026")
    set_run_font(run, size=10)


def add_front_matter(doc):
    add_heading(doc, "KATA PENGANTAR", 1)
    add_body(
        doc,
        "Buku panduan ini disusun sebagai acuan operasional penggunaan Website Unit Layanan Terpadu (ULT) FKIP Universitas Lampung dari sudut pandang pengelola layanan. Pembaruan dokumen dilakukan agar isi panduan selaras dengan antarmuka dan alur sistem terbaru, khususnya pada area daftar permohonan, detail layanan, workflow admin, setup layanan dokumen, nomor surat, dan modul pendukung administrasi.",
    )
    add_body(
        doc,
        "Melalui buku ini, pengguna diharapkan dapat memahami langkah penggunaan sistem secara lebih terarah sejak tahap login sampai pengelolaan modul yang lebih lanjut. Tangkapan layar yang digunakan pada setiap bab dipilih dari aset kerja portal admin agar uraian yang disajikan lebih mudah diikuti ketika praktik penggunaan dilakukan.",
    )
    add_body(
        doc,
        "Dokumen ini juga berfungsi sebagai pelengkap dokumentasi produk hasil pengembangan, sehingga dapat dipakai sebagai referensi saat orientasi pengguna internal, peninjauan tugas akhir, maupun evaluasi operasional layanan pada sistem.",
    )

    doc.add_page_break()
    add_field_note(
        doc,
        "DAFTAR ISI",
        'Daftar isi dapat diperbarui otomatis melalui fitur "Update Field" pada Microsoft Word setelah peninjauan akhir dokumen.',
    )

    doc.add_page_break()
    add_field_note(
        doc,
        "DAFTAR GAMBAR",
        'Daftar gambar dapat diperbarui otomatis melalui fitur "Update Field" pada Microsoft Word setelah caption final diperiksa.',
    )


def resolve_image_dimensions(path: Path):
    with Image.open(path) as image:
        width_px, height_px = image.size

    aspect_ratio = width_px / max(height_px, 1)
    content_width = PAGE_WIDTH_CM - LEFT_MARGIN_CM - RIGHT_MARGIN_CM
    content_height = PAGE_HEIGHT_CM - TOP_MARGIN_CM - BOTTOM_MARGIN_CM
    max_width = min(13.4, content_width - 0.2)
    max_height = min(15.2, content_height - 1.2)

    if aspect_ratio >= 4.5:
        target_width = 10.4
    elif aspect_ratio >= 2.8:
        target_width = 11.2
    elif aspect_ratio >= 1.6:
        target_width = 12.4
    elif aspect_ratio >= 1.0:
        target_width = 11.6
    elif aspect_ratio >= 0.72:
        target_width = 9.8
    else:
        target_width = 8.4

    target_width = min(target_width, max_width)
    target_height = target_width / max(aspect_ratio, 0.01)

    if target_height > max_height:
        scale = max_height / target_height
        target_width *= scale
        target_height = max_height

    return Cm(round(target_width, 2)), Cm(round(target_height, 2))


def add_figure(doc, filename, caption):
    candidates = [
        ASSETS / filename,
        ROOT / "public/example" / filename,
    ]
    path = next((candidate for candidate in candidates if candidate.exists()), None)
    if path is None:
        add_body(doc, f"[Gambar tidak ditemukan: {filename}]")
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    width, height = resolve_image_dimensions(path)
    p.add_run().add_picture(str(path), width=width, height=height)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(4)
    run = cap.add_run(caption)
    set_run_font(run, size=8, italic=True)


def add_chapters(doc):
    for chapter in CHAPTERS:
        add_heading(doc, chapter["title"], 1)
        for paragraph in chapter.get("intro", []):
            add_body(doc, paragraph)

        for section in chapter.get("sections", []):
            add_heading(doc, section["title"], 2)
            for paragraph in section.get("body", []):
                add_body(doc, paragraph)
            if section.get("figure"):
                add_figure(doc, section["figure"][0], section["figure"][1])
            if section.get("steps"):
                add_steps_intro(doc, section["title"])
                add_steps(doc, section["steps"])
            if section.get("bullets"):
                if section.get("steps") or section.get("figure"):
                    add_bullets_intro(doc, section["title"])
                add_bullets(doc, section["bullets"])


def main():
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc)
    doc.add_page_break()
    add_front_matter(doc)
    add_chapters(doc)
    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    main()
