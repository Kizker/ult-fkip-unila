import docx
import re
import os

doc_clean_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
doc_hl_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'

foreign_phrases = [
    "Design", "blueprint", "Use Case", "timeline", "Role-Based Access Control", 
    "audit trail", "layered architecture", "loosely coupled", "Client & Interface Layer", 
    "Progressive Web App", "Public Portal", "Student Portal", "Staff/Admin Portal", 
    "Signer Portal", "Security & Routing Gatekeeper", "Router", "Content Security Policy",
    "middleware", "Anti-IDOR", "HtmlSanitizer", "Cross-Site Scripting", "Application Core", 
    "Controllers Core", "DocumentAssembler", "HtmlToOpenXMLParser", "Database & Storage Layer", 
    "MySQL Database", "Private Storage Disk", "Document Assembly Engine", "Word Templates", 
    "placeholder", "ZipArchive", "PDF Converter", "form"
]

# Sort by length descending to avoid partial matches
foreign_phrases.sort(key=len, reverse=True)

def italicize_text_in_paragraph(paragraph, phrases):
    # This function is tricky because replacing text in runs while keeping other formatting is hard.
    # An easier way for paragraphs that only have homogeneous formatting (or just Highlighting)
    # is to clear the paragraph and rebuild it with the regex splits.
    # Since we JUST created these paragraphs, they have uniform formatting! (Times New Roman 12, Justified, etc.)
    # Exception: The highlight color in the HL version.
    
    # Let's check if the paragraph contains any of our phrases
    text = paragraph.text
    has_phrase = False
    for phrase in phrases:
        if re.search(r'\b' + re.escape(phrase) + r'\b', text, re.IGNORECASE):
            has_phrase = True
            break
            
    if not has_phrase:
        return False
        
    # Get the highlight color if any (we check the first run)
    hl_color = None
    if len(paragraph.runs) > 0:
        hl_color = paragraph.runs[0].font.highlight_color
        
    # Create a regex pattern for all phrases
    pattern = r'\b(' + '|'.join(re.escape(p) for p in phrases) + r')\b'
    
    # Split the text
    parts = re.split(pattern, text, flags=re.IGNORECASE)
    
    # Clear the paragraph
    paragraph.clear()
    
    from docx.shared import Pt
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if hl_color is not None:
            run.font.highlight_color = hl_color
            
        # Check if this part is one of the phrases
        # We need case-insensitive match against the list
        is_foreign = False
        for phrase in phrases:
            if part.lower() == phrase.lower():
                is_foreign = True
                break
                
        if is_foreign:
            run.italic = True
            
    return True

def process_doc(path):
    doc = docx.Document(path)
    
    # The paragraphs we modified are around indices 760 to 835
    # We can just check all paragraphs because `italicize_text_in_paragraph` only affects homogeneous paragraphs safely
    # Wait! If we apply this to the WHOLE document, it will wipe out bold/italic/equation formatting in other paragraphs!
    # So we MUST restrict it ONLY to the exact paragraphs we injected.
    
    # We know the injected texts start with specific strings.
    injected_starts = [
        "Pada tahap perancangan (Design)",
        "a. Penyusunan Instrumen",
        "b. Pemilihan Teknologi",
        "c. Perancangan Format",
        "d. Pemodelan Sistem",
        "Diagram Use Case memvisualisasikan",
        "Pada tataran operasional, Aktor",
        "Keseluruhan ekosistem interaksi fungsional",
        "Diagram arsitektur sistem pada",
        "Secara visual, lapisan teratas",
        "Lapisan kedua, Security",
        "Pada lapisan ketiga, Laravel",
        "Lapisan keempat merupakan",
        "Lapisan paling bawah adalah"
    ]
    
    count = 0
    for p in doc.paragraphs:
        for start in injected_starts:
            if p.text.startswith(start):
                if italicize_text_in_paragraph(p, foreign_phrases):
                    count += 1
                break
                
    print(f"Processed {count} paragraphs in {os.path.basename(path)}")
    doc.save(path)

process_doc(doc_clean_path)
process_doc(doc_hl_path)
