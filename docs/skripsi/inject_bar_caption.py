import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_clean_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
doc_hl_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'

def find_image_paragraph(doc):
    # Search for the paragraph that contains the image just before the text "Berdasarkan hasil rekapitulasi penilaian..."
    # The image is the Bar Chart.
    target_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if 'Berdasarkan hasil rekapitulasi penilaian dari ketiga ahli' in p.text:
            target_idx = i
            break
            
    if target_idx != -1:
        # Search backwards from target_idx to find the image paragraph
        for i in range(target_idx - 1, target_idx - 5, -1):
            for run in doc.paragraphs[i].runs:
                if '<w:drawing' in run.element.xml:
                    return i
    return -1

def add_caption_to_doc(file_path):
    doc = docx.Document(file_path)
    img_idx = find_image_paragraph(doc)
    
    if img_idx == -1:
        print(f"Could not find image paragraph in {file_path}")
        return
        
    print(f"File: {file_path}, Image at index: {img_idx}")
    
    # We will use the paragraph immediately after the image for the caption
    caption_p = doc.paragraphs[img_idx + 1]
    
    # Just to make sure we don't overwrite important text
    if len(caption_p.text.strip()) > 0 and 'Berdasarkan' in caption_p.text:
        # We need to insert a new paragraph
        caption_p = caption_p.insert_paragraph_before()
        
    caption_p.clear()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.line_spacing = 1.0
    
    # Add text
    run = caption_p.add_run("Gambar 4.26. Grafik Persentase Hasil Validasi Ahli.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.save(file_path)
    print("Caption added successfully.")

add_caption_to_doc(doc_clean_path)
add_caption_to_doc(doc_hl_path)
