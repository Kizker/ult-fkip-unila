import docx

doc_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
doc = docx.Document(doc_path)

for p in doc.paragraphs:
    if 'Tahap Pengem' in p.text or 'Tahap Implem' in p.text:
        print(f"Heading: '{p.text}'")
        
print("--- Check XML for graphic ---")
for i, p in enumerate(doc.paragraphs):
    if 'graphic' in p._p.xml or 'chart' in p._p.xml:
        print(f"P[{i}] has graphic. Length of text: {len(p.text)}")
