import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_average_row(doc_path):
    doc = docx.Document(doc_path)
    
    target_table = None
    for table in doc.tables:
        try:
            if 'Nama Validator' in table.cell(0,0).text and 'Bidang' in table.cell(0,1).text:
                target_table = table
                break
        except:
            pass
            
    if not target_table:
        print(f"Table not found in {doc_path}")
        return
        
    # Check if average row already exists
    last_row_text = target_table.rows[-1].cells[0].text.lower()
    if 'rata' in last_row_text:
        print(f"Average row already exists in {doc_path}")
        return

    # Calculate averages
    num_rows = len(target_table.rows) - 1 # exclude header
    
    q_sums = [0] * 10
    total_sum = 0
    perc_sum = 0.0
    
    for i in range(1, num_rows + 1):
        row = target_table.rows[i]
        for j in range(10):
            try:
                q_sums[j] += int(row.cells[j+2].text.strip())
            except:
                pass
        try:
            total_sum += int(row.cells[12].text.strip())
        except:
            pass
        try:
            p_str = row.cells[13].text.strip().replace('%', '')
            perc_sum += float(p_str)
        except:
            pass
            
    # Add new row
    new_row = target_table.add_row()
    cells = new_row.cells
    
    # Merge first two cells
    cells[0].merge(cells[1])
    cells[0].text = "Rata-rata Keseluruhan"
    
    for j in range(10):
        avg_q = q_sums[j] / num_rows
        cells[j+2].text = f"{avg_q:.2f}"
        
    avg_total = total_sum / num_rows
    cells[12].text = f"{avg_total:.2f}"
    
    avg_perc = perc_sum / num_rows
    cells[13].text = f"{avg_perc:.2f}%"
    
    # Format the new row
    for i in range(14):
        for p in cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                
    doc.save(doc_path)
    print(f"Added average row to {doc_path}")

add_average_row(r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx')
add_average_row(r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx')
