import docx
from pathlib import Path

def main():
    path = Path(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx")
    if not path.exists():
        print("Not found.")
        return
        
    doc = docx.Document(path)
    # Let's extract paragraphs from 625 to the end of the document to make sure we don't miss anything.
    start_idx = 625
    end_idx = len(doc.paragraphs)
    
    out_path = Path(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\extracted_bab4.txt")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("=== EXTRACTING ALL PARAGRAPHS FROM P 625 TO THE END ===\n")
        for idx in range(start_idx, end_idx):
            p = doc.paragraphs[idx]
            f.write(f"P {idx} ({p.style.name}): {p.text}\n")
    print(f"Extracted {end_idx - start_idx} paragraphs to {out_path}")

if __name__ == "__main__":
    main()
