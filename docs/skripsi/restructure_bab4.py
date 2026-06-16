import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import copy

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def format_paragraph(p, doc, style_name="Normal", alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing_before=0, spacing_after=0, line_spacing=1.5):
    try:
        p.style = doc.styles[style_name]
    except Exception:
        pass
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(spacing_before)
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = line_spacing

def add_run_formatted(p, text, bold=False, italic=False, font_size=12, highlight=None):
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(font_size)
    r.bold = bold
    r.italic = italic
    if highlight:
        r.font.highlight_color = highlight
    return r

# Define all naskah Chapter IV runs centrally to prevent duplication and ensure strict italicization of foreign words.
# Format: (text_str, bold_bool, italic_bool)
NASKAH_BAB_4 = [
    # Spacer
    {"type": "p", "runs": "", "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    # Heading A
    {"type": "p", "runs": [("A. Hasil Penelitian (Model Pengembangan ADDIE)", True, False)], "style": "Heading 2", "align": WD_ALIGN_PARAGRAPH.CENTER, "before": 12, "after": 6, "line": 1.5},
    
    # Intro
    {"type": "p", "runs": [
        ("Pengembangan Website Unit Layanan Terpadu (ULT) FKIP Universitas Lampung dilakukan dengan menerapkan model pengembangan ", False, False),
        ("ADDIE", True, False),
        (" (", False, False),
        ("Analysis, Design, Development, Implementation, ", False, True),
        ("dan ", False, False),
        ("Evaluation", False, True),
        ("). Model ini dipilih karena strukturnya yang sistematis, adaptif, dan berorientasi pada pemecahan masalah riil. Produk akhir yang dihasilkan berupa aplikasi web monolitik yang dibangun menggunakan ", False, False),
        ("framework", False, True),
        (" Laravel 12 dengan ", False, False),
        ("runtime", False, True),
        (" PHP 8.4+, basis data MySQL, dan antarmuka dinamis berbasis ", False, False),
        ("TailwindCSS", False, True),
        (" dan ", False, False),
        ("Alpine.js", False, True),
        (". Website ini dirancang untuk mendigitalisasi dan mengotomatisasi seluruh alur pengajuan dokumen dan surat-menyurat mahasiswa yang sebelumnya dilakukan secara manual dan terdistribusi. Berikut dijabarkan proses pengembangan sistem untuk masing-masing tahap secara mendetail.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    # 1. Tahap Analisis
    {"type": "p", "runs": [("1. Tahap Analisis (", True, False), ("Analysis", True, True), (")", True, False)], "style": "Heading 3", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 12, "after": 6, "line": 1.5},
    
    {"type": "p", "runs": [
        ("Pada tahap analisis, peneliti melakukan investigasi mendalam terhadap kebutuhan awal sistem melalui dua metode utama, yaitu wawancara pra-penelitian terstruktur dengan staf ULT dan administrator jurusan FKIP Unila, serta penyebaran angket analisis kebutuhan digitalisasi pelayanan kepada mahasiswa aktif. Hasil wawancara pra-penelitian dengan staf ULT mengonfirmasi keberadaan sejumlah kendala operasional ", False, False),
        ("pain points", False, True),
        (" yang signifikan pada alur administrasi pelayanan mahasiswa saat ini. Kendala-kendala utama tersebut meliputi: (1) tidak adanya sistem pencatatan terpusat, sehingga linimasa status dokumen mahasiswa menjadi sepenuhnya tidak transparan; (2) tingginya risiko kerusakan fisik atau kehilangan berkas penting akibat penumpukan dokumen di loket pelayanan; dan (3) waktu tunggu persetujuan pejabat fakultas/jurusan yang tidak terukur karena alur koordinasi manual yang terdistribusi dan sering mengalami hambatan.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("Untuk memperkuat temuan kualitatif tersebut, peneliti menyebarkan angket analisis kebutuhan kepada mahasiswa FKIP Unila. Hasil angket secara kuantitatif menunjukkan tingkat urgensi digitalisasi yang sangat tinggi. Mayoritas mahasiswa (87,5%) menyatakan bahwa pemindahan loket fisik ke platform digital adalah kebutuhan mendesak guna memangkas antrean loket dan memberikan kepastian penyelesaian dokumen secara transparan. Berdasarkan temuan empiris di atas, peneliti menyimpulkan bahwa transformasi digital pelayanan akademik melalui pengembangan website ULT FKIP Unila dengan fitur-fitur seperti pelacakan status ", False, False),
        ("timeline", False, True),
        (", penyimpanan berkas terenkripsi, penomoran surat otomatis, dan tanda tangan elektronik pimpinan adalah langkah strategis dan esensial.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    # 2. Tahap Desain
    {"type": "p", "runs": [("2. Tahap Desain (", True, False), ("Design", True, True), (")", True, False)], "style": "Heading 3", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 12, "after": 6, "line": 1.5},
    
    {"type": "p", "runs": [
        ("Tahap desain difokuskan pada penyusunan cetak biru (", False, False),
        ("blueprint", False, True),
        (") arsitektur teknologi informasi, pemetaan basis data, ", False, False),
        ("flowchart", False, True),
        (" alur kerja persuratan, dan rancangan antarmuka visual (", False, False),
        ("wireframe", False, True),
        ("). Guna memastikan keamanan dan ketertelusuran hak akses data, peneliti merancang hak akses pengguna berbasis ", False, False),
        ("Role-Based Access Control", False, True),
        (" (", False, False),
        ("RBAC", False, True),
        (") yang terbagi secara ketat ke dalam empat portal utama yang saling terintegrasi, yaitu:", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("1. ", True, False),
        ("Public Portal", True, True),
        (": Antarmuka publik yang dapat diakses secara bebas tanpa autentikasi login. Portal ini berfungsi sebagai ", False, False),
        ("landing page", False, True),
        (" utama yang menyajikan katalog seluruh layanan akademik, syarat-syarat berkas pengajuan, papan pengumuman resmi fakultas, artikel blog, dan CMS pengelola halaman publik.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("2. ", True, False),
        ("Student Portal", True, True),
        (": Dasbor khusus mahasiswa untuk mengajukan permohonan layanan akademik secara mandiri. Mahasiswa dapat mengisi formulir dinamis sesuai jenis layanan, mengunggah berkas persyaratan secara privat ke ", False, False),
        ("private storage disk", False, True),
        (" Laravel, melacak status pengajuan secara langsung melalui linimasa (", False, False),
        ("timeline", False, True),
        (") yang ", False, False),
        ("auditable", False, True),
        (", menerima notifikasi revisi, dan mengunduh dokumen hasil secara aman.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("3. ", True, False),
        ("Admin/Staff Portal", True, True),
        (": Dasbor operasional staf ULT FKIP yang berfungsi sebagai ", False, False),
        ("gatekeeper", False, True),
        (" alur kerja. Staf dapat memverifikasi berkas mahasiswa, meneruskan pengajuan kepada pejabat yang berwenang, menginput nomor surat resmi, dan mengelola ", False, False),
        ("template", False, True),
        (" dokumen dinamis berbasis PHPWord.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("4. ", True, False),
        ("Signer Portal", True, True),
        (": Portal khusus untuk pimpinan fakultas, ketua jurusan, atau ketua program studi untuk memverifikasi dan menandatangani permohonan surat secara elektronik/digital sebelum dokumen dirakit (", False, False),
        ("assemble", False, True),
        (") secara otomatis oleh sistem.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("Selain arsitektur portal, pada tahap ini dirancang pula skema relasi ", False, False),
        ("database", False, True),
        (" MySQL yang menampung entitas pengguna, peran, pengajuan layanan, riwayat status, nomor surat, dan dokumen hasil. ", False, False),
        ("Flowchart", False, True),
        (" alur kerja juga disiapkan untuk mendefinisikan transisi status permohonan dari status Draft, Diperiksa ULT, Disetujui Pimpinan, Perakitan Dokumen (", False, False),
        ("Assembly", False, True),
        ("), hingga Selesai yang siap diunduh.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("Guna menjembatani rancangan konseptual ini ke bentuk fisik, peneliti telah mengembangkan rancangan antarmuka visual interaktif tingkat tinggi (", False, False),
        ("high-fidelity UI/UX design", False, True),
        (") secara komprehensif menggunakan platform Figma untuk seluruh portal pelayanan ULT (Public, Student, Admin, Signer). Rancangan visual lengkap serta kode akses QR untuk menguji alur prototipe interaktif ini didokumentasikan secara terpisah pada ", False, False),
        ("Lampiran D. Dokumentasi Desain UI/UX Website ULT FKIP Unila", True, False),
        (".", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    # 3. Tahap Pengembangan
    {"type": "p", "runs": [("3. Tahap Pengembangan (", True, False), ("Development", True, True), (")", True, False)], "style": "Heading 3", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 12, "after": 6, "line": 1.5},
    
    {"type": "p", "runs": [
        ("Tahap pengembangan merealisasikan rancangan desain ke dalam bentuk kode program fungsional menggunakan ", False, False),
        ("framework", False, True),
        (" Laravel 12 dan ", False, False),
        ("runtime", False, True),
        (" PHP 8.4+. Seluruh berkas unggahan sensitif disimpan di dalam ", False, False),
        ("private storage disk", False, True),
        (" yang dilindungi ", False, False),
        ("middleware", False, True),
        (" otorisasi guna mencegah serangan IDOR. Setelah sistem selesai dikembangkan, peneliti melakukan pengujian kelayakan melalui uji validasi produk yang melibatkan sembilan validator ahli yang terbagi ke dalam tiga kelompok keahlian. Setiap kelompok terdiri dari tiga orang ahli yang menilai instrumen kelayakan menggunakan kuesioner skala Likert 1-5. Data penilaian dianalisis secara kuantitatif menggunakan persentase kelayakan (kriteria Arikunto) dan koefisien validitas isi Aiken's V. Berikut dijabarkan analisis detail evaluasi dan tindak lanjut perbaikan dari masing-masing validator secara mendalam dan terpisah:", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    # Ahli Materi
    {"type": "p", "runs": [("a) Analisis Evaluasi Kelayakan Ahli Materi (Buku Panduan)", True, False)], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 6, "after": 3, "line": 1.5},
    
    {"type": "p", "runs": [
        ("Validasi ahli materi difokuskan pada penilaian kesesuaian isi, keruntutan penyajian, kejelasan instruksi, kelengkapan informasi, dan akurasi buku panduan penggunaan Website ULT FKIP Unila. Validasi dilakukan oleh tiga ahli materi akademis, yaitu Validator Ahli Materi 1, Validator Ahli Materi 2, dan Validator Ahli Materi 3. Instrumen penilaian terdiri dari 22 butir yang dikelompokkan ke dalam empat aspek utama. Berdasarkan rekapitulasi data penilaian dari ketiga validator ahli materi, diperoleh total skor sebesar 315 dari skor maksimal 330, yang menunjukkan rata-rata persentase kelayakan materi sebesar ", False, False),
        ("95,45%", True, False),
        (" (kriteria 'Sangat Valid') dengan rerata koefisien validitas isi Aiken's V mencapai ", False, False),
        ("0,94", True, False),
        (" (tingkat validitas isi yang sangat tinggi). Hasil analisis detail dari masing-masing validator dijabarkan sebagai berikut:", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("1. ", True, False),
        ("Validator Ahli Materi 1", True, False),
        (" memberikan skor total 106 dari 110 (Kelayakan: 96,36%) dengan kesimpulan 'Layak digunakan dengan revisi'. Beliau memberikan komentar kualitatif krusial: 'Tambah penjelasan, cek kalimat, pastikan ", False, False),
        ("screenshot", False, True),
        (" jelas'. Masukan ini menyoroti perlunya penjelasan lebih mendalam pada buku panduan terkait alur tindak lanjut revisi berkas yang diajukan mahasiswa. Peneliti merespons dengan menyunting tata bahasa teks instruksi agar lebih dinamis, menambahkan paragraf penjelas status penolakan berkas, dan menaikkan resolusi gambar tangkapan layar dasbor.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("2. ", True, False),
        ("Validator Ahli Materi 2", True, False),
        (" memberikan skor total 104 dari 110 (Kelayakan: 94,55%) dengan keputusan kelayakan akhir 'Layak digunakan tanpa revisi' dan tanpa catatan perbaikan. Beliau menilai keruntutan penyajian materi pada buku panduan sudah sangat sistematis dan terstruktur dengan baik sesuai pembagian peran pengguna (mahasiswa, verifikator ULT, dan pimpinan penandatangan).", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("3. ", True, False),
        ("Validator Ahli Materi 3", True, False),
        (" memberikan skor total 105 dari 110 (Kelayakan: 95,45%) dengan kesimpulan 'Layak digunakan tanpa revisi' disertai saran tambahan: 'Beri pembatas header/content/footer & gambar'. Masukan ini dinilai peneliti sangat baik untuk meningkatkan kegunaan visual buku panduan. Peneliti menindaklanjutinya dengan memberikan garis pembatas visual (", False, False),
        ("border", False, True),
        (") warna abu-abu netral tipis berukuran 1 pt pada sekeliling tangkapan layar antarmuka sistem agar visualisasi batas aplikasi lebih tegas bagi pembaca yang membedakan bagian ", False, False),
        ("header, content, ", False, True),
        ("dan ", False, False),
        ("footer", False, True),
        (".", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    # Ahli Media
    {"type": "p", "runs": [("b) Analisis Evaluasi Kelayakan Ahli Media (Tampilan dan UI/UX)", True, False)], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 6, "after": 3, "line": 1.5},
    
    {"type": "p", "runs": [
        ("Validasi ahli media bertujuan untuk menguji kelayakan aspek visual, keselarasan warna, tata letak antarmuka dasbor, navigasi, ergonomi, dan pengalaman pengguna (", False, False),
        ("UI/UX", False, True),
        (") Website ULT FKIP Unila. Validasi dilakukan oleh tiga ahli media praktisi teknologi informasi, yaitu Validator Ahli Media 1, Validator Ahli Media 2, dan Validator Ahli Media 3. Penilaian mencakup 14 butir kuesioner yang terbagi ke dalam aspek Tampilan Visual Website dan UI/UX Website. Hasil rekapitulasi penilaian dari ketiga validator ahli media menunjukkan total skor sebesar 196 dari skor maksimal 210, yang menghasilkan rerata persentase kelayakan media mencapai ", False, False),
        ("93,33%", True, False),
        (" (kriteria 'Sangat Valid') dengan rerata koefisien validitas isi Aiken's V sebesar ", False, False),
        ("0,92", True, False),
        (" (validitas isi sangat tinggi). Penjelasan evaluasi detail per ahli media dipaparkan sebagai berikut:", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("1. ", True, False),
        ("Validator Ahli Media 1", True, False),
        (" memberikan total skor 65 dari 70 (Kelayakan: 92,86%) dengan keputusan kelayakan akhir 'Layak digunakan tanpa revisi' disertai komentar positif: 'Bagus, bisa digunakan untuk penelitian'. Beliau menilai estetika visual website sangat modern, tipografi bersih dengan kontras yang seimbang, serta layout yang ergonomis.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("2. ", True, False),
        ("Validator Ahli Media 2", True, False),
        (" memberikan skor total 65 dari 70 (Kelayakan: 92,86%) dengan keputusan kelayakan 'Layak digunakan tanpa revisi' disertai catatan perbaikan visual: 'Warna sebaiknya jangan gradasi'. Validator Ahli Media 2 menilai penggunaan efek gradasi warna latar belakang yang terlalu mencolok pada dasbor portal mahasiswa berpotensi memicu kelelahan mata (", False, False),
        ("eye strain", False, True),
        (") jika diakses dalam durasi lama. Tindak lanjut dari peneliti adalah merancang ulang UI dasbor dengan mengganti latar belakang gradasi menjadi warna solid netral pastel (", False, False),
        ("Slate-100", False, True),
        (" dan ", False, False),
        ("Teal-50", False, True),
        (") demi meningkatkan kenyamanan mata.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("3. ", True, False),
        ("Validator Ahli Media 3", True, False),
        (" memberikan total skor 66 dari 70 (Kelayakan: 94,29%) dengan kesimpulan 'Layak digunakan dengan revisi' beserta komentar: 'Beri icon pada setiap tile layanan'. Masukan ini bertujuan untuk meningkatkan fungsionalitas visual antarmuka agar pengguna dapat mengidentifikasi jenis permohonan secara instan. Peneliti langsung merespons dengan mengintegrasikan pustaka ", False, False),
        ("SVG icons", False, True),
        (" yang dinamis dan representatif di setiap kartu (", False, False),
        ("card/tile", False, True),
        (") layanan akademik pada dasbor mahasiswa.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    # Ahli Sistem
    {"type": "p", "runs": [("c) Analisis Evaluasi Kelayakan Ahli Sistem (Fungsionalitas dan Keamanan)", True, False)], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 6, "after": 3, "line": 1.5},
    
    {"type": "p", "runs": [
        ("Validasi ahli sistem difokuskan pada pengujian aspek fungsionalitas (", False, False),
        ("correctness", False, True),
        ("), keandalan proses bisnis (", False, False),
        ("reliability/error handling", False, True),
        ("), keamanan akses data (", False, False),
        ("security", False, True),
        ("), dan kinerja dasar sistem (", False, False),
        ("performance", False, True),
        ("). Uji validasi sistem dilakukan oleh tiga praktisi dan akademisi rekayasa perangkat lunak, yaitu Validator Ahli Sistem 1, Validator Ahli Sistem 2, dan Validator Ahli Sistem 3. Penilaian mencakup 22 butir kuesioner yang terbagi ke dalam empat dimensi kelayakan sistem. Hasil rekapitulasi penilaian dari ketiga validator ahli sistem menunjukkan total skor sebesar 289 dari skor maksimal 330, yang menghasilkan rerata persentase kelayakan sistem sebesar ", False, False),
        ("87,58%", True, False),
        (" (kriteria 'Sangat Valid') dengan rerata koefisien validitas isi Aiken's V mencapai ", False, False),
        ("0,85", True, False),
        (" (validitas isi tinggi). Penjelasan detail evaluasi dan perbaikan dari masing-masing ahli sistem dijabarkan sebagai berikut:", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("1. ", True, False),
        ("Validator Ahli Sistem 1", True, False),
        (" memberikan total skor 105 dari 110 (Kelayakan: 95,45%) dengan kesimpulan 'Layak digunakan tanpa revisi' dan komentar 'Perbaikan saran'. Validator Ahli Sistem 1 menilai bahwa fungsionalitas sistem secara umum (autentikasi peran RBAC, form input dinamis, dan pengunggahan berkas) telah berjalan stabil tanpa adanya cacat kode ", False, False),
        ("malfunction", False, True),
        (" yang berarti.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("2. ", True, False),
        ("Validator Ahli Sistem 2", True, False),
        (" memberikan skor total 93 dari 110 (Kelayakan: 84,55%) dengan keputusan akhir 'Layak digunakan dengan revisi' disertai komentar teknis yang sangat krusial: 'Login SSO, perhatikan CSP (unsafe-eval/inline), atur CSP default'. Masukan ini menyoroti celah keamanan potensial pada header HTTP Laravel yang rentan terhadap serangan injeksi ", False, False),
        ("script", False, True),
        (" jahat luar (", False, False),
        ("Cross-Site Scripting / XSS", False, True),
        (") jika ", False, False),
        ("Content Security Policy", False, True),
        (" (", False, False),
        ("CSP", False, True),
        (") default tidak disetel secara ketat. Menindaklanjuti hal tersebut, peneliti melakukan revisi mendalam dengan memperketat ", False, False),
        ("middleware", False, True),
        (" CSP default Laravel untuk memblokir ", False, False),
        ("unsafe-eval", False, True),
        (" dan ", False, False),
        ("unsafe-inline", False, True),
        (" secara total, serta mengenkripsi nama berkas pengajuan mahasiswa di ", False, False),
        ("private storage disk", False, True),
        (". Selain itu, peneliti juga menyusun cetak biru (", False, False),
        ("blueprint", False, True),
        (") skema database terenkripsi untuk persiapan integrasi basis data terpadu resmi Universitas Lampung dengan sistem autentikasi ", False, False),
        ("Single Sign-On", False, True),
        (" (", False, False),
        ("SSO", False, True),
        (") pada pengembangan tahap selanjutnya.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("3. ", True, False),
        ("Validator Ahli Sistem 3", True, False),
        (" memberikan skor total 91 dari 110 (Kelayakan: 82,73%) dengan keputusan 'Layak digunakan tanpa revisi' dan memberikan komentar: 'Tampilkan tombol kembali ke atas'. Masukan ini bertujuan untuk meningkatkan kenyamanan navigasi pengguna (", False, False),
        ("usability", False, True),
        ("). Peneliti merespons dengan mengimplementasikan tombol melayang kembali ke atas (", False, False),
        ("Back to Top", False, True),
        (") menggunakan perpaduan Alpine.js dan TailwindCSS di landing page portal publik.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    # Closing dev
    {"type": "p", "runs": [
        ("Secara kumulatif, rekapitulasi kelayakan dari sembilan validator ahli di atas membuktikan bahwa Website ULT FKIP Unila beserta dokumen panduannya telah memenuhi standar kualitas yang prima dan sangat layak untuk diuji coba di lapangan. Rekapitulasi hasil penilaian validasi akhir dari ketiga kelompok ahli tersebut disajikan secara lengkap pada Tabel 4.1 berikut.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    # Table 4.1
    {"type": "p", "runs": [("Tabel 4.1. Rekapitulasi Hasil Validasi Ahli (Materi, Media, Sistem)", True, False)], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.CENTER, "before": 12, "after": 3, "line": 1.5},
    {"type": "t_val"},
    {"type": "p", "runs": [("Sumber: Data primer olahan peneliti (2026)", False, True)], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.CENTER, "before": 3, "after": 12, "line": 1.5},
    
    # 4. Tahap Implementasi
    {"type": "p", "runs": [("4. Tahap Implementasi (", True, False), ("Implementation", True, True), (")", True, False)], "style": "Heading 3", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 12, "after": 6, "line": 1.5},
    
    {"type": "p", "runs": [
        ("Tahap implementasi bertujuan untuk mengukur tingkat kepraktisan dan kemudahan penggunaan operasional Website ULT FKIP Unila saat digunakan oleh pengguna sasaran. Pengujian dilakukan melalui uji coba terbatas yang melibatkan 18 responden civitas akademika FKIP Unila yang dipilih secara sengaja menggunakan teknik ", False, False),
        ("purposive sampling", False, True),
        (" (Sugiyono, 2013). Teknik pengambilan sampel ini dipilih dengan menetapkan kriteria khusus, yaitu responden harus merupakan pengguna aktif yang secara langsung terlibat atau membutuhkan pelayanan administrasi surat-menyurat di lingkungan fakultas. Responden dikelompokkan secara seimbang untuk mewakili seluruh corak disiplin ilmu dan peran operasional, yang terdiri atas: 12 mahasiswa tingkat akhir perwakilan rumpun jurusan di FKIP (jurusan Pendidikan Bahasa dan Sastra [PBS] = 3 orang, Pendidikan Ilmu Pengetahuan [PIP] = 3 orang, Pendidikan Ilmu Pengetahuan Sosial [PIPS] = 3 orang, dan Pendidikan Matematika dan IPA [PMIPA] = 3 orang), 3 Admin Program Studi yang mengurusi administrasi persuratan jurusan, serta 3 staf Unit Layanan Terpadu (ULT) selaku verifikator gerbang utama tingkat fakultas.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("Setiap responden diminta untuk menjalankan skenario penggunaan utama sistem (meliputi pengajuan layanan dari sisi mahasiswa, verifikasi berkas dan penomoran dari sisi staf ULT/admin prodi, tanda tangan elektronik pimpinan dari portal signer, hingga pengunduhan hasil dokumen surat resmi). Setelah menjalankan alur kerja sistem, seluruh responden mengisi kuesioner uji kepraktisan kuantitatif yang berisi 12 butir instrumen menggunakan skala Likert 1-5 (skor 5 = Sangat Setuju, skor 1 = Sangat Tidak Setuju) serta memberikan kesimpulan subjektif kelayakan akhir pada Bagian F kuesioner.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("Hasil analisis data kuantitatif dari ke-18 responden tersebut menunjukkan tingkat kepraktisan sistem yang sangat memuaskan. Rata-rata skor total yang diperoleh dari kuesioner adalah ", False, False),
        ("55,28", True, False),
        (" dari skor maksimal 60,00, yang menghasilkan persentase kepraktisan rata-rata sebesar ", False, False),
        ("92,13%", True, False),
        (". Berdasarkan kriteria Arikunto, hasil persentase ini menempatkan Website ULT FKIP Unila dalam kategori ", False, False),
        ("Sangat Praktis", True, True),
        (" untuk langsung diimplementasikan dalam kegiatan operasional sehari-hari. Evaluasi butir instrumen memperlihatkan skor yang sangat stabil di atas 4,40, dengan skor tertinggi berada pada butir A3 (Kemudahan mempelajari penggunaan sistem) yang meraih skor rata-rata 4,83 dari 5,00. Hal ini membuktikan bahwa antarmuka sistem sangat intuitif sehingga pengguna sasaran dapat beradaptasi dan menggunakan sistem secara mandiri dengan sangat cepat tanpa memerlukan pelatihan teknis khusus. Rekapitulasi matriks data kuantitatif hasil skor kuesioner uji kepraktisan responden secara lengkap dapat dilihat pada ", False, False),
        ("Lampiran C. Hasil Uji Kepraktisan Pengguna", True, False),
        (".", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("Selain data kuantitatif kuesioner, pada lembar angket Bagian F (Kesimpulan Responden), responden diminta memberikan pendapat kelayakan subjektif akhir produk secara langsung. Rekapitulasi distribusi pilihan kesimpulan akhir tersebut menunjukkan mayoritas responden (94,44% atau 17 orang) menyimpulkan bahwa sistem ini berkategori 'Sangat Praktis' (8 orang) and 'Praktis' (9 orang) untuk langsung digunakan dalam kegiatan operasional sehari-hari. Hanya 1 responden (5,56% yaitu Responden 3) yang memberikan kesimpulan 'Cukup Praktis'. Hal ini membuktikan antusiasme yang tinggi dari para responden pengguna terhadap efektivitas sistem.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("Meskipun memperoleh respon positif, terdapat masukan kualitatif yang berharga dari responden untuk penyempurnaan sistem ULT. Sebanyak 6 dari 18 responden menuliskan komentar dan saran perbaikan pada kolom yang disediakan. Peneliti telah menganalisis dan menindaklanjuti seluruh saran tersebut demi mencapai kesempurnaan website ULT pada fase pengembangan akhir. Rekapitulasi tanggapan kualitatif dari responden beserta tindakan perbaikan yang telah dikerjakan oleh peneliti dijabarkan di bawah ini:", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("1. ", True, False),
        ("Responden 1 (Admin Program Studi)", True, False),
        (": Menyarankan penambahan fitur visualisasi jadwal seminar mahasiswa mirip Airtable. Peneliti menindaklanjutinya dengan merancang modul kalender penjadwalan dinamis terintegrasi di portal Admin agar verifikator dapat melacak jadwal ujian mahasiswa secara periodik.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("2. ", True, False),
        ("Responden 2 (Admin Program Studi)", True, False),
        (": Menyarankan penomoran surat resmi otomatis terhubung langsung dengan pola penomoran resmi fakultas. Peneliti merespons dengan menyusun modul regex generator penomoran dinamis di portal staf ULT yang dapat menyesuaikan dengan format penomoran unit masing-masing program studi secara otomatis.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("3. ", True, False),
        ("Responden 3 (Admin Program Studi)", True, False),
        (": Menilai sistem mudah dipelajari, segera berjalan, dan sangat bermanfaat untuk memperlancar kelancaran proses administrasi persuratan mahasiswa maupun staf. Komentar ini menjadi bukti kemudahan adaptasi sistem di lingkungan operasional prodi.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("4. ", True, False),
        ("Responden 13 (Mahasiswa PMIPA)", True, False),
        (": Menyatakan kekagumannya terhadap desain antarmuka (", False, False),
        ("user interface", False, True),
        (") yang sederhana tetapi tetap menarik perhatian (", False, False),
        ("simple but eye-catching", False, True),
        ("), serta mengapresiasi kejelasan dan kelengkapan informasi katalog layanan yang disajikan di portal publik.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("5. ", True, False),
        ("Responden 14 (Mahasiswa PMIPA)", True, False),
        (": Memberikan masukan agar kepadatan teks diatur dengan baik agar tidak menimbulkan kesan menumpuk dan melelahkan mata. Peneliti menindaklanjuti saran ini pada visualisasi portal mahasiswa dengan menerapkan spasi baris 1.5 secara konsisten, meningkatkan ", False, False),
        ("padding", False, True),
        (" antar-komponen dasbor, dan mengoptimalkan ruang kosong (", False, False),
        ("white space", False, True),
        (") pada tata letak halaman.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("6. ", True, False),
        ("Responden 17 (Staf Unit Layanan Terpadu)", True, False),
        (": Menyarankan penambahan modul pelaporan berbasis grafik pelacakan data mingguan/bulanan dan penambahan modul layanan kepegawaian (dosen/staf). Peneliti menyambut baik saran pengembangan skala besar ini dengan menyusun rencana perluasan sistem di Bab V serta merancang modul visualisasi grafik data transaksi menggunakan pustaka Chart.js untuk fase pengembangan berikutnya.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    # 5. Tahap Evaluasi
    {"type": "p", "runs": [("5. Tahap Evaluasi (", True, False), ("Evaluation", True, True), (")", True, False)], "style": "Heading 3", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 12, "after": 6, "line": 1.5},
    
    {"type": "p", "runs": [
        ("Tahap evaluasi merupakan langkah terakhir yang melakukan evaluasi akhir sumatif terhadap kegunaan, keandalan, kelayakan, dan efektivitas operasional Website ULT FKIP Unila. Evaluasi dilakukan dengan menyinkronkan seluruh umpan balik dari validator ahli dan data uji coba kepraktisan oleh pengguna sasaran. Hasil evaluasi sumatif membuktikan bahwa transisi alur kerja dari loket fisik konvensional ke platform digital berbasis web berhasil mengatasi kendala keterlambatan alur birokrasi, penumpukan berkas loket, dan ketidakpastian persetujuan dokumen. Dengan rata-rata persentase kelayakan ahli sebesar ", False, False),
        ("91,95%", True, False),
        (" (Sangat Valid) dan persentase kepraktisan pengguna mencapai ", False, False),
        ("92,13%", True, False),
        (" (Sangat Praktis), sistem ini dinyatakan sangat layak dan andal untuk segera diimplementasikan secara resmi untuk mendukung digitalisasi pelayanan publik di lingkungan Fakultas Keguruan dan Ilmu Pendidikan Universitas Lampung.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    # Spacer
    {"type": "p", "runs": "", "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    # Heading B
    {"type": "p", "runs": [("B. Pembahasan Hasil Penelitian", True, False)], "style": "Heading 2", "align": WD_ALIGN_PARAGRAPH.CENTER, "before": 12, "after": 6, "line": 1.5},
    
    {"type": "p", "runs": [
        ("Hasil pengembangan Website ULT FKIP Unila ini membuktikan secara empiris bahwa penerapan model R&D dengan sintaks ADDIE mampu menghasilkan produk teknologi manajemen pelayanan kampus yang terarah, efisien, aman, dan tepat guna bagi seluruh civitas akademika FKIP Unila. Melalui analisis kebutuhan awal pada tahap ", False, False),
        ("Analysis", False, True),
        (", ditemukan bahwa kendala utama layanan administrasi di FKIP Unila adalah lambatnya ", False, False),
        ("turn-around time", False, True),
        (" pengurusan dokumen, hilangnya transparansi status dokumen, serta kerentanan fisik berkas yang diajukan. Sistem yang dikembangkan ini secara terarah menutup ", False, False),
        ("pain point", False, True),
        (" tersebut dengan memindahkan seluruh proses fisik ke dalam sistem berbasis web yang dinamis dan aman.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    # Keunggulan
    {"type": "p", "runs": [("1. Eksplorasi Keunggulan Produk Website ULT FKIP Unila", True, False)], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 6, "after": 3, "line": 1.5},
    
    {"type": "p", "runs": [
        ("Website ULT FKIP Unila memiliki sejumlah keunggulan teknis dan akademis yang membedakannya dengan platform administrasi konvensional. Keunggulan-keunggulan utama tersebut dijabarkan di bawah ini:", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("a) ", True, False),
        ("Integrasi Empat Portal Berbasis RBAC Ketat", True, False),
        (": Keberadaan portal khusus (Public, Student, Admin, Signer) memastikan setiap pengguna berinteraksi hanya dengan fungsi dan data yang relevan dengan kewenangannya. Model ini memblokir akses ilegal antarportal dan menjamin keamanan transaksi data.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("b) ", True, False),
        ("Otomatisasi Perakitan Dokumen Dinamis (OpenXML)", True, False),
        (": Integrasi pustaka PHPWord untuk merakit dokumen Word (.docx) langsung dari basis data menggunakan template dinamis terbukti sangat efisien. Logika parser HTML-to-OpenXML dinamis terpusat pada file XML mentah (`document.xml`) menjamin bahwa run teks yang dimasukkan dari WYSIWYG editor (Tiptap Editor) pada formulir mahasiswa tetap mewarisi gaya visual visual bawaan template (font, ukuran, warna) tanpa mengalami degradasi layout visual orisinal.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("c) ", True, False),
        ("Keamanan Data Tingkat Tinggi dan Proteksi Otorisasi", True, False),
        (": Penggunaan ", False, False),
        ("private storage disk", False, True),
        (" Laravel menjamin berkas berkas persyaratan sensitif mahasiswa tidak dapat diakses langsung secara publik tanpa login. Dilindungi oleh middleware otorisasi anti-IDOR yang ketat, platform berhasil menutup risiko ", False, False),
        ("bypass endpoint", False, True),
        (" berkas mahasiswa lain.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("d) ", True, False),
        ("Jejak Transaksi Digital Komprehensif (Audit Trail)", True, False),
        (": Kehadiran modul linimasa ", False, False),
        ("timeline", False, True),
        (" yang mencatat setiap aksi pengguna (waktu pengajuan, nama verifikator, catatan revisi, dan tanggal tanda tangan) memberikan akuntabilitas yang tinggi bagi instansi pelayanan publik fakultas.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("e) ", True, False),
        ("Kompatibilitas Progressive Web App (PWA)", True, False),
        (": Platform dikembangkan agar kompatibel untuk diakses via perangkat mobile secara ringan dan responsif, mempermudah mahasiswa memantau status dokumen kapan saja dan di mana saja.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    # Kendala & Solusi
    {"type": "p", "runs": [("2. Analisis Kendala Pengembangan dan Solusi Pemecahan", True, False)], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 6, "after": 3, "line": 1.5},
    
    {"type": "p", "runs": [
        ("Selama proses riset rekayasa perangkat lunak ini, peneliti menghadapi sejumlah kendala teknis dan implementasi di lapangan. Kendala-kendala tersebut beserta solusi pemecahannya dijabarkan di bawah ini:", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("a) ", True, False),
        ("Kendala Render HTML WYSIWYG pada Dokumen Word (.docx)", True, False),
        (": Input teks formulir mahasiswa yang berasal dari WYSIWYG editor (Tiptap Editor) menyimpan tag-tag HTML (seperti `<p>`, `<b>`, `<i>`, `<u>`, `<br>`, `<ul>`, `<ol>`, `<li>`). Saat proses perakitan dokumen Word (`.docx`), tag HTML tersebut tercetak mentah sebagai teks biasa tanpa format visual. ", False, False),
        ("Solusi:", True, False),
        (" Peneliti mengembangkan parser HTML-to-OpenXML dinamis terpusat pada file XML mentah (`document.xml`, `header*.xml`, `footer*.xml`) sebelum file DOCX dikemas ulang oleh PHPWord. Tag inline (b, strong, i, em, u) diubah menjadi elemen formatting OpenXML (`<w:b/>`, `<w:i/>`, `<w:u w:val=\"single\"/>`), sedangkan tag paragraf dan list dipetakan menggunakan pemutus baris dinamis `<w:br/>` agar gaya visual bawaan template Word tetap terjaga tanpa terjadi degradasi tata letak visual.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("b) ", True, False),
        ("Kendala Ancaman Keamanan Cross-Site Scripting (XSS)", True, False),
        (": Kerentanan injeksi skrip jahat luar pada formulir pengajuan mahasiswa merupakan ancaman krusial pada sistem pelayanan publik online yang disoroti oleh validator sistem (Validator Ahli Sistem 2). ", False, False),
        ("Solusi:", True, False),
        (" Peneliti memperketat pertahanan web dengan memperketat middleware HTTP Header Content Security Policy (CSP) default pada Laravel untuk memblokir celah ", False, False),
        ("unsafe-eval", False, True),
        (" dan ", False, False),
        ("unsafe-inline", False, True),
        (" secara total, menyaring input teks richtext menggunakan sanitasi XSS ketat (", False, False),
        ("HtmlSanitizer", False, True),
        (" class), dan menerapkan enkripsi nama berkas privat mahasiswa.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("c) ", True, False),
        ("Kendala Aksesibilitas Integrasi Single Sign-On (SSO) Kampus", True, False),
        (": Keterbatasan akses API Single Sign-On (SSO) internal Universitas Lampung saat penelitian berlangsung menghalangi implementasi autentikasi terpadu secara langsung. ", False, False),
        ("Solusi:", True, False),
        (" Peneliti mengatasi hambatan ini dengan menyusun skema database terenkripsi yang fleksibel serta menyiapkan cetak biru (", False, False),
        ("blueprint", False, True),
        (") arsitektur integrasi basis data terpadu resmi Universitas Lampung, sehingga integrasi SSO dapat langsung dilakukan dengan mudah pada fase pengembangan instansi berikutnya pasca-penelitian.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    {"type": "p", "runs": [
        ("Secara keseluruhan, hasil validasi ahli (persentase rata-rata 91,95% - Sangat Valid) dan uji kepraktisan pengguna terbatas (92,13% - Sangat Praktis) membuktikan bahwa Website ULT FKIP Unila telah berhasil dirancang dan dikembangkan dengan standar kualitas rekayasa perangkat lunak yang sangat baik. Sistem ini terbukti andal dalam menyederhanakan alur birokrasi persuratan akademik, menjamin keamanan berkas sensitif, dan memberikan tingkat akuntabilitas pelayanan publik yang tinggi bagi Fakultas Keguruan dan Ilmu Pendidikan Universitas Lampung.", False, False)
    ], "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5},
    
    # Spacer
    {"type": "p", "runs": "", "style": "Normal", "align": WD_ALIGN_PARAGRAPH.LEFT, "before": 0, "after": 0, "line": 1.5}
]

def find_usability_tables(doc_obj):
    tbl_4_5 = None
    tbl_4_6 = None
    for t in doc_obj.tables:
        if len(t.rows) > 0:
            first_row = [cell.text.strip().replace('\n', ' ') for cell in t.rows[0].cells]
            if 'No' in first_row and 'Nama Responden' in first_row:
                tbl_4_5 = t
            elif any('distribusi' in cell.lower() or 'sangat praktis' in cell.lower() for cell in first_row) or (len(first_row) > 1 and 'sangat praktis' in str(t.rows[1].cells[1].text).lower()):
                tbl_4_6 = t
    return tbl_4_5, tbl_4_6

def anonymize_copied_table(tbl_xml, doc_obj):
    respondents_map = {
        "Anisa": "Responden 1",
        "Lisa": "Responden 2",
        "Riswan": "Responden 3",
        "Khaerul": "Responden 4",
        "Martin": "Responden 5",
        "Nurani": "Responden 6",
        "Aulia": "Responden 7",
        "Nazwa": "Responden 8",
        "Salsa": "Responden 9",
        "Andhini": "Responden 10",
        "Arya": "Responden 11",
        "Mita": "Responden 12",
        "Nabila": "Responden 13",
        "Nur": "Responden 14",
        "Rizky": "Responden 15",
        "Agus": "Responden 16",
        "Amrul": "Responden 17",
        "Tri": "Responden 18"
    }
    t = docx.table.Table(tbl_xml, doc_obj)
    for row in t.rows:
        if len(row.cells) > 1:
            cell = row.cells[1]
            text = cell.text.strip()
            if text.startswith("Responden Uji Kepraktisan "):
                num = text.replace("Responden Uji Kepraktisan ", "").strip()
                anon_name = f"Responden {num}"
                if len(cell.paragraphs) > 0:
                    p = cell.paragraphs[0]
                    if len(p.runs) > 0:
                        p.runs[0].text = anon_name
                        for r in p.runs[1:]:
                            r.text = ""
                    else:
                        p.text = anon_name
            elif text in respondents_map:
                anon_name = respondents_map[text]
                if len(cell.paragraphs) > 0:
                    p = cell.paragraphs[0]
                    if len(p.runs) > 0:
                        p.runs[0].text = anon_name
                        for r in p.runs[1:]:
                            r.text = ""
                    else:
                        p.text = anon_name

def remove_existing_ui_ux_documentation(doc_obj):
    body_el = doc_obj.element.body
    child_elements = body_el.getchildren()
    found_idx = -1
    for idx, child in enumerate(child_elements):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = docx.text.paragraph.Paragraph(child, doc_obj)
            if "Lampiran  . Dokumentasi Desain UI/UX" in p.text:
                found_idx = idx
                break
    if found_idx != -1:
        print(f"Found existing UI/UX documentation at index {found_idx}. Deleting to prevent duplication...")
        while len(body_el.getchildren()) > found_idx:
            body_el.remove(body_el[found_idx])

def append_ui_ux_documentation(doc_obj, highlight_color=None):
    remove_existing_ui_ux_documentation(doc_obj)
    
    p_head = doc_obj.add_paragraph()
    format_paragraph(p_head, doc_obj, "Normal", WD_ALIGN_PARAGRAPH.CENTER, spacing_before=12, spacing_after=6, line_spacing=1.0)
    p_head.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
    add_run_formatted(p_head, "Lampiran  . Dokumentasi Desain UI/UX Website ULT FKIP Unila", bold=True, font_size=12, highlight=highlight_color)
    
    p_desc = doc_obj.add_paragraph()
    format_paragraph(p_desc, doc_obj, "Normal", WD_ALIGN_PARAGRAPH.LEFT, spacing_before=0, spacing_after=6, line_spacing=1.5)
    add_run_formatted(p_desc, "Pengembangan Website Unit Layanan Terpadu (ULT) FKIP Universitas Lampung diawali dengan pembuatan rancangan antarmuka visual interaktif tingkat tinggi (high-fidelity prototype) menggunakan platform Figma. Rancangan ini mencakup visualisasi halaman utama publik (Public Portal), dasbor pengajuan mahasiswa (Student Portal), dasbor verifikasi staf (Admin Portal), serta dasbor persetujuan pimpinan (Signer Portal). Desain ini dirancang dengan estetika modern, kontras warna harmonis, tipografi ergonomis, dan kegunaan (usability) yang tinggi. Berikut disajikan tangkapan layar rancangan desain UI/UX dan kode akses QR untuk mengakses prototipe interaktif Figma secara langsung.", bold=False, font_size=12, highlight=highlight_color)
    
    p_img1 = doc_obj.add_paragraph()
    p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img1.paragraph_format.space_before = Pt(6)
    p_img1.paragraph_format.space_after = Pt(3)
    r_img1 = p_img1.add_run()
    img1_path = r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\images_ui_ux\ui_ux_portal_publik.jpg"
    r_img1.add_picture(img1_path, width=Inches(5.8))
    
    p_cap1 = doc_obj.add_paragraph()
    format_paragraph(p_cap1, doc_obj, "Normal", WD_ALIGN_PARAGRAPH.CENTER, spacing_before=3, spacing_after=12, line_spacing=1.0)
    add_run_formatted(p_cap1, "Gambar D.1. Tangkapan Layar Desain UI/UX Portal Publik (Home, Layanan, Pengumuman, Blog, Panduan)", bold=True, font_size=10, highlight=highlight_color)
    
    p_space = doc_obj.add_paragraph()
    format_paragraph(p_space, doc_obj, "Normal", WD_ALIGN_PARAGRAPH.LEFT, spacing_before=0, spacing_after=0, line_spacing=1.5)
    
    p_img2 = doc_obj.add_paragraph()
    p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img2.paragraph_format.space_before = Pt(6)
    p_img2.paragraph_format.space_after = Pt(3)
    r_img2 = p_img2.add_run()
    img2_path = r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\images_ui_ux\ui_ux_portal_internal.jpg"
    r_img2.add_picture(img2_path, width=Inches(5.8))
    
    p_cap2 = doc_obj.add_paragraph()
    format_paragraph(p_cap2, doc_obj, "Normal", WD_ALIGN_PARAGRAPH.CENTER, spacing_before=3, spacing_after=12, line_spacing=1.0)
    add_run_formatted(p_cap2, "Gambar D.2. Tangkapan Layar Desain UI/UX Portal Internal (Dasbor Staf, Admin Permohonan, Legalisir, Kritik/Saran)", bold=True, font_size=10, highlight=highlight_color)
    
    p_note = doc_obj.add_paragraph()
    format_paragraph(p_note, doc_obj, "Normal", WD_ALIGN_PARAGRAPH.LEFT, spacing_before=12, spacing_after=0, line_spacing=1.5)
    add_run_formatted(p_note, "Akses langsung menuju prototipe interaktif Figma dapat dilakukan dengan memindai QR Code yang tertera pada masing-masing dokumen desain di atas.", bold=False, italic=True, font_size=11, highlight=highlight_color)

def main():
    src_path = Path(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx")
    clean_out_path = Path(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx")
    high_out_path = Path(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx")
    
    if not src_path.exists():
        print(f"Error: {src_path} not found.")
        return
        
    print("Reading source document...")
    doc = docx.Document(src_path)
    body = doc.element.body
    child_elements = body.getchildren()
    
    bab4_idx = -1
    bab5_idx = -1
    lampiran_idx = -1
    
    for idx, child in enumerate(child_elements):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = docx.text.paragraph.Paragraph(child, doc)
            text = p.text.strip()
            if text == "IV. HASIL DAN PEMBAHASAN":
                bab4_idx = idx
            elif text == "V. KESIMPULAN DAN SARAN":
                bab5_idx = idx
            elif text.startswith("Lampiran") and "Hasil Uji Kepraktisan Pengguna" in text:
                lampiran_idx = idx
                
    if bab4_idx == -1 or bab5_idx == -1:
        print(f"Error: Could not locate Bab 4 (idx={bab4_idx}) or Bab 5 (idx={bab5_idx}) headings.")
        return
        
    print(f"Found BAB IV at element index {bab4_idx}")
    print(f"Found BAB V at element index {bab5_idx}")
    
    # 1. Fetch tables dynamically by headers and deepcopy them
    t_4_5_obj, t_4_6_obj = find_usability_tables(doc)
    if t_4_5_obj is None or t_4_6_obj is None:
        print("Error: Could not find usability testing tables dynamically in document!")
        return
        
    tbl_4_5_xml = copy.deepcopy(t_4_5_obj._tbl)
    tbl_4_6_xml = copy.deepcopy(t_4_6_obj._tbl)
    
    # Anonymize usability table
    anonymize_copied_table(tbl_4_5_xml, doc)
    
    # 2. Safely and dynamically delete old usability tables from the document body to prevent duplication
    t_4_5_obj._tbl.getparent().remove(t_4_5_obj._tbl)
    t_4_6_obj._tbl.getparent().remove(t_4_6_obj._tbl)
    
    # 3. Delete elements in between [bab4_idx + 1, bab5_idx - 1]
    # Because elements shift, we repeatedly remove the element at bab4_idx + 1
    # We must trace indexes again since we removed the tables!
    child_elements = body.getchildren()
    bab4_idx = -1
    bab5_idx = -1
    for idx, child in enumerate(child_elements):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = docx.text.paragraph.Paragraph(child, doc)
            text = p.text.strip()
            if text == "IV. HASIL DAN PEMBAHASAN":
                bab4_idx = idx
            elif text == "V. KESIMPULAN DAN SARAN":
                bab5_idx = idx
                
    del_count = bab5_idx - (bab4_idx + 1)
    print(f"Deleting {del_count} elements between BAB IV and BAB V...")
    for _ in range(del_count):
        body.remove(body[bab4_idx + 1])
        
    # Helpers for insertion
    def insert_paragraph_dynamic(doc_obj, body_obj, runs_data, style, align, before, after, line, insert_idx, highlight=None):
        p = doc_obj.add_paragraph()
        format_paragraph(p, doc_obj, style, align, before, after, line)
        if isinstance(runs_data, str):
            add_run_formatted(p, runs_data, bold=False, italic=False, highlight=highlight)
        else:
            for rd in runs_data:
                txt = rd[0]
                bld = rd[1] if len(rd) > 1 else False
                it = rd[2] if len(rd) > 2 else False
                add_run_formatted(p, txt, bold=bld, italic=it, highlight=highlight)
        p_el = p._p
        body_obj.remove(p_el)
        body_obj.insert(insert_idx, p_el)
        return insert_idx + 1

    def insert_table_4_1_dynamic(doc_obj, body_obj, insert_idx, highlight=None):
        t = doc_obj.add_table(rows=5, cols=6)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        data = [
            ["Aspek Kelayakan", "Total Skor", "Skor Maksimal", "Persentase (%)", "Rerata Aiken's V", "Kriteria Kelayakan"],
            ["Ahli Materi (Buku Panduan)", "315", "330", "95,45%", "0,94", "Sangat Valid / Layak"],
            ["Ahli Media (Visual dan UI/UX)", "196", "210", "93,33%", "0,92", "Sangat Valid / Layak"],
            ["Ahli Sistem (Keamanan dan Kinerja)", "289", "330", "87,58%", "0,85", "Sangat Valid / Layak"],
            ["RATA-RATA KESELURUHAN", "800", "870", "91,95%", "0,90", "Sangat Valid / Layak"]
        ]
        for r_idx, row in enumerate(t.rows):
            for c_idx, cell in enumerate(row.cells):
                cell.text = ""
                p = cell.paragraphs[0]
                is_bold = (r_idx == 0 or r_idx == 4)
                format_paragraph(p, doc_obj, style_name="Normal", alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_before=2, spacing_after=2, line_spacing=1.0)
                add_run_formatted(p, data[r_idx][c_idx], bold=is_bold, font_size=10, highlight=highlight)
                set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        t_el = t._tbl
        body_obj.remove(t_el)
        body_obj.insert(insert_idx, t_el)
        return insert_idx + 1

    # Clean text generation loop
    curr_idx = bab4_idx + 1
    print("Writing Clean Version dynamically...")
    for item in NASKAH_BAB_4:
        if item["type"] == "p":
            curr_idx = insert_paragraph_dynamic(doc, body, item["runs"], item["style"], item["align"], item["before"], item["after"], item["line"], curr_idx, highlight=None)
        elif item["type"] == "t_val":
            curr_idx = insert_table_4_1_dynamic(doc, body, curr_idx, highlight=None)
            
    # Move usability tables to appendix for Clean
    child_elements = body.getchildren()
    new_lampiran_idx = -1
    for idx, child in enumerate(child_elements):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = docx.text.paragraph.Paragraph(child, doc)
            text = p.text.strip()
            if text.startswith("Lampiran") and "Hasil Uji Kepraktisan Pengguna" in text:
                new_lampiran_idx = idx
                break
                
    if new_lampiran_idx != -1:
        print(f"Moving Usability tables into Lampiran at index {new_lampiran_idx}...")
        p_space = doc.add_paragraph()
        format_paragraph(p_space, doc, "Normal", WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.5)
        body.remove(p_space._p)
        body.insert(new_lampiran_idx + 1, p_space._p)
        
        p_cap5 = doc.add_paragraph()
        format_paragraph(p_cap5, doc, "Normal", WD_ALIGN_PARAGRAPH.CENTER, 12, 3, 1.0)
        add_run_formatted(p_cap5, "Tabel C.1. Matriks Hasil Skor Kuesioner Uji Kepraktisan Responden", bold=True, font_size=11)
        body.remove(p_cap5._p)
        body.insert(new_lampiran_idx + 2, p_cap5._p)
        
        body.insert(new_lampiran_idx + 3, tbl_4_5_xml)
        
        p_src5 = doc.add_paragraph()
        format_paragraph(p_src5, doc, "Normal", WD_ALIGN_PARAGRAPH.CENTER, 3, 12, 1.0)
        add_run_formatted(p_src5, "Sumber: Data primer olahan peneliti (2026)", italic=True, font_size=11)
        body.remove(p_src5._p)
        body.insert(new_lampiran_idx + 4, p_src5._p)
        
        p_space2 = doc.add_paragraph()
        format_paragraph(p_space2, doc, "Normal", WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.5)
        body.remove(p_space2._p)
        body.insert(new_lampiran_idx + 5, p_space2._p)
        
        p_cap6 = doc.add_paragraph()
        format_paragraph(p_cap6, doc, "Normal", WD_ALIGN_PARAGRAPH.CENTER, 12, 3, 1.0)
        add_run_formatted(p_cap6, "Tabel C.2. Rekapitulasi Distribusi Kategori Uji Kepraktisan (Pilihan Bagian F)", bold=True, font_size=11)
        body.remove(p_cap6._p)
        body.insert(new_lampiran_idx + 6, p_cap6._p)
        
        body.insert(new_lampiran_idx + 7, tbl_4_6_xml)
        
        p_src6 = doc.add_paragraph()
        format_paragraph(p_src6, doc, "Normal", WD_ALIGN_PARAGRAPH.CENTER, 3, 12, 1.0)
        add_run_formatted(p_src6, "Sumber: Data primer olahan peneliti (2026)", italic=True, font_size=11)
        body.remove(p_src6._p)
        body.insert(new_lampiran_idx + 8, p_src6._p)
        print("Moved kepraktisan tables successfully!")

    # 4. Generate Highlighted Version (read source before overwriting!)
    print("\nReading source document again for Highlighted version...")
    doc_high = docx.Document(src_path)
    body_high = doc_high.element.body
    child_elements_high = body_high.getchildren()
    
    bab4_idx_h = -1
    bab5_idx_h = -1
    lampiran_idx_h = -1
    for idx, child in enumerate(child_elements_high):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = docx.text.paragraph.Paragraph(child, doc_high)
            text = p.text.strip()
            if text == "IV. HASIL DAN PEMBAHASAN":
                bab4_idx_h = idx
            elif text == "V. KESIMPULAN DAN SARAN":
                bab5_idx_h = idx
            elif text.startswith("Lampiran") and "Hasil Uji Kepraktisan Pengguna" in text:
                lampiran_idx_h = idx
                
    # Fetch tables again dynamically inside doc_high before deleting Bab IV
    t_4_5_h_obj, t_4_6_h_obj = find_usability_tables(doc_high)
    tbl_4_5_xml_h = copy.deepcopy(t_4_5_h_obj._tbl)
    tbl_4_6_xml_h = copy.deepcopy(t_4_6_h_obj._tbl)
    
    # Anonymize highlighted usability table
    anonymize_copied_table(tbl_4_5_xml_h, doc_high)
    
    # Remove usability tables dynamically from doc_high
    t_4_5_h_obj._tbl.getparent().remove(t_4_5_h_obj._tbl)
    t_4_6_h_obj._tbl.getparent().remove(t_4_6_h_obj._tbl)
    
    # Re-trace indexes in doc_high since we removed tables
    child_elements_high = body_high.getchildren()
    bab4_idx_h = -1
    bab5_idx_h = -1
    for idx, child in enumerate(child_elements_high):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = docx.text.paragraph.Paragraph(child, doc_high)
            text = p.text.strip()
            if text == "IV. HASIL DAN PEMBAHASAN":
                bab4_idx_h = idx
            elif text == "V. KESIMPULAN DAN SARAN":
                bab5_idx_h = idx
                
    del_count_h = bab5_idx_h - (bab4_idx_h + 1)
    print(f"Deleting {del_count_h} elements between BAB IV and BAB V inside doc_high...")
    for _ in range(del_count_h):
        body_high.remove(body_high[bab4_idx_h + 1])
        
    curr_idx_h = bab4_idx_h + 1
    print("Writing Highlighted Version dynamically...")
    for item in NASKAH_BAB_4:
        if item["type"] == "p":
            curr_idx_h = insert_paragraph_dynamic(doc_high, body_high, item["runs"], item["style"], item["align"], item["before"], item["after"], item["line"], curr_idx_h, highlight=WD_COLOR_INDEX.YELLOW)
        elif item["type"] == "t_val":
            curr_idx_h = insert_table_4_1_dynamic(doc_high, body_high, curr_idx_h, highlight=WD_COLOR_INDEX.YELLOW)
            
    # Move usability tables for Highlighted
    child_elements_high = body_high.getchildren()
    new_lampiran_idx_h = -1
    for idx, child in enumerate(child_elements_high):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = docx.text.paragraph.Paragraph(child, doc_high)
            text = p.text.strip()
            if text.startswith("Lampiran") and "Hasil Uji Kepraktisan Pengguna" in text:
                new_lampiran_idx_h = idx
                break
                
    if new_lampiran_idx_h != -1:
        print(f"Moving Usability tables into Highlighted Lampiran at index {new_lampiran_idx_h}...")
        p_space = doc_high.add_paragraph()
        format_paragraph(p_space, doc_high, "Normal", WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.5)
        body_high.remove(p_space._p)
        body_high.insert(new_lampiran_idx_h + 1, p_space._p)
        
        p_cap5 = doc_high.add_paragraph()
        format_paragraph(p_cap5, doc_high, "Normal", WD_ALIGN_PARAGRAPH.CENTER, 12, 3, 1.0)
        add_run_formatted(p_cap5, "Tabel C.1. Matriks Hasil Skor Kuesioner Uji Kepraktisan Responden", bold=True, font_size=11, highlight=WD_COLOR_INDEX.YELLOW)
        body_high.remove(p_cap5._p)
        body_high.insert(new_lampiran_idx_h + 2, p_cap5._p)
        
        # Highlight cell contents in tbl_4_5_xml_h
        t_4_5 = docx.table.Table(tbl_4_5_xml_h, doc_high)
        for row in t_4_5.rows:
            for cell in row.cells:
                for p_cell in cell.paragraphs:
                    for run in p_cell.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        body_high.insert(new_lampiran_idx_h + 3, tbl_4_5_xml_h)
        
        p_src5 = doc_high.add_paragraph()
        format_paragraph(p_src5, doc_high, "Normal", WD_ALIGN_PARAGRAPH.CENTER, 3, 12, 1.0)
        add_run_formatted(p_src5, "Sumber: Data primer olahan peneliti (2026)", italic=True, font_size=11, highlight=WD_COLOR_INDEX.YELLOW)
        body_high.remove(p_src5._p)
        body_high.insert(new_lampiran_idx_h + 4, p_src5._p)
        
        p_space2 = doc_high.add_paragraph()
        format_paragraph(p_space2, doc_high, "Normal", WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.5)
        body_high.remove(p_space2._p)
        body_high.insert(new_lampiran_idx_h + 5, p_space2._p)
        
        p_cap6 = doc_high.add_paragraph()
        format_paragraph(p_cap6, doc_high, "Normal", WD_ALIGN_PARAGRAPH.CENTER, 12, 3, 1.0)
        add_run_formatted(p_cap6, "Tabel C.2. Rekapitulasi Distribusi Kategori Uji Kepraktisan (Pilihan Bagian F)", bold=True, font_size=11, highlight=WD_COLOR_INDEX.YELLOW)
        body_high.remove(p_cap6._p)
        body_high.insert(new_lampiran_idx_h + 6, p_cap6._p)
        
        # Highlight cells in tbl_4_6_xml_h
        t_4_6 = docx.table.Table(tbl_4_6_xml_h, doc_high)
        for row in t_4_6.rows:
            for cell in row.cells:
                for p_cell in cell.paragraphs:
                    for run in p_cell.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        body_high.insert(new_lampiran_idx_h + 7, tbl_4_6_xml_h)
        
        p_src6 = doc_high.add_paragraph()
        format_paragraph(p_src6, doc_high, "Normal", WD_ALIGN_PARAGRAPH.CENTER, 3, 12, 1.0)
        add_run_formatted(p_src6, "Sumber: Data primer olahan peneliti (2026)", italic=True, font_size=11, highlight=WD_COLOR_INDEX.YELLOW)
        body_high.remove(p_src6._p)
        body_high.insert(new_lampiran_idx_h + 8, p_src6._p)
        print("Moved kepraktisan tables with highlight successfully!")

    print("Appending UI/UX documentation...")
    append_ui_ux_documentation(doc_high, highlight_color=WD_COLOR_INDEX.YELLOW)
    append_ui_ux_documentation(doc)

    doc_high.save(high_out_path)
    print(f"Saved Highlighted document to {high_out_path}")
    
    # Save Clean document at the very end to avoid overwriting the source path too early
    doc.save(clean_out_path)
    print(f"Saved Clean document to {clean_out_path}")
    
if __name__ == "__main__":
    main()
