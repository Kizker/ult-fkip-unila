import docx
from docx.enum.text import WD_COLOR_INDEX

def inject_citations(doc_path, is_highlighted=False):
    doc = docx.Document(doc_path)
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        
        # Paragraph 644 equivalent
        if text.startswith('Jumlah dan kualifikasi validator ahli ditetapkan') and 'ahli materi = 3' in text:
            p.text = text # Reset the text just in case to clear formatting
            run = p.runs[0] if p.runs else p.add_run()
            run.text = text
            # Append citation
            added_text = " Penggunaan keahlian multi-disiplin ini sejalan dengan standar evaluasi kelayakan sistem (Nieveen, 1999) yang mengharuskan produk dievaluasi dari sisi relevance (materi) dan consistency (sistem/media)."
            added_run = p.add_run(added_text)
            if is_highlighted:
                added_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                
        # Paragraph 647 equivalent
        elif text.startswith('Jumlah responden uji coba ditetapkan') and 'mahasiswa = 12' in text:
            # We want to insert text before "Kriteria inklusi responden:"
            parts = text.split('Kriteria inklusi responden:')
            if len(parts) == 2:
                base_text = parts[0].strip()
                # If "Teknik pemilihan responden menggunakan purposive sampling." is at the end of base_text
                if base_text.endswith('purposive sampling.'):
                    p.text = '' # Clear and rebuild
                    p.add_run(base_text)
                    
                    added_text = " Purposive sampling adalah teknik penentuan sampel dengan pertimbangan tertentu agar sampel yang dipilih relevan dengan kebutuhan pengujian (Sugiyono, 2013). Ukuran sampel sebanyak 18 pengguna ini dinilai sangat memadai untuk evaluasi kegunaan sistem, mengingat pengujian dengan 15-20 pengguna sudah mampu mengungkap lebih dari 95% masalah usability utama pada sebuah perangkat lunak (Faulkner, 2003; Weinger et al., 2010). "
                    added_run = p.add_run(added_text)
                    if is_highlighted:
                        added_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        
                    p.add_run('Kriteria inklusi responden:')
                    
    doc.save(doc_path)
    print(f"Successfully injected citations into {doc_path}")

if __name__ == '__main__':
    inject_citations(r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx', is_highlighted=False)
    inject_citations(r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx', is_highlighted=True)
