import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

img_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram\10_wireframe_beranda.png'

def process(doc_path):
    doc = docx.Document(doc_path)
    
    for p in doc.paragraphs:
        if 'GAMBAR_WIREFRAME_UTAMA_BELUM_ADA' in p.text:
            p.text = '' # clear placeholder
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(5.5))
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            
        if 'Gambar [NOMOR_GAMBAR]. Wireframe Beranda Utama' in p.text:
            p.text = p.text.replace('Gambar [NOMOR_GAMBAR].', 'Gambar 4.5.')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(12)
            # Ensure text is 10pt or 12pt (standard is 12pt for caption? Or 10pt?)
            # Actually just keep the run formatting by not overriding the whole text if possible, but replacing p.text replaces all runs.
            # We can just apply the font explicitly.
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10) # captions are usually 10pt or 11pt
                
    # Also fix the text in the first paragraph below it
    for p in doc.paragraphs:
        if 'Pada tampilan beranda utama sistem (Public Portal)' in p.text:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.line_spacing = 1.5
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)

    doc.save(doc_path)
    print(f"Injected wireframe into {os.path.basename(doc_path)}")

process(r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx')
process(r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx')
