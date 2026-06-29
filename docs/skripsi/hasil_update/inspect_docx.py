import docx

def inspect_doc(filename):
    try:
        doc = docx.Document(filename)
        for i in range(1050, min(1250, len(doc.paragraphs))):
            p = doc.paragraphs[i]
            text = p.text.strip()
            print(f"[{i}] {text[:60]} (Style: {p.style.name})")
    except Exception as e:
        print(f"Error: {e}")

inspect_doc("001_Skripsi_Andricha Dea Mitra_Clean.docx")
