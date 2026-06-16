import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def inject_diagrams(doc_path):
    doc = docx.Document(doc_path)
    
    diagrams = {
        "Gambar  4": r"c:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram\01_diagram_use_case.jpg",
        "Gambar  5": r"c:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram\02_diagram_arsitektur_sistem.jpg",
        "Gambar  6": r"c:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram\03_diagram_flowchart_dokumen.jpg",
        "Gambar  7": r"c:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram\04_diagram_erd_database.jpg",
        "Gambar  8": r"c:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram\05_diagram_sequence_parser.jpg"
    }

    # First, let's find the indices of the caption paragraphs.
    # Since we will modify the document, finding all indices first and then processing backwards 
    # (or finding one by one dynamically) is safer.
    
    for key, img_path in diagrams.items():
        if not os.path.exists(img_path):
            print(f"Warning: {img_path} not found!")
            continue

        # Find caption dynamically
        idx = -1
        for i, p in enumerate(doc.paragraphs):
            txt = p.text.strip()
            # Match strictly starting with the key (e.g., "Gambar 4", "Gambar  4")
            # Replace double spaces with single to be safe
            if txt.replace("  ", " ").startswith(key.replace("  ", " ")):
                idx = i
                break
                
        if idx != -1:
            caption_p = doc.paragraphs[idx]
            
            # Check if paragraph immediately before is an image paragraph
            if idx > 0:
                prev_p = doc.paragraphs[idx - 1]
                xml = prev_p._p.xml
                is_img_p = 'w:drawing' in xml or 'pic:pic' in xml or 'v:imagedata' in xml or 'v:shape' in xml
                
                # We should remove existing image to avoid duplicates
                if is_img_p and prev_p.text.strip() == "":
                    print(f"[{key}] Removing existing old image paragraph at index {idx - 1}")
                    prev_p._element.getparent().remove(prev_p._element)

            # Now insert the new high-res image
            print(f"[{key}] Inserting image {os.path.basename(img_path)} before caption")
            img_p = caption_p.insert_paragraph_before()
            run = img_p.add_run()
            # width=Inches(5.5) fits well within the 4-4-3-3 margins
            run.add_picture(img_path, width=Inches(5.5))
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            print(f"[{key}] Caption not found in document.")

    doc.save(doc_path)
    print(f"Successfully updated diagrams in {doc_path}")

if __name__ == '__main__':
    inject_diagrams(r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx')
    inject_diagrams(r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx')
