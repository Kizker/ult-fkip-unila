import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_highlight(run, is_highlighted):
    if is_highlighted:
        rPr = run._r.get_or_add_rPr()
        highlight = OxmlElement('w:highlight')
        highlight.set(qn('w:val'), 'yellow')
        rPr.append(highlight)

def insert_p_before(ref_p, text, style='Paragraph', is_highlighted=False, bold=False):
    new_p = ref_p.insert_paragraph_before(style=style)
    if text:
        run = new_p.add_run(text)
        if bold:
            run.font.bold = True
        add_highlight(run, is_highlighted)
    return new_p

def add_summary_paragraph(doc_path, is_highlighted=False):
    doc = docx.Document(doc_path)
    
    idx_instrumen = -1
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text == 'Instrumen Penelitian' or text == '3.6 Instrumen Penelitian':
            idx_instrumen = i
            break
            
    if idx_instrumen == -1:
        print(f"Error finding Instrumen Penelitian in {doc_path}")
        return

    p_inst = doc.paragraphs[idx_instrumen]
    
    # We want to insert the summary before 3.6 Instrumen Penelitian
    summary_text = "Berdasarkan rincian tersebut, secara keseluruhan subjek dalam penelitian ini berjumlah 27 orang yang dipilih melalui teknik purposive sampling. Tahap awal melibatkan 9 validator ahli yang bertugas mengevaluasi kelayakan aspek materi, media, dan sistem sebelum prototipe diujicobakan. Setelah dinyatakan valid, tahap selanjutnya melibatkan 18 responden pengguna dari berbagai elemen (mahasiswa, staf ULT, dan admin jurusan) untuk mengukur tingkat kepraktisan sistem di lapangan. Pelibatan berbagai unsur civitas akademika ini bertujuan agar umpan balik yang diperoleh bersifat komprehensif, representatif, dan relevan dengan alur birokrasi pelayanan dokumen di Fakultas Keguruan dan Ilmu Pendidikan Universitas Lampung."
    
    insert_p_before(p_inst, summary_text, 'Paragraph', is_highlighted)
    insert_p_before(p_inst, "", 'Normal')

    doc.save(doc_path)
    print(f"Successfully added summary paragraph to {doc_path}")

if __name__ == '__main__':
    add_summary_paragraph(r'001_Skripsi_Andricha Dea Mitra_Clean.docx', is_highlighted=False)
    add_summary_paragraph(r'001_Skripsi_Andricha Dea Mitra_Highlighted.docx', is_highlighted=True)
