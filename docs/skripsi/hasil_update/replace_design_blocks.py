import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_highlight(run, is_highlighted):
    if is_highlighted:
        rPr = run._r.get_or_add_rPr()
        highlight = OxmlElement('w:highlight')
        highlight.set(qn('w:val'), 'yellow')
        rPr.append(highlight)

def insert_p_before(ref_p, text, style='Paragraph', is_highlighted=False):
    new_p = ref_p.insert_paragraph_before(style=style)
    if text:
        run = new_p.add_run(text)
        add_highlight(run, is_highlighted)
    return new_p

def replace_design_blocks(doc_path, is_highlighted):
    doc = docx.Document(doc_path)
    
    # Identify indices of paragraphs we want to clear.
    # We will use text matching to find the start of each block.
    
    p_flowchart_start = -1
    p_erd_start = -1
    p_sequence_start = -1
    
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt.startswith("Diagram flowchart pada Gambar 6 merinci siklus hidup"):
            p_flowchart_start = i
        elif txt.startswith("Diagram ERD pada Gambar 7 menjabarkan delapan entitas"):
            p_erd_start = i
        elif txt.startswith("Diagram sequence pada Gambar 8 memvisualisasikan urutan interaksi"):
            p_sequence_start = i
            
    if p_flowchart_start == -1 or p_erd_start == -1 or p_sequence_start == -1:
        print(f"Could not find all blocks in {doc_path}")
        return

    # Clear Flowchart block (starts at p_flowchart_start, next 8 paragraphs)
    for i in range(p_flowchart_start, p_flowchart_start + 9):
        doc.paragraphs[i].text = ""
        
    # Clear ERD block (starts at p_erd_start, next 4 paragraphs)
    for i in range(p_erd_start, p_erd_start + 5):
        doc.paragraphs[i].text = ""

    # Clear Sequence block (starts at p_sequence_start, next 7 paragraphs)
    for i in range(p_sequence_start, p_sequence_start + 8):
        doc.paragraphs[i].text = ""
        
    # Inject Flowchart paragraphs
    t_flowchart = doc.paragraphs[p_flowchart_start]
    pA = "Diagram flowchart pada Gambar 6 merinci siklus hidup dokumen secara terstruktur tanpa menggunakan skema penomoran prosedural yang kaku. Mahasiswa mengawali proses dengan melakukan pengisian formulir dinamis menggunakan teks editor kaya (WYSIWYG Tiptap Editor) untuk menghasilkan draf awal dokumen. Pengajuan tersebut selanjutnya diteruskan ke portal Admin Jurusan untuk melewati tahapan verifikasi kelengkapan berkas pendukung secara seksama. Setelah dokumen dinyatakan valid pada tahap awal, sistem akan secara otomatis menerbitkan nomor surat resmi sesuai dengan pola penomoran unit program studi untuk meminimalisasi kesalahan administratif."
    pB = "Validasi lanjutan kemudian dieksekusi oleh Staf ULT sebagai representasi gerbang birokrasi fakultas untuk meninjau keabsahan keseluruhan data secara digital. Berkas yang telah disetujui fakultas akan masuk ke portal Pejabat untuk ditinjau secara substantif melalui pembubuhan tanda tangan elektronik pimpinan yang sah. Proses persetujuan tersebut secara otomatis memicu perakitan dokumen di mana modul parser mengolah tag-tag HTML mahasiswa menjadi format OpenXML yang sejalan dengan gaya bawaan template Word asli. Dokumen final tersebut kemudian diamankan ke dalam direktori privat sistem bersamaan dengan perekaman jejak audit, sebelum akhirnya mahasiswa menerima notifikasi penyelesaian untuk mengunduh berkas fisik secara personal."
    insert_p_before(t_flowchart, pA, 'Paragraph', is_highlighted)
    insert_p_before(t_flowchart, pB, 'Paragraph', is_highlighted)

    # Inject ERD paragraphs
    t_erd = doc.paragraphs[p_erd_start]
    pC = "Diagram ERD pada Gambar 7 memvisualisasikan arsitektur delapan entitas utama basis data MySQL yang menopang keseluruhan operasional sistem secara fungsional. Entitas pengguna (users) memegang peran sentral sebagai tabel master profil yang menyimpan kredensial dasar dan berelasi langsung dengan sistem kontrol akses berbasis peran. Di sisi operasional transaksional, entitas permohonan (requests) berfungsi sebagai pusat penyimpanan riwayat persuratan yang mencatat status permohonan, kode pelacakan pelapor, serta melacak tahapan dokumen secara presisi."
    pD = "Katalog jenis layanan persuratan akademik fakultas dikelola secara terpusat melalui entitas layanan (services) yang memuat konfigurasi rute template dokumen fisik. Relasi antartabel transaksional ini dioptimalkan lebih lanjut melalui perancangan arsitektur indeks pencarian pada berbagai kolom kunci primer maupun kunci pendatang. Optimalisasi pengindeksan tersebut memegang kontribusi krusial untuk mempercepat waktu respons eksekusi pencarian data, terutama saat administrator fakultas perlu memuat ribuan rekam jejak audit dokumen secara bersamaan."
    insert_p_before(t_erd, pC, 'Paragraph', is_highlighted)
    insert_p_before(t_erd, pD, 'Paragraph', is_highlighted)

    # Inject Sequence paragraphs
    t_seq = doc.paragraphs[p_sequence_start]
    pE = "Diagram sequence pada Gambar 8 memvisualisasikan dinamika pertukaran belasan pesan terstruktur yang melibatkan interaksi berkesinambungan antara antarmuka, pengendali, serta mesin perakit dokumen. Rangkaian instruksi tersebut bermula ketika Staf ULT mengeksekusi pemicu perakitan pada antarmuka portal operasional yang seketika meneruskan permintaan ke objek pengendali utama. Pengendali utama ini kemudian sepenuhnya mendelegasikan beban tugas eksekusi perakitan kepada layanan perakit dokumen untuk menangani keseluruhan logika manipulasi naskah Word."
    pF = "Kelas layanan perakit dokumen mengawali tugasnya dengan memuat salinan fail templat orisinal beserta ekstraksi data formulir permohonan sebelum meminta instruksi translasi teks dari modul pemilah tag HTML. Modul pemilah tersebut secara cermat menelusuri elemen model dokumen dan mereproduksi pemformatan gaya visual agar presisi dengan pengaturan bawaan tanpa merusak struktur rancangan aslinya. Setelah tahapan injeksi XML selesai, sistem secara otomatis mengemas ulang arsip berkas tersebut dan menyimpannya di dalam penyimpanan terenkripsi sebelum memancarkan notifikasi keberhasilan kepada layar staf."
    insert_p_before(t_seq, pE, 'Paragraph', is_highlighted)
    insert_p_before(t_seq, pF, 'Paragraph', is_highlighted)

    doc.save(doc_path)
    print(f"Saved {doc_path} with new paragraph blocks.")

if __name__ == '__main__':
    replace_design_blocks(r'001_Skripsi_Andricha Dea Mitra_Clean.docx', is_highlighted=False)
    replace_design_blocks(r'001_Skripsi_Andricha Dea Mitra_Highlighted.docx', is_highlighted=True)
