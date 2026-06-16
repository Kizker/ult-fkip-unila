import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_highlight(run, is_highlighted):
    if is_highlighted:
        rPr = run._r.get_or_add_rPr()
        highlight = OxmlElement('w:highlight')
        highlight.set(qn('w:val'), 'yellow')
        rPr.append(highlight)

def insert_p_before(ref_p, text, style='Paragraph', is_highlighted=False, bold=False):
    new_p = ref_p.insert_paragraph_before(style=style)
    if text:
        run = new_p.add_run(text)
        if bold:
            run.font.bold = True
        add_highlight(run, is_highlighted)
    return new_p

def apply_revisi(doc_path, is_highlighted=False):
    doc = docx.Document(doc_path)
    
    idx_subjek = -1
    idx_instrumen = -1
    idx_analisis = -1
    proc_texts = []
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if '3.5 Subjek Penelitian dan Sampling' in text:
            idx_subjek = i
        elif text == 'Instrumen Penelitian' or text == '3.6 Instrumen Penelitian':
            idx_instrumen = i
        elif text == 'Teknik Analisis Data' or text == '3.7 Teknik Analisis Data':
            idx_analisis = i
        
        # Capture the procedural texts
        if text.startswith('Setting uji coba dilakukan dalam kondisi terkontrol'):
            proc_texts.append(text)
        elif text.startswith('Urutan pelaksanaan uji coba:'):
            proc_texts.append(text)

    if idx_instrumen == -1 or idx_subjek == -1 or idx_analisis == -1:
        print(f"Error finding headings in {doc_path}: subjek={idx_subjek}, instrumen={idx_instrumen}, analisis={idx_analisis}")
        return

    # Move procedural text to right before "Teknik Analisis Data"
    p_analisis = doc.paragraphs[idx_analisis]
    for pt in proc_texts:
        insert_p_before(p_analisis, pt, 'Paragraph', is_highlighted)

    # Clear old Subjek text
    for i in range(idx_subjek + 1, idx_instrumen):
        doc.paragraphs[i].text = ''

    # Insert new Subjek text before Instrumen
    p_inst = doc.paragraphs[idx_instrumen]
    
    par1 = "Populasi dalam penelitian ini adalah seluruh civitas akademika Fakultas Keguruan dan Ilmu Pendidikan (FKIP) Universitas Lampung yang terdiri dari staf, administrator jurusan, dan mahasiswa aktif. Mengingat penelitian ini menggunakan pendekatan Research and Development (R&D), penentuan subjek penelitian dibagi ke dalam dua kelompok utama, yaitu kelompok validator ahli untuk uji kelayakan, dan responden pengguna untuk uji kepraktisan (implementasi)."
    par2 = "Teknik pengambilan sampel yang digunakan dalam penelitian ini adalah purposive sampling. Purposive sampling adalah teknik penentuan sampel dengan pertimbangan atau kriteria tertentu yang relevan dengan tujuan penelitian (Sugiyono, 2013). Berdasarkan pendekatan tersebut, rincian penetapan subjek penelitian adalah sebagai berikut:"
    
    insert_p_before(p_inst, par1, 'Paragraph', is_highlighted)
    insert_p_before(p_inst, par2, 'Paragraph', is_highlighted)
    insert_p_before(p_inst, "1. Validator Ahli", 'Paragraph', is_highlighted, bold=True)
    insert_p_before(p_inst, "Subjek validator ahli ditetapkan sebanyak 9 (sembilan) orang yang terbagi ke dalam tiga kategori untuk menguji prototipe awal sistem, yaitu:", 'Paragraph', is_highlighted)
    
    # List paragraphs
    insert_p_before(p_inst, "Ahli Materi (3 orang): Berperan menilai kesesuaian konten layanan, kelengkapan syarat dokumen, dan alur administrasi.", 'List Paragraph', is_highlighted)
    insert_p_before(p_inst, "Ahli Media (3 orang): Berperan menilai aspek desain antarmuka (UI), konsistensi navigasi, dan aspek usability.", 'List Paragraph', is_highlighted)
    insert_p_before(p_inst, "Ahli Sistem (3 orang): Berperan menilai kelayakan fungsional sistem, kestabilan fitur inti, keamanan, dan konsistensi alur berbasis role sebelum produk diuji oleh pengguna lapangan.", 'List Paragraph', is_highlighted)

    insert_p_before(p_inst, "2. Responden Uji Coba (Uji Kepraktisan)", 'Paragraph', is_highlighted, bold=True)
    insert_p_before(p_inst, "Jumlah sampel responden pengguna ditetapkan sebanyak 18 (delapan belas) orang. Ukuran sampel ini dinilai memadai, mengacu pada teori pengujian kegunaan sistem informasi (Faulkner, 2003), di mana pengujian dengan 15-20 pengguna sudah mampu mengungkap lebih dari 95% masalah usability utama pada sebuah perangkat lunak.", 'Paragraph', is_highlighted)
    insert_p_before(p_inst, "Kriteria pemilihan 18 responden tersebut ditetapkan sebagai berikut:", 'Paragraph', is_highlighted)
    
    insert_p_before(p_inst, "Mahasiswa (12 orang): Merupakan mahasiswa aktif FKIP Unila yang pernah atau sedang membutuhkan layanan administrasi. Sampel diambil secara proporsional sebagai perwakilan dari empat rumpun ilmu yang ada (PBS, PIP, PIPS, dan PMIPA).", 'List Paragraph', is_highlighted)
    insert_p_before(p_inst, "Staf ULT (3 orang): Merupakan staf pegawai administrasi yang secara langsung bertugas mengelola dan memproses permohonan dokumen mahasiswa di loket layanan ULT.", 'List Paragraph', is_highlighted)
    insert_p_before(p_inst, "Administrator Jurusan/Prodi (3 orang): Merupakan admin tingkat program studi yang memahami alur persetujuan dokumen sebelum diproses oleh pihak fakultas.", 'List Paragraph', is_highlighted)
    
    insert_p_before(p_inst, "", 'Normal')
    p_tabel_title = insert_p_before(p_inst, "Tabel 3.2 Subjek Penelitian dan Sampling", 'Normal', is_highlighted, bold=True)
    p_tabel_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Table
    table = doc.add_table(rows=8, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Move table before p_inst
    p_inst._p.addprevious(table._tbl)
    
    headers = ['No.', 'Kelompok Subjek', 'Kategori Subjek', 'Jumlah', 'Keterangan']
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                add_highlight(r, is_highlighted)
    
    data = [
        ['1', 'Validator Ahli', 'Ahli Materi', '3', 'Dosen/Praktisi Administrasi'],
        ['', '', 'Ahli Media', '3', 'Dosen/Praktisi UI/UX'],
        ['', '', 'Ahli Sistem', '3', 'Praktisi Rekayasa Perangkat Lunak'],
        ['2', 'Responden Uji Coba', 'Mahasiswa', '12', 'Perwakilan Rumpun Ilmu FKIP'],
        ['', '', 'Staf ULT', '3', 'Staf Operasional Pelayanan'],
        ['', '', 'Admin Jurusan', '3', 'Admin Verifikator Program Studi'],
        ['Total', '', '', '27', '']
    ]
    
    for row_idx, row_data in enumerate(data):
        for col_idx, val in enumerate(row_data):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = val
            for p in cell.paragraphs:
                if col_idx == 0 or col_idx == 3 or val == 'Total':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    if val == 'Total' or (row_idx == 6 and col_idx == 3):
                        r.font.bold = True
                    add_highlight(r, is_highlighted)
                    
    # Merge cells for "Total" row and others
    table.cell(1, 0).merge(table.cell(3, 0)) # No. 1
    table.cell(1, 1).merge(table.cell(3, 1)) # Validator Ahli
    table.cell(4, 0).merge(table.cell(6, 0)) # No. 2
    table.cell(4, 1).merge(table.cell(6, 1)) # Responden
    table.cell(7, 0).merge(table.cell(7, 2)) # Total merge cols 1-3
    
    widths = [0.4, 2.0, 1.5, 0.6, 2.5]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)

    insert_p_before(p_inst, "", 'Normal')

    doc.save(doc_path)
    print(f"Successfully applied revisi to {doc_path}")

if __name__ == '__main__':
    apply_revisi(r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx', is_highlighted=False)
    apply_revisi(r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx', is_highlighted=True)
