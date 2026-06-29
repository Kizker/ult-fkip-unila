import docx
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH

def process_images_in_docx(filename, output_filename, is_highlighted):
    doc = docx.Document(filename)
    
    img_counter = 37
    
    # Identify paragraphs that have tags
    # Since we'll be inserting new paragraphs, we should collect the original paragraph elements
    # and then process them.
    target_elements = []
    
    start_idx = -1
    end_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if "4.2 Pembahasan Hasil Penelitian" in text:
            start_idx = i
        elif text.startswith("KESIMPULAN DAN SARAN"):
            end_idx = i
            if start_idx != -1:
                break
                
    if start_idx == -1 or end_idx == -1:
        print(f"Failed to find bounds in {filename}.")
        return

    for i in range(start_idx, end_idx):
        if "[GAMBAR:" in doc.paragraphs[i].text:
            target_elements.append(doc.paragraphs[i]._element)
            
    for p_elem in target_elements:
        # Find the paragraph wrapper for this element
        # (This avoids index issues since we're tracking the underlying XML element)
        paragraphs = list(doc.paragraphs)
        try:
            idx = [p._element for p in paragraphs].index(p_elem)
        except ValueError:
            continue
            
        p = paragraphs[idx]
        text = p.text
        matches = re.findall(r'\[GAMBAR:\s*(.+?)\]', text)
        
        if not matches:
            continue
            
        new_text = text
        captions = []
        for match in matches:
            caption_title = match.strip()
            replacement = f"(Lihat Gambar {img_counter})"
            new_text = new_text.replace(f"[GAMBAR: {match}]", replacement)
            captions.append((img_counter, caption_title))
            img_counter += 1
            
        # Update paragraph text and keep highlight if needed
        p.clear()
        run = p.add_run(new_text)
        if is_highlighted:
            run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
            
        # We want to insert AFTER this paragraph.
        # Since we are iterating, we can just insert BEFORE the next paragraph.
        next_p = paragraphs[idx + 1]
        
        for cap_num, cap_title in captions:
            # We insert in forward order before next_p, so they end up in order AFTER p.
            ph_p = next_p.insert_paragraph_before(f"[Placeholder Gambar {cap_num}]", "Normal")
            ph_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            cap_title_clean = cap_title + "." if not cap_title.endswith(".") else cap_title
            cap_p = next_p.insert_paragraph_before(f"Gambar {cap_num}. {cap_title_clean}", "Caption")
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Spacer paragraph
            next_p.insert_paragraph_before("", "Normal")

    doc.save(output_filename)
    print(f"Successfully processed images for {output_filename}")

process_images_in_docx(
    "001_Skripsi_Andricha Dea Mitra_Clean.docx",
    "001_Skripsi_Andricha Dea Mitra_Clean.docx",
    is_highlighted=False
)

process_images_in_docx(
    "001_Skripsi_Andricha Dea Mitra_Highlighted.docx",
    "001_Skripsi_Andricha Dea Mitra_Highlighted.docx",
    is_highlighted=True
)
