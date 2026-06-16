import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

img1 = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram\11_wireframe_beranda.png'
img2 = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram\12_wireframe_student.png'
img3 = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram\13_wireframe_admin.png'

def process(doc_path):
    doc = docx.Document(doc_path)
    
    # 1. Update the existing first wireframe
    for i, p in enumerate(doc.paragraphs):
        if 'Gambar 4.5.' in p.text and 'Wireframe Beranda' in p.text:
            # the picture is in the previous paragraph
            pic_p = doc.paragraphs[i-1]
            pic_p.clear()
            run = pic_p.add_run()
            run.add_picture(img1, width=Inches(5.5))
            
    # 2. Find the paragraph describing the public portal
    target_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if 'Pada tampilan beranda utama sistem (Public Portal)' in p.text:
            target_idx = i
            break
            
    if target_idx != -1:
        # We need to insert AFTER target_idx
        # docx doesn't have a direct `insert_paragraph_after`, but we can use `insert_paragraph_before` on the next paragraph.
        # So we insert before target_idx + 1
        
        next_p = doc.paragraphs[target_idx + 1]
        
        # Insert Image 2
        p_img2 = next_p.insert_paragraph_before()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img2.paragraph_format.space_before = Pt(12)
        p_img2.paragraph_format.space_after = Pt(12)
        r_img2 = p_img2.add_run()
        r_img2.add_picture(img2, width=Inches(5.5))
        
        # Insert Caption 2
        p_cap2 = next_p.insert_paragraph_before('Gambar 4.6. Wireframe Dasbor Mahasiswa (Student Portal) Sistem ULT FKIP Unila.')
        p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap2.paragraph_format.space_before = Pt(0)
        p_cap2.paragraph_format.space_after = Pt(12)
        for r in p_cap2.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
            
        # Insert Desc 2
        p_desc2 = next_p.insert_paragraph_before('Pada portal dasbor mahasiswa (Student Portal), tata letak didesain menggunakan format dua kolom interaktif tanpa menu navigasi sisi (sidebar). Bagian atas menampilkan pahlawan visual ringkas (hero header) berisi informasi profil pengguna dan tombol aksi cepat. Kolom utama di sebelah kiri difokuskan pada daftar aktivitas pengajuan permohonan terbaru secara mendetail, sedangkan kolom kanan (side panel) difungsikan sebagai agregasi visual indikator status permohonan yang memberikan ringkasan kuantitatif harian. Desain ini memastikan mahasiswa dapat memantau progres permohonan secara mandiri, efisien, dan transparan.')
        p_desc2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_desc2.paragraph_format.first_line_indent = Inches(0.5)
        p_desc2.paragraph_format.line_spacing = 1.5
        for r in p_desc2.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            
        # Insert Image 3
        p_img3 = next_p.insert_paragraph_before()
        p_img3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img3.paragraph_format.space_before = Pt(12)
        p_img3.paragraph_format.space_after = Pt(12)
        r_img3 = p_img3.add_run()
        r_img3.add_picture(img3, width=Inches(5.5))
        
        # Insert Caption 3
        p_cap3 = next_p.insert_paragraph_before('Gambar 4.7. Wireframe Dasbor Operasional (Admin Portal) Sistem ULT FKIP Unila.')
        p_cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap3.paragraph_format.space_before = Pt(0)
        p_cap3.paragraph_format.space_after = Pt(12)
        for r in p_cap3.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
            
        # Insert Desc 3
        p_desc3 = next_p.insert_paragraph_before('Antarmuka dasbor operasional bagi jajaran staf dan administrator ULT (Admin Portal) mengusung tata letak tiga bagian dengan penambahan menu navigasi tersemat vertikal (sidebar) di sisi kiri layar. Struktur hierarki ini memberikan fleksibilitas akses yang cepat ke berbagai modul administratif tingkat tinggi seperti peninjauan antrian permohonan, manajemen inventori layanan, pengaturan templat dokumen elektronik, serta hak akses peran pengguna. Area konten utama dipertahankan menggunakan pola dua kolom yang mengutamakan rincian antrian operasional aktif di kolom primer, bersandingan dengan indikator kinerja (KPI) status kumulatif pada kolom sekunder.')
        p_desc3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_desc3.paragraph_format.first_line_indent = Inches(0.5)
        p_desc3.paragraph_format.line_spacing = 1.5
        for r in p_desc3.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)

    doc.save(doc_path)
    print(f"Successfully processed {os.path.basename(doc_path)}")

process(r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx')
process(r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx')
