import docx

def update_pplg_curriculum(filename, output_filename):
    doc = docx.Document(filename)
    
    # Target elements
    heading_p = None
    para_1 = None
    table_1 = None
    table_2 = None
    para_2 = None
    
    # Find the heading
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading 2') and "Analisis Kurikulum Pendidikan Teknologi Informasi" in p.text:
            heading_p = p
            para_1 = doc.paragraphs[i+1]
            break
            
    if not heading_p:
        print(f"Heading not found in {filename}!")
        return

    # Find the tables that come after heading_p
    # python-docx doesn't store tables and paragraphs sequentially in a single list
    # But we can find the tables based on their content, since they were recently added
    for t in doc.tables:
        if "Capaian Pembelajaran" in t.cell(0,0).text:
            table_1 = t
        if "Kedalaman Materi" in t.cell(0,0).text:
            table_2 = t
            
    # Find para_2 (the closing paragraph)
    for i, p in enumerate(doc.paragraphs):
        if "Pemetaan struktur matriks analisis kurikulum" in p.text:
            para_2 = p
            break
            
    # Update Heading
    heading_p.text = "Analisis Kurikulum Dasar-dasar Pengembangan Perangkat Lunak dan Gim (PPLG)"
    
    # Update Para 1
    para_1.text = "Pengembangan Website Unit Layanan Terpadu (ULT) di Fakultas Keguruan dan Ilmu Pendidikan Universitas Lampung ini turut dianalisis potensinya sebagai sumber belajar kontekstual bagi pendidikan vokasi. Berdasarkan pra-penelitian di SMK Negeri 9 Bandar Lampung, analisis kurikulum dieksekusi spesifik pada mata pelajaran Dasar-dasar Pengembangan Perangkat Lunak dan Gim (PPLG). Analisis ini bertujuan untuk membuktikan kesesuaian antara tahapan rekayasa perangkat lunak sistem nyata (Website ULT) dengan Capaian Pembelajaran (CP) dan Tujuan Pembelajaran (TP) yang dituntut dalam pedoman kejuruan tingkat SMK. Rincian penjabaran relevansi pengembangan produk tata kelola administrasi institusi ini terhadap materi kejuruan diuraikan pada tabel berikut."
    
    # Update Table 1
    row_cells = table_1.rows[1].cells
    row_cells[0].text = "Pada akhir pembelajaran PPLG, peserta didik mampu memahami proses bisnis industri perangkat lunak, menerapkan metodologi pengembangan sistem informasi, mengimplementasikan kerangka kerja pemrograman web modern, serta menyusun dokumentasi proyek rekayasa perangkat lunak secara terstruktur dan komprehensif."
    row_cells[1].text = "Peserta didik mampu:\n1. Mengidentifikasi kebutuhan klien dan merumuskan alur proses bisnis antarmuka web (C4).\n2. Menerapkan arsitektur kerangka kerja pemrograman spesifik (Laravel dan Tailwind CSS) (C3).\n3. Merancang integrasi tata kelola basis data relasional ke dalam sistem operasional (C5).\n4. Menunjukkan etika profesional dan tanggung jawab kerja sama tim manajemen proyek (A4).\n5. Memproduksi perangkat lunak fungsional teruji beserta dokumentasi algoritma teknis (C6)."
    
    for p in row_cells[0].paragraphs: p.alignment = 3 # JUSTIFY
    for p in row_cells[1].paragraphs: p.alignment = 3 # JUSTIFY

    # Update Table 2 (Rows 1-4)
    data = [
        ("Arsitektur Antarmuka (UI/UX) dan Front-end", "1. Desain tata letak antarmuka web responsif menggunakan Tailwind CSS.\n2. Pemenuhan standar aksesibilitas interaksi pengguna via Progressive Web App (PWA).\n3. Pemisahan fungsionalitas dasbor spesifik untuk ragam entitas pengguna."),
        ("Manajemen Basis Data & Keamanan Sistem", "1. Implementasi sistem pangkalan data relasional terstruktur (MySQL).\n2. Proteksi tautan dan modul enkripsi penyimpanan repositori sistem fail lokal.\n3. Penerapan penyaringan serangan siber injeksi skrip berbahaya (Cross-Site Scripting)."),
        ("Logika Bisnis (Back-end) dan Algoritma", "1. Integrasi pustaka konversi dokumen terpusat (HTML-to-OpenXML).\n2. Ekstraksi dan penggabungan variabel pendaftar ke dalam templat resmi.\n3. Pembangkitan kode pelacakan dan nomor identifikasi elektronik terotomatisasi."),
        ("Manajemen Proyek dan Dokumentasi (Audit)", "1. Pencatatan log aktivitas tanggal waktu pemrosesan transaksi (Audit Trail).\n2. Pemodelan tahapan alur persetujuan terstruktur pada sebuah sistem operasional.\n3. Pembuatan rekapitulasi data pangkalan evaluasi durasi penanganan layanan.")
    ]
    
    for i, (kedalaman, keluasan) in enumerate(data):
        cells = table_2.rows[i+1].cells
        cells[0].text = kedalaman
        cells[1].text = keluasan
        for p in cells[0].paragraphs: p.alignment = 3
        for p in cells[1].paragraphs: p.alignment = 3

    # Update Para 2
    if para_2:
        para_2.text = "Pemetaan struktur analisis kurikulum mata pelajaran PPLG tersebut secara faktual menjustifikasi bahwa produk Website ULT tidak sekadar berfungsi sebagai solusi administratif internal kampus semata. Penyelarasan antara kewajiban pemenuhan pencapaian tujuan kejuruan vokasi dengan arsitektur teknologi terbukti menjadikan proyek pengembangan sistem ini sebagai media pembelajaran kontekstual yang ideal. Kerangka pemetaan ini secara akademis memastikan bahwa proses rancang bangun teknologi tata kelola administrasi dapat langsung diadaptasi oleh peserta didik SMK sebagai rujukan proyek riil yang berstandar dunia industri."
        para_2.alignment = 3 # JUSTIFY

    doc.save(output_filename)
    print(f"Successfully updated {output_filename}")

update_pplg_curriculum("001_Skripsi_Andricha Dea Mitra_Clean.docx", "001_Skripsi_Andricha Dea Mitra_Clean.docx")
update_pplg_curriculum("001_Skripsi_Andricha Dea Mitra_Highlighted.docx", "001_Skripsi_Andricha Dea Mitra_Highlighted.docx")
