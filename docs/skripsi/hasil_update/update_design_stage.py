import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
import os

files_to_process = [
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx',
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
]

def add_paragraph_after(paragraph, text, is_highlighted, is_bold=False):
    p = paragraph.insert_paragraph_before() # We will insert before the current one, wait, it's easier to insert after by using p.insert_paragraph_before on the NEXT paragraph.
    # Actually, we can use insert_paragraph_before on the target paragraph and then we advance the target.
    pass

# Better approach: 
# iterate through doc.paragraphs, find the target, and we can just modify or insert.
def insert_content(file_path, is_highlighted_doc):
    doc = docx.Document(file_path)
    
    # Text content to insert
    intro_text = "Pada tahap perancangan (Design) dalam model pengembangan ADDIE, fokus utama peneliti dialihkan dari perumusan masalah konseptual menuju perancangan solusi teknis. Berdasarkan model Four-D yang diadaptasi dari Thiagarajan (1974), tahap perancangan ini meliputi empat langkah sistematis utama, yaitu:"
    
    a_title = "a. Constructing Criterion-Referenced Tests"
    a_text = "Penyusunan tes acuan patokan (Criterion-Referenced Tests) merupakan langkah awal yang menjadi jembatan operasional antara tahapan analisis dan tahapan desain. Pada langkah ini, disusun instrumen evaluasi produk yang spesifik dan terukur untuk menilai tingkat kelayakan serta kepraktisan Website ULT FKIP Unila yang akan dikembangkan. Instrumen yang disusun meliputi kisi-kisi uji validasi ahli yang terbagi atas tiga ranah kompetensi pokok: ahli materi untuk memvalidasi substansi alur birokrasi persuratan, ahli media untuk menilai struktur dan hierarki visual antarmuka (UI/UX), serta ahli sistem informasi untuk mengevaluasi fungsionalitas teknis, integrasi database, dan keamanan siber aplikasi. Selain itu, disusun pula instrumen kuesioner uji kepraktisan skala terbatas (limited trial) bagi 18 orang responden akhir (mahasiswa, admin program studi, dan staf ULT) sebagai tolak ukur ketercapaian fungsionalitas produk dalam menyelesaikan kendala administrasi dokumen yang ada. Penggunaan tes acuan patokan ini krusial untuk memastikan bahwa rancangan sistem nantinya senantiasa berpedoman kuat pada solusi yang telah dirumuskan di awal."
    
    b_title = "b. Media Selection"
    b_text = "Langkah pemilihan media difokuskan pada penentuan teknologi, perangkat keras, dan perangkat lunak yang paling sesuai untuk menerjemahkan model konsepsional ke dalam bentuk fisik sistem. Berdasarkan hasil analisis kebutuhan (requirement analysis), sistem administrasi terpadu membutuhkan media yang fleksibel, dapat diakses kapan saja dan dari mana saja (ubiquitous), serta tidak memerlukan proses instalasi rumit bagi pengguna akhir. Oleh sebab itu, media utama yang dipilih adalah aplikasi berbasis peramban (web-based application) berformat monolitik yang dibangun menggunakan framework Laravel versi 12 dengan bahasa pemrograman PHP 8.4+.\n\nUntuk lapisan antarmuka pengguna (front-end), dipilih kerangka kerja TailwindCSS dipadukan dengan Alpine.js yang dikompilasi menggunakan Vite, sehingga menghasilkan responsivitas dan kecepatan akses yang optimal layaknya Progressive Web App (PWA). Pilihan pangkalan data jatuh pada MySQL untuk menjamin integritas data relasional bervolume tinggi, didukung oleh private local storage disk pada Laravel guna memastikan berkas-berkas akademik mahasiswa yang sensitif tetap terlindungi dan terisolasi dari akses publik. Lebih lanjut, teknologi pemroses dokumen secara khusus melibatkan pustaka phpoffice/phpword dan dompdf/dompdf yang digabungkan bersama Tiptap Editor sebagai editor teks kaya (WYSIWYG), memungkinkan ekstraksi serta perakitan dokumen Word dan PDF secara dinamis tanpa merusak susunan visual templat orisinalnya."
    
    c_title = "c. Format Selection"
    c_text = "Pemilihan format (Format Selection) sangat erat kaitannya dengan pendefinisian antarmuka sistem (user interface format) dan bagaimana informasi disajikan serta diatur pembagiannya agar intuitif bagi pengguna. Untuk mengakomodasi arsitektur role-based access control (RBAC), format antarmuka dipecah menjadi lima portal akses utama dengan hak otonomi spesifik, yakni:\n1. Public Portal (Tampilan Publik): Format antarmuka halaman depan (landing page) yang didesain secara informatif, menayangkan profil singkat ULT, katalog prosedur layanan persuratan, portal berita atau pengumuman fakultas, serta sistem Content Management System (CMS) publik untuk mempermudah tamu umum.\n2. Authentication System (Sistem Autentikasi): Tampilan layar terpusat untuk proses pendaftaran (register) akun baru, validasi masuk (login), dan pemulihan kata sandi (reset password) yang dibangun di atas pondasi Laravel Breeze (Blade).\n3. Student Portal (Dasbor Mahasiswa): Format antarmuka privat yang ringkas dan bebas distraksi, dirancang khusus untuk mahasiswa agar dapat mengajukan permohonan berbagai jenis layanan surat, mengisi formulir isian data dinamis (dynamic input fields), mengunggah berkas persyaratan terisolasi, serta melacak tahapan proses layanan (auditable timeline).\n4. Admin/Staff Portal (Dasbor Operasional): Antarmuka manajemen inti bagi petugas pintu gerbang (Gatekeeper) ULT. Format halamannya diatur dengan paradigma data grid untuk memfasilitasi tugas-tugas administratif berat seperti pengecekan permohonan, penomoran surat otomatis, pengelolaan templat layanan, dan perakitan (assembly) keluaran akhir dokumen fisik.\n5. Signer Portal (Portal Tanda Tangan): Format halaman minimalis khusus bagi para pejabat struktural fakultas maupun jurusan untuk melakukan telaah substantif surat, verifikasi keabsahan data, serta melampirkan tanda tangan persetujuan elektronik dengan aman."
    
    d_title = "d. Initial Design"
    d_text = "Langkah perancangan awal (Initial Design) adalah tahap finalisasi purwarupa dari segala bentuk konseptual arsitektur sistem. Rancangan ini menggabungkan seluruh komponen pemodelan struktural dan perilaku (behavioral) untuk memastikan aplikasi dibangun dengan presisi teknis. Desain awal ini diwujudkan melalui serangkaian cetak biru visual berupa pemodelan fungsional, arsitektural, alur kerja, relasi data, diagram aktivitas, wireframe kerangka antarmuka dasar, hingga antarmuka pengguna final (UI/UX). Penjabaran masing-masing elemen perancangan tersebut diuraikan sebagai berikut:"

    act_title = "Perancangan Aktivitas Pengguna Menggunakan Diagram Activity"
    act_text = "Diagram aktivitas (Activity Diagram) merinci lebih lanjut alur kerja (workflow) transaksional dan operasional yang telah divisualisasikan dalam diagram Use Case. Pemodelan ini secara spesifik memetakan runtutan aksi dan keputusan dinamis antara aktor pengguna (mahasiswa, staf ULT, dan pejabat) saat berinteraksi secara sekuensial dengan sistem ULT FKIP Unila."
    act_img_cap = "Gambar [NOMOR_GAMBAR]. Diagram Activity Pengajuan Layanan Dokumen ULT FKIP Unila."
    act_text2 = "Secara konseptual, aktivitas berawal dari mahasiswa yang mengautentikasi diri ke dalam Student Portal untuk memilih jenis layanan surat akademik yang relevan dengan kebutuhannya. Sistem kemudian akan memuat formulir dinamis sesuai dengan basis data layanan, mengharuskan mahasiswa mengisi detail data teks maupun mengunggah berkas-berkas persyaratan yang diamankan secara privat. Setelah formulir dikirim, transisi kendali berpindah ke dasbor Staf ULT yang akan melakukan verifikasi berkas awal. Bila terdapat kekeliruan dokumen, staf dapat memberikan umpan balik penolakan sehingga mahasiswa diwajibkan untuk merevisi permohonan. Sebaliknya, bila berkas tervalidasi lengkap, sistem meneruskan status layanan kepada Pejabat Fakultas atau Jurusan di portal terpisah (Signer Portal) untuk menelaah kesesuaian akademis. Pada tahapan kritis ini, pejabat berhak membubuhkan tanda tangan elektronik yang secara langsung menstimulasi sistem (melalui DocumentAssemblerService) untuk merakit dan menggabungkan data pemohon ke dalam berkas .docx siap unduh. Proses aktivitas ini ditutup ketika mahasiswa menerima notifikasi digital bahwa dokumen fisik maupun elektronik mereka telah selesai dan siap diunduh lewat akun masing-masing."

    wf_title = "Perancangan Kerangka Tampilan Utama (Wireframe)"
    wf_text = "Rancangan kerangka tampilan (wireframe) dibuat sebagai kerangka tulang punggung awal antarmuka (low-fidelity design). Wireframe berfungsi untuk merancang dan memastikan peletakan struktur elemen navigasi, hirarki penataan konten, dan tata letak hierarki informasi tanpa melibatkan distraksi elemen visual seperti pewarnaan maupun tipografi."
    wf_img_cap = "Gambar [NOMOR_GAMBAR]. Wireframe Beranda Utama (Main View) Sistem ULT FKIP Unila."
    wf_text2 = "Pada tampilan beranda utama sistem (Public Portal), wireframe disusun menggunakan tata letak berbasis kolom tunggal terpusat (centered layout) guna memusatkan perhatian pengunjung pada elemen inti. Area header (navigasi atas) memuat logo institusi Fakultas Keguruan dan Ilmu Pendidikan bersandingan dengan tautan menu navigasi utama seperti 'Beranda', 'Katalog Layanan', 'Berita', dan tombol 'Masuk / Daftar'. Tepat di bawahnya, terdapat pahlawan visual (hero section) besar yang menampung slogan layanan terpadu serta tombol Call-to-Action (CTA) berukuran dominan untuk memandu pengguna baru agar segera masuk ke portal mahasiswa. Bergerak ke bawah, tata letak dipecah menjadi pola grid asimetris untuk menyajikan kartu-kartu panduan alur permohonan secara berurutan, memberikan kejelasan visual tahap demi tahap (step-by-step guide). Bagian kaki (footer) dialokasikan untuk informasi kontak institusi, tautan regulasi, serta informasi hak cipta. Secara keseluruhan, pemetaan visual beresolusi rendah ini sukses memenuhi standar kesederhanaan akses tanpa menghilangkan kedalaman informasi esensial."

    ui_title = "Perancangan Antarmuka Pengguna Akhir (Final UI/UX Design)"
    ui_text = "Tahap perancangan ini bermuara pada hasil desain interaktif berkualitas tinggi (high-fidelity design) dari sisi User Interface (UI) dan User Experience (UX). Desain antarmuka akhir ini merepresentasikan wujud nyata aplikasi yang sudah mengimplementasikan kerangka kerja TailwindCSS dan panduan desain standar modern dengan proporsi estetika yang ramah pengguna."
    ui_img_cap = "Gambar [NOMOR_GAMBAR]. Desain Antarmuka Akhir (UI/UX) Sistem ULT FKIP Unila."
    ui_text2 = "Implementasi UI/UX menerapkan skema warna dominan biru korporat Universitas Lampung yang berpadu dengan aksen warna netral (abu-abu dan putih) guna menanamkan kesan profesional, kredibel, dan institusional yang bersih. Elemen interaktif seperti tombol pengiriman data, panel navigasi sisi (sidebar), maupun indikator status permohonan dokumen (badges) diberikan peningkat interaktivitas visual layaknya hover-effect dan transisi halus dari Alpine.js agar pengguna merasakan pengalaman berselancar yang dinamis. Pendekatan perancangan mobile-first juga diprioritaskan sehingga tampilan UI dapat beradaptasi dan menyesuaikan ruang kosong (whitespace) secara proporsional, baik ketika diakses melalui resolusi monitor layar lebar maupun gawai pintar berlayar kecil. Rancangan antarmuka yang presisi ini memastikan agar mahasiswa hingga pejabat tinggi fakultas tidak mengalami ambiguitas kognitif saat memproses dokumen mereka secara digital di dalam Web ULT FKIP Unila."

    def format_para(p, is_high, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(12), space_before=Pt(0), line_spacing=1.5):
        p.alignment = align
        p.paragraph_format.space_after = space_after
        p.paragraph_format.space_before = space_before
        p.paragraph_format.line_spacing = line_spacing
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            if is_high:
                r.font.highlight_color = WD_COLOR_INDEX.YELLOW

    def insert_p(ref_p, text, is_high=is_highlighted_doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        if text.startswith('[GAMBAR'):
            # It's an image placeholder
            new_p = ref_p.insert_paragraph_before()
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            new_p.paragraph_format.space_after = Pt(0)
            new_p.paragraph_format.space_before = Pt(12)
            run = new_p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = docx.shared.RGBColor(255, 0, 0) # Red to indicate missing
            return new_p

        new_p = ref_p.insert_paragraph_before()
        new_p.alignment = align
        new_p.paragraph_format.space_after = Pt(0)
        new_p.paragraph_format.line_spacing = 1.5
        
        # Determine if there are italics needed (e.g. English words inside parentheses)
        # We will just do a simple run addition. 
        # But for titles like "a. Constructing Criterion-Referenced Tests", the english words should be italicized.
        
        if text in [a_title, b_title, c_title, d_title]:
            run = new_p.add_run(text.split(' ', 1)[0] + ' ') # "a. "
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if is_high: run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
            run2 = new_p.add_run(text.split(' ', 1)[1]) # "Constructing..."
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(12)
            run2.font.italic = True
            if is_high: run2.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif text.startswith('Gambar ['):
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            new_p.paragraph_format.space_after = Pt(12) # Jarak 3 spasi (1.5 * 2) after image caption? we'll use 12pt
            new_p.paragraph_format.line_spacing = 1.0
            run = new_p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10) # Caption size
            if is_high: run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            # Handle newlines inside text
            for part in text.split('\n'):
                if part.strip() == '': continue
                run = new_p.add_run(part + '\n')
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                if is_high: run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
            # Remove trailing newline from last run
            if new_p.runs:
                new_p.runs[-1].text = new_p.runs[-1].text.rstrip('\n')
        
        return new_p

    # Find Tahap Desain intro
    intro_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith('Pada tahap perancangan (Design) dalam model pengembangan ADDIE, fokus utama peneliti dialihkan'):
            intro_idx = i
            break
            
    if intro_idx != -1:
        target_p = doc.paragraphs[intro_idx]
        target_p.text = "" # Clear the old text
        
        # We will insert everything BEFORE target_p, and then delete target_p
        insert_p(target_p, intro_text)
        insert_p(target_p, a_title)
        insert_p(target_p, a_text)
        insert_p(target_p, b_title)
        insert_p(target_p, b_text)
        insert_p(target_p, c_title)
        insert_p(target_p, c_text)
        insert_p(target_p, d_title)
        insert_p(target_p, d_text)
        
        p = target_p._element
        p.getparent().remove(p)

    # Find the end of Sequence Diagram description to insert Activity, Wireframe, UIUX
    # The last paragraph before "Guna menjembatani rancangan konseptual" is the end of Sequence diagram
    insert_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith('Guna menjembatani rancangan konseptual ini ke bentuk fisik'):
            insert_idx = i
            break

    if insert_idx != -1:
        target_p2 = doc.paragraphs[insert_idx]
        
        insert_p(target_p2, act_title)
        insert_p(target_p2, act_text)
        insert_p(target_p2, "[GAMBAR_ACTIVITY_DIAGRAM_BELUM_ADA]")
        insert_p(target_p2, act_img_cap)
        insert_p(target_p2, act_text2)

        insert_p(target_p2, wf_title)
        insert_p(target_p2, wf_text)
        insert_p(target_p2, "[GAMBAR_WIREFRAME_UTAMA_BELUM_ADA]")
        insert_p(target_p2, wf_img_cap)
        insert_p(target_p2, wf_text2)

        insert_p(target_p2, ui_title)
        insert_p(target_p2, ui_text)
        insert_p(target_p2, "[GAMBAR_UIUX_FINAL_BELUM_ADA]")
        insert_p(target_p2, ui_img_cap)
        insert_p(target_p2, ui_text2)

    doc.save(file_path)
    print(f"Processed: {file_path}")

for f in files_to_process:
    insert_content(f, "Highlighted" in f)

