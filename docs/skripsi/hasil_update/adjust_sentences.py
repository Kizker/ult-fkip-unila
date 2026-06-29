import docx
import sys
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding='utf-8')

def rewrite_doc(filename, is_highlighted=False):
    print(f"Processing {filename}...")
    doc = docx.Document(filename)
    
    start_idx = -1
    end_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == '1. Aspek Komposisi Warna (Validasi Media)':
            start_idx = i
        if start_idx != -1 and p.text.strip().startswith('Penambahan ornamen interaksi ini pada hakikatnya'):
            end_idx = i
            break
            
    if start_idx == -1 or end_idx == -1:
        print(f"Could not find target paragraphs in {filename}. start={start_idx}, end={end_idx}")
        return
        
    insert_after_p = doc.paragraphs[start_idx - 1]
    
    to_delete = []
    idx = start_idx
    while idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        if idx <= end_idx:
            to_delete.append(p)
            idx += 1
        elif p.text.strip() == '':
            to_delete.append(p)
            idx += 1
            if idx > end_idx + 3:
                break
        else:
            break
            
    for p in to_delete:
        p_elem = p._element
        if p_elem.getparent() is not None:
            p_elem.getparent().remove(p_elem)
            
    data = [
        {
            "title": "1. Aspek Komposisi Warna (Validasi Media)",
            "intro": "Evaluasi pertama dari validator ahli media berfokus pada aspek komposisi warna yang digunakan pada antarmuka utama sistem informasi. Kesan pertama pengguna saat mengakses portal sangat dipengaruhi oleh kombinasi warna yang disajikan di layar. Oleh karena itu, pemilihan skema warna harus mempertimbangkan faktor kenyamanan visual sekaligus merepresentasikan identitas institusi secara profesional.",
            "awal_ss": "[MASUKKAN SCREENSHOT: UI Beranda dengan warna gradasi mencolok]",
            "awal_cap": "Gambar 9. Desain Awal (Pra-Validasi) Antarmuka Utama",
            "awal_expl": "Rancangan antarmuka awal yang diperlihatkan pada Gambar 9 menggunakan skema warna gradasi mencolok pada bagian latar belakang. Pakar media mengkritisi pilihan desain ini karena intensitas warna yang terlalu tinggi berpotensi mendistraksi fokus visual pengguna dari informasi utama. Desain tersebut dinilai dapat menyebabkan kelelahan mata jika pengguna mengakses halaman dalam durasi yang cukup lama tanpa jeda istirahat visual.",
            "akhir_ss": "[MASUKKAN SCREENSHOT: UI Beranda flat design biru yang elegan]",
            "akhir_cap": "Gambar 10. Desain Akhir (Pasca-Validasi) Antarmuka Utama",
            "akhir_expl": "Sebagai tindak lanjut korektif, desain antarmuka yang tersaji pada Gambar 10 direstrukturisasi secara menyeluruh menggunakan pendekatan flat design dengan palet warna solid biru. Pemilihan warna ini jauh lebih elegan dan secara langsung selaras dengan identitas visual (corporate identity) Fakultas Keguruan dan Ilmu Pendidikan. Selain itu, elemen antarmuka yang sebelumnya mencolok kini telah dilembutkan untuk menciptakan keseimbangan kontras antara teks dan latar belakang.",
            "summary": "Perubahan arsitektur warna ini terbukti mampu merevolusi atmosfer visual saat sistem pertama kali diakses oleh entitas pengguna. Keputusan untuk beralih pada palet warna yang lebih redup tidak hanya menyelesaikan masalah kelelahan mata, tetapi juga memperkuat penekanan pada fitur-fitur esensial layanan. Pada akhirnya, tata warna yang baru ini berhasil membangun impresi platform layanan akademik yang modern, kredibel, dan berfokus penuh pada kenyamanan operasional."
        },
        {
            "title": "2. Aspek Tata Letak/Spasi Tabel (Validasi Media/Sistem)",
            "intro": "Penilaian selanjutnya menyoroti tata letak dan pengaturan spasi pada elemen tabel yang menampung data operasional sistem. Tabel merupakan komponen krusial bagi staf ULT untuk memantau, memverifikasi, dan mengelola ratusan data permohonan mahasiswa setiap harinya. Struktur tabel yang optimal sangat dibutuhkan agar proses pemindaian informasi dapat berlangsung secara cepat tanpa membebani daya kognitif staf pelaksana.",
            "awal_ss": "[MASUKKAN SCREENSHOT: Tabel Dasbor Staf ULT yang padat/sempit]",
            "awal_cap": "Gambar 11. Desain Awal (Pra-Validasi) Tata Letak Tabel",
            "awal_expl": "Matriks tabel operasional pada desain pra-validasi sebagaimana yang disajikan pada Gambar 11 memiliki pengaturan baris yang sangat padat dan sempit. Validator media dan sistem secara kompak menyoroti kepadatan baris ini karena berpotensi memicu keletihan visual yang serius akibat penumpukan data. Kondisi jarak antarbaris yang minim ini sangat menyulitkan proses pemindaian data (eye-scanning) secara berulang saat staf mencari spesifikasi dokumen tertentu dari tumpukan antrean.",
            "akhir_ss": "[MASUKKAN SCREENSHOT: Tabel Dasbor Staf ULT yang renggang dan rapi]",
            "akhir_cap": "Gambar 12. Desain Akhir (Pasca-Validasi) Tata Letak Tabel",
            "akhir_expl": "Mengakomodasi kritik konstruktif tersebut, desain perbaikan pada Gambar 12 menunjukkan adanya ekspansi white space dengan memperbesar padding sel dan mengatur spasi antarbaris menjadi 1,5. Modifikasi tata letak ini memberikan ruang bernapas pada teks sehingga setiap sel data dapat terpisahkan secara visual dengan sangat jelas. Melalui penataan ulang proporsi spasi ini, tingkat keterbacaan data tabular meningkat secara drastis tanpa harus mengurangi kapasitas informasi yang disajikan pada satu layar.",
            "summary": "Restrukturisasi elemen visual pada tabel operasional ini membawa implikasi besar terhadap alur kerja administrasi harian. Peningkatan kejelasan struktur teks secara langsung berkontribusi pada penurunan beban kognitif staf ULT saat berhadapan dengan volume data yang masif. Hasil penyempurnaan ini pada gilirannya meminimalisasi potensi human error, memastikan setiap berkas permohonan mahasiswa dapat divalidasi dengan tingkat akurasi dan kecepatan yang jauh lebih optimal."
        },
        {
            "title": "3. Aspek Fitur Ikon Navigasi (Validasi Media)",
            "intro": "Aspek ketiga yang dievaluasi oleh pakar berpusat pada ketersediaan fitur penanda visual atau ikon navigasi pada halaman direktori layanan. Direktori layanan merupakan gerbang utama bagi mahasiswa untuk mengenali rincian administrasi yang disediakan oleh pihak fakultas. Ketersediaan elemen semiotik yang representatif amat vital untuk membimbing pengguna baru agar tidak merasa kebingungan saat pertama kali mengeksplorasi modul layanan.",
            "awal_ss": "[MASUKKAN SCREENSHOT: Katalog layanan yang hanya berisi teks]",
            "awal_cap": "Gambar 13. Desain Awal (Pra-Validasi) Katalog Layanan",
            "awal_expl": "Desain awal dari katalog layanan akademik yang dapat diamati pada Gambar 13 hanya mengandalkan susunan teks murni tanpa adanya elemen grafis pendamping. Ketiadaan penanda visual ini dinilai oleh validator sangat menghambat proses identifikasi fungsi sistem oleh pengguna akhir secara cepat. Pengguna secara terpaksa harus membaca setiap judul layanan satu per satu, yang mana pola interaksi ini sangat tidak efisien dan melanggar prinsip navigasi intuitif.",
            "akhir_ss": "[MASUKKAN SCREENSHOT: Katalog layanan dilengkapi ikon]",
            "akhir_cap": "Gambar 14. Desain Akhir (Pasca-Validasi) Katalog Layanan",
            "akhir_expl": "Solusi teknis untuk mengatasi kelemahan tersebut ditunjukkan pada Gambar 14 yang diwujudkan melalui pengintegrasian ikon vektor spesifik pada setiap kartu layanan (service tile). Kehadiran elemen grafis ini langsung berfungsi sebagai jangkar kognitif yang memecah dominasi teks pada antarmuka pengguna. Dengan pendekatan ikonografi yang tepat, setiap jenis layanan kini memiliki representasi visual yang unik sehingga dapat dikenali hanya melalui pemindaian sekilas.",
            "summary": "Integrasi penanda visual ini terbukti mampu menjembatani kesenjangan komunikasi antara kompleksitas sistem dan pemahaman awal pengguna. Pendekatan ini memangkas waktu orientasi mahasiswa secara signifikan karena mereka tidak lagi bergantung sepenuhnya pada kemampuan literasi membaca untuk mencari menu yang tepat. Inovasi kecil namun strategis ini mengukuhkan kapabilitas sistem dalam menyediakan pengalaman pengguna yang responsif, terarah, dan ramah bagi berbagai lapisan mahasiswa."
        },
        {
            "title": "4. Aspek Fitur Pendukung Interaksi (Validasi Sistem)",
            "intro": "Evaluasi terakhir dari pakar sistem menyentuh aspek ketersediaan fitur pendukung interaksi pada halaman yang memiliki struktur konten bergulir panjang. Pengalaman menggulir layar (scrolling) merupakan aktivitas fundamental yang tak terhindarkan saat pengguna membaca informasi atau prosedur yang mendetail. Ketersediaan fasilitas navigasi balik yang taktis sangat krusial untuk mencegah rasa frustrasi ketika pengguna berniat kembali ke titik awal halaman.",
            "awal_ss": "[MASUKKAN SCREENSHOT: Halaman panjang tanpa tombol navigasi balik]",
            "awal_cap": "Gambar 15. Desain Awal (Pra-Validasi) Navigasi Halaman",
            "awal_expl": "Kondisi halaman informasi dengan gulir panjang (long-scroll pages) yang diilustrasikan pada Gambar 15 belum terfasilitasi oleh pintasan navigasi terintegrasi. Pakar sistem mengkritisi inefisiensi mobilitas pada kerangka ini, mengingat absennya tombol pengarah cepat menuju area navigasi atas layar. Tanpa keberadaan fitur akselerasi ini, pengguna terbebani keharusan melakukan gerakan gulir manual yang berulang-ulang, sebuah proses yang memboroskan waktu dan energi.",
            "akhir_ss": "[MASUKKAN SCREENSHOT: Halaman dengan tombol floating Back to Top]",
            "akhir_cap": "Gambar 16. Desain Akhir (Pasca-Validasi) Navigasi Halaman",
            "akhir_expl": "Sebagai bentuk penyelesaian terhadap temuan tersebut, eksekusi perbaikan pada Gambar 16 diwujudkan melalui penambahan tombol melayang (floating button) bertuliskan 'Back to Top'. Tombol interaktif ini ditempatkan secara strategis pada sudut layar sehingga senantiasa mudah diakses kapan pun pengguna membutuhkannya. Implementasi mekanik gulir otomatis ini secara langsung mengoptimalkan kelancaran mobilitas vertikal di seluruh halaman sistem tanpa terkecuali.",
            "summary": "Penambahan ornamen interaksi ini pada hakikatnya merepresentasikan kepedulian pengembang terhadap aspek mikromobilitas di dalam ekosistem aplikasi. Penghapusan rutinitas menggulir yang menjemukan tersebut tidak sekadar mempercepat transisi halaman, tetapi secara substansial melipatgandakan keluwesan sistem secara keseluruhan. Langkah penyempurnaan ini menjamin agar seluruh interaksi pengguna senantiasa terjaga dalam koridor yang praktis, terkendali, dan terbebas dari hambatan navigasi struktural."
        }
    ]
    
    current_element = insert_after_p._element
    
    def add_p(text, bold=False, highlight=False, center=False):
        nonlocal current_element
        new_p = doc.add_paragraph()
        if center:
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text:
            run = new_p.add_run(text)
            if bold:
                run.bold = True
            if highlight:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        current_element.addnext(new_p._element)
        current_element = new_p._element
        return new_p
        
    for item in data:
        add_p(item['title'], bold=True, highlight=is_highlighted)
        add_p(item['intro'], highlight=is_highlighted)
        add_p("")
        add_p(item['awal_ss'], highlight=is_highlighted, center=True)
        add_p(item['awal_cap'], highlight=is_highlighted, center=True)
        add_p(item['awal_expl'], highlight=is_highlighted)
        add_p("")
        add_p(item['akhir_ss'], highlight=is_highlighted, center=True)
        add_p(item['akhir_cap'], highlight=is_highlighted, center=True)
        add_p(item['akhir_expl'], highlight=is_highlighted)
        add_p(item['summary'], highlight=is_highlighted)
        add_p("")
        
    doc.save(filename)
    print(f"Successfully processed {filename}")

rewrite_doc(r'C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx', False)
rewrite_doc(r'C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx', True)
