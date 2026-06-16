import docx
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
import os

diagrams = [
    {
        "file": r"c:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram\06_activity_diagram_autentikasi.png",
        "caption": "Gambar 4.1. Diagram Activity Autentikasi Pengguna",
        "desc": "Gambar 4.1 mengilustrasikan alur aktivitas autentikasi pengguna pada Web ULT FKIP Unila. Mahasiswa baru diwajibkan mengisi formulir registrasi yang akan divalidasi dan dikirimkan tautan verifikasi melalui email. Pengguna yang telah memiliki akun terverifikasi dapat langsung memasukkan email dan kata sandi pada halaman login. Sistem akan memvalidasi kredensial tersebut terhadap basis data; jika valid, sistem menciptakan sesi login dan mengarahkan pengguna ke dasbor masing-masing sesuai dengan perannya (role), sedangkan jika gagal, sistem akan menampilkan pesan peringatan."
    },
    {
        "file": r"c:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram\07_activity_diagram_pengajuan.png",
        "caption": "Gambar 4.2. Diagram Activity Pengajuan Layanan Akademik",
        "desc": "Gambar 4.2 memetakan proses utama pengajuan layanan persuratan oleh mahasiswa. Dimulai dari Student Portal, mahasiswa memilih jenis layanan spesifik, mengisi formulir berbasis variabel dinamis yang disediakan oleh sistem, serta mengunggah berkas-berkas prasyarat. Sistem secara otomatis melakukan validasi input. Permohonan yang tervalidasi kemudian disimpan ke dalam basis data dengan status awal 'Menunggu Verifikasi', dan sebuah notifikasi proaktif diteruskan kepada staf ULT untuk ditindaklanjuti."
    },
    {
        "file": r"c:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram\08_activity_diagram_verifikasi.png",
        "caption": "Gambar 4.3. Diagram Activity Verifikasi dan Pemrosesan Dokumen",
        "desc": "Gambar 4.3 menjabarkan alur kerja operasional oleh Staf ULT dan Pejabat struktural. Staf ULT berperan sebagai verifikator pertama yang mengecek kesesuaian berkas. Jika terdapat ketidaksesuaian, permohonan akan ditolak dengan catatan khusus untuk diperbaiki mahasiswa. Jika valid, sistem mendeteksi apakah dokumen memerlukan tanda tangan elektronik tingkat pimpinan. Jika ya, draf elektronik diteruskan ke Signer Portal agar pejabat dapat mereview dan membubuhkan tanda tangan. Pada tahap final, mesin assembly merakit keseluruhan data dan tanda tangan ke dalam format dokumen legal (.docx), lalu mengubah status permohonan menjadi 'Selesai'."
    },
    {
        "file": r"c:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram\09_activity_diagram_manajemen.png",
        "caption": "Gambar 4.4. Diagram Activity Manajemen Layanan dan Templat",
        "desc": "Gambar 4.4 menunjukkan fungsionalitas manajemen layanan yang dioperasikan oleh Admin ULT. Admin memiliki wewenang untuk meracik layanan persuratan baru melalui antarmuka dasbor tanpa intervensi langsung pada kode sumber aplikasi (source code). Proses ini mencakup pengaturan penamaan layanan, pendefinisian bidang isian dinamis (dynamic inputs) yang wajib diisi mahasiswa, serta pengunggahan templat master dokumen berformat Word (.docx). Sistem lalu akan memvalidasi format dokumen tersebut dan segera memperbarui katalog layanan secara seketika (real-time) di halaman publik."
    }
]

def process_doc(filepath, is_highlighted):
    doc = docx.Document(filepath)
    
    start_idx = -1
    end_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        if 'Diagram aktivitas (Activity Diagram) merinci lebih lanjut alur kerja' in p.text:
            start_idx = i
        if 'Perancangan Kerangka Tampilan Utama (Wireframe)' in p.text:
            end_idx = i
            break
            
    if start_idx == -1 or end_idx == -1:
        print(f"Bounds not found in {filepath}. Start: {start_idx}, End: {end_idx}")
        return
        
    for i in range(start_idx + 1, end_idx):
        doc.paragraphs[i].clear()
        
    target_p = doc.paragraphs[end_idx]
    
    for item in diagrams:
        if not os.path.exists(item["file"]):
            print("Image not found:", item["file"])
            continue
            
        img_p = target_p.insert_paragraph_before()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.paragraph_format.space_before = Pt(12)
        img_p.paragraph_format.space_after = Pt(6)
        img_p.paragraph_format.line_spacing = 1.0
        run_img = img_p.add_run()
        run_img.add_picture(item["file"], width=Inches(4.8)) # Adjusted slightly for the plantuml
        
        cap_p = target_p.insert_paragraph_before()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_before = Pt(0)
        cap_p.paragraph_format.space_after = Pt(12)
        cap_p.paragraph_format.line_spacing = 1.0
        run_cap = cap_p.add_run(item["caption"])
        run_cap.font.name = 'Times New Roman'
        run_cap.font.size = Pt(11)
        run_cap.font.bold = False 
        if is_highlighted:
            run_cap.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
        desc_p = target_p.insert_paragraph_before()
        desc_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        desc_p.paragraph_format.space_before = Pt(0)
        desc_p.paragraph_format.space_after = Pt(12)
        desc_p.paragraph_format.line_spacing = 1.5
        desc_p.paragraph_format.first_line_indent = Cm(1.27)
        run_desc = desc_p.add_run(item["desc"])
        run_desc.font.name = 'Times New Roman'
        run_desc.font.size = Pt(12)
        if is_highlighted:
            run_desc.font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.save(filepath)
    print("Cleaned and updated", filepath)

files = [
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx',
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
]

for f in files:
    process_doc(f, "Highlighted" in f)
