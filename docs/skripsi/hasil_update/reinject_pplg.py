import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

def reinject_pplg(filename, output_filename):
    doc = docx.Document(filename)
    
    # Find start and end paragraphs
    start_idx = -1
    end_idx = -1
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if p.style.name.startswith('Heading 2') and ("Analisis Kurikulum Pendidikan Teknologi Informasi" in text or "Analisis Kurikulum Dasar-dasar" in text):
            start_idx = i
        if p.style.name.startswith('Heading 2') and "Keterkaitan dengan Bidang Pendidikan" in text:
            end_idx = i
            break

    if start_idx == -1 or end_idx == -1:
        print(f"Bounds not found in {filename}!")
        return

    print(f"Found bounds in {filename}: {start_idx} to {end_idx}")

    # Delete all elements between start_idx and end_idx
    # Since python-docx doesn't easily expose the sequential elements (paragraphs vs tables),
    # we need to remove the XML nodes.
    start_p = doc.paragraphs[start_idx]
    end_p = doc.paragraphs[end_idx]
    
    body = start_p._p.getparent()
    in_range = False
    elements_to_remove = []
    
    for elem in body:
        if elem == start_p._p:
            in_range = True
        if in_range:
            elements_to_remove.append(elem)
        if elem == end_p._p:
            break
            
    # Keep the end_p, remove everything from start_p to just before end_p
    for elem in elements_to_remove[:-1]:
        body.remove(elem)

    # Now insert the new content before end_p
    target_p = end_p
    
    def insert_p(text, style='Normal', align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False):
        new_p = target_p.insert_paragraph_before(text, style)
        if style == 'Normal':
            new_p.alignment = align
            if bold and new_p.runs:
                new_p.runs[0].bold = True
        return new_p

    # Insert Heading 2
    insert_p("Analisis Kurikulum Dasar-dasar Pengembangan Perangkat Lunak dan Gim (PPLG)", style='Heading 2')
    
    # Insert Para 1
    insert_p("Pengembangan Website Unit Layanan Terpadu (ULT) di Fakultas Keguruan dan Ilmu Pendidikan Universitas Lampung ini turut dianalisis potensinya sebagai sumber belajar kontekstual bagi pendidikan vokasi. Berdasarkan pra-penelitian di SMK Negeri 9 Bandar Lampung, analisis kurikulum dieksekusi spesifik pada mata pelajaran Dasar-dasar Pengembangan Perangkat Lunak dan Gim (PPLG). Analisis ini bertujuan untuk membuktikan kesesuaian antara tahapan rekayasa perangkat lunak sistem nyata (Website ULT) dengan Capaian Pembelajaran (CP) dan Tujuan Pembelajaran (TP) yang dituntut dalam pedoman kejuruan tingkat SMK. Rincian penjabaran relevansi pengembangan produk tata kelola administrasi institusi ini terhadap materi kejuruan diuraikan pada tabel berikut.")
    
    insert_p("", style='Normal')
    
    # Caption Table 1
    caption1 = insert_p("Tabel 1. Analisis Kurikulum Dasar-dasar PPLG", style='Normal', align=WD_ALIGN_PARAGRAPH.LEFT)
    for r in caption1.runs:
        r.bold = True
        
    # Table 1
    table1 = doc.add_table(rows=2, cols=2)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Capaian Pembelajaran'
    hdr_cells[1].text = 'Tujuan Pembelajaran'
    for c in hdr_cells:
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True

    row_cells = table1.rows[1].cells
    row_cells[0].text = "Pada akhir pembelajaran PPLG, peserta didik mampu memahami proses bisnis industri perangkat lunak, menerapkan metodologi pengembangan sistem informasi, mengimplementasikan kerangka kerja pemrograman web modern, serta menyusun dokumentasi proyek rekayasa perangkat lunak secara terstruktur dan komprehensif."
    row_cells[1].text = "Peserta didik mampu:\n1. Mengidentifikasi kebutuhan klien dan merumuskan alur proses bisnis antarmuka web (C4).\n2. Menerapkan arsitektur kerangka kerja pemrograman spesifik (Laravel dan Tailwind CSS) (C3).\n3. Merancang integrasi tata kelola basis data relasional ke dalam sistem operasional (C5).\n4. Menunjukkan etika profesional dan tanggung jawab kerja sama tim manajemen proyek (A4).\n5. Memproduksi perangkat lunak fungsional teruji beserta dokumentasi algoritma teknis (C6)."
    
    for c in row_cells:
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    target_p._p.addprevious(table1._tbl)
    
    insert_p("", style='Normal')
    
    # Caption Table 2
    caption2 = insert_p("Tabel 2. Pemetaan Kedalaman dan Keluasan Materi PPLG Berbasis Proyek ULT", style='Normal', align=WD_ALIGN_PARAGRAPH.LEFT)
    for r in caption2.runs:
        r.bold = True
        
    # Table 2
    table2 = doc.add_table(rows=5, cols=2)
    table2.style = 'Table Grid'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Kedalaman Materi'
    hdr_cells2[1].text = 'Keluasan Materi'
    for c in hdr_cells2:
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                
    data = [
        ("Arsitektur Antarmuka (UI/UX) dan Front-end", "1. Desain tata letak antarmuka web responsif menggunakan Tailwind CSS.\n2. Pemenuhan standar aksesibilitas interaksi pengguna via Progressive Web App (PWA).\n3. Pemisahan fungsionalitas dasbor spesifik untuk ragam entitas pengguna."),
        ("Manajemen Basis Data & Keamanan Sistem", "1. Implementasi sistem pangkalan data relasional terstruktur (MySQL).\n2. Proteksi tautan dan modul enkripsi penyimpanan repositori sistem fail lokal.\n3. Penerapan penyaringan serangan siber injeksi skrip berbahaya (Cross-Site Scripting)."),
        ("Logika Bisnis (Back-end) dan Algoritma", "1. Integrasi pustaka konversi dokumen terpusat (HTML-to-OpenXML).\n2. Ekstraksi dan penggabungan variabel pendaftar ke dalam templat resmi.\n3. Pembangkitan kode pelacakan dan nomor identifikasi elektronik terotomatisasi."),
        ("Manajemen Proyek dan Dokumentasi (Audit)", "1. Pencatatan log aktivitas tanggal waktu pemrosesan transaksi (Audit Trail).\n2. Pemodelan tahapan alur persetujuan terstruktur pada sebuah sistem operasional.\n3. Pembuatan rekapitulasi data pangkalan evaluasi durasi penanganan layanan.")
    ]
    
    for i, (kedalaman, keluasan) in enumerate(data):
        cells = table2.rows[i+1].cells
        cells[0].text = kedalaman
        cells[1].text = keluasan
        for c in cells:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
    target_p._p.addprevious(table2._tbl)
    
    insert_p("", style='Normal')
    
    # Para 2
    insert_p("Pemetaan struktur analisis kurikulum mata pelajaran PPLG tersebut secara faktual menjustifikasi bahwa produk Website ULT tidak sekadar berfungsi sebagai solusi administratif internal kampus semata. Penyelarasan antara kewajiban pemenuhan pencapaian tujuan kejuruan vokasi dengan arsitektur teknologi terbukti menjadikan proyek pengembangan sistem ini sebagai media pembelajaran kontekstual yang ideal. Kerangka pemetaan ini secara akademis memastikan bahwa proses rancang bangun teknologi tata kelola administrasi dapat langsung diadaptasi oleh peserta didik SMK sebagai rujukan proyek riil yang berstandar dunia industri.")

    insert_p("", style='Normal')

    doc.save(output_filename)
    print(f"Successfully updated {output_filename}")

reinject_pplg("001_Skripsi_Andricha Dea Mitra_Clean.docx", "001_Skripsi_Andricha Dea Mitra_Clean.docx")
reinject_pplg("001_Skripsi_Andricha Dea Mitra_Highlighted.docx", "001_Skripsi_Andricha Dea Mitra_Highlighted.docx")
