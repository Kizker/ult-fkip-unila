import docx

def find_text(filename):
    doc = docx.Document(filename)
    found = False
    for p in doc.paragraphs:
        if "Kedalaman Materi" in p.text:
            print("Found in paragraph:", p.text)
            found = True
            
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if "Kedalaman Materi" in c.text:
                    print("Found in table cell:", c.text)
                    found = True
                    
    if not found:
        print("Not found anywhere!")

find_text("001_Skripsi_Andricha Dea Mitra_Clean.docx")
