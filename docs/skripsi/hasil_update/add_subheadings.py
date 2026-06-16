import docx
import copy
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

files = [
    "c:/laragon/www/ult-fkip-unila/docs/skripsi/hasil_update/001_Skripsi_Andricha Dea Mitra_Clean.docx",
    "c:/laragon/www/ult-fkip-unila/docs/skripsi/hasil_update/001_Skripsi_Andricha Dea Mitra_Highlighted.docx"
]

def add_headings(filepath, is_highlighted):
    doc = docx.Document(filepath)
    
    # Cari paragraf referensi
    ref_p = None
    for p in doc.paragraphs:
        if "Perancangan Fungsionalitas Menggunakan Diagram Use Case" in p.text:
            ref_p = p
            break
            
    if not ref_p:
        print(f"Reference not found in {filepath}")
        return

    # Kumpulkan tempat-tempat yang akan disisipkan judul baru
    insert_points = []
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith("Setelah interaksi fungsional didefinisikan, langkah berikutnya adalah merancang"):
            insert_points.append((p, "Perancangan Arsitektur Sistem"))
        elif text.startswith("Perancangan alur kerja persuratan ditujukan untuk memetakan"):
            insert_points.append((p, "Perancangan Alur Kerja Persuratan (Flowchart)"))

    for target_p, heading_text in insert_points:
        new_p = target_p.insert_paragraph_before(heading_text)
        new_p.style = ref_p.style
        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Copy numbering properties
        if ref_p._p.pPr is not None and ref_p._p.pPr.numPr is not None:
            if new_p._p.get_or_add_pPr().numPr is None:
                new_p._p.pPr.append(copy.deepcopy(ref_p._p.pPr.numPr))
                
        # Format text
        for run in new_p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            # Karena ini heading level kecil, mungkin italic atau bold tidak diminta, tapi di list awal ngga bold.
            if is_highlighted:
                from docx.enum.text import WD_COLOR_INDEX
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.save(filepath)

add_headings(files[0], False)
add_headings(files[1], True)
print("Penambahan sub-subjudul selesai.")
