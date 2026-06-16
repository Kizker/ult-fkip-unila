import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def revert_subjek(doc_path):
    doc = docx.Document(doc_path)
    
    # 1. Delete the stuff we inserted earlier in 3.5 and 3.6
    # Find 3.5 Subjek Penelitian
    idx_subjek = -1
    idx_instrumen = -1
    idx_analisis = -1
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if '3.5 Subjek Penelitian dan Sampling' in text:
            idx_subjek = i
        elif text == 'Instrumen Penelitian' or text == '3.6 Instrumen Penelitian':
            idx_instrumen = i
        elif text == 'Teknik Analisis Data' or text == '3.7 Teknik Analisis Data':
            idx_analisis = i
            
    # Clear 3.5
    if idx_subjek != -1 and idx_instrumen != -1:
        for i in range(idx_subjek + 1, idx_instrumen):
            doc.paragraphs[i].text = ''
            
    # Remove procedural text in 3.6
    if idx_analisis != -1:
        # Search backwards from analisis for the procedural texts and clear them
        for i in range(idx_analisis - 1, idx_instrumen, -1):
            text = doc.paragraphs[i].text.strip()
            if text.startswith('Setting uji coba dilakukan') or text.startswith('Urutan pelaksanaan uji coba:'):
                doc.paragraphs[i].text = ''
                
    # Also delete any tables between subjek and instrumen
    # In python-docx, table elements are not in doc.paragraphs list. 
    # But wait, earlier we deleted the ghost table using XML element removal.
    # The new table we inserted earlier is currently just before idx_instrumen.
    # Let's remove any table between 3.5 and 3.6.
    to_delete_tbls = []
    start_del = False
    for child in doc._body._body:
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl
        from docx.text.paragraph import Paragraph
        if isinstance(child, CT_P):
            text = Paragraph(child, doc).text.strip()
            if '3.5 Subjek Penelitian dan Sampling' in text:
                start_del = True
            elif 'Instrumen Penelitian' in text:
                start_del = False
        elif isinstance(child, CT_Tbl):
            if start_del:
                to_delete_tbls.append(child)
                
    for tbl in to_delete_tbls:
        tbl.getparent().remove(tbl)

    # 2. Insert the restored original text
    p_inst = doc.paragraphs[idx_instrumen]
    
    def add_p(text, style='Paragraph', bold=False):
        new_p = p_inst.insert_paragraph_before(style=style)
        if text:
            run = new_p.add_run(text)
            if bold:
                run.font.bold = True
        return new_p

    add_p('Subjek penelitian ditentukan berdasarkan kebutuhan evaluasi produk website Unit Layanan Terpadu (ULT) FKIP Universitas Lampung, yaitu validator ahli (expert judgement) dan responden uji coba (pengguna) pada tahap implementasi. Pemilihan subjek ini ditujukan agar kualitas produk dapat dinilai dari sisi kelayakan isi atau layanan, kelayakan media atau teknis, serta kemudahan penggunaan berdasarkan pengalaman pengguna setelah mencoba sistem.')
    add_p('Validator Ahli', 'Paragraph', bold=True)
    add_p('Validator ahli terdiri dari tiga kategori, yaitu: (1) ahli materi (content expert), (2) ahli media (media/technical expert), dan (3) ahli sistem (system expert). Ahli materi berperan menilai kesesuaian konten layanan, alur administrasi, dan keterpahaman informasi layanan. Ahli media berperan menilai aspek desain antarmuka, konsistensi navigasi, dan aspek media atau usability. Ahli sistem berperan menilai kelayakan teknis sistem, kestabilan fungsi inti, konsistensi alur berbasis role, serta kesesuaian implementasi fitur dengan spesifikasi sebelum produk diuji oleh pengguna.')
    add_p('Jumlah dan kualifikasi validator ahli ditetapkan sebanyak 9 (sembilan) orang, dengan rincian: ahli materi = 3 orang; ahli media = 3 orang; ahli sistem = 3 orang. Pemilihan validator dilakukan secara purposive berdasarkan relevansi keahlian dengan objek yang dinilai, pengalaman atau kompetensi pada bidangnya, serta kesediaan memberikan penilaian dan saran perbaikan terhadap produk. Penggunaan keahlian multi-disiplin ini sejalan dengan standar evaluasi kelayakan sistem (Nieveen, 1999) yang mengharuskan produk dievaluasi dari sisi relevance (materi) dan consistency (sistem/media).')
    
    add_p('Responden Uji Coba (Uji Kepraktisan atau Angket)', 'Paragraph', bold=True)
    add_p('Responden uji coba merupakan pengguna sasaran yang mencoba website ULT pada uji coba terbatas (limited trial/small group). Responden dibagi menjadi tiga kelompok pengguna utama, yaitu mahasiswa, staf Unit Layanan Terpadu (ULT), dan Administrator jurusan. Komposisi responden ditetapkan agar mewakili peran pengguna yang berinteraksi langsung dengan fitur inti (pengajuan layanan, verifikasi atau proses, pelacakan status, dan riwayat layanan).')
    add_p('Jumlah responden uji coba ditetapkan sebanyak 18 (delapan belas) orang, dengan rincian: mahasiswa = 12 orang; staf ULT = 3 orang; Administrator jurusan = 3 orang. Teknik pemilihan responden menggunakan purposive sampling. Purposive sampling adalah teknik penentuan sampel dengan pertimbangan tertentu agar sampel yang dipilih relevan dengan kebutuhan pengujian (Sugiyono, 2013). Ukuran sampel 18 orang ini dinilai sangat memadai untuk evaluasi kegunaan sistem, mengingat pengujian dengan 15-20 pengguna sudah mampu mengungkap lebih dari 95% masalah usability utama pada sebuah perangkat lunak (Faulkner, 2003; Weinger et al., 2010). Kriteria inklusi responden sebagai berikut:')
    
    add_p('Mahasiswa: mahasiswa aktif FKIP dari perwakilan rumpun PBS, PIP, PIPS, dan PMIPA; pernah atau berpotensi mengurus layanan administrasi; mampu mengoperasikan laptop atau ponsel dan browser; bersedia mengikuti skenario tugas sampai selesai dan mengisi angket.')
    add_p('Staf ULT: staf yang terlibat pelayanan ULT atau administrasi; memahami alur layanan; bersedia mencoba sistem sesuai peran dan memberi masukan.')
    add_p('Administrator jurusan: admin atau verifikator jurusan yang terlibat verifikasi atau proses layanan; memahami alur administrasi; bersedia mencoba sistem sesuai peran dan memberi masukan.')
    add_p('Umum: bersedia memberikan persetujuan (informed consent) dan data dijaga kerahasiaannya.')
    
    add_p('Setting uji coba dilakukan dalam kondisi terkontrol agar dapat direplikasi. Responden menggunakan perangkat yang tersedia seperti laptop, komputer atau ponsel dengan browser modern dan koneksi internet yang memadai. Durasi uji coba: 30 menit. Sebelum pengisian angket kepraktisan, responden diberikan pengarahan singkat dan panduan tugas (task scenario) untuk memastikan semua responden mencoba fungsi inti yang sama.')
    add_p('Urutan pelaksanaan uji coba: (1) briefing singkat dan pembagian akun sesuai peran; (2) pelaksanaan skenario tugas minimal (login sesuai peran, mahasiswa mengajukan layanan dan mengunggah dokumen, staf atau admin memverifikasi dan memproses permohonan, pengguna melacak status layanan, dan pengguna melihat riwayat layanan); (3) pengisian angket uji kepraktisan (angket) setelah mencoba sistem; (4) pencatatan komentar/kendala untuk bahan revisi. Jumlah responden (N) pada tahap ini digunakan dalam perhitungan skor rata-rata uji kepraktisan, sehingga harus konsisten dengan lampiran instrumen dan pelaporan hasil.')
    
    add_p('')
    ptitle = add_p('Tabel 3.2 Subjek Penelitian dan Sampling', 'Normal', bold=True)
    ptitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    p_inst._p.addprevious(table._tbl)
    
    headers = ['Kelompok Subjek', 'Peran dalam Evaluasi', 'Jumlah (n)', 'Teknik Pemilihan & Kriteria Ringkas']
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold = True
            
    data = [
        ['Validator Ahli Materi', 'Menilai kesesuaian konten layanan, alur administrasi, dan keterpahaman informasi.', '3', 'Purposive; kompetensi pada layanan administrasi; memahami SOP/alur layanan; bersedia memberi masukan.'],
        ['Validator Ahli Media', 'Menilai desain antarmuka, konsistensi navigasi, dan kualitas fungsionalitas UI/UX sistem.', '3', 'Purposive; kompetensi pengembangan website atau UI/UX; memahami aspek teknis dan usability.'],
        ['Validator Ahli Sistem', 'Menilai kelayakan teknis sistem, kestabilan, konsistensi alur berbasis role, serta keamanan.', '3', 'Purposive; kompetensi rekayasa perangkat lunak dan pengujian fungsional.'],
        ['Mahasiswa', 'Mengikuti uji coba terbatas dan mengisi angket kepraktisan setelah mencoba sistem.', '12', 'Purposive; mahasiswa aktif FKIP 4 rumpun ilmu; mampu mengoperasikan web browser; mengikuti skenario.'],
        ['Staf ULT', 'Mengikuti uji coba terbatas sesuai peran pengelola layanan dan mengisi angket kepraktisan.', '3', 'Purposive; staf yang terlibat pelayanan ULT/administrasi; memahami alur layanan.'],
        ['Administrator Jurusan', 'Mengikuti uji coba terbatas sesuai peran verifikator layanan dan mengisi angket kepraktisan.', '3', 'Purposive; admin verifikator jurusan terlibat verifikasi; memahami alur administrasi.']
    ]
    
    for row_idx, row_data in enumerate(data):
        for col_idx, val in enumerate(row_data):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = val
            if col_idx == 2:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    widths = [1.5, 2.5, 0.7, 2.5]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)

    add_p('Sumber: Adaptasi dari rancangan evaluasi penelitian pengembangan (R&D) dan disesuaikan oleh peneliti dengan konteks pengembangan website ULT FKIP Universitas Lampung.', 'Normal')
    add_p('')
    
    doc.save(doc_path)
    print(f"Successfully restored {doc_path}")

if __name__ == '__main__':
    revert_subjek(r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx')
    revert_subjek(r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx')
