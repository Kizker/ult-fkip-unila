import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_waktu_tempat(doc_path, is_highlighted=False):
    doc = docx.Document(doc_path)
    
    # Find the target paragraph "Prosedur Penelitian"
    target_p = None
    for p in doc.paragraphs:
        if p.style.name == 'Heading 2' and 'Prosedur Penelitian' in p.text:
            target_p = p
            break
            
    if not target_p:
        print(f"Could not find 'Prosedur Penelitian' in {doc_path}")
        return
        
    def add_highlight(run):
        if is_highlighted:
            rPr = run._r.get_or_add_rPr()
            highlight = OxmlElement('w:highlight')
            highlight.set(qn('w:val'), 'yellow')
            rPr.append(highlight)
            
    def insert_p_before(text, style_name):
        new_p = target_p.insert_paragraph_before(text, style=style_name)
        if text:
            # Re-apply text through run to handle highlight
            new_p.text = ''
            run = new_p.add_run(text)
            add_highlight(run)
        return new_p

    # Insert contents
    # Heading 2
    insert_p_before('Waktu dan Tempat Penelitian', 'Heading 2')
    
    # Heading 3 - Waktu
    insert_p_before('Waktu Penelitian', 'Heading 3')
    
    # Narrative
    narasi = (
        "Penelitian ini dilaksanakan melalui serangkaian tahapan yang terstruktur dan sistematis. "
        "Berdasarkan rekam jejak pengembangan sistem, tahapan awal berupa persiapan dan revisi pasca seminar "
        "proposal dimulai pada akhir bulan Januari 2026. Selanjutnya, tahapan pembuatan perangkat lunak (website) "
        "berlangsung secara komprehensif dari bulan Februari hingga bulan April 2026. "
        "Setelah sistem web dinyatakan siap secara fungsional, tahapan uji coba terbatas kepada para ahli (validator) "
        "dan pengguna lapangan (mahasiswa, admin program studi, dan staf ULT) dilaksanakan secara intensif pada bulan Mei 2026. "
        "Secara paralel, proses penyusunan dan penulisan karya ilmiah (skripsi) dikerjakan mulai dari bulan Maret hingga proses penyelesaian pada bulan Juni 2026. "
        "Rincian penjadwalan waktu pelaksanaan kegiatan penelitian ini disajikan secara lengkap pada Tabel 3.1 berikut."
    )
    insert_p_before(narasi, 'Paragraph')
    
    # Table Title (Normal style but centered, bold? Let's use 1.0 spacing as per Unila)
    # The user's Unila style uses 'Normal' for table title, usually single spaced. We'll use Normal.
    tbl_title = insert_p_before('Tabel 3.1 Waktu Pelaksanaan Penelitian', 'Normal')
    tbl_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Clear runs and add bold run
    tbl_title.text = ''
    run = tbl_title.add_run('Tabel 3.1 Waktu Pelaksanaan Penelitian')
    run.font.bold = True
    add_highlight(run)
    
    # Create Table
    # Table headers: No. | Kegiatan Penelitian | Jan | Feb | Mar | Apr | Mei | Jun
    table = doc.add_table(rows=5, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Move table to before target_p
    # Element insertion: target_p._p.addprevious(table._tbl)
    target_p._p.addprevious(table._tbl)
    
    headers = ['No.', 'Kegiatan Penelitian', 'Jan', 'Feb', 'Mar', 'Apr', 'Mei', 'Jun']
    data = [
        ['1', 'Persiapan & Revisi Proposal', 'v', 'v', 'v', '', '', ''],
        ['2', 'Pembuatan Website', 'v', 'v', 'v', 'v', '', ''],
        ['3', 'Uji Coba Terbatas', '', '', '', '', 'v', ''],
        ['4', 'Penulisan Skripsi', '', '', 'v', 'v', 'v', 'v']
    ]
    
    # Fill headers
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                add_highlight(r)
                
    # Fill data
    for row_idx, row_data in enumerate(data):
        for col_idx, text in enumerate(row_data):
            cell = table.cell(row_idx + 1, col_idx)
            val = '✔️' if text == 'v' else text
            cell.text = val
            for p in cell.paragraphs:
                # Center all columns except "Kegiatan Penelitian"
                if col_idx != 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    add_highlight(r)
                    
    # Adjust column widths manually since Word tables need explicit widths to look good
    widths = [0.4, 2.5, 0.5, 0.5, 0.5, 0.5, 0.5, 0.5]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)

    # Empty paragraph after table for spacing
    insert_p_before('', 'Normal')
    
    # Heading 3 - Tempat
    insert_p_before('Tempat Penelitian', 'Heading 3')
    
    # Narrative Tempat
    narasi_tempat = (
        "Penelitian pengembangan ini dilaksanakan di lingkungan Unit Layanan Terpadu (ULT) "
        "Fakultas Keguruan dan Ilmu Pendidikan (FKIP) Universitas Lampung, yang berlokasi di Jl. Prof. Dr. Ir. Sumantri Brojonegoro No.1, "
        "Gedong Meneng, Kec. Rajabasa, Kota Bandar Lampung, Lampung 35141."
    )
    insert_p_before(narasi_tempat, 'Paragraph')
    
    # Save document
    doc.save(doc_path)
    print(f"Successfully updated {doc_path}")

if __name__ == '__main__':
    clean_path = r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
    highlighted_path = r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
    
    add_waktu_tempat(clean_path, is_highlighted=False)
    add_waktu_tempat(highlighted_path, is_highlighted=True)
