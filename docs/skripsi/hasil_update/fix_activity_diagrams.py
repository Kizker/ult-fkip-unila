import urllib.request
import urllib.parse
import base64
import json
import os
import docx
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

diagrams = [
    {
        "file": r"c:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram\06_activity_diagram_autentikasi.jpg",
        "search_cap": "Diagram Activity Autentikasi Pengguna",
        "new_cap": "Gambar 4.1. Diagram Activity Autentikasi Pengguna",
        "desc": "Gambar 4.1 mengilustrasikan alur aktivitas autentikasi pengguna pada Web ULT FKIP Unila. Mahasiswa baru diwajibkan mengisi formulir registrasi yang akan divalidasi dan dikirimkan tautan verifikasi melalui email. Pengguna yang telah memiliki akun terverifikasi dapat langsung memasukkan email dan kata sandi pada halaman login. Sistem akan memvalidasi kredensial tersebut terhadap basis data; jika valid, sistem menciptakan sesi login dan mengarahkan pengguna ke dasbor masing-masing sesuai dengan perannya (role), sedangkan jika gagal, sistem akan menampilkan pesan peringatan.",
        "code": """flowchart TD
    subgraph Pengguna
        A((Mulai)) --> B{Sudah punya akun?}
        B -- Belum --> C[Isi Form Registrasi]
        C --> D[Submit Form]
        B -- Sudah --> E[Buka Halaman Login]
        E --> F[Input Email & Password]
        F --> G[Submit Login]
    end
    subgraph Sistem
        D --> H[Simpan Data & Kirim Email]
        H --> I[Verifikasi Email]
        I --> E
        
        G --> J{Validasi Kredensial}
        J -- Gagal --> K[Tampilkan Error]
        K --> E
        J -- Berhasil --> L[Buat Sesi Login]
    end
    L --> M((Selesai))"""
    },
    {
        "file": r"c:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram\07_activity_diagram_pengajuan.jpg",
        "search_cap": "Diagram Activity Pengajuan Layanan Akademik",
        "new_cap": "Gambar 4.2. Diagram Activity Pengajuan Layanan Akademik",
        "desc": "Gambar 4.2 memetakan proses utama pengajuan layanan persuratan oleh mahasiswa. Dimulai dari Student Portal, mahasiswa memilih jenis layanan spesifik, mengisi formulir berbasis variabel dinamis yang disediakan oleh sistem, serta mengunggah berkas-berkas prasyarat. Sistem secara otomatis melakukan validasi input. Permohonan yang tervalidasi kemudian disimpan ke dalam basis data dengan status awal 'Menunggu Verifikasi', dan sebuah notifikasi proaktif diteruskan kepada staf ULT untuk ditindaklanjuti.",
        "code": """flowchart TD
    subgraph Mahasiswa
        A((Mulai)) --> B[Login Student Portal]
        B --> C[Buka Buat Permohonan]
        C --> D[Pilih Layanan]
        D --> E[Isi Formulir Dinamis]
        E --> F[Unggah Berkas Prasyarat]
        F --> G[Kirim Permohonan]
    end
    subgraph Sistem
        G --> H{Validasi Input}
        H -- Tidak Valid --> I[Tampilkan Peringatan]
        I --> E
        H -- Valid --> J[Simpan Status: Menunggu]
        J --> K[Kirim Notifikasi ke Staf]
    end
    K --> L((Selesai))"""
    },
    {
        "file": r"c:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram\08_activity_diagram_verifikasi.jpg",
        "search_cap": "Diagram Activity Verifikasi dan Pemrosesan Dokumen",
        "new_cap": "Gambar 4.3. Diagram Activity Verifikasi dan Pemrosesan Dokumen",
        "desc": "Gambar 4.3 menjabarkan alur kerja operasional oleh Staf ULT dan Pejabat struktural. Staf ULT berperan sebagai verifikator pertama yang mengecek kesesuaian berkas. Jika terdapat ketidaksesuaian, permohonan akan ditolak dengan catatan khusus untuk diperbaiki mahasiswa. Jika valid, sistem mendeteksi apakah dokumen memerlukan tanda tangan elektronik tingkat pimpinan. Jika ya, draf elektronik diteruskan ke Signer Portal agar pejabat dapat mereview dan membubuhkan tanda tangan. Pada tahap final, mesin assembly merakit keseluruhan data dan tanda tangan ke dalam format dokumen legal (.docx), lalu mengubah status permohonan menjadi 'Selesai'.",
        "code": """flowchart TD
    subgraph Staf ULT
        A((Mulai)) --> B[Terima Notifikasi]
        B --> C[Buka Detail Permohonan]
        C --> D{Verifikasi Valid?}
        D -- Tidak --> E[Catatan Penolakan]
    end
    subgraph Sistem
        E --> F[Status: Ditolak/Revisi]
    end
    subgraph Staf ULT
        D -- Ya --> H{Butuh Tanda Tangan?}
        H -- Tidak --> I[Input Nomor Surat]
    end
    subgraph Pejabat Signer
        H -- Ya --> J[Review Draf Surat]
        J --> K[Berikan Tanda Tangan]
    end
    subgraph Sistem
        K --> M[Sistem Catat Persetujuan]
        I --> M
        M --> N[Assembly Engine: Render DOCX/PDF]
        N --> O[Status: Selesai]
        O --> P((Selesai))
    end"""
    },
    {
        "file": r"c:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram\09_activity_diagram_manajemen.jpg",
        "search_cap": "Diagram Activity Manajemen Layanan dan Templat",
        "new_cap": "Gambar 4.4. Diagram Activity Manajemen Layanan dan Templat",
        "desc": "Gambar 4.4 menunjukkan fungsionalitas manajemen layanan yang dioperasikan oleh Admin ULT. Admin memiliki wewenang untuk meracik layanan persuratan baru melalui antarmuka dasbor tanpa intervensi langsung pada kode sumber aplikasi (source code). Proses ini mencakup pengaturan penamaan layanan, pendefinisian bidang isian dinamis (dynamic inputs) yang wajib diisi mahasiswa, serta pengunggahan templat master dokumen berformat Word (.docx). Sistem lalu akan memvalidasi format dokumen tersebut dan segera memperbarui katalog layanan secara seketika (real-time) di halaman publik.",
        "code": """flowchart TD
    subgraph Admin ULT
        A((Mulai)) --> B[Buka Menu Layanan]
        B --> C[Tambah/Edit Layanan]
        C --> D[Konfigurasi Syarat]
        D --> E[Atur Variabel Input]
        E --> F[Unggah Template .docx]
        F --> G[Simpan Konfigurasi]
    end
    subgraph Sistem
        G --> H{Validasi Template}
        H -- Error --> I[Tampilkan Peringatan]
        I --> F
        H -- Valid --> J[Simpan Layanan]
        J --> K[Update Katalog Publik]
    end
    K --> L((Selesai))"""
    }
]

# 1. Generate Monochrome Images
theme_vars = {
    'primaryColor': '#ffffff',
    'primaryTextColor': '#000000',
    'primaryBorderColor': '#000000',
    'lineColor': '#000000',
    'secondaryColor': '#ffffff',
    'tertiaryColor': '#ffffff',
    'fontFamily': 'Times New Roman',
    'clusterBkg': '#ffffff',
    'clusterBorder': '#000000',
    'edgeLabelBackground': '#ffffff'
}
theme_param = urllib.parse.quote(json.dumps(theme_vars))

for item in diagrams:
    b64 = base64.b64encode(item["code"].encode('utf-8')).decode('ascii')
    url = f'https://mermaid.ink/img/{b64}?theme=base&themeVariables={theme_param}'
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    try:
        with urllib.request.urlopen(req) as response, open(item["file"], 'wb') as out_file:
            out_file.write(response.read())
        print(f"Generated {os.path.basename(item['file'])}")
    except Exception as e:
        print(f"Failed to generate {os.path.basename(item['file'])}: {e}")

# 2. Process Word Documents
def process_doc(filepath, is_highlighted):
    doc = docx.Document(filepath)
    
    for item in diagrams:
        idx = -1
        for i, p in enumerate(doc.paragraphs):
            if item["search_cap"] in p.text and 'Gambar' in p.text:
                idx = i
                break
                
        if idx == -1:
            print(f"Caption '{item['search_cap']}' not found in {filepath}")
            continue
            
        img_p = doc.paragraphs[idx-1]
        cap_p = doc.paragraphs[idx]
        desc_p = doc.paragraphs[idx+1]
        
        img_p.clear()
        cap_p.clear()
        desc_p.clear()
        
        # Apply formatting
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.paragraph_format.space_before = Pt(12)
        img_p.paragraph_format.space_after = Pt(6)
        run_img = img_p.add_run()
        run_img.add_picture(item["file"], width=Inches(4.2)) # Resized to 4.2 inches for better scale
        
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_before = Pt(0)
        cap_p.paragraph_format.space_after = Pt(12)
        cap_p.paragraph_format.line_spacing = 1.0
        run_cap = cap_p.add_run(item["new_cap"])
        run_cap.font.name = 'Times New Roman'
        run_cap.font.size = Pt(11) # Standard caption size
        if is_highlighted:
            run_cap.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
        desc_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        desc_p.paragraph_format.space_before = Pt(0)
        desc_p.paragraph_format.space_after = Pt(12)
        desc_p.paragraph_format.line_spacing = 1.5
        desc_p.paragraph_format.first_line_indent = Cm(1.27) # Proper indent
        run_desc = desc_p.add_run(item["desc"])
        run_desc.font.name = 'Times New Roman'
        run_desc.font.size = Pt(12)
        if is_highlighted:
            run_desc.font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.save(filepath)
    print("Updated", filepath)

files = [
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx',
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
]

for f in files:
    process_doc(f, "Highlighted" in f)
