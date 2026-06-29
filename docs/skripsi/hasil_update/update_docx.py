import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def update_doc(filename, output_filename, is_highlighted):
    doc = docx.Document(filename)
    
    # Find "Pembahasan Hasil Penelitian" and "KESIMPULAN DAN SARAN"
    start_idx = -1
    end_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.endswith("Pembahasan Hasil Penelitian") and p.style.name == 'Heading 2':
            start_idx = i
        elif text.startswith("KESIMPULAN DAN SARAN") and p.style.name == 'Heading 1':
            end_idx = i
            if start_idx != -1:
                break
                
    if start_idx == -1 or end_idx == -1:
        print(f"Failed to find bounds in {filename}. Start: {start_idx}, End: {end_idx}")
        return
        
    print(f"Found bounds in {filename}: {start_idx} to {end_idx}")
    
    # Delete paragraphs between start_idx and end_idx - 1
    for i in range(start_idx, end_idx):
        p = doc.paragraphs[start_idx]
        p._element.getparent().remove(p._element)
        
    # Content to insert
    content = [
        ("Heading 2", "4.2 Pembahasan Hasil Penelitian"),
        ("Normal", "Penelitian ini mengadopsi model desain instruksional ADDIE sebagai kerangka kerja filosofis sekaligus sistematis dalam memecahkan permasalahan administrasi akademik. Pendekatan ini dipilih secara terencana untuk memastikan setiap fase pengembangan produk selalu dikawal oleh proses evaluasi empiris yang terukur. Rangkaian panjang penelitian berawal dari observasi mendalam terhadap inefisiensi loket konvensional fakultas yang sangat bergantung pada peredaran kertas fisik. Iterasi desain antarmuka digital yang kompleks pada akhirnya lahir sebagai jawaban mutlak atas kemandekan sistem birokrasi lama tersebut."),
        ("Normal", "Produk akhir dari siklus pengembangan ini berhasil mencapai tingkat kelayakan ahli sebesar 91,95 persen dengan predikat Sangat Valid. Angka kepraktisan pengguna turut menyentuh 92,13 persen yang secara empiris menegaskan predikat Sangat Praktis dalam uji coba lapangan. Capaian kuantitatif tersebut membuktikan penciptaan sebuah solusi komprehensif yang secara holistik merevolusi tata kelola administrasi kampus lokal. Pembenahan struktural ini secara nyata menghasilkan arsitektur layanan unggul alih-alih sekadar meluncurkan purwarupa aplikasi biasa."),
        ("Heading 3", "4.2.1 Makna Kepraktisan dan Korelasi Landasan Teori"),
        ("Normal", "Pencapaian skor kepraktisan sebesar 92,13 persen merepresentasikan penerimaan adaptasi teknologi yang masif di kalangan sivitas akademika Universitas Lampung. Skor tinggi ini memanifestasikan penurunan beban kognitif pengguna secara drastis saat mengakses layanan persuratan akademik. Mahasiswa tidak perlu lagi memahami alur birokrasi berbelit akibat kehadiran sistem cerdas yang memandu setiap langkah pengajuan dokumen. Interaksi manusia dan komputer terbangun secara natural seiring dengan berkurangnya kerumitan proses administratif harian."),
        ("Normal", "Temuan empiris ini beresonansi kuat dengan teori User Experience yang digagas oleh Maulana et al. (2023) mengenai urgensi perancangan antarmuka dalam mereduksi friksi interaksi pengguna. Kajian efisiensi digitalisasi oleh Sukorini et al. (2024) turut terkonfirmasi melalui percepatan durasi pemrosesan dokumen di tingkat fakultas maupun program studi. Implementasi website Unit Layanan Terpadu ini secara sah melahirkan paradigma baru dalam tata kelola administrasi kampus lokal yang mengedepankan efisiensi radikal. Kemudahan akses informasi menjadi pilar utama dalam membangun fondasi ekosistem pendidikan tinggi berbasis teknologi pintar."),
        ("Heading 3", "4.2.2 Transformasi Sistem Konvensional Menuju Transparansi Layanan Terpadu"),
        ("Normal", "Ketiadaan transparansi dokumen mahasiswa pada birokrasi konvensional berhasil diselesaikan secara tuntas melalui implementasi fitur Audit Trail digital. Mahasiswa sebelumnya selalu menghadapi ketidakpastian informasi mengenai posisi berkas fisik mereka yang sering tertumpuk tanpa status kejelasan di meja petugas [GAMBAR: Pelacakan Status Dokumen - Kondisi Sebelum]. Pelacakan status permohonan kini dapat dipantau secara langsung melalui antarmuka dasbor interaktif yang senantiasa merekam setiap perubahan kronologis secara akurat [GAMBAR: Pelacakan Status Dokumen - Kondisi Sesudah]. Fitur pelacakan progresif ini sukses mendobrak budaya komunikasi pasif satu arah yang selama ini sangat merugikan pihak sivitas akademika."),
        ("Normal", "Mekanisme pelacakan ini beroperasi secara seketika mengikuti rentetan siklus pengajuan, peninjauan staf, penandatanganan pimpinan, hingga penyelesaian dokumen akhir. Visibilitas proses birokrasi ini memberikan dampak psikologis yang sangat positif dengan menurunkan tingkat kecemasan mahasiswa secara signifikan setiap kali mengajukan permohonan. Ekosistem digital ini secara fundamental merombak tata kelola konvensional yang tertutup menjadi sebuah bentuk birokrasi institusi pendidikan yang transparan secara radikal. Budaya kerja berbasis keterbukaan data berhasil menggantikan pola kerja lama yang bergantung pada dokumen cetak nirjejak."),
        ("Heading 3", "4.2.3 Keunggulan Sistem"),
        ("Normal", "Arsitektur Role-Based Access Control berhasil memisahkan tanggung jawab operasional melalui empat portal terdedikasi yang meliputi area publik, mahasiswa, staf, dan pimpinan penandatangan. Sistem konvensional sebelumnya mencampuradukkan wewenang pengelolaan dokumen tanpa batas demarkasi yang jelas antar entitas pengguna akademik [GAMBAR: Hierarki Hak Akses Pengguna - Kondisi Sebelum]. Pendekatan portal terpisah ini menjamin pendelegasian tugas akademik berjalan tepat sasaran sesuai tingkat otorisasi masing-masing peran pengguna [GAMBAR: Hierarki Hak Akses Pengguna - Kondisi Sesudah]. Pimpinan fakultas memegang kendali penuh atas persetujuan akhir tanpa harus terlibat dalam proses administrasi tingkat bawah."),
        ("Normal", "Inovasi perakitan dokumen dinamis berbasis standar OpenXML memungkinkan injeksi data mahasiswa secara terprogram ke dalam templat surat resmi tanpa merusak tata letak visual bawaan. Petugas layanan konvensional sebelumnya harus menyalin teks secara manual ke dalam aplikasi pengolah kata yang rentan terhadap kesalahan pengetikan [GAMBAR: Proses Perakitan Dokumen - Kondisi Sebelum]. Generator dokumen otomatis ini sanggup memelihara gaya tipografi orisinal sekaligus mempercepat durasi penerbitan surat secara eksponensial [GAMBAR: Proses Perakitan Dokumen - Kondisi Sesudah]. Otomatisasi ini mengeliminasi beban kerja berulang staf kependidikan dalam menyusun draf administrasi akademik mahasiswa secara keseluruhan."),
        ("Normal", "Infrastruktur pelindungan data dibangun menggunakan lapisan keamanan penyimpanan privat yang terintegrasi penuh dengan mekanisme mitigasi Insecure Direct Object Reference. Arsip persuratan akademik pada era prakomputerisasi rentan mengalami kerusakan fisik serta kebocoran informasi akibat ketiadaan brankas penyimpanan terisolasi [GAMBAR: Keamanan Penyimpanan Dokumen - Kondisi Sebelum]. Lapisan autentikasi modern ini memblokir seluruh upaya akses ilegal pihak ketiga terhadap berkas sensitif mahasiswa secara definitif [GAMBAR: Keamanan Penyimpanan Dokumen - Kondisi Sesudah]. Skema enkripsi jalur akses secara masif mempertebal benteng pertahanan peladen web dari potensi peretasan massal dokumen privat."),
        ("Normal", "Ekosistem pencatatan jejak audit beroperasi di belakang layar untuk merekam seluruh jejak aktivitas transaksional setiap aktor di dalam sistem secara permanen. Pengecekan riwayat penanganan surat pada sistem pencatatan manual memerlukan waktu berjam-jam akibat pencarian buku register yang sangat memakan tenaga fisik [GAMBAR: Riwayat Transaksi Layanan - Kondisi Sebelum]. Tabel jejak audit digital ini menyajikan rekam jejak waktu nyata yang tidak dapat dimanipulasi demi menjamin akuntabilitas seluruh pemangku kepentingan [GAMBAR: Riwayat Transaksi Layanan - Kondisi Sesudah]. Pemantauan riwayat persetujuan memberikan landasan evaluasi kinerja staf ULT secara objektif bersandarkan data rekaman operasional peladen harian."),
        ("Normal", "Kompabilitas Progressive Web App menawarkan fleksibilitas akses luar biasa yang menyerupai pengalaman menggunakan aplikasi bawaan sistem operasi pada perangkat bergerak. Mahasiswa sebelumnya harus membuka peramban web dan mengetikkan alamat portal secara berulang kali untuk sekadar mengecek pembaruan status permohonan mereka [GAMBAR: Aksesibilitas Sistem Mobile - Kondisi Sebelum]. Instalasi instan pada layar utama gawai mahasiswa memastikan notifikasi serta fitur layanan persuratan dapat dijangkau kapan saja tanpa friksi [GAMBAR: Aksesibilitas Sistem Mobile - Kondisi Sesudah]. Kenyamanan akses antarmuka gawai menegaskan komitmen pengembang terhadap paradigma desain responsif tingkat lanjut demi memaksimalkan kepuasan pengguna akhir."),
        ("Heading 3", "4.2.4 Evaluasi, Kendala, dan Resolusi Teknis"),
        ("Normal", "Degradasi format perenderan teks Hypertext Markup Language bawaan editor Tiptap menuju format dokumen Microsoft Word menjadi kendala teknis paling krusial pada fase awal pengembangan. Teks tebal, miring, dan daftar bernomor dari antarmuka web sebelumnya tercetak mentah beserta label kodenya saat diunduh sebagai berkas fisik keluaran [GAMBAR: Render Format Teks Dokumen - Kondisi Sebelum]. Pengembangan parser dinamis khusus yang mampu menerjemahkan kode web ke dalam standar OpenXML berhasil memulihkan integritas visual dokumen hasil rakitan [GAMBAR: Render Format Teks Dokumen - Kondisi Sesudah]. Pemecahan masalah konversi format ini sukses mengangkat kualitas dokumen surat keluar hingga setara dengan pengetikan manual standar format universitas."),
        ("Normal", "Tindak lanjut evaluasi tim ahli mencakup mitigasi kerentanan serangan Cross-Site Scripting melalui pengetatan Content Security Policy pada lapisan kepala permintaan peladen. Formulir pengisian data mahasiswa pada iterasi awal belum memiliki tembok penyaring karakter berbahaya yang rentan dieksploitasi oleh pihak tidak bertanggung jawab untuk menyisipkan skrip asing [GAMBAR: Validasi Keamanan Input Form - Kondisi Sebelum]. Perbaikan lapis ganda berupa validasi ketat dan penyaringan konten web menjamin keamanan data akademik tetap terproteksi secara absolut dari injeksi skrip eksternal [GAMBAR: Validasi Keamanan Input Form - Kondisi Sesudah]. Sistem pencegahan intrusi terpasang sangat kuat di setiap modul formulir sebagai penjaga gerbang utama pencegah masuknya manipulasi data kotor."),
        ("Normal", "Saran perbaikan pakar media langsung diakomodasi melalui transformasi komposisi warna antarmuka web dari palet gradasi yang mencolok menuju gaya desain datar yang jauh lebih profesional. Tampilan visual beranda sebelumnya mengeksploitasi efek transisi warna berlebihan yang justru mendistraksi fokus pembaca terhadap esensi penyampaian informasi layanan akademik kampus [GAMBAR: Komposisi Warna Antarmuka Web - Kondisi Sebelum]. Penggunaan warna solid korporat kampus dengan tingkat kontras yang proporsional berhasil menaikkan nilai keeleganan visual portal pelayanan terpadu [GAMBAR: Komposisi Warna Antarmuka Web - Kondisi Sesudah]. Revitalisasi elemen grafis antarmuka sanggup menciptakan harmoni visual tata letak komponen yang sejalan dengan pakem interaksi ramah pengguna."),
        ("Normal", "Responden uji coba lapangan merekomendasikan integrasi modul kalender penjadwalan dinamis dan penambahan generator ekspresi reguler khusus untuk penomoran surat program studi secara otomatis. Staf fakultas sebelumnya harus memasukkan rangkaian kode surat secara manual dan menebak jadwal ketersediaan pimpinan yang sering kali berujung pada bentrok agenda [GAMBAR: Penomoran Surat dan Penjadwalan - Kondisi Sebelum]. Sistem penomoran cerdas beserta kalender interaktif kini mengotomatiskan rutinitas administrasi harian sekaligus mencegah risiko tumpang tindih waktu penandatanganan dokumen antar pimpinan [GAMBAR: Penomoran Surat dan Penjadwalan - Kondisi Sesudah]. Pengurangan beban memori operasional para pemangku kepentingan terwujud sepenuhnya berkat penerapan komputasi cerdas dari rangkaian algoritma peladen mandiri."),
        ("Normal", "Optimalisasi rasio ruang kosong dan penerapan jarak tepi 1.5 dieksekusi demi mencegah kelelahan visual para pengguna yang harus berlama-lama menatap dasbor sistem operasional. Barisan teks pada antarmuka tabel versi purwarupa tampak terlalu rapat serta saling berdempetan yang sangat membebani daya akomodasi mata para pemangku kepentingan [GAMBAR: Tata Letak Ruang Kosong Tabel - Kondisi Sebelum]. Penataan ulang ruang negatif antarkomponen desain web sukses menciptakan ruang bernapas visual yang memanjakan kenyamanan baca seluruh strata pengguna sistem [GAMBAR: Tata Letak Ruang Kosong Tabel - Kondisi Sesudah]. Proporsi anatomi teks pada perangkat pelihat berukuran kecil hingga resolusi layar besar kini sukses menempati hierarki pembacaan ergonomis tingkat prima.")
    ]
    
    # We will insert new paragraphs before start_idx (which is now the KESIMPULAN DAN SARAN paragraph)
    target_p = doc.paragraphs[start_idx]
    
    for style, text in content:
        new_p = target_p.insert_paragraph_before(text, style)
        if style == "Normal":
            # Justify is standard, but the instructions said "Rata Kiri (Left) hanya diterapkan pada naskah isi (Normal) yang sebelumnya berformat Justify atau default (None)."
            # Let's align to Left.
            new_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if is_highlighted:
                # Add yellow highlight to all runs in the paragraph
                for run in new_p.runs:
                    run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
                    
    # Ensure there is an empty line before KESIMPULAN DAN SARAN
    target_p.insert_paragraph_before("", "Normal")

    doc.save(output_filename)
    print(f"Successfully saved {output_filename}")

update_doc(
    "001_Skripsi_Andricha Dea Mitra_Clean.docx", 
    "001_Skripsi_Andricha Dea Mitra_Clean.docx", 
    is_highlighted=False
)
update_doc(
    "001_Skripsi_Andricha Dea Mitra_Highlighted.docx", 
    "001_Skripsi_Andricha Dea Mitra_Highlighted.docx", 
    is_highlighted=True
)
