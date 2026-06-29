import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from generate_full_text import content_structure

def update_doc(filename, output_filename, is_highlighted):
    doc = docx.Document(filename)
    
    # Find bounds
    start_idx = -1
    end_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if "4.2 Pembahasan Hasil Penelitian" in text:
            start_idx = i
        elif text.startswith("KESIMPULAN DAN SARAN") and p.style.name == 'Heading 1':
            end_idx = i
            if start_idx != -1:
                break
                
    if start_idx == -1 or end_idx == -1:
        print(f"Failed to find bounds in {filename}. Start: {start_idx}, End: {end_idx}")
        return
        
    print(f"Found bounds in {filename}: {start_idx} to {end_idx}")
    
    # Delete paragraphs between start_idx and end_idx - 1
    for i in range(start_idx, end_idx):
        p = doc.paragraphs[start_idx]
        p._element.getparent().remove(p._element)
        
    # We will insert new paragraphs before start_idx (which is now KESIMPULAN DAN SARAN)
    target_p = doc.paragraphs[start_idx]
    
    for style, text in content_structure:
        new_p = target_p.insert_paragraph_before(text, style if style != "Placeholder" else "Normal")
        
        if style == "Placeholder":
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style == "Caption":
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style == "Normal":
            new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Match standard academic style (space after 12pt, or keep spacer approach)
            if is_highlighted:
                # Add yellow highlight
                for run in new_p.runs:
                    run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
                    
        # Add a spacer paragraph after every element to ensure distance between paragraphs
        # The user loved the spacer: "dan setiap antar paragraph diberikan jarak"
        if style != "Placeholder": 
            spacer = target_p.insert_paragraph_before("", "Normal")
            spacer.paragraph_format.space_after = Pt(0)
            spacer.paragraph_format.space_before = Pt(0)

    # Ensure there is an empty line before KESIMPULAN DAN SARAN
    target_p.insert_paragraph_before("", "Normal")

    doc.save(output_filename)
    print(f"Successfully saved {output_filename}")

try:
    update_doc(
        "001_Skripsi_Andricha Dea Mitra_Clean.docx", 
        "001_Skripsi_Andricha Dea Mitra_Clean.docx", 
        is_highlighted=False
    )
    update_doc(
        "001_Skripsi_Andricha Dea Mitra_Highlighted.docx", 
        "001_Skripsi_Andricha Dea Mitra_Highlighted.docx", 
        is_highlighted=True
    )
except Exception as e:
    print(f"An error occurred: {e}")
