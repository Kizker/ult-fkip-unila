import os
import shutil
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def insert_paragraph_after(element, text, style=None):
    new_p = OxmlElement("w:p")
    # Add basic run and text
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    new_p.append(r)
    # Insert right after the given element
    element.addnext(new_p)
    return new_p

def process_docx(file_path):
    print(f"Processing {file_path}...")
    doc = Document(file_path)

    print("Step 1: Fix Headings & Page Breaks & Find/Replace")
    for p in doc.paragraphs:
        # Find and replace
        for run in p.runs:
            if "Ringkasan" in run.text:
                run.text = run.text.replace("Ringkasan", "Rekapitulasi")

        # Heading 1 format
        if p.style.name.startswith('Heading 1') or p.style.name.startswith('heading 1'):
            p.paragraph_format.space_before = Pt(85)
            p.paragraph_format.space_after = Pt(48)

        # Alignment fix for Covers, Captions
        text_stripped = p.text.strip()
        if text_stripped.startswith("Gambar ") or text_stripped.startswith("Tabel "):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Cover titles (all caps)
        if text_stripped in [
            "PENGEMBANGAN SISTEM INFORMASI BERBASIS WEBSITE UNTUK", 
            "PENGELOLAAN DATA DAN INFORMASI DI PROGRAM STUDI PPG FKIP",
            "UNIVERSITAS LAMPUNG"
        ]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Remove manual page breaks
        for run in p.runs:
            br_elements = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')
            for br in br_elements:
                if br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
                    run._element.remove(br)

    print("Step 2: Fix Italics in Tables")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.font.italic is None:
                            run.font.italic = False

    print("Step 3: Replace specific tables with narratives")
    body = doc._body._body
    elements = list(body)
    
    tables_to_remove = []
    
    for i, el in enumerate(elements):
        if el.tag.endswith('p'):
            p_text = "".join(node.text for node in el.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t') if node.text)
            
            if p_text.startswith("Tabel 4.2") or p_text.startswith("Tabel 4.3") or p_text.startswith("Tabel 4.4"):
                tables_to_remove.append(el)
                if i + 1 < len(elements) and elements[i+1].tag.endswith('tbl'):
                    tables_to_remove.append(elements[i+1])
                    if p_text.startswith("Tabel 4.2"): # Only insert once for the series
                        insert_paragraph_after(elements[i+1], "Berdasarkan pengujian, jabaran detail hasil pengujian dari 9 validator ahli (3 Ahli Materi, 3 Ahli Media, 3 Ahli Sistem) dilakukan secara mendalam per individu ahli (Margaretha, Eko, Putut, Rafiqa, Daniel, Dwi, Ghea, Radinal, Rahmad). Hasil menunjukkan skor kuantitatif validasi Materi mencapai 95,45%, Media 93,33%, dan Sistem 87,58%. Keputusan kelayakan beserta komentar kualitatif dari ahli digunakan sebagai dasar tindak lanjut perbaikan riil yang dilakukan peneliti pada sistem.")
                        
            elif p_text.startswith("Tabel 4.5") or p_text.startswith("Tabel 4.6"):
                tables_to_remove.append(el)
                if i + 1 < len(elements) and elements[i+1].tag.endswith('tbl'):
                    tables_to_remove.append(elements[i+1])
                    if p_text.startswith("Tabel 4.5"):
                        insert_paragraph_after(elements[i+1], "Tabel rekapitulasi data kepraktisan dapat dilihat pada bagian Lampiran. Dari 18 responden yang terdiri atas mahasiswa, admin prodi, dan staf ULT, diperoleh rerata skor kuesioner kepraktisan sebesar 55,28 dari total 60,00 atau mencapai 92,13% yang masuk dalam kategori Sangat Praktis. Berdasarkan distribusi kategori Bagian F kuesioner, mayoritas responden menyatakan sistem sangat praktis. Narasi tanggapan kualitatif dari 6 responden (Anisa, Lisa, Riswan, Nabila, Nur, Amrul) turut digunakan sebagai acuan tindak lanjut perbaikan sistem yang telah dikerjakan secara tuntas oleh peneliti.")

    for el in tables_to_remove:
        el.getparent().remove(el)

    doc.save(file_path)
    print(f"Saved to {file_path}")

if __name__ == '__main__':
    src = r"c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx"
    tmp_out = r"c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\tmp_clean.docx"
    shutil.copy2(src, tmp_out) # Backup
    
    process_docx(src)
    
    highlight_src = r"c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx"
    shutil.copy2(highlight_src, highlight_src + ".bak") # Backup
    process_docx(highlight_src)
