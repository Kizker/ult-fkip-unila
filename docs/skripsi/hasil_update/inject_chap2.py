import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

def inject_chap2(filename, output_filename):
    doc = docx.Document(filename)
    
    # Find insertion point
    target_p = None
    for p in doc.paragraphs:
        if p.style.name.startswith('Heading 2') and "Keterkaitan dengan Bidang Pendidikan" in p.text:
            target_p = p
            break
            
    if not target_p:
        print(f"Target paragraph not found in {filename}!")
        return

    # Helper to add paragraph before
    def insert_p(text, style='Normal', align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False):
        new_p = target_p.insert_paragraph_before(text, style)
        if style == 'Normal':
            new_p.alignment = align
            if bold:
                new_p.runs[0].bold = True
        return new_p

    # Insert Heading 2
    insert_p("Analisis Kurikulum Pendidikan Teknologi Informasi", style='Heading 2')
    
    # Insert Para 1
    insert_p("Pengembangan Website Unit Layanan Terpadu (ULT) di Fakultas Keguruan dan Ilmu Pendidikan Universitas Lampung ini dianalisis relevansinya terhadap struktur kompetensi Kurikulum Program Studi Pendidikan Teknologi Informasi. Analisis kurikulum dieksekusi berdasarkan penjabaran Capaian Pembelajaran Lulusan (CPL) yang berorientasi pada profil keahlian pengembang sistem informasi institusional. Analisis tersebut bertujuan untuk membuktikan kesesuaian antara kemampuan merancang bangun arsitektur perangkat lunak dengan standar penyelesaian masalah birokrasi kampus. Rincian penjabaran pemenuhan target kurikulum penciptaan rekayasa teknologi ini secara komprehensif dideskripsikan pada tabel berikut.")
    
    # Spacer
    insert_p("", style='Normal')
    
    # Caption Table 1
    caption1 = insert_p("Tabel 1. Analisis Kurikulum Pengembangan Sistem", style='Normal', align=WD_ALIGN_PARAGRAPH.LEFT)
    for r in caption1.runs:
        r.bold = True
        
    # Table 1
    table1 = doc.add_table(rows=2, cols=2)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Capaian Pembelajaran'
    hdr_cells[1].text = 'Tujuan Pembelajaran'
    for c in hdr_cells:
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True

    row_cells = table1.rows[1].cells
    row_cells[0].text = "Pada akhir pengerjaan tugas akhir, mahasiswa mampu mengimplementasikan pengetahuan rekayasa perangkat lunak dan arsitektur basis data, menganalisis kebutuhan tata kelola sistem institusi, menunjukkan profesionalitas pengembangan teknologi, serta merancang platform layanan akademik digital secara komprehensif dan terpadu."
    row_cells[1].text = "Mahasiswa mampu:\n1. Menganalisis masalah birokrasi dan merumuskan spesifikasi arsitektur antarmuka sistem (C4).\n2. Menerapkan kerangka kerja peranti lunak spesifik seperti Laravel dan Tailwind CSS (C3).\n3. Merancang integrasi basis data relasional beserta modul pengamanan privasi berkas (C5).\n4. Menunjukkan kedisiplinan dan ketelitian tanggung jawab dalam menyusun logika kode pemrograman (A4).\n5. Memproduksi sistem perakitan dokumen elektronik mandiri berbasis algoritma OpenXML (C6)."
    
    for c in row_cells:
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    target_p._p.addprevious(table1._tbl)
    
    # Spacer
    insert_p("", style='Normal')
    
    # Caption Table 2
    caption2 = insert_p("Tabel 2. Pemetaan Kedalaman dan Keluasan Materi Sistem", style='Normal', align=WD_ALIGN_PARAGRAPH.LEFT)
    for r in caption2.runs:
        r.bold = True
        
    # Table 2
    table2 = doc.add_table(rows=5, cols=2)
    table2.style = 'Table Grid'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Kedalaman Materi'
    hdr_cells2[1].text = 'Keluasan Materi'
    for c in hdr_cells2:
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                
    data = [
        ("Arsitektur Antarmuka (UI/UX)", "1. Desain tata letak responsif menggunakan TailwindCSS.\n2. Pemenuhan kriteria aksesibilitas seluler via Progressive Web App (PWA).\n3. Pemisahan dasbor operasional spesifik untuk entitas Mahasiswa, Staf, dan Pimpinan."),
        ("Manajemen Basis Data & Keamanan", "1. Implementasi sistem basis data relasional MySQL.\n2. Proteksi tautan unduhan dan enkripsi sistem fail lokal.\n3. Modul penyaringan ancaman injeksi skrip berbahaya lintas situs (Cross-Site Scripting)."),
        ("Mesin Perakitan Dokumen (Document Assembly)", "1. Integrasi parser konversi HTML-to-OpenXML terpusat.\n2. Ekstraksi otomatis variabel profil pendaftar ke dalam templat dokumen Word korporat.\n3. Pembangkitan kode nomor identifikasi surat persetujuan elektronik."),
        ("Pelacakan dan Matriks Riwayat (Audit Trail)", "1. Pencatatan log aktivitas tanggal waktu pemrosesan dokumen harian petugas.\n2. Visualisasi indikator tahapan alur persetujuan administratif berjenjang.\n3. Rekapitulasi pangkalan evaluasi durasi penanganan layanan bagi pimpinan.")
    ]
    
    for i, (kedalaman, keluasan) in enumerate(data):
        cells = table2.rows[i+1].cells
        cells[0].text = kedalaman
        cells[1].text = keluasan
        for c in cells:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
    target_p._p.addprevious(table2._tbl)
    
    # Spacer
    insert_p("", style='Normal')
    
    # Para 2
    insert_p("Pemetaan struktur matriks analisis kurikulum tersebut secara langsung menjustifikasi relevansi antara spesifikasi kompetensi teknis mahasiswa dengan tuntutan penciptaan karya cipta teknologi terapan. Penyelarasan antara kewajiban pemenuhan pencapaian tujuan pembelajaran dengan proses rancang bangun arsitektur sistem terbukti mampu melegitimasi kedalaman penguasaan keilmuan rekayasa perangkat lunak. Kerangka pemetaan korelasi ini secara akademis memastikan bahwa produk teknologi layanan administrasi yang dihasilkan telah memenuhi kualifikasi standar keahlian program studi pendidikan teknologi informasi modern.")

    # Spacer
    insert_p("", style='Normal')

    doc.save(output_filename)
    print(f"Successfully saved {output_filename}")

inject_chap2("001_Skripsi_Andricha Dea Mitra_Clean.docx", "001_Skripsi_Andricha Dea Mitra_Clean.docx")
inject_chap2("001_Skripsi_Andricha Dea Mitra_Highlighted.docx", "001_Skripsi_Andricha Dea Mitra_Highlighted.docx")
