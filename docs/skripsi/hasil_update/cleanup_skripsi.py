import docx
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Inches

def cleanup_document(doc_path):
    doc = docx.Document(doc_path)
    start_del = False
    to_delete = []

    # 1. Remove ghost paragraphs and ghost table
    for child in doc._body._body:
        if isinstance(child, CT_P):
            p = Paragraph(child, doc)
            text = p.text.strip()
            if text == '3.5 Subjek Penelitian dan Sampling':
                start_del = True
                continue
            if start_del and text.startswith('Populasi dalam penelitian ini'):
                start_del = False
                # break but we want to continue checking other things if needed
            if start_del:
                to_delete.append(child)
        elif isinstance(child, CT_Tbl):
            if start_del:
                to_delete.append(child)

    for el in to_delete:
        el.getparent().remove(el)
        
    # 2. Fix the tables (autofit and widths)
    for table in doc.tables:
        try:
            row_text = ' | '.join(cell.text.strip()[:10] for cell in table.rows[0].cells)
            
            # Gantt Chart (Tabel 3.1)
            if 'Jan | Feb' in row_text:
                table.autofit = False
                widths = [0.4, 2.5, 0.5, 0.5, 0.5, 0.5, 0.5, 0.5]
                for row in table.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = Inches(width)
                        
            # Subjek Tabel (Tabel 3.2)
            elif 'Kelompok S | Kategori S' in row_text:
                table.autofit = False
                widths = [0.4, 2.0, 1.5, 0.6, 2.5]
                for row in table.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = Inches(width)
        except Exception:
            pass

    # 3. Add explicit bullet points to the list paragraphs if they don't have them
    for p in doc.paragraphs:
        if p.style.name == 'List Paragraph':
            text = p.text.strip()
            if text.startswith('Ahli Materi') or text.startswith('Ahli Media') or text.startswith('Ahli Sistem') or \
               text.startswith('Mahasiswa (12') or text.startswith('Staf ULT') or text.startswith('Administrator'):
                if not text.startswith('•'):
                    p.text = '• ' + text
                    
    doc.save(doc_path)
    print(f"Successfully cleaned up {doc_path}")

if __name__ == '__main__':
    cleanup_document(r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx')
    cleanup_document(r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx')
