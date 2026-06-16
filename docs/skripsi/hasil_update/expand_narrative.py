import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX

files = [
    "c:/laragon/www/ult-fkip-unila/docs/skripsi/hasil_update/001_Skripsi_Andricha Dea Mitra_Clean.docx",
    "c:/laragon/www/ult-fkip-unila/docs/skripsi/hasil_update/001_Skripsi_Andricha Dea Mitra_Highlighted.docx"
]

narasi_5 = [
    "Diagram arsitektur sistem pada Gambar 5 mengadopsi model berlapis (layered architecture) yang membagi komponen sistem secara vertikal menjadi lima lapisan utama. Pembagian ini bertujuan untuk memisahkan tanggung jawab masing-masing komponen, mulai dari antarmuka interaksi pengguna hingga pada lapisan mesin perakit dokumen tingkat rendah, sehingga menghasilkan sistem yang loosely coupled, mudah dipelihara, dan memiliki tingkat keamanan yang terisolasi dengan baik.",
    "Secara visual, lapisan teratas adalah Client & Interface Layer yang dirancang agar kompatibel dengan teknologi Progressive Web App (PWA). Lapisan ini menaungi empat portal utama: Public Portal sebagai sarana informasi umum, Student Portal yang menjadi dasbor utama mahasiswa untuk mengajukan permohonan dokumen, Staff/Admin Portal untuk manajemen operasional oleh staf, dan Signer Portal yang didedikasikan khusus untuk pejabat dalam memberikan persetujuan elektronik. Pengguna akan berinteraksi langsung dengan antarmuka ini sebelum permintaan diteruskan ke sistem.",
    "Lapisan kedua, Security & Routing Gatekeeper, bertindak sebagai perisai pelindung utama sistem. Lapisan ini memuat Laravel Router yang dikonfigurasi dengan Content Security Policy (CSP) untuk mencegah eksekusi skrip berbahaya. Selain itu, terdapat Spatie RBAC (Role-Based Access Control) sebagai middleware otorisasi yang secara dinamis memvalidasi hak akses pengguna berdasarkan peran mereka. Fitur Anti-IDOR juga disematkan untuk memproteksi akses berkas agar hanya dapat diunduh oleh pemilik sah, serta HtmlSanitizer yang secara ketat membersihkan setiap input berformat teks kaya (WYSIWYG) dari potensi serangan Cross-Site Scripting (XSS).",
    "Pada lapisan ketiga, Laravel 12 Application Core memegang kendali penuh atas logika bisnis aplikasi. Controllers Core, khususnya RequestAdminController, menjadi pusat pengatur lalu lintas data permohonan persuratan akademik. Komponen ini berkoordinasi langsung dengan DocumentAssembler sebagai fasilitator orkestrasi perakitan dokumen, serta HtmlToOpenXMLParser yang bertugas secara spesifik untuk membedah dan menerjemahkan tag-tag HTML dari editor teks ke dalam format XML mentah tanpa merusak gaya bawaan.",
    "Lapisan keempat merupakan Database & Storage Layer yang menangani persistensi data. Di dalamnya, MySQL Database berfungsi untuk menyimpan data operasional yang terstruktur seperti formulir pengguna, profil, dan log rekam jejak (audit trail). Secara paralel, Private Storage Disk bertindak sebagai ruang penyimpanan terenkripsi yang sepenuhnya terisolasi dari akses publik, dikhususkan untuk mengamankan fisik berkas lampiran maupun dokumen hasil perakitan berformat .docx dan PDF.",
    "Lapisan paling bawah adalah Document Assembly Engine yang ditenagai oleh pustaka PHPOffice. Lapisan ini bertanggung jawab mengeksekusi proses perakitan tingkat rendah. Word Templates berisi placeholder statis digunakan sebagai landasan surat. Selanjutnya, modul PhpWord & ZipArchive bekerja memanipulasi struktur XML mentah secara langsung di dalam arsip dokumen, menginjeksi data mahasiswa yang telah diproses. Terakhir, PDF Converter menghasilkan luaran dokumen dengan format siap cetak yang dapat diunduh langsung melalui antarmuka."
]

narasi_6 = [
    "Diagram flowchart pada Gambar 6 menampilkan visualisasi alur siklus hidup permohonan dokumen secara berurutan (sequential) dari tahap inisiasi hingga dokumen siap diunduh. Diagram ini disusun vertikal dari titik terminal 'MULAI' hingga 'SELESAI', menggunakan simbol persegi panjang untuk merepresentasikan blok operasi atau tindakan, dan simbol belah ketupat untuk menggambarkan titik pengambilan keputusan (decision point) krusial di dalam alur birokrasi.",
    "Fungsi dan interaksi sistem dimulai dengan blok 'Pengisian Form Dinamis' oleh mahasiswa melalui Student Portal. Pada tahap ini, mahasiswa menginputkan data permohonan menggunakan isian teks kaya (Rich Text HTML) yang difasilitasi oleh antarmuka Tiptap Editor. Setelah data dikirim, alur bergerak menuju blok 'Pengecekan Pertama Berkas' yang dilakukan oleh Admin Jurusan. Di sinilah letak titik pengambilan keputusan ('Layak?'); jika berkas atau data dinilai kurang lengkap, alur akan terputus dan berbelok ke status 'Ditolak/Draft' untuk direvisi kembali oleh mahasiswa, atau proses langsung dihentikan menuju terminal penolakan.",
    "Apabila permohonan dinyatakan valid dan lengkap, sistem akan mengambil alih secara otomatis pada blok 'Pemberian Nomor Surat Otomatis'. Modul ini menerbitkan kode unik dan nomor surat resmi yang menyesuaikan dengan standar format penomoran di masing-masing program studi. Alur kemudian berlanjut ke blok 'Review & Validasi ULT', di mana staf unit layanan terpadu meninjau kembali keabsahan dan keakuratan seluruh data secara digital sebelum dinaikkan ke level pimpinan fakultas.",
    "Tahap birokrasi tertinggi berada pada blok 'Verifikasi & TTD Elektronik'. Pejabat yang berwenang menggunakan Signer Portal untuk memberikan persetujuan final secara substantif sekaligus membubuhkan tanda tangan elektronik. Persetujuan ini secara instan memicu blok komputasi inti, yaitu 'Perakitan Dokumen Otomatis'. Pada tahap ini, mesin perakit (Assembly Core) dan parser HTML-to-OpenXML bekerja di balik layar menerjemahkan tag-tag pemformatan Tiptap, mengkloning atribut font bawaan (run asli), lalu menyuntikkan data ke dalam templat dokumen (.docx) sehingga menghasilkan keluaran yang presisi tanpa distorsi visual.",
    "Dua blok terakhir mendeskripsikan proses pasca-perakitan. Pada blok 'Penyimpanan Berkas & Audit Trail', sistem secara terprogram memindahkan berkas akhir ke dalam direktori Private Disk yang terisolasi, serta mencatat secara permanen setiap aksi yang telah dilakukan ke dalam log database sebagai bukti jejak audit digital. Akhirnya, pada blok 'Unduh Berkas Privat', mahasiswa menerima notifikasi bahwa dokumen telah selesai dan dapat mengunduhnya dengan aman berkat perlindungan middleware anti-IDOR yang memastikan integritas kepemilikan berkas."
]

narasi_7 = [
    "Diagram Entity-Relationship (ERD) pada Gambar 7 memvisualisasikan arsitektur dan relasi antara delapan entitas utama di dalam basis data MySQL yang menopang keseluruhan sistem. Setiap entitas direpresentasikan sebagai sebuah tabel persegi panjang dengan tajuk gelap, di mana baris-baris di bawahnya menjabarkan daftar atribut (kolom) beserta spesifikasi tipe datanya, seperti BIGINT, VARCHAR, TEXT, BOOLEAN, maupun TIMESTAMP. Garis-garis solid yang menghubungkan tabel-tabel ini menunjukkan struktur relasi fisik berupa foreign key constraints.",
    "Entitas users, roles, beserta tabel pivot model_has_roles membentuk sebuah ekosistem manajemen identitas yang solid untuk mendukung fungsionalitas Role-Based Access Control (RBAC). Entitas users menyimpan kredensial pengguna, sementara roles mendefinisikan level akses, yang kemudian dipetakan secara dinamis melalui tabel pivot. Tabel services berperan sebagai katalog sentral yang menampung informasi jenis layanan persuratan beserta jalur referensi lokasi templat fisik dokumen (template_path).",
    "Sebagai pusat operasional transaksional, entitas requests bertindak untuk merekam setiap permohonan yang masuk, menyimpan status alur, kode pelacakan (tracking_code), dan menerbitkan nomor surat (issue_number). Mengingat setiap layanan memiliki kebutuhan isian yang berbeda-beda, sistem mengimplementasikan entitas request_inputs berformat pasangan kunci-nilai (field_name dan field_value) yang berelasi One-to-Many dengan requests. Mekanisme dinamis ini memastikan skema database tidak perlu direstrukturisasi setiap kali terdapat penambahan bidang isian baru. Selain itu, entitas request_documents secara fungsional dikhususkan untuk menangani metadata berkas, meliputi jalur file fisik, tipe, dan ukuran file lampiran.",
    "Alur kerja relasional ERD ini menjamin integritas data dari awal hingga akhir. Saat mahasiswa membuat permohonan, rekam induk (parent record) tercipta di tabel requests dengan status awal. Data borang dinamis terhubung masuk berantai ke request_inputs, dan seluruh lampiran diindeks di request_documents. Segala aktivitas perubahan status yang dipicu oleh intervensi staf maupun pejabat akan selalu memicu perekaman baris data baru di entitas audit_trails. Entitas audit log ini menyimpan tindakan riwayat aksi (action) beserta alamat IP aktor bersangkutan, berfungsi ganda sebagai lapisan pengamanan historis untuk memastikan akuntabilitas operasional sistem administrasi."
]

narasi_8 = [
    "Diagram sequence pada Gambar 8 menyajikan visualisasi interaksi dinamis berorientasi waktu (time-oriented) yang mendeskripsikan secara spesifik bagaimana objek-objek sistem berkolaborasi dalam menyelesaikan proses perakitan dokumen tingkat rendah. Antarmuka diagram tersusun atas elemen lifeline vertikal yang merepresentasikan kelima entitas atau aktor yang berpartisipasi, direpresentasikan oleh kotak hitam di posisi teratas. Garis-garis panah horizontal bersiklus menggambarkan urutan spesifik pengiriman pesan (message passing) atau pemanggilan metode (method calls) yang mengalir kronologis dari atas ke bawah.",
    "Komponen yang terlibat secara menyeluruh meliputi aktor Staf ULT / User sebagai inisiator pemicu, RequestController sebagai antarmuka pengatur rute permintaan pada level kerangka kerja aplikasi, DocumentAssemblerService yang memegang mandat utama orkestrasi perakitan dan ekstraksi templat, kelas utilitas independen HtmlToOpenXMLParser sebagai mesin murni penerjemah tag HTML kaya ke dalam blok XML, serta Private Storage sebagai representasi infrastruktur tujuan penyimpanan fisik berkas keluaran.",
    "Skenario perakitan dokumen dimulai pada siklus langkah 1, saat staf ULT menekan tombol perakitan dari antarmuka dasbor yang mengirimkan sinyal clickAssemble(requestId) menuju RequestController. Pada langkah 2, pengendali mendelegasikan perintah komputasi transaksional ini seutuhnya ke kelas DocumentAssemblerService lewat instruksi assembleDocument. Objek service ini kemudian mempersiapkan lingkungan kerja di langkah 3 dengan memuat fail templat dokumen mentah dan mengekstraksi formulir berisikan data dari basis data (loadTemplateAndFormData). Pada titik delegasi kritis di langkah 4, service mengoper beban terjemahan teks kaya (WYSIWYG) seutuhnya kepada HtmlToOpenXMLParser dengan mengeksekusi metode writeHtmlToWordRun.",
    "HtmlToOpenXMLParser lantas mengambil alih proses eksekusi mendalam di langkah 5 secara rekursif (traverseHtmlAndInsertRuns) untuk membedah satu demi satu elemen DOM HTML, seperti tag paragraf, tebal, maupun miring. Guna mempertahankan nilai estetik dan gaya visual orisinal templat Word, parser tidak memformat teks dari nol melainkan mengeksekusi metode cloneRunPropertiesAndFormat (langkah 6) untuk menduplikasi atribut font pewarisan dari placeholder aslinya, menghindari degradasi dokumen. Setelah seluruh format OpenXML beserta atribut silsilahnya berhasil dirajut (langkah 7), himpunan instruksi simpul dikembalikan kepada DocumentAssemblerService.",
    "Pada siklus penyelesaian di langkah 8, DocumentAssemblerService merajut kembali seluruh struktur berkas XML yang telah dimodifikasi menjadi satu kesatuan arsip biner ZIP berbentuk fail .docx yang utuh (compileAndZipArchive). Berkas fisik ini selanjutnya ditransfer ke objek Private Storage pada langkah 9 melalui perintah putFileInPrivateDisk untuk dikarantina dan diamankan dari akses anonim. Siklus diakhiri dengan mekanisme sinyal balasan (return value) yang merambat mundur; status keberhasilan persisten dari penyimpanan di langkah 10, diteruskan sebagai respons HTTP sukses menuju kontroler (langkah 11), dan bermuara kembali pada layar peramban aktor utama sebagai pemanggilan fungsi redirectWithSuccessToast (langkah 12), mengkonfirmasi kepada pengguna bahwa rantai proses perakitan dokumen digital sistem telah paripurna."
]

def delete_p(p):
    if p._element.getparent() is not None:
        p._element.getparent().remove(p._element)
    p._p = p._element = None

def process_doc(filepath, is_highlighted):
    doc = docx.Document(filepath)
    paragraphs = doc.paragraphs
    
    # Flags to start deleting text after Gambar 5 & 6
    del_5 = False
    del_6 = False
    
    for p in paragraphs:
        text = p.text.strip()
        
        if text.startswith("Gambar  5. Diagram"):
            del_5 = True
            continue
        if text.startswith("Gambar  6. Flowchart"):
            del_6 = True
            continue
            
        if text.startswith("Perancangan alur kerja persuratan") and del_5:
            del_5 = False
            
        if text.startswith("Perancangan Struktur Data") and del_6:
            del_6 = False
            
        if del_5 or del_6:
            if text:
                delete_p(p)
                continue
                
        if text.startswith("Aktivitas perancangan basis data memetakan") or \
           text.startswith("Diagram ERD pada Gambar 7") or \
           text.startswith("Katalog jenis layanan persuratan akademik"):
            delete_p(p)
            
        if text.startswith("Diagram sequence pada Gambar 8") or \
           text.startswith("Kelas layanan perakit dokumen mengawali") or \
           text.startswith("Perancangan interaksi dinamis memodelkan urutan pesan"):
            delete_p(p)
            
        # Clean empty paragraphs between sections
        # if not text and not p.runs:
        #     delete_p(p)

    def insert_after(caption_text, naras_list):
        # Refresh paragraphs list
        current_paragraphs = doc.paragraphs
        for i, p in enumerate(current_paragraphs):
            if p.text.strip().startswith(caption_text):
                # find the next valid paragraph to insert before
                if i + 1 < len(current_paragraphs):
                    target_p = current_paragraphs[i+1]
                    for text in naras_list:
                        new_p = target_p.insert_paragraph_before("")
                        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        new_p.paragraph_format.first_line_indent = Cm(1.27)
                        new_p.paragraph_format.line_spacing = 1.5
                        new_p.paragraph_format.space_after = Pt(0)
                        new_p.paragraph_format.space_before = Pt(0)
                        
                        run = new_p.add_run(text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        
                        if is_highlighted:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    break

    insert_after("Gambar  5. Diagram", narasi_5)
    insert_after("Gambar  6. Flowchart", narasi_6)
    insert_after("Gambar  7. Diagram Hubungan", narasi_7)
    insert_after("Gambar  8. Alur Perakitan", narasi_8)

    doc.save(filepath)

process_doc(files[0], False)
process_doc(files[1], True)
print("Penulisan narasi selesai.")
