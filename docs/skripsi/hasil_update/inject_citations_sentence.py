import docx
from docx.enum.text import WD_COLOR_INDEX

def modify_citations_sentence(doc_path, is_highlighted=False):
    doc = docx.Document(doc_path)
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        
        # Paragraph 644 equivalent
        if 'kualifikasi validator ahli ditetapkan' in text and 'ahli materi = 3' in text:
            p.text = '' # Clear
            
            base_text1 = "Jumlah dan kualifikasi validator ahli ditetapkan sebagai berikut: ahli materi = 3 orang; ahli media = 3 orang; ahli sistem = 3 orang. Pemilihan validator dilakukan secara purposive berdasarkan relevansi keahlian dengan objek yang dinilai, pengalaman atau kompetensi pada bidangnya, serta kesediaan memberikan penilaian dan saran perbaikan terhadap produk. Total validator ahli adalah 9 orang. "
            p.add_run(base_text1)
            
            added_text = "Menurut Nieveen (1999), penggunaan keahlian multi-disiplin ini sejalan dengan standar evaluasi kelayakan sistem yang mengharuskan produk dievaluasi dari sisi relevance (materi) dan consistency (sistem/media)."
            added_run = p.add_run(added_text)
            if is_highlighted:
                added_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                
        # Paragraph 647 equivalent
        elif 'Jumlah responden uji coba ditetapkan' in text and 'mahasiswa = 12' in text:
            p.text = '' # Clear
            
            base_text1 = "Jumlah responden uji coba ditetapkan sebagai berikut: mahasiswa = 12 orang; staf ULT = 3 orang; Administrator jurusan = 3 orang. "
            p.add_run(base_text1)
            
            added_text1 = "Menurut Faulkner (2003) dan Weinger et al. (2010), ukuran sampel sebanyak 15-20 pengguna sudah terbukti memadai dan mampu mengungkap lebih dari 95% masalah usability utama pada sebuah perangkat lunak. "
            added_run1 = p.add_run(added_text1)
            if is_highlighted:
                added_run1.font.highlight_color = WD_COLOR_INDEX.YELLOW
                
            base_text2 = "Teknik pemilihan responden menggunakan purposive sampling. "
            p.add_run(base_text2)
            
            added_text2 = "Menurut Sugiyono (2013), purposive sampling adalah teknik penentuan sampel dengan pertimbangan tertentu agar sampel yang dipilih relevan dengan kebutuhan pengujian. "
            added_run2 = p.add_run(added_text2)
            if is_highlighted:
                added_run2.font.highlight_color = WD_COLOR_INDEX.YELLOW
                
            base_text3 = "Kriteria inklusi responden:"
            p.add_run(base_text3)
            
    doc.save(doc_path)
    print(f"Successfully modified citations to start of sentences in {doc_path}")

if __name__ == '__main__':
    modify_citations_sentence(r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx', is_highlighted=False)
    modify_citations_sentence(r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx', is_highlighted=True)
