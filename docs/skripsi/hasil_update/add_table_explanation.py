import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX

def insert_text_after_table(file_path, is_highlighted):
    doc = docx.Document(file_path)
    
    target_p = None
    for p in doc.paragraphs:
        if p.text.startswith('d. Initial Design') or p.text.startswith('Langkah perancangan awal (Initial Design)'):
            target_p = p
            break
            
    if not target_p:
        print(f"Target paragraph not found in {file_path}")
        return

    text = "Berdasarkan rincian Tabel 4.1 di atas, format antarmuka pada sistem dirancang dengan pendekatan modular yang memisahkan tanggung jawab (separation of concerns) ke dalam lingkungan visual yang spesifik. Pemisahan antarmuka menjadi lima portal berbeda ini tidak hanya menyederhanakan navigasi pengguna sesuai dengan wewenang (role), melainkan juga meminimalisasi distraksi visual dan kelebihan beban informasi (information overload). Secara praktikal, mahasiswa dapat berfokus murni pada pelacakan dokumen melalui Dasbor Mahasiswa, staf ULT dimudahkan dengan tampilan tabel data operasional di Dasbor Admin, sementara pejabat struktural disuguhkan tampilan minimalis dan ringkas di Signer Portal guna mempercepat birokrasi persetujuan dokumen. Struktur tata letak ini dikonstruksi secara adaptif agar optimal digunakan di berbagai ukuran perangkat, menjamin kontinuitas pelayanan administrasi yang tidak terputus."

    new_p = target_p.insert_paragraph_before()
    new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_p.paragraph_format.space_before = Pt(0)
    new_p.paragraph_format.space_after = Pt(12)
    new_p.paragraph_format.line_spacing = 1.5

    run = new_p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if is_highlighted:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.save(file_path)
    print(f"Added explanation to {file_path}")

files_to_process = [
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx',
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
]

for f in files_to_process:
    insert_text_after_table(f, "Highlighted" in f)
