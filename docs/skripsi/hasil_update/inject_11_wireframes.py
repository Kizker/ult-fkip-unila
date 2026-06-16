import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

docs = [
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx',
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
]

items = [
    {
        "img": "11_wireframe_beranda_balsamiq.png",
        "cap": "Gambar 4.5. Wireframe Beranda Utama (Landing Page).",
        "desc": "Pada tampilan beranda utama sistem (Public Portal), wireframe disusun menggunakan tata letak berbasis kolom tunggal terpusat (centered layout) guna memusatkan perhatian pengunjung pada elemen inti. Area header memuat logo institusi Fakultas Keguruan dan Ilmu Pendidikan bersandingan dengan tautan navigasi utama seperti 'Beranda', 'Katalog Layanan', 'Berita', dan tombol 'Masuk'. Tepat di bawahnya, terdapat pahlawan visual (hero section) besar yang menampung slogan layanan terpadu serta tombol aksi (call-to-action) untuk menjelajahi layanan. Bergerak ke bawah, tata letak menampilkan grid asimetris yang berisi daftar layanan terpopuler. Bagian kaki (footer) dialokasikan untuk informasi kontak institusi dan hak cipta."
    },
    {
        "img": "12_wireframe_katalog_balsamiq.png",
        "cap": "Gambar 4.6. Wireframe Katalog Layanan (Public Portal).",
        "desc": "Halaman Katalog Layanan dirancang untuk menyajikan informasi menyeluruh terkait ragam permohonan persuratan yang difasilitasi oleh fakultas. Susunan antarmuka mengadopsi pola grid kartu vertikal yang rapi, di mana setiap kartu mendeskripsikan secara spesifik satu jenis layanan, seperti 'Surat Aktif Kuliah' atau 'Cuti Akademik', lengkap beserta rincian lampiran prasyaratnya. Desain ini bertujuan untuk meminimalkan kebingungan administratif mahasiswa sebelum mengajukan dokumen."
    },
    {
        "img": "13_wireframe_berita_balsamiq.png",
        "cap": "Gambar 4.7. Wireframe Berita & Pengumuman (Public Portal).",
        "desc": "Halaman Berita dan Pengumuman mengusung tata letak daftar blok tunggal (list view) yang memprioritaskan pembaruan informasi terkini dari pihak dekanat maupun tata usaha. Setiap entri berita direpresentasikan dengan kotak pembungkus citra digital sederhana (thumbnail placeholder) berserta judul, cuplikan isi, dan tanggal publikasi. Modul ini dilengkapi dengan sistem penomoran halaman (pagination) di bagian bawah guna memperlancar penjelajahan arsip publikasi lawas tanpa membebani pemuatan halaman (page load)."
    },
    {
        "img": "14_wireframe_login_balsamiq.png",
        "cap": "Gambar 4.8. Wireframe Login & Register (Authentication System).",
        "desc": "Antarmuka sistem autentikasi diwujudkan dalam bentuk kotak kartu terpusat (centered card) di tengah layar. Desain ini secara sengaja melepaskan diri dari elemen navigasi lain agar fokus visual pengguna sepenuhnya tertuju pada pengisian kredensial. Elemen isian meliputi alamat surel dan kata sandi, kotak centang 'Ingat Saya', tombol aksi utama untuk masuk, serta tautan pengalihan jika pengguna lupa kata sandi. Pendekatan minimalis ini meningkatkan laju penyelesaian autentikasi tanpa hambatan kognitif."
    },
    {
        "img": "15_wireframe_student_balsamiq.png",
        "cap": "Gambar 4.9. Wireframe Beranda Mahasiswa (Student Portal).",
        "desc": "Pada portal dasbor mahasiswa, tata letak didesain menggunakan format dua kolom interaktif tanpa menu navigasi sisi vertikal (sidebar). Bagian atas menampilkan informasi profil ringkas dan tombol aksi utama untuk mengajukan layanan baru. Kolom utama di sebelah kiri difokuskan pada daftar aktivitas pengajuan permohonan terbaru secara mendetail, sedangkan kolom sekunder di sisi kanan (side panel) difungsikan sebagai agregasi visual status pengajuan dalam bentuk angka total dan kategori (menunggu, diproses, selesai)."
    },
    {
        "img": "16_wireframe_form_balsamiq.png",
        "cap": "Gambar 4.10. Wireframe Form Pengajuan Layanan (Student Portal).",
        "desc": "Halaman form pengajuan layanan mengadaptasi formulir masukan dinamis (dynamic inputs) berbasis tumpukan vertikal (stacked vertical layout). Mahasiswa akan melihat prasyarat data secara runtut dari atas ke bawah, mulai dari data teks pokok seperti Nomor Pokok Mahasiswa dan nomor telepon, hingga area interaktif khusus untuk mengunggah dokumen digital prasyarat berformat PDF atau gambar. Tombol konfirmasi di bagian bawah diformulasikan tegas untuk memvalidasi kelengkapan berkas sebelum dikirimkan."
    },
    {
        "img": "17_wireframe_timeline_balsamiq.png",
        "cap": "Gambar 4.11. Wireframe Riwayat & Pelacakan Dokumen (Timeline).",
        "desc": "Halaman riwayat dan pelacakan dirancang untuk menjamin prinsip transparansi birokrasi (birocratic transparency) kepada mahasiswa. Tata letak layar dibagi menjadi dua modul esensial; modul sebelah kiri memuat rincian data masukan serta pratinjau lampiran yang telah diajukan, sementara modul di sebelah kanan menampilkan visualisasi linimasa interaktif (auditable timeline). Visualisasi linimasa merunut pergerakan berkas dari status draf, persetujuan admin, penyematan tanda tangan pimpinan fakultas, hingga produk akhir diterbitkan."
    },
    {
        "img": "18_wireframe_admin_balsamiq.png",
        "cap": "Gambar 4.12. Wireframe Dasbor Utama Staff (Admin Portal).",
        "desc": "Antarmuka dasbor operasional bagi jajaran staf dan administrator ULT mengusung tata letak asimetris tiga bagian dengan penambahan menu navigasi tersemat vertikal (sidebar) di sisi kiri layar. Struktur hierarki ini memberikan fleksibilitas akses yang cepat ke berbagai modul administratif. Area konten utama mempertahankan pola dua kolom yang menonjolkan rincian antrian operasional harian yang membutuhkan peninjauan di kolom primer, dan indikator kinerja (KPI) status kumulatif di kolom sekunder."
    },
    {
        "img": "19_wireframe_review_balsamiq.png",
        "cap": "Gambar 4.13. Wireframe Detail Permohonan / Review (Admin Portal).",
        "desc": "Halaman peninjauan (review) oleh staf memuat sekumpulan panel analitik (analytic panels) yang menyajikan seluruh instrumen dokumen mahasiswa secara holistik tanpa harus mengunduh file satu persatu ke peramban lokal. Staf disediakan mekanisme umpan balik interaktif, format isian nomor surat fakultas secara terotomatisasi, serta kontrol untuk memutus status penolakan atau melanjutkan perakitan dokumen elektronik secara mandiri menggunakan protokol ekstraksi OpenXML."
    },
    {
        "img": "20_wireframe_template_balsamiq.png",
        "cap": "Gambar 4.14. Wireframe Manajemen Template (Admin Portal).",
        "desc": "Modul manajemen templat direpresentasikan melalui tabel variabel isian dinamis (dynamic input mapping table) serta area konfigurasi pengunggahan format master dokumen berekstensi Microsoft Word (.docx). Melalui halaman ini, administrator diizinkan menambahkan atau menghapus kolom variabel tambahan (seperti Nama Orang Tua, Alamat Lengkap, dll) serta menandai status kewajibannya, tanpa memerlukan modifikasi kode sumber (hardcode) pada arsitektur perangkat lunak dasarnya."
    },
    {
        "img": "21_wireframe_signer_balsamiq.png",
        "cap": "Gambar 4.15. Wireframe Antarmuka Verifikasi Pejabat (Signer Portal).",
        "desc": "Pada kotak masuk verifikasi pejabat (Signer Portal), antarmuka dititikberatkan secara eksklusif pada kemudahan eksekusi otoritas dan kecepatan pratinjau (fast-preview). Halaman menghilangkan elemen menu kompleks, digantikan dengan tumpukan tabel prioritas tugas persetujuan di sisi atas dan mesin pratinjau pembaca dokumen di bawahnya. Pejabat fakultas seperti Wakil Dekan dapat menginspeksi kelayakan narasi surat serta membubuhkan kriptografi tanda tangan elektronik (TTE) secara terintegrasi dalam satu bingkai kerja (single-frame workspace)."
    }
]

def update_doc(path):
    doc = docx.Document(path)
    # Find start and end
    start_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if 'Gambar 4.5. Wireframe Beranda Utama' in p.text:
            start_idx = i - 1  # the image paragraph
            break
            
    if start_idx == -1: 
        print("Could not find Gambar 4.5 in", path)
        return
    
    end_idx = -1
    for i in range(start_idx, len(doc.paragraphs)):
        if 'Perancangan Antarmuka Pengguna Akhir' in doc.paragraphs[i].text:
            end_idx = i
            break
            
    if end_idx != -1:
        # Clear the old paragraphs between start_idx and end_idx
        for i in range(start_idx, end_idx):
            doc.paragraphs[i].clear()
            
    target_p = doc.paragraphs[start_idx]
    
    # Insert new items
    for item in items:
        # Img
        p_img = target_p.insert_paragraph_before()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(12)
        p_img.paragraph_format.space_after = Pt(12)
        run = p_img.add_run()
        img_path = os.path.join(r'c:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram', item['img'])
        if os.path.exists(img_path):
            run.add_picture(img_path, width=Inches(5.5))
            
        # Caption
        p_cap = target_p.insert_paragraph_before(item['cap'])
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(0)
        p_cap.paragraph_format.space_after = Pt(12)
        r_cap = p_cap.add_run(item['cap'])
        r_cap.font.name = 'Times New Roman'
        r_cap.font.size = Pt(10)
            
        # Desc
        p_desc = target_p.insert_paragraph_before(item['desc'])
        p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_desc.paragraph_format.first_line_indent = Inches(0.5)
        p_desc.paragraph_format.line_spacing = 1.5
        r_desc = p_desc.add_run(item['desc'])
        r_desc.font.name = 'Times New Roman'
        r_desc.font.size = Pt(12)
            
    # Fix the next image number placeholder if it exists
    for p in doc.paragraphs:
        if '[NOMOR_GAMBAR]' in p.text and 'Desain Antarmuka Akhir' in p.text:
            p.text = '' # Clear the old text completely
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(12)
            r_cap2 = p.add_run('Gambar 4.16. Desain Antarmuka Akhir (UI/UX) Sistem ULT FKIP Unila.')
            r_cap2.font.name = 'Times New Roman'
            r_cap2.font.size = Pt(10)
            break

    doc.save(path)
    print("Saved", os.path.basename(path))

for d in docs:
    update_doc(d)
