import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color="D9EAD3"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill_color))
    tcPr.append(shading)

def add_highlight(run, is_highlighted):
    if is_highlighted:
        rPr = run._r.get_or_add_rPr()
        highlight = OxmlElement('w:highlight')
        highlight.set(qn('w:val'), 'yellow')
        rPr.append(highlight)

def process_doc(doc_path, is_highlighted=False):
    doc = docx.Document(doc_path)
    
    # 1. Find "Subjek Penelitian dan Sampling"
    target_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if 'Subjek Penelitian dan Sampling' in p.text and p.style.name.startswith('Heading'):
            target_idx = i
            break
            
    if target_idx == -1:
        print(f"Could not find 'Subjek Penelitian dan Sampling' in {doc_path}")
        return
        
    target_p = doc.paragraphs[target_idx]
    
    # 2. Find the start of junk to delete (after "Evaluation")
    start_del_idx = -1
    for i in range(target_idx - 1, -1, -1):
        p = doc.paragraphs[i]
        if 'Tahap Evaluation' in p.text or 'Evaluation (Evaluasi)' in p.text:
            start_del_idx = i + 1
            break
            
    if start_del_idx != -1:
        # Delete paragraphs between start_del_idx and target_idx
        for i in range(start_del_idx, target_idx):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)
            
    # Helper to insert
    def insert_p_before(text, style_name):
        new_p = target_p.insert_paragraph_before(text, style=style_name)
        if text:
            new_p.text = ''
            run = new_p.add_run(text)
            add_highlight(run, is_highlighted)
        return new_p
        
    # 3. Insert new section
    # Space before
    insert_p_before('', 'Normal')
    
    # Heading
    # The user wants "3.4 Tempat dan Waktu Penelitian", but Unila styles usually auto-number.
    # We'll just provide the text. If auto-numbering is on, it will add it.
    h = insert_p_before('3.4 Tempat dan Waktu Penelitian', 'Heading 2')
    
    # Narrative
    narasi = (
        "Penelitian ini dilaksanakan di Fakultas Keguruan dan Ilmu Pendidikan (FKIP) Universitas Lampung sebagai lokasi utama, "
        "dengan objek kajian berupa pengembangan website Unit Layanan Terpadu (ULT) yang ditujukan untuk mahasiswa, "
        "staf administrasi, dan Administrator jurusan. Pemilihan lokasi ini didasarkan pada permasalahan nyata, "
        "yaitu masih dominannya prosedur layanan manual yang berdampak pada keterlambatan, inefisiensi, dan rendahnya kepuasan pengguna. "
        "Selain itu, pra-penelitian juga dilakukan di SMK Negeri 9 Bandar Lampung pada Jurusan Sistem Informasi dan Jaringan (SIJA). "
        "Tujuannya adalah untuk memperoleh masukan tambahan mengenai relevansi pengembangan website dengan pembelajaran berbasis proyek "
        "pada mata pelajaran Dasar-dasar Pengembangan Perangkat Lunak dan Gim (PPLG). Dengan demikian, penelitian ini tidak hanya "
        "berkontribusi terhadap digitalisasi layanan administrasi perguruan tinggi, tetapi juga dapat dimanfaatkan sebagai sumber belajar "
        "kontekstual di pendidikan vokasi.\n\n"
        "Penelitian ini dilaksanakan melalui serangkaian tahapan yang terstruktur dan sistematis berdasarkan kerangka kerja pengembangan ADDIE. "
        "Mengacu pada riwayat histori pengembangan sistem, tahapan awal berupa Analysis dan Design dimulai pada akhir bulan Januari hingga bulan Maret 2026. "
        "Tahapan pembuatan perangkat lunak (Development) berlangsung secara intensif dari awal bulan Maret hingga awal bulan Mei 2026. "
        "Setelah sistem dinyatakan matang secara fungsional, uji coba terbatas (Implementation) dilaksanakan pada bulan Mei 2026, "
        "disusul oleh proses evaluasi hasil uji coba (Evaluation). Secara paralel, penyusunan draf laporan penulisan karya ilmiah "
        "dikerjakan bertahap mulai dari bulan Maret hingga bulan Juni 2026. Rincian penjadwalan pelaksanaan kegiatan penelitian ini disajikan pada Tabel 3.1 berikut."
    )
    insert_p_before(narasi, 'Paragraph')
    
    # Table Title
    insert_p_before('', 'Normal') # spacing
    tbl_title = insert_p_before('Tabel 3.1 Waktu Pelaksanaan Penelitian', 'Normal')
    tbl_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl_title.text = ''
    run = tbl_title.add_run('Tabel 3.1 Waktu Pelaksanaan Penelitian')
    run.font.bold = True
    add_highlight(run, is_highlighted)
    
    # Create Table
    table = doc.add_table(rows=7, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    target_p._p.addprevious(table._tbl)
    
    headers = ['No.', 'Kegiatan Penelitian', 'Jan', 'Feb', 'Mar', 'Apr', 'Mei', 'Jun']
    data = [
        ['1', 'Analysis', True, True, False, False, False, False],
        ['2', 'Design', False, True, True, False, False, False],
        ['3', 'Development', False, False, True, True, True, False],
        ['4', 'Implementation', False, False, False, False, True, False],
        ['5', 'Evaluation', False, False, False, False, True, True],
        ['6', 'Penulisan Karya Ilmiah', False, False, True, True, True, True]
    ]
    
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                add_highlight(r, is_highlighted)
                
    for row_idx, row_data in enumerate(data):
        for col_idx, val in enumerate(row_data):
            cell = table.cell(row_idx + 1, col_idx)
            if col_idx < 2:
                cell.text = str(val)
                for p in cell.paragraphs:
                    if col_idx == 0:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        add_highlight(r, is_highlighted)
                        if col_idx == 1 and val != 'Penulisan Karya Ilmiah':
                            r.font.italic = True
            else:
                if val:
                    set_cell_background(cell, "A9D08E") # light green
                
    widths = [0.4, 2.5, 0.5, 0.5, 0.5, 0.5, 0.5, 0.5]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)

    insert_p_before('', 'Normal')
    
    doc.save(doc_path)
    print(f"Successfully updated {doc_path}")

if __name__ == '__main__':
    clean_path = r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
    highlighted_path = r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
    
    process_doc(clean_path, is_highlighted=False)
    process_doc(highlighted_path, is_highlighted=True)
