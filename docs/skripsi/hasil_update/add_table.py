import docx
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX

def set_cell_background(cell, color_hex):
    # Set cell background color
    tcProperties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    tcProperties.append(shading)

def set_table_borders(table):
    # Set borders for the table using OpenXML
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_table_to_doc(file_path, is_highlighted):
    doc = docx.Document(file_path)
    
    # Find insertion point: the paragraph starting with 'd. Initial Design'
    target_p = None
    for p in doc.paragraphs:
        if p.text.startswith('d. Initial Design') or p.text.startswith('d. ') and 'Initial Design' in p.text:
            target_p = p
            break
            
    if not target_p:
        # Fallback to the text itself
        for p in doc.paragraphs:
            if p.text.startswith('Langkah perancangan awal (Initial Design)'):
                target_p = p
                break

    if not target_p:
        print(f"Target paragraph not found in {file_path}")
        return

    # Create the title paragraph
    title_p = target_p.insert_paragraph_before()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(6)
    title_p.paragraph_format.line_spacing = 1.0
    run = title_p.add_run("Tabel 4.1 Format Antarmuka Sistem Web ULT FKIP Unila")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = True
    if is_highlighted:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Table data
    table_data = [
        ("GROUP", "Public Portal (Tampilan Publik)"),
        ("1.", "Beranda Utama (Landing Page)", "• Header: Logo institusi, menu navigasi utama (Beranda, Katalog Layanan, Berita), dan tombol Masuk/Daftar.\n• Konten: Hero section besar dengan slogan layanan, kartu langkah-langkah alur permohonan.\n• Footer: Informasi kontak, tautan penting, dan hak cipta."),
        ("2.", "Katalog Layanan", "• Konten: Daftar lengkap layanan persuratan fakultas beserta rincian syarat dokumen yang dibutuhkan."),
        ("3.", "Berita & Pengumuman", "• Konten: Daftar artikel berita terbaru dan pengumuman resmi fakultas dengan penomoran halaman (pagination)."),
        ("GROUP", "Authentication System (Sistem Autentikasi)"),
        ("1.", "Login & Register", "• Konten: Form input kredensial (email dan kata sandi), logo fakultas, tombol aksi utama, dan tautan lupa kata sandi diletakkan di tengah halaman (centered card)."),
        ("GROUP", "Student Portal (Dasbor Mahasiswa)"),
        ("1.", "Beranda Mahasiswa", "• Sidebar: Menu navigasi privat (Dasbor, Buat Permohonan, Riwayat Layanan).\n• Konten: Ringkasan statistik permohonan aktif, notifikasi, dan tombol pintasan pengajuan layanan."),
        ("2.", "Form Pengajuan Layanan", "• Header: Judul jenis layanan terkait.\n• Konten: Formulir isian dinamis (dynamic inputs) sesuai jenis layanan dan area unggah berkas prasyarat dengan validasi ekstensi file."),
        ("3.", "Riwayat & Pelacakan (Timeline)", "• Konten: Tabel daftar permohonan, status proses (Draft, Menunggu Verifikasi, Sedang Diproses, Selesai), dan linimasa pelacakan (auditable timeline) untuk setiap dokumen."),
        ("GROUP", "Admin/Staff Portal (Dasbor Operasional)"),
        ("1.", "Dasbor Utama Staff", "• Sidebar: Menu navigasi manajemen (Permohonan Masuk, Daftar Layanan, Template Dokumen, Manajemen Pengguna).\n• Konten: Tabel data grid berisi permohonan masuk dengan filter pencarian dan tombol aksi (Review, Setujui, Tolak)."),
        ("2.", "Detail Permohonan (Review)", "• Konten: Pratinjau data isian mahasiswa, pratinjau berkas lampiran, kolom pembuatan nomor surat, serta tombol perakitan dokumen (Assembly)."),
        ("3.", "Manajemen Template", "• Konten: Form pengaturan nama layanan, manajemen variabel input dinamis, dan area unggah master template Word (.docx)."),
        ("GROUP", "Signer Portal (Portal Tanda Tangan)"),
        ("1.", "Antarmuka Verifikasi Pejabat", "• Konten: Tampilan ringkas daftar dokumen yang butuh persetujuan, pratinjau draf surat, dan tombol aksi untuk membubuhkan tanda tangan elektronik (Digital Signature).")
    ]

    # Create table (inserted at the end of the document initially, we will move it)
    # python-docx insert_table_before doesn't exist, we must add table to doc then move xml
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Set column widths
    table.columns[0].width = Cm(1.0)
    table.columns[1].width = Cm(4.5)
    table.columns[2].width = Cm(10.5)

    set_table_borders(table)

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'No.'
    hdr_cells[1].text = 'Halaman'
    hdr_cells[2].text = 'Format Antarmuka'
    
    for cell in hdr_cells:
        set_cell_background(cell, 'D9D9D9') # Light gray
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
                r.font.bold = True
                if is_highlighted:
                    r.font.highlight_color = WD_COLOR_INDEX.YELLOW

    for row_data in table_data:
        row_cells = table.add_row().cells
        if row_data[0] == "GROUP":
            # Merge all three cells
            row_cells[0].merge(row_cells[1])
            row_cells[0].merge(row_cells[2])
            row_cells[0].text = row_data[1]
            set_cell_background(row_cells[0], 'F2F2F2') # Very light gray
            for p in row_cells[0].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)
                    r.font.bold = True
                    if is_highlighted:
                        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            row_cells[0].text = row_data[0]
            row_cells[1].text = row_data[1]
            row_cells[2].text = row_data[2]
            
            for i, cell in enumerate(row_cells):
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_after = Pt(0)
                    for r in p.runs:
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(10)
                        if is_highlighted:
                            r.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Move table to be after title_p
    # title_p._p is the xml element of the title paragraph.
    # We want to insert the table xml right after title_p xml
    # Wait, the table is currently at the very end of the document.
    body = doc._body._body
    tbl_xml = table._tbl
    body.remove(tbl_xml)
    title_p._p.addnext(tbl_xml)
    
    # We also need to add a single paragraph break after the table so it doesn't collide with the next text.
    # Actually, we can just insert an empty paragraph after the table.
    spacer_p = docx.oxml.OxmlElement('w:p')
    pPr = docx.oxml.OxmlElement('w:pPr')
    spacing = docx.oxml.OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '240') # 12 pt
    pPr.append(spacing)
    spacer_p.append(pPr)
    tbl_xml.addnext(spacer_p)

    doc.save(file_path)
    print(f"Added table to {file_path}")

files_to_process = [
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx',
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
]

for f in files_to_process:
    add_table_to_doc(f, "Highlighted" in f)
