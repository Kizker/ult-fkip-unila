import docx

doc = docx.Document()
p = doc.add_paragraph("This is the target paragraph.")

# Add table
table = doc.add_table(rows=2, cols=2)
table.cell(0,0).text = "A"
table.cell(0,1).text = "B"
table.cell(1,0).text = "C"
table.cell(1,1).text = "D"

# Move table before paragraph
p._p.addprevious(table._element)

doc.save("test_table_move.docx")
print("Saved test_table_move.docx")
