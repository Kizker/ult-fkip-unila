import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_highlight(run, is_highlighted):
    if is_highlighted:
        rPr = run._r.get_or_add_rPr()
        highlight = OxmlElement('w:highlight')
        highlight.set(qn('w:val'), 'yellow')
        rPr.append(highlight)

def fix_section_412(doc_path, is_highlighted):
    doc = docx.Document(doc_path)
    
    # 1. Clean up ALL empty paragraphs that have numbering/bullet styles
    for p in list(doc.paragraphs):
        txt = p.text.strip()
        if not txt:
            pPr = p._p.pPr
            if pPr is not None:
                numPr = pPr.numPr
                if numPr is not None:
                    # It's an empty numbered/bulleted list item, delete it!
                    p._element.getparent().remove(p._element)

    # Re-iterate to find Gambar 4 and 5 blocks
    idx_g4 = -1
    idx_g5 = -1
    
    # Text to match
    g4_intro = "Diagram use case pada Gambar 4 secara visual menggambarkan interaksi logis"
    g5_intro = "Struktur berlapis pada Gambar 5 terbagi secara hierarkis"
    
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt.startswith(g4_intro):
            idx_g4 = i
        elif txt.startswith(g5_intro):
            idx_g5 = i

    if idx_g4 != -1:
        # Delete the intro and the 5 bullet items
        # Be careful, removing elements shifts indices. We'll just clear text for now, 
        # and remove their xml element so they disappear entirely.
        # We need the reference to insert before. We will insert before idx_g4, then delete idx_g4 to idx_g4+5.
        target_g4 = doc.paragraphs[idx_g4]
        
        p4_A = "Diagram use case pada Gambar 4 secara visual menggambarkan interaksi logis antara lima aktor utama civitas akademika dengan sistem fungsional Web ULT FKIP Unila tanpa menggunakan hierarki penomoran yang kaku. Aktor Mahasiswa bertindak sebagai pemohon utama yang mengajukan layanan persuratan akademik secara mandiri dengan akses penuh untuk pengajuan, pelacakan linimasa, hingga pengunduhan berkas akhir. Di tingkat verifikasi awal, Aktor Administrator Jurusan memegang kendali vital untuk meninjau kesesuaian lampiran pengajuan mahasiswa sebelum menerbitkan nomor registrasi surat otomatis."
        p4_B = "Pada level manajemen fakultas, Aktor Staf ULT bertindak selaku pengelola utama operasional yang bertanggung jawab atas validasi tingkat lanjut, pengelolaan templat dokumen, serta eksekusi perakitan berkas. Hasil verifikasi staf ini kemudian dilanjutkan kepada Aktor Pejabat selaku pimpinan fakultas atau jurusan untuk dilakukan peninjauan substantif dan legalisasi melalui pembubuhan tanda tangan elektronik. Keseluruhan ekosistem interaksi ini senantiasa diawasi oleh Aktor Admin Utama yang mengendalikan hak akses pengguna dan memantau rekaman jejak audit demi menjaga stabilitas serta akuntabilitas platform."
        
        # Insert
        nA = target_g4.insert_paragraph_before(p4_A, style='Paragraph')
        if is_highlighted:
            for r in nA.runs: add_highlight(r, True)
            
        nB = target_g4.insert_paragraph_before(p4_B, style='Paragraph')
        if is_highlighted:
            for r in nB.runs: add_highlight(r, True)
            
        # Remove old elements (intro + 5 bullets)
        for i in range(idx_g4, idx_g4 + 6):
            if i < len(doc.paragraphs):
                p_del = doc.paragraphs[i]
                p_del._element.getparent().remove(p_del._element)

    # Need to find g5 again because indices shifted
    idx_g5 = -1
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt.startswith(g5_intro):
            idx_g5 = i

    if idx_g5 != -1:
        target_g5 = doc.paragraphs[idx_g5]
        
        p5_A = "Struktur berlapis pada Gambar 5 terbagi secara hierarkis ke dalam lima lapisan operasional terpadu guna menjamin keamanan data akademik sensitif sekaligus memperlancar pemrosesan persuratan. Lapisan terluar diisi oleh Client and Interface Layer yang mengandalkan teknologi Progressive Web App sebagai gerbang interaksi visual bagi mahasiswa, staf, dan pimpinan secara responsif. Lapis pertahanan sistem dilanjutkan oleh Security and Routing Gatekeeper yang secara ketat menangani penjaluran akses, validasi otorisasi peran, penyaringan input kotor, serta proteksi kerentanan serangan peretasan dari luar."
        p5_B = "Logika bisnis aplikasi sepenuhnya dikendalikan oleh lapisan Laravel Application Core yang memuat serangkaian pengendali proses perakitan dan modul penerjemah sintaks teks kaya. Seluruh data operasional maupun relasional selanjutnya dititipkan ke dalam Database and Storage Layer yang secara fisik memisahkan penyimpanan konfigurasi tabel dengan direktori penyimpanan lampiran berkas mentah. Rangkaian arsitektur ini berpuncak pada modul Document Assembly Engine yang beroperasi secara laten sebagai mesin utama untuk memanipulasi dokumen XML dan memproduksi luaran berkas persuratan siap cetak."
        
        nA = target_g5.insert_paragraph_before(p5_A, style='Paragraph')
        if is_highlighted:
            for r in nA.runs: add_highlight(r, True)
            
        nB = target_g5.insert_paragraph_before(p5_B, style='Paragraph')
        if is_highlighted:
            for r in nB.runs: add_highlight(r, True)
            
        # Remove old elements (intro + 5 bullets)
        for i in range(idx_g5, idx_g5 + 6):
            if i < len(doc.paragraphs):
                p_del = doc.paragraphs[i]
                p_del._element.getparent().remove(p_del._element)

    doc.save(doc_path)
    print(f"Fixed numbering and expanded text for {doc_path}")

if __name__ == '__main__':
    fix_section_412(r'001_Skripsi_Andricha Dea Mitra_Clean.docx', is_highlighted=False)
    fix_section_412(r'001_Skripsi_Andricha Dea Mitra_Highlighted.docx', is_highlighted=True)
