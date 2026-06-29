import docx
import sys
import os

def replace_section(doc, heading, teks_list, next_headings):
    start_idx = -1
    end_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        if heading in p.text:
            start_idx = i
            break
            
    if start_idx == -1:
        print(f"Heading not found: {heading}")
        return False
        
    for i in range(start_idx + 1, len(doc.paragraphs)):
        if any(nx in doc.paragraphs[i].text for nx in next_headings):
            end_idx = i
            break
            
    if end_idx == -1:
        end_idx = len(doc.paragraphs)
        
    print(f"Replacing '{heading}' from index {start_idx} to {end_idx}")
    
    style = None
    insert_idx = -1
    for i in range(start_idx + 1, end_idx):
        if doc.paragraphs[i].text.strip():
            if not style:
                style = doc.paragraphs[i].style
                insert_idx = i
            doc.paragraphs[i].text = ""

    if insert_idx == -1:
        insert_idx = end_idx
    
    for teks in teks_list:
        p = doc.paragraphs[end_idx].insert_paragraph_before(teks)
        if style:
            p.style = style
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
        
    return True

def add_new_section(doc, after_heading, next_heading, new_heading, teks_list):
    start_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if after_heading in p.text:
            start_idx = i
            break
            
    if start_idx == -1:
        print(f"After heading not found: {after_heading}")
        return False
        
    end_idx = -1
    for i in range(start_idx + 1, len(doc.paragraphs)):
        if next_heading in doc.paragraphs[i].text:
            end_idx = i
            break
            
    if end_idx == -1:
        print(f"Next heading not found: {next_heading}")
        return False
        
    print(f"Inserting new section '{new_heading}' before index {end_idx}")
    
    style = None
    for i in range(start_idx + 1, end_idx):
        if doc.paragraphs[i].text.strip():
            style = doc.paragraphs[i].style
            break
            
    h_style = doc.paragraphs[start_idx].style
    
    doc.paragraphs[end_idx].insert_paragraph_before("")
    
    hp = doc.paragraphs[end_idx].insert_paragraph_before(new_heading)
    hp.style = h_style
    
    doc.paragraphs[end_idx].insert_paragraph_before("")
    
    for teks in teks_list:
        p = doc.paragraphs[end_idx].insert_paragraph_before(teks)
        if style:
            p.style = style
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
        
    return True

def apply_revisions(doc_path):
    print(f"\nProcessing {doc_path}...")
    doc = docx.Document(doc_path)
    
    teks_421 = [
        "Pencapaian skor kepraktisan sebesar 92,13% pada uji coba lapangan mengindikasikan lebih dari sekadar keberhasilan fungsional sebuah sistem perangkat lunak. Tingginya angka ini secara empiris mengonfirmasi teori User Experience (Maulana et al., 2023) yang menekankan urgensi perancangan antarmuka untuk meminimalkan beban kognitif (cognitive load) bagi pengguna pemula. Antarmuka yang dikembangkan, khususnya pada portal mahasiswa (Student Portal), terbukti mampu memandu proses pengajuan dokumen tanpa memerlukan kurva pembelajaran yang curam. Keselarasan fungsionalitas aplikasi dengan mentalitas pengguna (mental model) menghasilkan interaksi yang sangat intuitif, sehingga hambatan-hambatan adaptasi teknologi yang sering memicu frustrasi dapat direduksi secara signifikan.",
        "Lebih jauh, capaian ini melahirkan sebuah paradigma baru dalam ekosistem tata kelola administratif di lingkungan ULT FKIP Universitas Lampung. Keberhasilan implementasi sistem ini menegaskan bahwa digitalisasi bukan sebatas rutinitas memindahkan borang fisik menjadi format elektronik, melainkan sebuah bentuk business process reengineering (rekayasa ulang proses bisnis) yang fundamental. Pembaruan ini secara efektif memangkas hierarki birokrasi yang konvensional, mempercepat laju koordinasi lintas program studi hingga ke tingkat fakultas, serta mendefinisikan ulang standar pelayanan publik yang terukur, responsif, dan berorientasi penuh pada pengguna akhir."
    ]
    
    teks_422 = [
        "Pemecahan masalah ketiadaan transparansi layanan yang selama ini menjadi keluhan laten di kalangan mahasiswa berhasil direalisasikan melalui implementasi fitur Audit Trail dan linimasa pelacakan (tracking timeline) yang terintegrasi secara komprehensif. Melalui fitur ini, setiap pergerakan permohonan memiliki rekam jejak digital yang presisi, mencatat secara pasti cap waktu (timestamp) eksekusi beserta entitas pengguna yang bertanggung jawab. Siklus status dokumen dirancang dengan transisi yang terstruktur; dimulai dari status \"Diajukan\" ketika mahasiswa pertama kali menyerahkan berkas sistemik. Sistem kemudian memperbarui status menjadi \"Direview\" saat staf ULT atau admin program studi melakukan verifikasi kelengkapan preskriptif. Apabila permohonan membutuhkan otorisasi pejabat berwenang, alur dilanjutkan ke tahap \"Dalam Signing/Penandatanganan\". Siklus ini bermuara pada status \"Selesai\" ketika dokumen telah dirakit dan siap diunduh secara privat, atau diubah menjadi \"Revisi/Ditolak\" lengkap dengan anotasi catatan manakala didapati adanya ketidaksesuaian prasyarat.",
        "Visibilitas real-time pada setiap tahap pergerakan dokumen ini secara mutlak mengeliminasi fenomena \"black box\" birokrasi, di mana sebelumnya pemohon sama sekali buta terhadap laju progres dokumen mereka. Mahasiswa kini tidak perlu lagi membuang waktu untuk hadir secara fisik sekadar mempertanyakan status permohonan. Lebih esensial lagi, transparansi pelacakan ini menciptakan pertanggungjawaban (accountability) yang mengikat bagi seluruh pemangku kepentingan. Dengan demikian, risiko manajerial fundamental seperti dokumen yang menumpuk di atas meja, berkas yang terselip, hingga dokumen yang hilang secara fisik dapat diatasi hingga pada titik yang tidak lagi relevan dalam tata kelola layanan."
    ]
    
    teks_evaluasi = [
        "Sebagai pilar pamungkas dari metodologi ADDIE, tahap evaluasi ini menyajikan cerminan reflektif terhadap keseluruhan ikhtiar pemecahan masalah. Menilik kembali diskursus urgensi pada analisis kebutuhan (Bab I), operasional administrasi di Fakultas Keguruan dan Ilmu Pendidikan (FKIP) Universitas Lampung dalam waktu yang lama terjerat oleh inefisiensi mekanis. Mahasiswa terus dihadapkan pada tumpukan borang manual yang memicu antrean panjang, birokrasi layanan yang nir-transparansi, serta kerentanan struktural yang memfasilitasi terjadinya kasus berkas akademik yang terselip, rusak, maupun hilang. Kumpulan patologi administratif inilah yang menginisiasi perlunya intervensi teknologi berskala besar.",
        "Mengartikulasikan problematika masa lalu, rumusan cetak biru pada tahapan desain (Bab III) dan proses rekayasa sistem (Development) pada tahapan ini (Bab IV) sukses mengejawantahkan solusi yang komprehensif. Perancangan kode sumber aplikasi menjadi instrumen transisi sistem operasional, di mana mekanisme Role-Based Access Control (RBAC) menyekat wewenang portal dan sinkronisasi pustaka OpenXML mengotomatisasi perakitan dokumen Word. Hasil pengembangan ini sukses meleburkan prosedur-prosedur layanan lintas prodi dan jurusan yang semula terpisah ke dalam satu wadah portal tunggal yang tersentralisasi.",
        "Laju transisi sistem konvensional menuju sistem cerdas tidak terlepas dari dinamika lapangan pada tahap implementasi (Implementation). Tantangan manajerial di fase ini—mulai dari kelambatan respons pengguna, kebutuhan adaptasi konfigurasi peladen, hingga kesibukan jadwal staf dan pimpinan fakultas—telah menuntut strategi \"jemput bola\" dan pendampingan terfokus. Meskipun berhadapan dengan turbulensi eksekusi lapangan, uji coba mampu terlaksana secara persisten, sukses membuktikan ketahanan sistem Audit Trail serta validasi privasi dokumen yang solid di bawah operasional lingkungan pengguna nyata.",
        "Sebagai akumulasi keberhasilan dari keseluruhan proses penelitian, evaluasi sumatif sistem terbuktikan dengan indeks Validitas dari pakar yang mencatat angka presisi 91,95% (Sangat Valid), serta tingkat Kepraktisan lapangan sebesar 92,13% (Sangat Praktis). Tingginya persentase tersebut melampaui justifikasi sebagai sebatas laporan kuantitatif. Angka-angka ini berdiri sebagai pembuktian mutlak bahwa sistem layanan elektronik ULT FKIP Unila—secara akademis, operasional, dan teknologis—telah menuntaskan akar persoalan inefisiensi layanan masa lalu, dan mendeklarasikan kesiapan absolut untuk diimplementasikan secara holistik dan permanen."
    ]
    
    teks_425 = [
        "Dalam merealisasikan produk pengembangan dari lingkungan lokal ke tahap deployment pada peladen awan (cloud server), penelitian ini berhadapan dengan sejumlah kendala teknis terkait konfigurasi infrastruktur. Salah satu tantangan utama pada peladen hosting berbasis cPanel adalah ketidaksesuaian antara konfigurasi bawaan peladen dengan spesifikasi lingkungan PHP yang diwajibkan oleh framework Laravel 12. Di samping itu, pengadopsian arsitektur penyimpanan privat (Private Storage) menuntut pengamanan direktori ekstra agar file dokumen penting tidak terekspos secara publik. Sebagai resolusi teknis, dilakukan penyelarasan pustaka lingkungan pada versi PHP 8.4+ dan manipulasi tautan simbolik (symlink) struktur direktori. Pendekatan ini memastikan direktori dokumen terisolasi dari public root, sehingga pengunduhan output selalu ditegakkan melalui middleware otorisasi. Hasil dari perombakan arsitektural ini berhasil menjamin data persisten tetap terenkripsi dan imun dari eksploitasi kerentanan Insecure Direct Object Reference (IDOR).",
        "Dari dimensi manajemen lapangan, tantangan substansial muncul saat mengeksekusi tahapan uji kepraktisan. Ritme kesibukan agenda operasional dari pihak fakultas, staf ULT, serta mobilitas perkuliahan mahasiswa berdampak langsung pada kelambatan pengumpulan respons pengujian dan pengisian instrumen kuesioner. Keterbatasan slot waktu yang dimiliki oleh responden menuntut penerapan strategi yang efisien, guna memastikan validitas data tidak terkompromi oleh proses pengumpulan yang berkepanjangan.",
        "Untuk mengatasi stagnasi pengujian tersebut, diimplementasikan strategi eksekusi \"jemput bola\" berbasis persuasi aktif. Peneliti melaksanakan visitasi langsung ke ruang-ruang staf ULT dan admin prodi, menginterupsi rutinitas harian dengan pendampingan teknis yang efisien. Pada segmentasi pengujian mahasiswa, uji coba dikonsolidasikan dalam sesi paralel berbekal desain skenario instruksi yang dikondensasi (compact). Skenario singkat ini dikonstruksi tanpa menghilangkan titik-titik interaksi krusial dari fitur aplikasi, sehingga waktu pengujian dapat diminimalkan sembari tetap menyerap impresi kepraktisan secara autentik dan komprehensif."
    ]

    replace_section(doc, "Tahap Evaluasi (Evaluation)", teks_evaluasi, ["4.2", "Pembahasan Hasil Penelitian"])
    replace_section(doc, "4.2.1 Implikasi Kepraktisan", teks_421, ["4.2.2", "Transparansi"])
    
    idx_422 = -1
    for i, p in enumerate(doc.paragraphs):
        if "4.2.2 Transparansi" in p.text:
            idx_422 = i
            break
            
    if idx_422 != -1:
        print("Handling 4.2.2 specially to preserve images...")
        count_cleared = 0
        style = None
        for i in range(idx_422 + 1, len(doc.paragraphs)):
            text = doc.paragraphs[i].text.strip()
            if "Gambar" in text or "Placeholder" in text:
                continue
            if "4.2.3" in text or "BAB V" in text or "KESIMPULAN" in text:
                break
                
            if len(text) > 50:
                if not style: style = doc.paragraphs[i].style
                doc.paragraphs[i].text = ""
                count_cleared += 1
                
        for teks in teks_422:
            p = doc.paragraphs[idx_422 + 1].insert_paragraph_before(teks)
            if style: p.style = style
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    add_new_section(doc, "4.2", "KESIMPULAN DAN SARAN", "4.2.5 Kendala Pelaksanaan Lapangan dan Solusi Eksekusi", teks_425)

    doc.save(doc_path)
    print(f"Saved {doc_path}")

if __name__ == "__main__":
    file1 = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
    file2 = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
    apply_revisions(file1)
    apply_revisions(file2)
