import docx
from pathlib import Path

def main():
    path = Path(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\update skripsi terakhir-seminar proposal\001_Skripsi_Andricha Dea Mitra.docx")
    if not path.exists():
        print("Not found.")
        return
        
    doc = docx.Document(path)
    print("Total paragraphs:", len(doc.paragraphs))
    print("Total tables:", len(doc.tables))
    
    # Search for percentage values in the document paragraphs
    print("\n=== Searching for values ===")
    queries = ["95.45", "93.33", "87.58", "91.95", "92.13", "purposive sampling", "Aiken"]
    for q in queries:
        found = False
        for i, p in enumerate(doc.paragraphs):
            if q.lower() in p.text.lower():
                print(f"P {i}: '{p.text[:120]}...'")
                found = True
        if not found:
            print(f"'{q}' not found in paragraphs.")
            
    # Let's inspect Table 12 first row and some cells
    if len(doc.tables) > 12:
        t = doc.tables[12]
        print("\n=== Table 12 contents ===")
        for r_idx, row in enumerate(t.rows):
            vals = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            print(f"Row {r_idx}: {vals}")

if __name__ == "__main__":
    main()
