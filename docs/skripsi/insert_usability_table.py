import sys
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

# Set stdout to UTF-8
sys.stdout.reconfigure(encoding='utf-8')

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def format_paragraph(p, doc, style_name="Normal", alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing_before=0, spacing_after=0, line_spacing=1.5):
    try:
        p.style = doc.styles[style_name]
    except Exception:
        pass
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(spacing_before)
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = line_spacing

def add_run_formatted(p, text, bold=False, italic=False, font_size=12):
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(font_size)
    r.bold = bold
    r.italic = italic
    return r

def main():
    path = Path(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx")
    if not path.exists():
        print("Not found.")
        return
        
    doc = docx.Document(path)
    body = doc.element.body
    
    # Step 1: Find the target paragraph under "Tahap Implementasi"
    target_p = None
    target_idx = -1
    
    print("Searching for the target paragraph under 'Tahap Implementasi'...")
    for idx, child in enumerate(body.getchildren()):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = docx.text.paragraph.Paragraph(child, doc)
            txt = p.text.strip()
            # We match the paragraph that starts with "Selain data kuantitatif kuesioner"
            if txt.startswith("Selain data kuantitatif kuesioner") and "Bagian F" in txt:
                target_p = p
                target_idx = idx
                break
                
    if target_p is None:
        print("Error: Target paragraph not found!")
        return
        
    print(f"Found target paragraph at element index {target_idx}")
    
    # Step 2: Update the text of the target paragraph to add reference to Table 4.2
    new_text = (
        "Selain data kuantitatif kuesioner, pada lembar angket Bagian F (Kesimpulan Responden), "
        "responden diminta memberikan pendapat kelayakan subjektif akhir produk secara langsung. "
        "Rekapitulasi distribusi pilihan kesimpulan akhir tersebut menunjukkan mayoritas responden "
        "(94,44% atau 17 orang) menyimpulkan bahwa sistem ini berkategori 'Sangat Praktis' (8 orang) dan "
        "'Praktis' (9 orang) untuk langsung digunakan dalam kegiatan operasional sehari-hari. Hanya 1 "
        "responden (5,56% yaitu Responden 3) yang memberikan kesimpulan 'Cukup Praktis'. Hal ini membuktikan "
        "antusiasme yang tinggi dari para responden pengguna terhadap efektivitas sistem. Rekapitulasi hasil "
        "distribusi kategori uji kepraktisan tersebut disajikan secara lengkap pada Tabel 4.2 berikut."
    )
    
    # We clear runs and set the text
    target_p.text = ""
    add_run_formatted(target_p, new_text, bold=False, italic=False, font_size=12)
    format_paragraph(target_p, doc, "Normal", WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.5)
    print("Updated paragraph text and formatting successfully.")
    
    # Step 3: Insert Judul Tabel (Tabel 4.2)
    p_title = doc.add_paragraph()
    format_paragraph(p_title, doc, "Normal", WD_ALIGN_PARAGRAPH.CENTER, spacing_before=12, spacing_after=3, line_spacing=1.0)
    add_run_formatted(p_title, "Tabel 4.2. Rekapitulasi Distribusi Kategori Uji Kepraktisan Pengguna", bold=True, font_size=11)
    
    body.remove(p_title._p)
    body.insert(target_idx + 1, p_title._p)
    print("Inserted Table Title paragraph.")
    
    # Step 4: Insert Table 4.2
    table = doc.add_table(rows=7, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    table_data = [
        ["No", "Kategori Kepraktisan (Pilihan Bagian F)", "Jumlah Responden", "Persentase (%)"],
        ["1", "Sangat Praktis (81% - 100%)", "8", "44,44%"],
        ["2", "Praktis (61% - 80%)", "9", "50,00%"],
        ["3", "Cukup Praktis (41% - 60%)", "1", "5,56%"],
        ["4", "Kurang Praktis (21% - 40%)", "0", "0,00%"],
        ["5", "Tidak Praktis (0% - 20%)", "0", "0,00%"],
        ["", "Total", "18", "100,00%"]
    ]
    
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.text = ""
            p_cell = cell.paragraphs[0]
            
            # Formatting cells: Bold for headers and totals
            is_bold = (r_idx == 0 or r_idx == 6)
            
            # Alignments: Left for category, Center for other columns
            align = WD_ALIGN_PARAGRAPH.CENTER
            if c_idx == 1 and r_idx != 0 and r_idx != 6:
                align = WD_ALIGN_PARAGRAPH.LEFT
                
            format_paragraph(p_cell, doc, style_name="Normal", alignment=align, spacing_before=2, spacing_after=2, line_spacing=1.0)
            add_run_formatted(p_cell, table_data[r_idx][c_idx], bold=is_bold, font_size=10)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            
    body.remove(table._tbl)
    body.insert(target_idx + 2, table._tbl)
    print("Inserted Table 4.2 data grid.")
    
    # Step 5: Insert Sumber Tabel
    p_src = doc.add_paragraph()
    format_paragraph(p_src, doc, "Normal", WD_ALIGN_PARAGRAPH.CENTER, spacing_before=3, spacing_after=12, line_spacing=1.0)
    add_run_formatted(p_src, "Sumber: Data primer olahan peneliti (2026)", italic=True, font_size=10)
    
    body.remove(p_src._p)
    body.insert(target_idx + 3, p_src._p)
    print("Inserted Table Source paragraph.")
    
    # Save the updated Clean.docx
    doc.save(path)
    print(f"Clean document successfully updated and saved to: {path}")

if __name__ == "__main__":
    main()
