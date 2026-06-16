import docx
from pathlib import Path

def main():
    path = Path(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx")
    if not path.exists():
        print("Clean doc not found.")
        return
        
    doc = docx.Document(path)
    out_path = Path(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\bab3_aiken_output.txt")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("Printing paragraphs 670 to 715:\n")
        for idx in range(670, 715):
            if idx < len(doc.paragraphs):
                f.write(f"P {idx}: '{doc.paragraphs[idx].text}'\n")
    print("Successfully wrote bab3_aiken_output.txt")

if __name__ == "__main__":
    main()
