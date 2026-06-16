import docx
import re
import os

doc_clean_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
doc_hl_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'

foreign_phrases = [
    "Entity-Relationship", "ERD", "MySQL", "BIGINT", "VARCHAR", "TEXT", "BOOLEAN", "TIMESTAMP", 
    "foreign key constraints", "users", "roles", "model_has_roles", "Role-Based Access Control", "RBAC", 
    "services", "template_path", "requests", "tracking_code", "issue_number", "request_inputs", 
    "field_name", "field_value", "One-to-Many", "request_documents", "parent record", "audit_trails", 
    "audit log", "action", "IP", "sequence", "time-oriented", "lifeline", "message passing", "method calls", 
    "User", "RequestController", "DocumentAssemblerService", "HtmlToOpenXMLParser", "HTML", "XML", 
    "Private Storage", "clickAssemble", "requestId", "assembleDocument", "service", "loadTemplateAndFormData", 
    "WYSIWYG", "writeHtmlToWordRun", "traverseHtmlAndInsertRuns", "DOM HTML", "Word", "parser", 
    "cloneRunPropertiesAndFormat", "font", "placeholder", "OpenXML", "compileAndZipArchive", "ZIP", "file", 
    ".docx", "putFileInPrivateDisk", "return value", "HTTP", "redirectWithSuccessToast"
]

foreign_phrases.sort(key=len, reverse=True)

def italicize_text_in_paragraph(paragraph, phrases):
    text = paragraph.text
    has_phrase = False
    for phrase in phrases:
        if re.search(r'\b' + re.escape(phrase) + r'\b', text, re.IGNORECASE):
            has_phrase = True
            break
            
    if not has_phrase:
        return False
        
    hl_color = None
    if len(paragraph.runs) > 0:
        hl_color = paragraph.runs[0].font.highlight_color
        
    pattern = r'\b(' + '|'.join(re.escape(p) for p in phrases) + r')\b'
    parts = re.split(pattern, text, flags=re.IGNORECASE)
    
    paragraph.clear()
    
    from docx.shared import Pt
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if hl_color is not None:
            run.font.highlight_color = hl_color
            
        is_foreign = False
        for phrase in phrases:
            if part.lower() == phrase.lower():
                is_foreign = True
                break
                
        if is_foreign:
            run.italic = True
            
    return True

def process_doc(path):
    doc = docx.Document(path)
    
    injected_starts = [
        "Diagram Entity-Relationship (ERD) pada",
        "Entitas users, roles, beserta",
        "Sebagai pusat operasional transaksional",
        "Alur kerja relasional ERD ini",
        "Diagram sequence pada Gambar 8",
        "Komponen yang terlibat secara",
        "Skenario perakitan dokumen dimulai",
        "HtmlToOpenXMLParser lantas mengambil",
        "Pada siklus penyelesaian di langkah"
    ]
    
    count = 0
    for p in doc.paragraphs:
        for start in injected_starts:
            if p.text.startswith(start):
                if italicize_text_in_paragraph(p, foreign_phrases):
                    count += 1
                break
                
    print(f"Processed {count} paragraphs in {os.path.basename(path)}")
    doc.save(path)

process_doc(doc_clean_path)
process_doc(doc_hl_path)
