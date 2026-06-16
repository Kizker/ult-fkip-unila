import sys
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# Set stdout to UTF-8
sys.stdout.reconfigure(encoding='utf-8')

def format_paragraph(p, doc, style_name="Normal", alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing_before=0, spacing_after=0, line_spacing=1.5):
    try:
        p.style = doc.styles[style_name]
    except Exception:
        pass
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(spacing_before)
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = line_spacing

def add_run_formatted(p, text, bold=False, italic=False, font_size=12):
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(font_size)
    r.bold = bold
    r.italic = italic
    return r

def main():
    path = Path(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx")
    if not path.exists():
        print("Error: Clean.docx not found.")
        return
        
    doc = docx.Document(path)
    
    found_quant = False
    found_qual = False
    
    print("Searching for paragraphs to update...")
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        
        # 1. Match the quantitative paragraph
        if txt.startswith("Hasil analisis data kuantitatif dari ke-18 responden") and "Lampiran C" in txt:
            print(f"Found quantitative paragraph at index {i}")
            new_quant_text = (
                "Hasil analisis data kuantitatif dari ke-18 responden tersebut menunjukkan tingkat kepraktisan sistem yang sangat memuaskan. "
                "Rata-rata skor total yang diperoleh dari kuesioner adalah 55,28 dari skor maksimal 60,00, yang menghasilkan persentase kepraktisan rata-rata sebesar 92,13%. "
                "Berdasarkan kriteria Arikunto, hasil persentase ini menempatkan Website ULT FKIP Unila dalam kategori Sangat Praktis untuk langsung "
                "diimplementasikan dalam kegiatan operasional sehari-hari. Evaluasi butir instrumen memperlihatkan skor yang sangat stabil di atas 4,40, "
                "dengan skor tertinggi berada pada butir A3 (Kemudahan mempelajari penggunaan sistem) yang meraih skor rata-rata 4,83 dari 5,00. "
                "Hal ini membuktikan bahwa antarmuka sistem sangat intuitif sehingga pengguna sasaran dapat beradaptasi dan menggunakan sistem secara mandiri "
                "dengan sangat cepat tanpa memerlukan pelatihan teknis khusus. Adapun rekapitulasi tingkat kepraktisan berdasarkan persepsi pengguna disajikan secara lengkap pada Tabel 4.2."
            )
            p.text = ""
            add_run_formatted(p, new_quant_text, bold=False, italic=False, font_size=12)
            format_paragraph(p, doc, "Normal", WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.5)
            found_quant = True
            
        # 2. Match the qualitative paragraph
        elif txt.startswith("Selain data kuantitatif kuesioner") and "Tabel 4.2 berikut" in txt:
            print(f"Found qualitative paragraph at index {i}")
            new_qual_text = (
                "Selain data kuantitatif kuesioner, pada lembar angket Bagian F (Kesimpulan Responden), "
                "responden diminta memberikan pendapat kelayakan subjektif akhir produk secara langsung. "
                "Rekapitulasi distribusi pilihan kesimpulan akhir tersebut menunjukkan mayoritas responden "
                "(94,44% atau 17 orang) menyimpulkan bahwa sistem ini berkategori 'Sangat Praktis' (8 orang) dan "
                "'Praktis' (9 orang) untuk langsung digunakan dalam kegiatan operasional sehari-hari. Hanya 1 "
                "responden (5,56% yaitu Responden 3) yang memberikan kesimpulan 'Cukup Praktis'. Hal ini membuktikan "
                "antusiasme yang tinggi dari para responden pengguna terhadap efektivitas sistem. Distribusi kategori tingkat "
                "kepraktisan dari pilihan responden tersebut secara detail dapat dilihat pada Tabel 4.2 di bawah ini."
            )
            p.text = ""
            add_run_formatted(p, new_qual_text, bold=False, italic=False, font_size=12)
            format_paragraph(p, doc, "Normal", WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.5)
            found_qual = True
            
    if not found_quant:
        print("Warning: Quantitative paragraph not found or already updated.")
    if not found_qual:
        print("Warning: Qualitative paragraph not found or already updated.")
        
    doc.save(path)
    print("Clean.docx saved successfully.")

if __name__ == "__main__":
    main()
