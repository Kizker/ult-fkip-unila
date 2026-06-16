import docx
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

def count_mendeley_and_math(path):
    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    }
    
    with zipfile.ZipFile(path) as z:
        doc_xml = ET.fromstring(z.read('word/document.xml'))
        
    mendeley_fld = doc_xml.findall('.//w:fldSimple', ns)
    mendeley_instr = doc_xml.findall('.//w:instrText', ns)
    math_eq = doc_xml.findall('.//m:oMath', ns)
    
    # Mendeley fields usually have 'ADDIN ONEOBJ' or similar in instrText
    mendeley_count = len(mendeley_fld) + sum(1 for t in mendeley_instr if 'ADDIN' in (t.text or ''))
    math_count = len(math_eq)
    
    return mendeley_count, math_count

def verify_document(path_str, is_highlighted=False):
    path = Path(path_str)
    print(f"\n--- VERIFYING DOCUMENT: {path.name} ---")
    if not path.exists():
        print(f"Error: {path.name} does not exist!")
        return False
        
    doc = docx.Document(path)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    
    # 1. Verify Mendeley and Word Equations
    m_count, eq_count = count_mendeley_and_math(path)
    print(f"Mendeley field codes count: {m_count}")
    print(f"Word Equations count: {eq_count}")
    
    # 2. Check headings in Bab IV
    headings = []
    has_table_4_1 = False
    has_table_4_2 = False
    has_table_4_3 = False
    has_table_4_4 = False
    has_table_4_5 = False
    has_table_4_6 = False
    
    target_headings = [
        "Hasil Penelitian (Model Pengembangan ADDIE)",
        "1. Tahap Analisis (Analysis)",
        "2. Tahap Desain (Design)",
        "3. Tahap Pengembangan (Development)",
        "4. Tahap Implementasi (Implementation)",
        "5. Tahap Evaluasi (Evaluation)",
        "B. Pembahasan Hasil Penelitian"
    ]
    
    found_headings = {h: False for h in target_headings}
    
    for p in doc.paragraphs:
        txt = p.text.strip()
        for th in target_headings:
            if th.lower() in txt.lower():
                found_headings[th] = True
                
        # Check for deleted aspect tables references or headings
        if "Tabel 4.1" in txt:
            has_table_4_1 = True
        if "Tabel 4.2" in txt:
            has_table_4_2 = True
        if "Tabel 4.3" in txt:
            has_table_4_3 = True
        if "Tabel 4.4" in txt:
            has_table_4_4 = True
            
    print("Headings check:")
    for h, found in found_headings.items():
        print(f"  - '{h}': {'FOUND' if found else 'NOT FOUND'}")
        
    # Check for highlight in highlighted version
    if is_highlighted:
        highlight_found = False
        for p in doc.paragraphs:
            if "tahap analisis" in p.text.lower() or "tahap desain" in p.text.lower():
                for r in p.runs:
                    if r.font.highlight_color is not None:
                        highlight_found = True
                        break
        print(f"Highlight check (Yellow expected): {'FOUND' if highlight_found else 'NOT FOUND'}")
        
    return all(found_headings.values())

def main():
    clean_path = r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx"
    high_path = r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx"

    
    c_ok = verify_document(clean_path, is_highlighted=False)
    h_ok = verify_document(high_path, is_highlighted=True)
    
    if c_ok and h_ok:
        print("\n=== VERIFICATION SUCCESSFUL: 100% PASS! ===")
    else:
        print("\n=== VERIFICATION FAILED! ===")

if __name__ == "__main__":
    main()
