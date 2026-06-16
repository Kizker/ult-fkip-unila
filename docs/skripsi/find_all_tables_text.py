import sys
import docx
from pathlib import Path

# Set stdout to UTF-8
sys.stdout.reconfigure(encoding='utf-8')

def main():
    path = Path(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx")
    doc = docx.Document(path)
    
    print(f"Total tables: {len(doc.tables)}")
    for idx, t in enumerate(doc.tables):
        first_row = []
        if len(t.rows) > 0:
            first_row = [c.text.strip().replace('\n', ' ')[:50] for c in t.rows[0].cells]
        print(f"Table {idx}: Rows={len(t.rows)}, Cols={len(t.columns)}, First Row={first_row[:4]}")
        
if __name__ == "__main__":
    main()
