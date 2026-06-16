import docx
import copy
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def merge_table():
    # 1. Grab table XML from standalone doc
    doc_standalone = docx.Document(r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\Matriks_Ahli_Validasi.docx')
    table_element_to_copy = copy.deepcopy(doc_standalone.tables[0]._element)
    
    # 2. Open main docs
    for doc_name in ['001_Skripsi_Andricha Dea Mitra_Clean.docx', '001_Skripsi_Andricha Dea Mitra_Highlighted.docx']:
        doc_path = rf'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\{doc_name}'
        doc = docx.Document(doc_path)
        
        # 3. Clean up previous broken insertions
        elements_to_remove = []
        for i, el in enumerate(doc._body._body):
            if el.tag.endswith('p'):
                p = docx.text.paragraph.Paragraph(el, doc)
                if 'Lampiran 10. Matriks Hasil Skor Kuesioner Uji Ahli Validasi' in p.text:
                    elements_to_remove.append(el)
                    if i+1 < len(doc._body._body) and doc._body._body[i+1].tag.endswith('tbl'):
                        elements_to_remove.append(doc._body._body[i+1])
                        
        for el in elements_to_remove:
            try:
                doc._body._body.remove(el)
            except:
                pass

        # 4. Find target paragraph
        target_p = None
        for p in doc.paragraphs:
            if 'Lampiran 11. Matriks Hasil Skor Kuesioner Uji Kepraktisan Responden' in p.text:
                target_p = p
                break
                
        if target_p:
            # Create heading
            new_heading = target_p.insert_paragraph_before("Lampiran 10. Matriks Hasil Skor Kuesioner Uji Ahli Validasi")
            new_heading.style = doc.styles['Normal']
            
            # Insert the copied table element
            # We must use deepcopy for each document to avoid cross-document references
            tbl_clone = copy.deepcopy(table_element_to_copy)
            target_p._p.addprevious(tbl_clone)
            
            # Add an empty paragraph for spacing
            empty_p_xml = docx.oxml.OxmlElement('w:p')
            target_p._p.addprevious(empty_p_xml)
            
            doc.save(doc_path)
            print(f"Successfully injected table into {doc_name}")

merge_table()
