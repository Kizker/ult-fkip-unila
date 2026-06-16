import docx
import os

doc_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
doc = docx.Document(doc_path)

in_impl = False
for p in doc.paragraphs:
    text = p.text.strip()
    if 'Tahap Implementasi (Implementation)' in text or '4.1.4' in text:
        in_impl = True
    elif in_impl and ('Tahap Evaluasi (Evaluation)' in text or '4.1.5' in text):
        in_impl = False
        break
    
    if in_impl:
        print(text)

for table in doc.tables:
    try:
        first_cell = table.cell(0,0).text.strip()
        if 'Kepraktisan' in first_cell or 'Responden' in first_cell or 'No' in first_cell:
            print(f"Table found, first cell: {first_cell}")
    except:
        pass
