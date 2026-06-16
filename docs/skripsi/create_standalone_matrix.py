import pandas as pd
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

file_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil uji ahli validator\Rekap_Uji_Validitas_Ahli.xlsx'
validators = []

for sheet in ['Validitas Materi', 'Validitas Media', 'Validitas Sistem']:
    df = pd.read_excel(file_path, sheet_name=sheet)
    names = df.columns[2:5]
    for name in names:
        scores = df[name].iloc[0:10].tolist()
        total = sum([float(x) for x in scores])
        perc = (total / 50.0) * 100
        validators.append({
            'Nama': name.strip(),
            'Bidang': sheet.replace('Validitas ', 'Ahli '),
            'Scores': [int(float(x)) for x in scores],
            'Total': int(total),
            'Persentase': f"{perc:.2f}%"
        })

doc = docx.Document()

# Setup landscape orientation for a wide table just to be safe
section = doc.sections[0]
new_width, new_height = section.page_height, section.page_width
section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height

doc.add_paragraph("Lampiran 10. Matriks Hasil Skor Kuesioner Uji Ahli Validasi")

table = doc.add_table(rows=1, cols=14)
table.style = 'Table Grid'
table.autofit = True

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Nama Validator'
hdr_cells[1].text = 'Bidang Keahlian'
for i in range(10):
    hdr_cells[i+2].text = f'Q{i+1}'
hdr_cells[12].text = 'Total'
hdr_cells[13].text = 'Persentase (%)'

for c in hdr_cells:
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(10)
            
for v in validators:
    row_cells = table.add_row().cells
    row_cells[0].text = v['Nama']
    row_cells[1].text = v['Bidang']
    for i, score in enumerate(v['Scores']):
        row_cells[i+2].text = str(score)
    row_cells[12].text = str(v['Total'])
    row_cells[13].text = v['Persentase']
    
    for i in range(14):
        for p in row_cells[i].paragraphs:
            if i >= 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(10)

doc.save(r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\Matriks_Ahli_Validasi.docx')
print("Created Matriks_Ahli_Validasi.docx")
