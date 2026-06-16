import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

with open(r'c:\laragon\www\ult-fkip-unila\docs\skripsi\daftar_pustaka_raw.txt', 'r', encoding='utf-8') as f:
    lines = [l.strip() for l in f.readlines() if l.strip() and l.strip() != 'DAFTAR PUSTAKA']

fixes = {
    'Ahmad, Z.': {'journal': 'Network and Distributed System Security (NDSS)', 'pages': '1-8'},
    'Aiken, L. R.': {'journal': 'Educational and Psychological Measurement', 'vol': '45(1)', 'pages': '131-142'},
    'Akbari, T. T.': {'pages': '52-65'},
    'Alisa, N.': {'journal': 'Journal of Governance and Public Administration (JoGaPA)'},
    'Arikunto, S. (2009)': {'publisher': 'Bumi Aksara, Jakarta', 'pages': '310 hlm'},
    'Arikunto, S. (2013)': {'publisher': 'Rineka Cipta, Jakarta', 'pages': '413 hlm'},
    'Branch, R. M.': {'publisher': 'Springer, Boston', 'pages': '175 hlm'},
    'Deshpande, Y.': {'journal': 'Web Engineering', 'vol': '1(1)', 'pages': '1-438'},
    'Henim, S. R.': {'journal': 'Jurnal Komputer Terapan'},
    'Johns, K.': {'publisher': 'World Bank Group, Washington, DC', 'pages': '150 hlm'},
    'Mustopa, A.': {'journal': 'Jurnal Ilmiah Informatika Komputer'},
    'Ridwan, M. Y.': {'journal': 'Jurnal Sistem Informasi'},
    'Sudjana. (2005)': {'publisher': 'Tarsito, Bandung', 'pages': '508 hlm'},
    'Sugiyono. (2013)': {'publisher': 'Alfabeta, Bandung', 'pages': '334 hlm'},
    'Widiawati, M.': {'journal': 'Jurnal Informatika', 'vol': '2', 'pages': '95-109'}
}

def parse_reference(ref_text):
    # Match author, year, title, and optional rest
    match = re.match(r'^([^\(]+?)\s+\((\d{4})\)\.\s+(.+?)\.(?:\s+(.+))?$', ref_text)
    if not match:
        return None
    author_part = match.group(1).strip()
    year = match.group(2).strip()
    title = match.group(3).strip()
    rest = match.group(4).strip() if match.group(4) else ''
    
    author_part = author_part.replace('&', 'dan')
    
    is_book = False
    if ' hlm' in rest.lower() or 'hlm.' in rest.lower() or 'jakarta' in rest.lower() or 'bandung' in rest.lower() or 'publisher' in rest.lower():
        is_book = True
    elif re.search(r'\d+\(\d+\)', rest) or 'Journal' in rest or 'Jurnal' in rest:
        is_book = False
    else:
        for k in fixes:
            if k in author_part:
                if 'publisher' in fixes[k]:
                    is_book = True
                else:
                    is_book = False
        if not rest and not is_book:
            is_book = True

    return {
        'original': ref_text,
        'authors': author_part,
        'year': year,
        'title': title,
        'rest': rest,
        'is_book': is_book
    }

parsed_refs = []
unparsed = []

for line in lines:
    if line.startswith('The above content shows') or 'http' in line and len(line) < 50:
        continue
    line = re.sub(r'https?://[^\s]+', '', line).strip()
    line = re.sub(r'www\.[^\s]+', '', line).strip()
    
    parsed = parse_reference(line)
    if parsed:
        parsed_refs.append(parsed)
    else:
        unparsed.append(line)

formatted_refs = []

for p in parsed_refs:
    auth = p['authors']
    year = p['year']
    title = p['title']
    rest = p['rest']
    is_book = p['is_book']
    
    for k, v in fixes.items():
        if k in auth:
            if 'Arikunto' in auth:
                if year == '2009' and '2009' in k:
                    rest = f"{v.get('publisher', '')}. {v.get('pages', '')}"
                    is_book = True
                elif year == '2013' and '2013' in k:
                    rest = f"{v.get('publisher', '')}. {v.get('pages', '')}"
                    is_book = True
            elif 'Sudjana' in auth and '2005' in k and year == '2005':
                rest = f"{v.get('publisher', '')}. {v.get('pages', '')}"
                is_book = True
            elif 'Sugiyono' in auth and '2013' in k and year == '2013':
                rest = f"{v.get('publisher', '')}. {v.get('pages', '')}"
                is_book = True
            else:
                if 'journal' in v:
                    rest = f"{v['journal']} {v.get('vol', '')}, {v.get('pages', rest)}"
                    is_book = False
                elif 'publisher' in v:
                    rest = f"{v['publisher']}. {v['pages']}"
                    is_book = True
                else:
                    if 'pages' in v:
                        rest += f", {v['pages']}"

    if is_book:
        words = title.split()
        title_case = ' '.join([w.capitalize() for w in words])
        p['formatted_title'] = title_case
        p['italic_title'] = True
    else:
        p['formatted_title'] = title.capitalize() if title else ''
        p['italic_title'] = False

    if not is_book:
        journal_match = re.match(r'^(.+?)(?:,\s*(\d+(?:\(\d+\))?.*?))?$', rest)
        if journal_match:
            p['journal_name'] = journal_match.group(1).strip()
            p['journal_vol'] = journal_match.group(2) if journal_match.group(2) else ''
        else:
            p['journal_name'] = rest
            p['journal_vol'] = ''
    else:
        p['rest'] = rest

    formatted_refs.append(p)

formatted_refs.sort(key=lambda x: x['authors'].lower())

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

p_title = doc.add_paragraph('DAFTAR PUSTAKA')
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.runs[0].bold = True
doc.add_paragraph('')

for ref in formatted_refs:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(24)
    
    p.add_run(f"{ref['authors']}. {ref['year']}. ")
    
    if ref['is_book']:
        p.add_run(f"{ref['formatted_title']}. ").italic = True
        p.add_run(f"{ref['rest'].rstrip('.')}.").italic = False
    else:
        p.add_run(f"{ref['formatted_title']}. ")
        p.add_run(f"{ref.get('journal_name', '')} ").italic = True
        p.add_run(f"{ref.get('journal_vol', '').rstrip('.')}.").italic = False

for u in unparsed:
    if len(u) > 10:
        p = doc.add_paragraph()
        p.add_run(f"[UNPARSED] {u}")

doc.save(r'c:\laragon\www\ult-fkip-unila\docs\skripsi\Daftar_Pustaka_Unila_2020.docx')
print('Unparsed items remaining:', len([u for u in unparsed if len(u)>10]))
