import pandas as pd
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

file_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil uji ahli validator\Rekap_Uji_Validitas_Ahli.xlsx'

validators = []

# Materi
df_materi = pd.read_excel(file_path, sheet_name='Validitas Materi')
names_materi = df_materi.columns[2:5]
for name in names_materi:
    scores = df_materi[name].iloc[0:10].tolist()
    total = sum([float(x) for x in scores])
    perc = (total / 50.0) * 100
    validators.append({
        'Nama': name.strip(),
        'Bidang': 'Ahli Materi',
        'Scores': [int(float(x)) for x in scores],
        'Total': int(total),
        'Persentase': f"{perc:.2f}%"
    })

# Media
df_media = pd.read_excel(file_path, sheet_name='Validitas Media')
names_media = df_media.columns[2:5]
for name in names_media:
    scores = df_media[name].iloc[0:10].tolist()
    total = sum([float(x) for x in scores])
    perc = (total / 50.0) * 100
    validators.append({
        'Nama': name.strip(),
        'Bidang': 'Ahli Media',
        'Scores': [int(float(x)) for x in scores],
        'Total': int(total),
        'Persentase': f"{perc:.2f}%"
    })

# Sistem
df_sistem = pd.read_excel(file_path, sheet_name='Validitas Sistem')
names_sistem = df_sistem.columns[2:5]
for name in names_sistem:
    scores = df_sistem[name].iloc[0:10].tolist()
    total = sum([float(x) for x in scores])
    perc = (total / 50.0) * 100
    validators.append({
        'Nama': name.strip(),
        'Bidang': 'Ahli Sistem',
        'Scores': [int(float(x)) for x in scores],
        'Total': int(total),
        'Persentase': f"{perc:.2f}%"
    })

def create_table_in_doc(doc_path, output_path):
    doc = docx.Document(doc_path)
    
    target_p = None
    for p in doc.paragraphs:
        if 'Lampiran 11. Matriks Hasil Skor Kuesioner Uji Kepraktisan Responden' in p.text:
            target_p = p
            break
            
    if not target_p:
        for p in doc.paragraphs:
            if 'Matriks Hasil Skor Kuesioner Uji Kepraktisan' in p.text:
                target_p = p
                break

    if target_p:
        # Create heading
        new_heading = target_p.insert_paragraph_before("Lampiran 10. Matriks Hasil Skor Kuesioner Uji Ahli Validasi")
        new_heading.style = doc.styles['Normal']
        
        # Table
        table = doc.add_table(rows=1, cols=14)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nama Validator'
        hdr_cells[1].text = 'Bidang Keahlian'
        for i in range(10):
            hdr_cells[i+2].text = f'Q{i+1}'
        hdr_cells[12].text = 'Total'
        hdr_cells[13].text = 'Persentase (%)'
        
        for c in hdr_cells:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    
        for v in validators:
            row_cells = table.add_row().cells
            row_cells[0].text = v['Nama']
            row_cells[1].text = v['Bidang']
            for i, score in enumerate(v['Scores']):
                row_cells[i+2].text = str(score)
            row_cells[12].text = str(v['Total'])
            row_cells[13].text = v['Persentase']
            
            for i in range(14):
                for p in row_cells[i].paragraphs:
                    if i >= 2:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(10)

        tbl_element = table._element
        doc._body._body.remove(tbl_element)
        target_p._p.addprevious(tbl_element)
        
        empty_p_xml = docx.oxml.OxmlElement('w:p')
        target_p._p.addprevious(empty_p_xml)
        
        doc.save(output_path)
        print(f"Added matrix to {output_path}")
    else:
        print(f"Could not find anchor paragraph in {doc_path}")

create_table_in_doc(
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx',
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
)
create_table_in_doc(
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx',
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
)
