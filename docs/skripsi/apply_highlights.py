import shutil
import sys
import docx
from docx.enum.text import WD_COLOR_INDEX
from pathlib import Path

# Set stdout to UTF-8
sys.stdout.reconfigure(encoding='utf-8')

def highlight_paragraph(p):
    """Highlight all runs in a paragraph with Yellow."""
    for r in p.runs:
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW

def highlight_table(table):
    """Highlight all text runs inside a table."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                highlight_paragraph(p)

def main():
    src_path = Path(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx")
    dest_path = Path(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx")
    
    print("Step 1: Performing binary copy from Clean.docx to Highlighted.docx...")
    if not src_path.exists():
        print(f"Error: Source file {src_path} not found!")
        return
        
    shutil.copy2(src_path, dest_path)
    print("Binary copy completed successfully.")
    
    print("\nStep 2: Loading Highlighted.docx with python-docx...")
    doc = docx.Document(dest_path)
    
    total_paragraphs = len(doc.paragraphs)
    total_tables = len(doc.tables)
    print(f"Loaded document. Paragraphs: {total_paragraphs}, Tables: {total_tables}")
    
    # Flags to track which sections we are currently inside
    in_bab3 = False
    in_bab4 = False
    in_bab5 = False
    in_daftar_pustaka = False
    in_lampiran = False
    
    in_subjek_sampling = False
    in_instrumen = False
    in_teknik_analisis = False
    
    print("\nStep 3: Processing paragraphs and applying highlights...")
    
    # We will iterate through paragraphs and set highlights based on structural sections
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        style = p.style.name
        
        # Check if the style is heading or if it's heading-like (all caps, etc.)
        is_heading = style.startswith("Heading") or (txt.isupper() and len(txt) > 3)
        
        # Section transition detection
        if "METODE PENELITIAN" in txt.upper() and is_heading:
            in_bab3 = True
            in_bab4 = False
            in_bab5 = False
            in_daftar_pustaka = False
            in_lampiran = False
            print(f"Entering BAB III at Paragraph {i}")
            continue
            
        elif "HASIL DAN PEMBAHASAN" in txt.upper() and is_heading:
            in_bab3 = False
            in_bab4 = True
            in_bab5 = False
            in_daftar_pustaka = False
            in_lampiran = False
            print(f"Entering BAB IV at Paragraph {i}")
            # We highlight heading of Bab IV as well
            highlight_paragraph(p)
            continue
            
        elif "KESIMPULAN DAN SARAN" in txt.upper() and is_heading:
            in_bab3 = False
            in_bab4 = False
            in_bab5 = True
            in_daftar_pustaka = False
            in_lampiran = False
            print(f"Entering BAB V at Paragraph {i}")
            highlight_paragraph(p)
            continue
            
        elif "DAFTAR PUSTAKA" in txt.upper() and is_heading:
            in_bab3 = False
            in_bab4 = False
            in_bab5 = False
            in_daftar_pustaka = True
            in_lampiran = False
            print(f"Entering DAFTAR PUSTAKA at Paragraph {i}")
            continue
            
        elif in_daftar_pustaka and ("VALIDATOR" in txt.upper() or "LAMPIRAN" in txt.upper() or "Ibu Etha" in txt or "Bapak Eko" in txt):
            # Transition from Daftar Pustaka to Lampiran
            in_bab3 = False
            in_bab4 = False
            in_bab5 = False
            in_daftar_pustaka = False
            in_lampiran = True
            print(f"Entering LAMPIRAN at Paragraph {i}")
            highlight_paragraph(p)
            continue
            
        # Specific sub-sections inside BAB III
        if in_bab3:
            # 1. Subjek Penelitian dan Sampling
            if "SUBJEK PENELITIAN DAN SAMPLING" in txt.upper() and is_heading:
                in_subjek_sampling = True
                in_instrumen = False
                in_teknik_analisis = False
                print(f"  Start Subjek Penelitian & Sampling at Paragraph {i}")
                
            # 2. Instrumen Penelitian
            elif "INSTRUMEN PENELITIAN" in txt.upper() and is_heading:
                in_subjek_sampling = False
                in_instrumen = True
                in_teknik_analisis = False
                print(f"  Start Instrumen Penelitian at Paragraph {i}")
                
            # 3. Teknik Pengumpulan Data
            elif "TEKNIK PENGUMPULAN DATA" in txt.upper() and is_heading:
                in_subjek_sampling = False
                in_instrumen = False
                in_teknik_analisis = False
                
            # 4. Teknik Analisis Data
            elif "TEKNIK ANALISIS DATA" in txt.upper() and is_heading:
                in_subjek_sampling = False
                in_instrumen = False
                # We start highlighting from Aiken's V formula which is inside this section
                in_teknik_analisis = True
                print(f"  Start Teknik Analisis Data (for highlighting) at Paragraph {i}")
                
            # Highlight execution inside BAB III
            if in_subjek_sampling:
                highlight_paragraph(p)
            elif in_instrumen:
                # Inside Instrumen, we highlight table captions, table titles, list of aspects, and descriptions of validator and kepraktisan questionnaires
                keywords_instrumen = ["kisi-kisi", "validasi", "ahli", "materi", "media", "sistem", "kepraktisan", "tabel 5", "tabel 6", "tabel 7", "tabel 8", "angket", "skala likert"]
                if any(kw in txt.lower() for kw in keywords_instrumen) or p.style.name.startswith("Caption") or p.style.name.startswith("Heading"):
                    highlight_paragraph(p)
            elif in_teknik_analisis:
                # Inside Teknik Analisis Data, we highlight starting from Aiken's V discussion to the end of BAB III
                keywords_analisis = ["aiken", "validitas", "rekapitulasi", "tabel 10", "tabel 12", "kriteria", "penafsiran", "rumus", "likert", "12 butir", "18 responden", "skor", "v ="]
                if any(kw in txt.lower() for kw in keywords_analisis) or p.style.name.startswith("Caption") or p.style.name.startswith("Heading") or i >= 681:
                    highlight_paragraph(p)
                    
        # Highlight execution inside BAB IV
        elif in_bab4:
            highlight_paragraph(p)
            
        # Highlight execution inside BAB V
        elif in_bab5:
            highlight_paragraph(p)
            
        # Highlight specific reference items inside DAFTAR PUSTAKA
        elif in_daftar_pustaka:
            if "aiken, l." in txt.lower() or "sugiyono" in txt.lower():
                highlight_paragraph(p)
                
        # Highlight execution inside LAMPIRAN
        elif in_lampiran:
            highlight_paragraph(p)

    print("\nStep 4: Processing tables and applying highlights...")
    
    # Highlight all tables that belong to BAB III (specific), BAB IV, or LAMPIRAN
    for idx, table in enumerate(doc.tables):
        # We can extract all text inside the table to search for keywords
        table_text = ""
        for row in table.rows:
            for cell in row.cells:
                table_text += cell.text + " "
                
        table_text_lower = table_text.lower()
        
        # Check if the table needs highlighting
        should_highlight = False
        
        # 1. LAMPIRAN tables (have respondents, UI/UX documentation, etc.)
        # These are usually at the end of the document, containing things like "Responden 1", "Sangat Praktis", etc.
        if "responden" in table_text_lower and ("matriks" in table_text_lower or "skor" in table_text_lower or "distribusi" in table_text_lower or "kategori" in table_text_lower):
            should_highlight = True
            loc = "LAMPIRAN"
        elif "tangkapan layar" in table_text_lower or "figma" in table_text_lower or "gambar d.1" in table_text_lower or "gambar d.2" in table_text_lower:
            should_highlight = True
            loc = "LAMPIRAN UI/UX"
        # 2. BAB IV tables (Tabel 4.1 Rekapitulasi Hasil Validasi Ahli & Tabel 4.2 Rekapitulasi Uji Kepraktisan)
        elif "rekapitulasi hasil validasi ahli" in table_text_lower or ("ahli materi" in table_text_lower and "ahli media" in table_text_lower and "ahli sistem" in table_text_lower and "91,95%" in table_text_lower):
            should_highlight = True
            loc = "BAB IV (Tabel 4.1)"
        elif "kategori kepraktisan (pilihan bagian f)" in table_text_lower and "jumlah responden" in table_text_lower:
            should_highlight = True
            loc = "Tabel Rekapitulasi Kepraktisan"
        # 3. BAB III tables
        elif "kisi-kisi" in table_text_lower or "aspek penilaian" in table_text_lower or "butir" in table_text_lower or "likert" in table_text_lower or "kriteria penafsiran" in table_text_lower:
            should_highlight = True
            loc = "BAB III"
        elif "subjek penelitian" in table_text_lower or "validator" in table_text_lower or "staf ult" in table_text_lower:
            should_highlight = True
            loc = "BAB III"
            
        if should_highlight:
            print(f"Highlighting Table {idx} located in {loc}...")
            highlight_table(table)
            
    print("\nStep 5: Saving the highlighted document...")
    doc.save(dest_path)
    print(f"Highlighted document saved successfully to: {dest_path}")

if __name__ == "__main__":
    main()
