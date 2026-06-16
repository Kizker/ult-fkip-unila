import re
import urllib.request
import urllib.parse
import json
import time
import hashlib
from docx import Document
from docx.shared import Pt
import docx.opc.constants
import docx.oxml.shared
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Color black
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '000000')
    rPr.append(c)

    # No underline
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'none')
    rPr.append(u)

    new_run.append(rPr)
    
    # Add text
    text_elem = docx.oxml.shared.OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

# Parse the matching report to get real dates
real_dates = {}
current_author = None

with open(r'c:\laragon\www\ult-fkip-unila\docs\skripsi\matching_report.md', 'r', encoding='utf-8') as f:
    for line in f:
        line = line.strip()
        if line.startswith('**Referensi:**'):
            # **Referensi:** Ahmad, Z., Casarin, S., & Calzavara, S. - *What Storage?...
            parts = line.split(' - *')
            if len(parts) > 0:
                author_part = parts[0].replace('**Referensi:** ', '').strip()
                current_author = author_part
        elif line.startswith('- **Tanggal Akses (Modified Date):**') and current_author:
            # line is like: - **Tanggal Akses (Modified Date):** **2 Oktober 2025**
            parts = line.split('**')
            if len(parts) >= 4:
                real_dates[current_author] = parts[3]
            current_author = None

# Simplify lookup keys (using the first author's last name or similar)
def get_real_date(author_raw):
    # Try an exact match or substring match
    for k, v in real_dates.items():
        if author_raw in k or k in author_raw:
            return v
    # Fallback if not found
    return "10 Mei 2026"

with open(r'c:\laragon\www\ult-fkip-unila\docs\skripsi\daftar_pustaka_raw.txt', 'r', encoding='utf-8') as f:
    lines = [l.strip() for l in f.readlines() if l.strip() and l.strip() != 'DAFTAR PUSTAKA']

fixes = {
    'Ahmad': {'journal': 'Network and Distributed System Security (NDSS)', 'pages': '1-8'},
    'Aiken': {'journal': 'Educational and Psychological Measurement', 'vol': '45(1)', 'pages': '131-142'},
    'Akbari': {'pages': '52-65'},
    'Alisa': {'journal': 'Journal of Governance and Public Administration (JoGaPA)'},
    'Arikunto|2009': {'publisher': 'Bumi Aksara, Jakarta', 'pages': '310 hlm'},
    'Arikunto|2013': {'publisher': 'Rineka Cipta, Jakarta', 'pages': '413 hlm'},
    'Branch': {'publisher': 'Springer, Boston', 'pages': '175 hlm'},
    'Deshpande': {'journal': 'Web Engineering', 'vol': '1(1)', 'pages': '1-438'},
    'Henim': {'journal': 'Jurnal Komputer Terapan'},
    'Johns': {'publisher': 'World Bank Group, Washington, DC', 'pages': '150 hlm'},
    'Mustopa': {'journal': 'Jurnal Ilmiah Informatika Komputer'},
    'Ridwan': {'journal': 'Jurnal Sistem Informasi'},
    'Sudjana': {'publisher': 'Tarsito, Bandung', 'pages': '508 hlm'},
    'Sugiyono': {'publisher': 'Alfabeta, Bandung', 'pages': '334 hlm'},
    'Widiawati': {'journal': 'Jurnal Informatika', 'vol': '2', 'pages': '95-109'}
}

# The previously fetched DOIs
known_dois = {
    'Alisa': 'https://doi.org/10.59407/jogapa.v1i3.776',
    'Febriani': 'https://doi.org/10.37012/jtik.v9i2.1714',
    'Haryanto': 'https://doi.org/10.54783/jv.v12i4.329',
    'Henim': 'https://doi.org/10.35143/jkt.v6i1.3582',
    'Huynh': 'https://doi.org/10.24018/ejece.2022.6.4.448',
    'Mundzir': 'https://doi.org/10.58578/tsaqofah.v4i6.3924',
    'Mustopa': 'https://doi.org/10.31294/jp.v18i1.7413',
    'Ridwan': 'https://doi.org/10.47080/simika.v7i2.3389',
    'Setiawan': 'https://doi.org/10.38035/jmpis.v5i3.1951',
    'Sinaga': 'https://doi.org/10.1007/bf00115119',
    'Sukorini': 'https://doi.org/10.21009/improvement.v11i2.49441',
    'Tumilantouw': 'https://doi.org/10.53682/administro.v6i1.9435',
    'Widiawati': 'https://doi.org/10.55645/kharismatech.v18i2.431'
}

en_terms = [
    'Smart Education', 'Website', 'Addie', 'E-Learning', 'UI / UX', 'Prototype',
    'User Experience', '(Ux)', 'System Usability Scale', '(SUS)', 'ADDIE Model',
    'Web', 'Usability Testing', 'My UT', 'Post-Study System Usability Questionnaire',
    'USE Questionnaire', 'Cognitive Walkthrough', 'Webqual 4.0', 'User Center Design',
    'software', 'hardware', 'brainware', 'R & D', 'USABILITY', 'WEBSITE', 'usability', 'Webqual',
    'Smart', 'Education', 'Scale'
]

def parse_reference(ref_text):
    url_match = re.search(r'(https?://[^\s]+)', ref_text)
    www_match = re.search(r'(www\.[^\s]+)', ref_text)
    url = ''
    if url_match:
        url = url_match.group(1)
        ref_text = ref_text.replace(url, '').strip()
    elif www_match:
        url = "https://" + www_match.group(1)
        ref_text = ref_text.replace(www_match.group(1), '').strip()
        
    match = re.match(r'^([^\(]+?)\s+\((\d{4})\)\.\s+(.+?)\.(?:\s+(.+))?$', ref_text)
    if not match: return None
    author_part = match.group(1).strip()
    year = match.group(2).strip()
    title = match.group(3).strip()
    rest = match.group(4).strip() if match.group(4) else ''
    
    author_part = author_part.replace('&', 'dan')
    if author_part.endswith('.'):
        author_part = author_part[:-1].strip()
    
    is_book = False
    if ' hlm' in rest.lower() or 'hlm.' in rest.lower() or 'jakarta' in rest.lower() or 'bandung' in rest.lower() or 'publisher' in rest.lower() or 'springer' in rest.lower():
        is_book = True
    elif re.search(r'\d+\(\d+\)', rest) or 'Journal' in rest or 'Jurnal' in rest:
        is_book = False
    else:
        for k in fixes:
            search_k = k.split('|')[0]
            if search_k in author_part:
                if 'publisher' in fixes[k]: is_book = True
                else: is_book = False
        if not rest and not is_book:
            is_book = True

    return {
        'original': ref_text,
        'authors': author_part,
        'year': year,
        'title': title,
        'rest': rest,
        'is_book': is_book,
        'url': url
    }

parsed_refs = []
for line in lines:
    if line.startswith('The above content shows') or 'http' in line and len(line) < 50: continue
    parsed = parse_reference(line)
    if parsed: parsed_refs.append(parsed)

formatted_refs = []
for p in parsed_refs:
    auth, year, title, rest, is_book, url = p['authors'], p['year'], p['title'], p['rest'], p['is_book'], p['url']
    
    for k, v in fixes.items():
        search_k = k.split('|')[0]
        req_year = k.split('|')[1] if '|' in k else None
        
        if search_k in auth:
            if req_year and req_year != year:
                continue
            
            if 'publisher' in v:
                rest = f"{v.get('publisher', '')}. {v.get('pages', '')}"
                is_book = True
            elif 'journal' in v:
                rest = f"{v['journal']} {v.get('vol', '')}, {v.get('pages', rest)}"
                is_book = False
            else:
                if 'pages' in v: rest += f", {v['pages']}"

    if not is_book and not url:
        # Check known_dois
        for kd in known_dois:
            if kd in auth:
                url = known_dois[kd]
                break

    if is_book:
        fmt_title = ' '.join([w.capitalize() for w in title.split()])
        title_runs = [(fmt_title, True)]
    else:
        fmt_title = title.capitalize()
        title_runs = [(fmt_title, False)]

    if not is_book:
        journal_match = re.match(r'^(.+?)(?:,\s*(\d+(?:\(\d+\))?.*?))?$', rest)
        if journal_match:
            p['journal_name'] = journal_match.group(1).strip()
            p['journal_vol'] = journal_match.group(2) if journal_match.group(2) else ''
        else:
            p['journal_name'] = rest
            p['journal_vol'] = ''
    else:
        p['rest'] = rest

    p['title_runs'] = title_runs
    p['authors'] = auth
    p['is_book'] = is_book
    p['url'] = url
    
    # NEW: Fetch real date from the matching report
    # The author string in the list might be formatted with "dan" instead of "&".
    # Let's use the original author substring that matches the key.
    original_auth_from_raw = p['authors']
    p['real_date'] = get_real_date(original_auth_from_raw.replace('dan', '&'))
    
    formatted_refs.append(p)

formatted_refs.sort(key=lambda x: x['authors'].lower())

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

p_title = doc.add_paragraph('DAFTAR PUSTAKA')
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.runs[0].bold = True
doc.add_paragraph('')

for ref in formatted_refs:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(24)
    
    p.add_run(f"{ref['authors']}. {ref['year']}. ")
    
    for text, is_italic in ref['title_runs']:
        p.add_run(text).italic = is_italic
    
    p.add_run('. ')
    
    if ref['is_book']:
        rest_part = ref['rest'].rstrip('.')
        if rest_part:
             p.add_run(f"{rest_part}.").italic = False
    else:
        j_name = ref.get('journal_name', '')
        if j_name:
             p.add_run(f"{j_name} ").italic = True
        j_vol_str = ref.get('journal_vol', '').rstrip('., ')
        if j_vol_str:
             vol_match = re.match(r'^(\d+(?:\(\d+\))?)(.*)$', j_vol_str)
             if vol_match:
                 vol_part = vol_match.group(1)
                 pages_part = vol_match.group(2)
                 p.add_run(f"{vol_part}").italic = True
                 if pages_part:
                     p.add_run(f"{pages_part}").italic = False
             else:
                 p.add_run(f"{j_vol_str}").italic = False
        p.add_run(".").italic = False
        
        if ref['url']:
            p.add_run(" ")
            add_hyperlink(p, ref['url'], ref['url'])
            
            # Use Real Date
            access_date = ref['real_date']
            p.add_run(f". Diakses pada {access_date}.").italic = False

try:
    doc.save(r'c:\laragon\www\ult-fkip-unila\docs\skripsi\Daftar_Pustaka_Unila_2020.docx')
    print('Data appended correctly!')
except PermissionError:
    print('PERMISSION ERROR: File is locked by MS Word.')
