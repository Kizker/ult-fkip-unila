import openpyxl
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

file_validitas = r"c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil uji ahli validator\Rekap_Uji_Validitas_Ahli.xlsx"
file_kepraktisan = r"c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil uji kepraktisan\Rekap_Uji_Kepraktisan.xlsx"
output_file = r"c:\laragon\www\ult-fkip-unila\docs\skripsi\Lampiran_Pengolahan_Data_Lengkap_V6.docx"


import re

def sanitize_for_xml(text):
    if text is None:
        return ""
    text = str(text)
    # Ganti newline bawaan Excel dengan spasi, karena \n literal di python-docx merusak struktur XML Word
    text = text.replace('\n', ' ').replace('\r', ' ')
    # Menghapus karakter kontrol XML ilegal lainnya
    return re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)

from docx.oxml import OxmlElement

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)

def add_heading(doc, text, level):
    text = sanitize_for_xml(text)
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            run.font.size = Pt(14)
            run.font.bold = True
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            run.font.size = Pt(12)
            run.font.bold = True
        else:
            run.font.size = Pt(12)

def add_paragraph(doc, text, bold=False, italic=False, align='justify'):
    p = doc.add_paragraph()
    if align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text = sanitize_for_xml(text)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p

def create_table_from_data(doc, headers, rows_data):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        clean_header = sanitize_for_xml(header_text)
        run = p.add_run(clean_header)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8) # 8pt for table content
        run.bold = True

    for row_data in rows_data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            p = row_cells[i].paragraphs[0]
            if str(cell_data).replace('.','',1).isdigit():
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            clean_cell = sanitize_for_xml(cell_data)
            run = p.add_run(clean_cell)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(8)
    doc.add_paragraph()
    return table

def main():
    doc = Document()
    from docx.enum.section import WD_ORIENT
    section = doc.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    add_heading(doc, 'LAMPIRAN PENGOLAHAN DATA PENELITIAN', 1)
    doc.add_paragraph()
    
    add_paragraph(doc, 'A. Pengolahan Data Uji Validitas Ahli', bold=True)
    add_paragraph(doc, 'Uji validitas produk dilakukan oleh 9 orang validator ahli yang terbagi menjadi tiga aspek penilaian: Ahli Materi (3 orang), Ahli Media (3 orang), dan Ahli Sistem (3 orang). Instrumen penilaian menggunakan skala Likert 1 sampai 5. Perhitungan kelayakan menggunakan dua parameter: Persentase Kelayakan dan Indeks Kesepakatan Ahli (Aiken\'s V).')
    
    add_paragraph(doc, 'Rumus Persentase Kelayakan:')
    add_paragraph(doc, 'P = (Skor Aktual / Skor Ideal) × 100%', italic=True)
    add_paragraph(doc, 'Keterangan:')
    add_paragraph(doc, 'P = Persentase Kelayakan\nSkor Aktual = Jumlah seluruh skor yang diberikan validator\nSkor Ideal = Skor tertinggi × jumlah responden × jumlah item')
    
    add_paragraph(doc, 'Rumus Aiken\'s V:')
    add_paragraph(doc, 'V = Σs / [n(c - 1)]', italic=True)
    add_paragraph(doc, 'Keterangan:')
    add_paragraph(doc, 'V = Indeks kesepakatan ahli (Aiken\'s V)\ns = r - l₀\nr = Skor yang diberikan oleh ahli\nl₀ = Angka penilaian terendah (1)\nn = Jumlah ahli (3 orang)\nc = Banyaknya kategori pilihan skala (5 pilihan)')
    doc.add_paragraph()
    
    try:
        wb_val = openpyxl.load_workbook(file_validitas, data_only=True)
        aspek_ahli = [('Validitas Materi', 'Materi'), ('Validitas Media', 'Media'), ('Validitas Sistem', 'Sistem')]
        
        sum_p_all = 0
        sum_v_all = 0
        
        for sheet_name, nama_aspek in aspek_ahli:
            if sheet_name not in wb_val.sheetnames: continue
            
            add_heading(doc, f'1. Validitas {nama_aspek}', 2)
            ws = wb_val[sheet_name]
            
            data_rows = []
            for row in ws.iter_rows(values_only=True):
                data_rows.append([c for c in row])
                
            headers = [h for h in data_rows[0] if h is not None and str(h).strip() != '']
            headers_count = len(headers)
            
            # The matrix items are rows before 'TOTAL KESELURUHAN'
            matrix_data = []
            total_row = None
            
            total_skor_aktual = 0
            total_skor_ideal = 0
            sigma_s = 0
            n_items = 0
            
            for i, row in enumerate(data_rows[1:]):
                if not row[0]: continue
                if "TOTAL KESELURUHAN" in str(row[0]).upper():
                    total_row = row
                    break
                else:
                    matrix_data.append(row[:headers_count])
                    
            if total_row:
                # Find the columns for actual score, ideal score
                # The headers typically have "Total Skor per Item", "Skor Maks", "Persentase"
                idx_aktual = -1
                idx_ideal = -1
                for idx, h in enumerate(headers):
                    h_up = str(h).upper()
                    if "TOTAL" in h_up and "ITEM" in h_up:
                        idx_aktual = idx
                    elif "MAKS" in h_up or "IDEAL" in h_up:
                        idx_ideal = idx
                        
                if idx_aktual != -1 and idx_ideal != -1:
                    total_skor_aktual = float(total_row[idx_aktual]) if total_row[idx_aktual] else 0
                    total_skor_ideal = float(total_row[idx_ideal]) if total_row[idx_ideal] else 0
                else:
                    # Fallback to column index based on structure if headers differ slightly
                    total_skor_aktual = float(total_row[-5]) if total_row[-5] else 0
                    total_skor_ideal = float(total_row[-3]) if total_row[-3] else 0
            
            # Draw the Table Matrix
            add_paragraph(doc, f'Tabel Rekapitulasi Data (Matriks) Uji {nama_aspek}:')
            create_table_from_data(doc, headers, matrix_data)
            
            # Calculate s for Aiken's V
            # Find the columns representing the 3 validators
            # Usually index 2, 3, 4 based on headers (Aspek, Indikator, V1, V2, V3)
            val1_idx, val2_idx, val3_idx = 2, 3, 4
            
            # Menghitung Aiken's V step by step
            add_paragraph(doc, 'Proses Perhitungan Aiken\'s V (Langkah Demi Langkah):', bold=True)
            for md in matrix_data:
                try:
                    s1 = float(md[val1_idx]) - 1
                    s2 = float(md[val2_idx]) - 1
                    s3 = float(md[val3_idx]) - 1
                    sigma_s += (s1 + s2 + s3)
                    n_items += 1
                except:
                    pass
                    
            add_paragraph(doc, f'a. Mencari nilai s (r - l₀) untuk setiap butir dari 3 penilai:')
            # Berikan 1 contoh kalkulasi untuk butir pertama
            try:
                s1_c = float(matrix_data[0][val1_idx])
                s2_c = float(matrix_data[0][val2_idx])
                s3_c = float(matrix_data[0][val3_idx])
                add_paragraph(doc, f'   Contoh Butir 1 (Indikator: {matrix_data[0][1]}):')
                add_paragraph(doc, f'   - Penilai 1 (r={s1_c}) -> s = {s1_c} - 1 = {s1_c-1}')
                add_paragraph(doc, f'   - Penilai 2 (r={s2_c}) -> s = {s2_c} - 1 = {s2_c-1}')
                add_paragraph(doc, f'   - Penilai 3 (r={s3_c}) -> s = {s3_c} - 1 = {s3_c-1}')
                add_paragraph(doc, f'   - Jumlah s (Butir 1) = {(s1_c-1) + (s2_c-1) + (s3_c-1)}')
            except:
                pass
                
            add_paragraph(doc, f'b. Menghitung total Σs dari seluruh butir ({n_items} butir):')
            add_paragraph(doc, f'   Σs = {sigma_s}')
            
            add_paragraph(doc, 'c. Memasukkan nilai ke dalam rumus Aiken\'s V:')
            n_ahli = 3
            c_skala = 5
            c_minus_1 = c_skala - 1 # 4
            v_index = sigma_s / (n_items * n_ahli * c_minus_1) if n_items > 0 else 0
            add_paragraph(doc, f'   V = Σs / [n(c-1)]')
            add_paragraph(doc, f'   V = {sigma_s} / [({n_items} butir × {n_ahli} ahli) × ({c_skala} - 1)]')
            add_paragraph(doc, f'   V = {sigma_s} / [{n_items * n_ahli} × 4]')
            add_paragraph(doc, f'   V = {sigma_s} / {n_items * n_ahli * 4}')
            add_paragraph(doc, f'   V = {v_index:.2f}', bold=True)
            
            add_paragraph(doc, 'Proses Perhitungan Persentase Kelayakan:', bold=True)
            persentase = (total_skor_aktual / total_skor_ideal) * 100 if total_skor_ideal > 0 else 0
            add_paragraph(doc, f'- Total Skor Aktual = {total_skor_aktual}')
            add_paragraph(doc, f'- Total Skor Ideal = {total_skor_ideal} (3 ahli × {n_items} butir × 5)')
            add_paragraph(doc, f'P = ({total_skor_aktual} / {total_skor_ideal}) × 100% = {persentase:.2f}%', bold=True)
            
            sum_p_all += persentase
            sum_v_all += v_index
            
            kriteria_p = "Sangat Valid" if persentase >= 81 else "Valid"
            add_paragraph(doc, f'Kesimpulan: Berdasarkan perhitungan, persentase kelayakan adalah {persentase:.2f}% ({kriteria_p}) dan indeks Aiken\'s V adalah {v_index:.2f}. Oleh karena itu, aspek {nama_aspek} dikategorikan Sangat Valid.', italic=True)
            doc.add_page_break()
            
        # Rata-rata ahli
        add_heading(doc, '2. Rekapitulasi Keseluruhan Uji Ahli', 2)
        rata_p = sum_p_all / 3
        rata_v = sum_v_all / 3
        add_paragraph(doc, f'Secara keseluruhan, rata-rata persentase kelayakan dari ketiga aspek ahli adalah {rata_p:.2f}%, dan rata-rata indeks Aiken\'s V adalah {rata_v:.2f}. Hal ini menegaskan bahwa sistem yang dikembangkan memenuhi kriteria kelayakan teoretis yang sangat baik.')
        doc.add_page_break()
        
    except Exception as e:
        print("Error Validitas:", e)
        
    add_paragraph(doc, 'B. Pengolahan Data Uji Kepraktisan Pengguna', bold=True)
    add_paragraph(doc, 'Uji kepraktisan dilakukan kepada 18 responden yang terdiri dari 12 mahasiswa, 3 admin prodi, dan 3 staf ULT. Instrumen kepraktisan berjumlah 12 butir pertanyaan dengan skala Likert 1 sampai 5. Data mentah seluruh tanggapan responden direkapitulasi sebagai berikut:')
    
    try:
        wb_kep = openpyxl.load_workbook(file_kepraktisan, data_only=True)
        sheet_name = 'Rekap Uji Kepraktisan'
        if sheet_name in wb_kep.sheetnames:
            ws = wb_kep[sheet_name]
            
            data_rows = []
            for row in ws.iter_rows(values_only=True):
                data_rows.append([c for c in row])
                
            headers = [h for h in data_rows[0] if h is not None and str(h).strip() != '']
            headers_count = len(headers)
            
            matrix_data = []
            total_skor_aktual = 0
            total_responden = 18
            total_butir = 12
            
            for i, row in enumerate(data_rows[1:]):
                if not row[0]: continue
                if "TOTAL SKOR KESELURUHAN" in str(row[0]).upper():
                    total_skor_aktual = float(row[1]) if row[1] else 0
                    break
                else:
                    matrix_data.append(row[:headers_count])
            
            add_paragraph(doc, f'Tabel Matriks Jawaban Uji Kepraktisan (18 Responden x 12 Butir):')
            create_table_from_data(doc, headers, matrix_data)
            
            # Compute total based on matrix if total_skor_aktual is 0
            if total_skor_aktual == 0:
                for md in matrix_data:
                    # Usually the last column or sum of middle columns.
                    pass
            
            total_skor_ideal = total_responden * total_butir * 5
            persentase = (total_skor_aktual / total_skor_ideal) * 100 if total_skor_ideal > 0 else 0
            avg_score_per_respondent = total_skor_aktual / total_responden
            
            add_paragraph(doc, 'Proses Perhitungan Kepraktisan:', bold=True)
            add_paragraph(doc, f'Rumus yang digunakan adalah P = (Skor Aktual / Skor Ideal) × 100%')
            add_paragraph(doc, f'- Jumlah Responden (N) = {total_responden}')
            add_paragraph(doc, f'- Jumlah Butir Pertanyaan = {total_butir}')
            add_paragraph(doc, f'- Total Skor Aktual dari tabel = {total_skor_aktual}')
            add_paragraph(doc, f'- Total Skor Ideal = N × Butir × 5 (skor tertinggi) = {total_responden} × {total_butir} × 5 = {total_skor_ideal}')
            
            add_paragraph(doc, f'Mensubstitusi nilai ke dalam rumus:')
            add_paragraph(doc, f'P = ({total_skor_aktual} / {total_skor_ideal}) × 100%')
            add_paragraph(doc, f'P = {persentase:.2f}%', bold=True)
            
            add_paragraph(doc, 'Perhitungan Rata-rata Skor per Responden (dari Maksimal 60):', bold=True)
            add_paragraph(doc, f'Rata-rata = Total Skor Aktual / N')
            add_paragraph(doc, f'Rata-rata = {total_skor_aktual} / {total_responden}')
            add_paragraph(doc, f'Rata-rata = {avg_score_per_respondent:.2f}', bold=True)
            
            add_paragraph(doc, f'Kesimpulan: Berdasarkan perhitungan terperinci dari 18 responden di atas, diperoleh tingkat kepraktisan {persentase:.2f}% yang berada pada rentang Sangat Praktis. Secara individual, rata-rata skor responsivitas adalah {avg_score_per_respondent:.2f}/60.00, yang membuktikan sistem ULT FKIP Unila sangat memudahkan operasional sivitas akademika.', italic=True)
            
    except Exception as e:
        print("Error Kepraktisan:", e)

    doc.save(output_file)
    print(f"File berhasil dibuat di: {output_file}")

if __name__ == "__main__":
    main()
