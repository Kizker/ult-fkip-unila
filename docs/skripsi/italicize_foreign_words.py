import re
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

def process_paragraph(p, pattern, highlight=False):
    modified = False
    
    # We will use a while loop that restarts if a run is split
    changed_in_pass = True
    while changed_in_pass:
        changed_in_pass = False
        
        # Re-fetch runs because we modify them
        runs = p.runs
        for i in range(len(runs)):
            run = runs[i]
            
            # Skip if already italic, or has no text
            if run.italic or not run.text:
                continue
                
            text = run.text
            match = pattern.search(text)
            
            if match:
                modified = True
                changed_in_pass = True
                
                start, end = match.span()
                before_text = text[:start]
                word_text = text[start:end]
                after_text = text[end:]
                
                # Replace text in current run
                run.text = before_text
                
                # Insert new run for word
                word_run = p.add_run(word_text)
                word_run.font.name = run.font.name
                word_run.font.size = run.font.size
                word_run.bold = run.bold
                word_run.underline = run.underline
                if run.font.color.rgb:
                    word_run.font.color.rgb = run.font.color.rgb
                
                word_run.italic = True
                if highlight:
                    word_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                
                # Insert new run for after text
                after_run = p.add_run(after_text)
                after_run.font.name = run.font.name
                after_run.font.size = run.font.size
                after_run.bold = run.bold
                after_run.underline = run.underline
                if run.font.color.rgb:
                    after_run.font.color.rgb = run.font.color.rgb
                after_run.italic = run.italic
                
                # Move elements in XML
                p_elem = p._element
                try:
                    run_idx = p_elem.index(run._element)
                    p_elem.insert(run_idx + 1, word_run._element)
                    p_elem.insert(run_idx + 2, after_run._element)
                except ValueError:
                    # element not found in direct children, could be inside an hyperlink or something.
                    # python-docx p.runs flattens runs inside hyperlinks!
                    # If it's inside a hyperlink, its parent is the hyperlink.
                    parent = run._element.getparent()
                    run_idx = parent.index(run._element)
                    parent.insert(run_idx + 1, word_run._element)
                    parent.insert(run_idx + 2, after_run._element)
                
                # Restart the loop
                break
                
    return modified

def main():
    print("Loading foreign words...")
    with open("foreign_words_to_italicize.txt", "r", encoding="utf-8") as f:
        target_words = [line.strip() for line in f if line.strip()]
        
    # Sort by length descending to match longest words first (e.g., 'digitalization' before 'digital')
    target_words = sorted(target_words, key=len, reverse=True)
    
    escaped_words = [re.escape(w) for w in target_words]
    pattern = re.compile(r'\b(' + '|'.join(escaped_words) + r')\b', re.IGNORECASE)
    
    input_file = r"c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx"
    output_clean = r"c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx"
    output_highlight = r"c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx"
    
    print("Processing Clean Document...")
    doc_clean = Document(input_file)
    
    mod_count = 0
    # Process paragraphs in body
    for p in doc_clean.paragraphs:
        if process_paragraph(p, pattern, highlight=False):
            mod_count += 1
            
    # Process paragraphs in tables
    for table in doc_clean.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if process_paragraph(p, pattern, highlight=False):
                        mod_count += 1
                        
    print("Processing Highlighted Document...")
    doc_hl = Document(input_file)
    
    # Process paragraphs in body
    for p in doc_hl.paragraphs:
        process_paragraph(p, pattern, highlight=True)
            
    # Process paragraphs in tables
    for table in doc_hl.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p, pattern, highlight=True)
    
    print(f"Modified {mod_count} paragraphs/cells in Document.")
    doc_clean.save(output_clean)
    doc_hl.save(output_highlight)
    print("Saved Clean and Highlighted Documents.")

if __name__ == "__main__":
    main()
