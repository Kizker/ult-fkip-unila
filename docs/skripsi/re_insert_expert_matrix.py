import pandas as pd
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

file_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil uji ahli validator\Rekap_Uji_Validitas_Ahli.xlsx'
validators = []

for sheet in ['Validitas Materi', 'Validitas Media', 'Validitas Sistem']:
    df = pd.read_excel(file_path, sheet_name=sheet)
    names = df.columns[2:5]
    for name in names:
        scores = df[name].iloc[0:10].tolist()
        total = sum([float(x) for x in scores])
        perc = (total / 50.0) * 100
        validators.append({
            'Nama': name.strip(),
            'Bidang': sheet.replace('Validitas ', 'Ahli '),
            'Scores': [int(float(x)) for x in scores],
            'Total': int(total),
            'Persentase': f"{perc:.2f}%"
        })

def fix_doc(doc_path, output_path):
    doc = docx.Document(doc_path)
    
    # 1. Clean up previous broken insertions
    elements_to_remove = []
    found_broken = False
    for i, el in enumerate(doc._body._body):
        if el.tag.endswith('p'):
            p = docx.text.paragraph.Paragraph(el, doc)
            if 'Lampiran 10. Matriks Hasil Skor Kuesioner Uji Ahli Validasi' in p.text:
                elements_to_remove.append(el)
                found_broken = True
                # also remove the next element if it is a table
                if i+1 < len(doc._body._body) and doc._body._body[i+1].tag.endswith('tbl'):
                    elements_to_remove.append(doc._body._body[i+1])
                    
    for el in elements_to_remove:
        try:
            doc._body._body.remove(el)
        except:
            pass

    # 2. Find target
    target_p = None
    for p in doc.paragraphs:
        if 'Lampiran 11. Matriks Hasil Skor Kuesioner Uji Kepraktisan Responden' in p.text:
            target_p = p
            break
            
    if target_p:
        # Create heading
        new_heading = target_p.insert_paragraph_before("Lampiran 10. Matriks Hasil Skor Kuesioner Uji Ahli Validasi")
        new_heading.style = doc.styles['Normal']
        
        # Create table xml element directly without adding it to the end of the document first
        # Actually, let's create it at the end, then move it. DO NOT call remove().
        table = doc.add_table(rows=1, cols=14)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nama Validator'
        hdr_cells[1].text = 'Bidang Keahlian'
        for i in range(10):
            hdr_cells[i+2].text = f'Q{i+1}'
        hdr_cells[12].text = 'Total'
        hdr_cells[13].text = 'Persentase (%)'
        
        for c in hdr_cells:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    
        for v in validators:
            row_cells = table.add_row().cells
            row_cells[0].text = v['Nama']
            row_cells[1].text = v['Bidang']
            for i, score in enumerate(v['Scores']):
                row_cells[i+2].text = str(score)
            row_cells[12].text = str(v['Total'])
            row_cells[13].text = v['Persentase']
            
            for i in range(14):
                for p in row_cells[i].paragraphs:
                    if i >= 2:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(10)

        # Move table
        target_p._p.addprevious(table._element)
        
        # Add spacing
        empty_p_xml = docx.oxml.OxmlElement('w:p')
        target_p._p.addprevious(empty_p_xml)
        
        doc.save(output_path)
        print(f"Fixed and added matrix to {output_path}")

fix_doc(
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx',
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
)
fix_doc(
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx',
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
)
