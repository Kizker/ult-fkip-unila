import sys
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# Set stdout to UTF-8
sys.stdout.reconfigure(encoding='utf-8')

def format_paragraph(p, doc, style_name="Normal", alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing_before=0, spacing_after=0, line_spacing=1.5):
    try:
        p.style = doc.styles[style_name]
    except Exception:
        pass
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(spacing_before)
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = line_spacing

def add_run_formatted(p, text, bold=False, italic=False, font_size=12):
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(font_size)
    r.bold = bold
    r.italic = italic
    return r

def main():
    path = Path(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx")
    if not path.exists():
        print("Error: Clean.docx not found.")
        return
        
    doc = docx.Document(path)
    body = doc.element.body
    
    # Step 1: Find indexes of "Kesimpulan" and "B. Saran"
    kesimpulan_idx = -1
    saran_idx = -1
    
    print("Locating Kesimpulan and B. Saran headings...")
    for idx, child in enumerate(body.getchildren()):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = docx.text.paragraph.Paragraph(child, doc)
            txt = p.text.strip()
            if txt == "Kesimpulan" and p.style.name.startswith("Heading"):
                kesimpulan_idx = idx
            elif txt == "B. Saran" and p.style.name.startswith("Heading"):
                saran_idx = idx
                break
                
    if kesimpulan_idx == -1 or saran_idx == -1:
        print(f"Error: Headings not found! Kesimpulan: {kesimpulan_idx}, Saran: {saran_idx}")
        return
        
    print(f"Found Kesimpulan heading at element index {kesimpulan_idx}")
    print(f"Found B. Saran heading at element index {saran_idx}")
    
    # Step 2: Delete old kesimpulan paragraphs between kesimpulan_idx and saran_idx
    # Since we are deleting elements, indexes of elements after the deleted ones will shift.
    # We repeatedly delete the element at kesimpulan_idx + 1.
    del_count = saran_idx - (kesimpulan_idx + 1)
    print(f"Deleting {del_count} old kesimpulan elements...")
    for _ in range(del_count):
        body.remove(body[kesimpulan_idx + 1])
    print("Old kesimpulan elements successfully removed.")
    
    # Step 3: Define the new 3 paragraphs of kesimpulan
    new_kesimpulan_paras = [
        # Paragraf 1: Pengembangan & Fungsionalitas
        "Website Unit Layanan Terpadu (ULT) FKIP Universitas Lampung telah berhasil dikembangkan dengan "
        "menerapkan model pengembangan ADDIE (Analysis, Design, Development, Implementation, dan Evaluation) secara "
        "sistematis. Produk akhir berupa aplikasi web monolitik berbasis Laravel 12 dan PHP 8.4+ yang memiliki empat portal "
        "terintegrasi, yaitu Public Portal, Student Portal, Admin/Staff Portal, dan Signer Portal. Website ini "
        "mengintegrasikan fitur-fitur pendukung rekayasa tingkat tinggi seperti pelacakan status pengajuan (auditable timeline), "
        "tanda tangan elektronik, penomoran otomatis, dan penyimpanan berkas secara privat (private storage) untuk melindungi "
        "kerahasiaan dokumen akademik mahasiswa. Seluruh arsitektur portal dan modul fungsional tersebut berhasil diwujudkan "
        "guna mengotomatisasi birokrasi loket pelayanan fisik yang sebelumnya berjalan manual dan terdistribusi menjadi terpusat "
        "pada platform digital yang aman.",
        
        # Paragraf 2: Validitas Ahli
        "Uji kelayakan produk yang melibatkan sembilan validator ahli dari tiga bidang keahlian (materi, media, dan sistem) "
        "menunjukkan bahwa Website ULT FKIP Unila memiliki tingkat validitas yang sangat tinggi. Penilaian kuantitatif dari "
        "para ahli tersebut menghasilkan rata-rata persentase kelayakan kumulatif sebesar 95,45% untuk aspek materi, 93,33% "
        "untuk aspek media, dan 87,58% untuk aspek sistem, dengan rata-rata keseluruhan mencapai 91,95% yang dikategorikan "
        "\"Sangat Valid\". Hasil ini diperkuat oleh analisis validitas isi menggunakan koefisien Aiken's V dengan rata-rata indeks "
        "mencapai 0,90 yang membuktikan bahwa setiap butir instrumen memiliki tingkat relevansi isi yang sangat memuaskan. "
        "Berbagai masukan kualitatif dari para ahli, termasuk penguatan Content Security Policy (CSP) untuk pencegahan serangan "
        "Cross-Site Scripting (XSS) dan penyesuaian kontras dasbor antarmuka, telah berhasil ditindaklanjuti secara tuntas "
        "demi menghasilkan sistem yang andal dan aman sebelum diuji coba secara langsung.",
        
        # Paragraf 3: Uji Kepraktisan Pengguna
        "Dari segi kepraktisan operasional di lapangan, uji coba terbatas yang melibatkan 18 responden civitas akademika "
        "FKIP Unila membuktikan bahwa Website ULT FKIP Unila sangat mudah digunakan dan efisien. Pengisian kuesioner uji "
        "kepraktisan menghasilkan rata-rata skor total sebesar 55,28 dari skor maksimal 60,00, atau setara dengan persentase "
        "kepraktisan sebesar 92,13% yang menempatkan sistem ini ke dalam kriteria \"Sangat Praktis\". Temuan kuantitatif "
        "tersebut selaras dengan penilaian subjektif pada lembar angket Bagian F, di mana sebanyak 94,44% responden menyatakan "
        "secara langsung bahwa sistem ini berkategori \"Sangat Praktis\" and \"Praktis\" untuk menunjang kebutuhan sehari-hari. "
        "Walaupun respon pengguna secara umum sangat positif, masukan kualitatif dari enam responden mengenai penataan padding "
        "halaman dan modul filter jadwal telah diimplementasikan oleh peneliti guna menyempurnakan kegunaan (usability) "
        "sistem pada fase final."
    ]
    
    # Step 4: Insert the 3 new paragraphs at kesimpulan_idx + 1
    insert_idx = kesimpulan_idx + 1
    
    for i, para_text in enumerate(new_kesimpulan_paras):
        p = doc.add_paragraph()
        format_paragraph(p, doc, "Normal", WD_ALIGN_PARAGRAPH.LEFT, spacing_before=6 if i > 0 else 12, spacing_after=6, line_spacing=1.5)
        # Indent the first line of paragraph (0.5 inch or ~1.27 cm is typical in indonesian thesis)
        p.paragraph_format.first_line_indent = Pt(36) # 0.5 inch indention in points
        
        add_run_formatted(p, para_text, bold=False, italic=False, font_size=12)
        
        p_el = p._p
        body.remove(p_el)
        body.insert(insert_idx, p_el)
        insert_idx += 1
        
    # Add a spacer paragraph at the very end of kesimpulan to keep spacing beautiful before B. Saran
    p_space = doc.add_paragraph()
    format_paragraph(p_space, doc, "Normal", WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.5)
    body.remove(p_space._p)
    body.insert(insert_idx, p_space._p)
    
    doc.save(path)
    print("Clean.docx kesimpulan section successfully restructured in paragraph format and saved.")

if __name__ == "__main__":
    main()
