import sys
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# Set stdout to UTF-8
sys.stdout.reconfigure(encoding='utf-8')

def format_paragraph(p, doc, style_name="Normal", alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing_before=0, spacing_after=0, line_spacing=1.5):
    try:
        p.style = doc.styles[style_name]
    except Exception:
        pass
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(spacing_before)
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = line_spacing

def add_run_formatted(p, text, bold=False, italic=False, font_size=12):
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(font_size)
    r.bold = bold
    r.italic = italic
    return r

def main():
    path = Path(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx")
    if not path.exists():
        print("Error: Clean.docx not found.")
        return
        
    doc = docx.Document(path)
    body = doc.element.body
    
    # 1. Update P 759 (Validator Ahli Sistem 2) in BAB IV
    print("Step 1: Refining Validator Ahli Sistem 2 paragraph...")
    found_val2 = False
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt.startswith("Validator Ahli Sistem 2 memberikan skor total 93") and "Login SSO" in txt:
            print(f"  Found Validator Ahli Sistem 2 paragraph at index {i}")
            # Replace the paragraph content to remove preparation blueprint of SSO, focusing only on CSP security
            refined_val2_text = (
                "Validator Ahli Sistem 2 memberikan skor total 93 dari 110 (Kelayakan: 84,55%) dengan keputusan akhir "
                "'Layak digunakan dengan revisi' disertai komentar teknis yang sangat krusial: 'Login SSO, perhatikan "
                "CSP (unsafe-eval/inline), atur CSP default'. Masukan ini menyoroti celah keamanan potensial pada header HTTP "
                "Laravel yang rentan terhadap serangan injeksi script jahat luar (Cross-Site Scripting / XSS) jika Content "
                "Security Policy (CSP) default tidak disetel secara ketat. Menindaklanjuti hal tersebut, peneliti melakukan "
                "revisi mendalam dengan memperketat middleware CSP default Laravel untuk memblokir unsafe-eval dan unsafe-inline "
                "secara total, serta mengenkripsi nama berkas pengajuan mahasiswa di private storage disk. Tindakan ini "
                "memastikan celah keamanan injeksi skrip jahat dari luar dapat diblokir secara total dan data privat mahasiswa "
                "terlindungi secara maksimal."
            )
            p.text = ""
            add_run_formatted(p, refined_val2_text, bold=False, italic=False, font_size=12)
            format_paragraph(p, doc, "Normal", WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.5)
            found_val2 = True
            break
            
    if not found_val2:
        print("  Warning: Validator Ahli Sistem 2 paragraph not found.")
        
    # 2. Find and delete P 812 (Kendala SSO) in BAB IV
    print("\nStep 2: Locating and deleting SSO Kendala paragraph...")
    sso_kendala_idx = -1
    for idx, child in enumerate(body.getchildren()):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = docx.text.paragraph.Paragraph(child, doc)
            txt = p.text.strip()
            if txt.startswith("Kendala Aksesibilitas Integrasi Single Sign-On (SSO) Kampus"):
                sso_kendala_idx = idx
                print(f"  Found SSO Kendala paragraph at element index {idx}")
                break
                
    if sso_kendala_idx != -1:
        body.remove(body[sso_kendala_idx])
        print("  Successfully deleted SSO Kendala paragraph.")
    else:
        print("  Warning: SSO Kendala paragraph not found.")
        
    # 3. Refine B. Saran in BAB V
    # Reload doc paragraphs since index changed
    print("\nStep 3: Refining Saran in BAB V...")
    sso_saran_idx = -1
    for idx, child in enumerate(body.getchildren()):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = docx.text.paragraph.Paragraph(child, doc)
            txt = p.text.strip()
            if txt.startswith("1. Integrasi Single Sign-On (SSO)"):
                sso_saran_idx = idx
                print(f"  Found SSO Saran paragraph at element index {idx}")
                break
                
    if sso_saran_idx != -1:
        # Delete P 831 (SSO Saran)
        body.remove(body[sso_saran_idx])
        print("  Deleted SSO Saran paragraph.")
        
        # Now we need to update the next two paragraphs numbers
        # Reload children to find the shifted paragraphs
        children_after_delete = body.getchildren()
        # Find next two paragraph elements
        saran1_p = None
        saran2_p = None
        found_saran_count = 0
        
        for idx in range(sso_saran_idx, len(children_after_delete)):
            child = children_after_delete[idx]
            tag = child.tag.split("}")[-1]
            if tag == "p":
                p = docx.text.paragraph.Paragraph(child, doc)
                txt = p.text.strip()
                if txt.startswith("2. Penambahan Modul Layanan Kepegawaian"):
                    saran1_p = p
                    found_saran_count += 1
                elif txt.startswith("3. Pengujian Keamanan Tingkat Lanjut"):
                    saran2_p = p
                    found_saran_count += 1
                    break
                    
        if saran1_p:
            new_saran1_text = (
                "1. Penambahan Modul Layanan Kepegawaian: Pengembangan modul kepegawaian (staf dan dosen) disarankan untuk "
                "ditambahkan pada fase berikutnya agar layanan administrasi yang disediakan tidak hanya terbatas pada "
                "mahasiswa, melainkan mencakup seluruh civitas akademika FKIP."
            )
            saran1_p.text = ""
            add_run_formatted(saran1_p, new_saran1_text, bold=False, italic=False, font_size=12)
            format_paragraph(saran1_p, doc, "Normal", WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.5)
            print("  Updated Kepegawaian Saran to number 1.")
            
        if saran2_p:
            new_saran2_text = (
                "2. Pengujian Keamanan Tingkat Lanjut: Perlu dilakukan pengujian keamanan web secara rutin (seperti "
                "Vulnerability Assessment dan Penetration Testing) secara berkala untuk mengevaluasi efektivitas "
                "header Content Security Policy (CSP) dan meminimalkan risiko serangan injeksi data di masa mendatang."
            )
            saran2_p.text = ""
            add_run_formatted(saran2_p, new_saran2_text, bold=False, italic=False, font_size=12)
            format_paragraph(saran2_p, doc, "Normal", WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.5)
            print("  Updated Security Saran to number 2.")
    else:
        print("  Warning: SSO Saran paragraph not found.")
        
    # Save document
    doc.save(path)
    print("\nClean.docx successfully updated without SSO details and saved.")

if __name__ == "__main__":
    main()
