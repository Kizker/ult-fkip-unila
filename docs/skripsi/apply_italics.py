import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

with open(r'c:\laragon\www\ult-fkip-unila\docs\skripsi\daftar_pustaka_raw.txt', 'r', encoding='utf-8') as f:
    lines = [l.strip() for l in f.readlines() if l.strip() and l.strip() != 'DAFTAR PUSTAKA']

fixes = {
    'Ahmad': {'journal': 'Network and Distributed System Security (NDSS)', 'pages': '1-8'},
    'Aiken': {'journal': 'Educational and Psychological Measurement', 'vol': '45(1)', 'pages': '131-142'},
    'Akbari': {'pages': '52-65'},
    'Alisa': {'journal': 'Journal of Governance and Public Administration (JoGaPA)'},
    'Arikunto|2009': {'publisher': 'Bumi Aksara, Jakarta', 'pages': '310 hlm'},
    'Arikunto|2013': {'publisher': 'Rineka Cipta, Jakarta', 'pages': '413 hlm'},
    'Branch': {'publisher': 'Springer, Boston', 'pages': '175 hlm'},
    'Deshpande': {'journal': 'Web Engineering', 'vol': '1(1)', 'pages': '1-438'},
    'Henim': {'journal': 'Jurnal Komputer Terapan'},
    'Johns': {'publisher': 'World Bank Group, Washington, DC', 'pages': '150 hlm'},
    'Mustopa': {'journal': 'Jurnal Ilmiah Informatika Komputer'},
    'Ridwan': {'journal': 'Jurnal Sistem Informasi'},
    'Sudjana': {'publisher': 'Tarsito, Bandung', 'pages': '508 hlm'},
    'Sugiyono': {'publisher': 'Alfabeta, Bandung', 'pages': '334 hlm'},
    'Widiawati': {'journal': 'Jurnal Informatika', 'vol': '2', 'pages': '95-109'}
}

en_terms = [
    'Smart Education', 'Website', 'Addie', 'E-Learning', 'UI / UX', 'Prototype',
    'User Experience', '(Ux)', 'System Usability Scale', '(SUS)', 'ADDIE Model',
    'Web', 'Usability Testing', 'My UT', 'Post-Study System Usability Questionnaire',
    'USE Questionnaire', 'Cognitive Walkthrough', 'Webqual 4.0', 'User Center Design',
    'software', 'hardware', 'brainware', 'R & D', 'USABILITY', 'WEBSITE', 'usability', 'Webqual',
    'Smart', 'Education', 'Scale'
]

def parse_reference(ref_text):
    match = re.match(r'^([^\(]+?)\s+\((\d{4})\)\.\s+(.+?)\.(?:\s+(.+))?$', ref_text)
    if not match: return None
    author_part = match.group(1).strip()
    year = match.group(2).strip()
    title = match.group(3).strip()
    rest = match.group(4).strip() if match.group(4) else ''
    
    author_part = author_part.replace('&', 'dan')
    if author_part.endswith('.'):
        author_part = author_part[:-1].strip()
    
    is_book = False
    if ' hlm' in rest.lower() or 'hlm.' in rest.lower() or 'jakarta' in rest.lower() or 'bandung' in rest.lower() or 'publisher' in rest.lower() or 'springer' in rest.lower():
        is_book = True
    elif re.search(r'\d+\(\d+\)', rest) or 'Journal' in rest or 'Jurnal' in rest:
        is_book = False
    else:
        for k in fixes:
            search_k = k.split('|')[0]
            if search_k in author_part:
                if 'publisher' in fixes[k]: is_book = True
                else: is_book = False
        if not rest and not is_book:
            is_book = True

    return {
        'original': ref_text,
        'authors': author_part,
        'year': year,
        'title': title,
        'rest': rest,
        'is_book': is_book
    }

parsed_refs = []
for line in lines:
    if line.startswith('The above content shows') or 'http' in line and len(line) < 50: continue
    line = re.sub(r'https?://[^\s]+', '', line).strip()
    line = re.sub(r'www\.[^\s]+', '', line).strip()
    parsed = parse_reference(line)
    if parsed: parsed_refs.append(parsed)

formatted_refs = []
for p in parsed_refs:
    auth, year, title, rest, is_book = p['authors'], p['year'], p['title'], p['rest'], p['is_book']
    
    for k, v in fixes.items():
        search_k = k.split('|')[0]
        req_year = k.split('|')[1] if '|' in k else None
        
        if search_k in auth:
            if req_year and req_year != year:
                continue
            
            if 'publisher' in v:
                rest = f"{v.get('publisher', '')}. {v.get('pages', '')}"
                is_book = True
            elif 'journal' in v:
                rest = f"{v['journal']} {v.get('vol', '')}, {v.get('pages', rest)}"
                is_book = False
            else:
                if 'pages' in v: rest += f", {v['pages']}"

    id_words = ['dalam', 'pada', 'di', 'dan', 'yang', 'untuk', 'dengan', 'dari', 'ke', 'sebagai', 'berbasis', 'sistem', 'informasi', 'pendidikan', 'pengembangan', 'pelayanan', 'kualitas', 'mahasiswa', 'perguruan', 'tinggi']
    words_lower = title.lower().split()
    is_id_title = any(w in words_lower for w in id_words)
    
    title_runs = []
    if is_book or not is_id_title:
        fmt_title = ' '.join([w.capitalize() for w in title.split()]) if is_book else title.capitalize()
        title_runs.append((fmt_title, True))
    else:
        fmt_title = title.capitalize()
        pattern = re.compile(r'(' + '|'.join(map(re.escape, en_terms)) + r')', re.IGNORECASE)
        parts = pattern.split(fmt_title)
        for part in parts:
            if not part: continue
            is_en = any(part.lower() == et.lower() for et in en_terms)
            title_runs.append((part, is_en))

    if not is_book:
        journal_match = re.match(r'^(.+?)(?:,\s*(\d+(?:\(\d+\))?.*?))?$', rest)
        if journal_match:
            p['journal_name'] = journal_match.group(1).strip()
            p['journal_vol'] = journal_match.group(2) if journal_match.group(2) else ''
        else:
            p['journal_name'] = rest
            p['journal_vol'] = ''
    else:
        p['rest'] = rest

    p['title_runs'] = title_runs
    p['authors'] = auth
    p['is_book'] = is_book  # CRITICAL BUG FIX!!!!
    formatted_refs.append(p)

formatted_refs.sort(key=lambda x: x['authors'].lower())

with open(r'c:\laragon\www\ult-fkip-unila\docs\skripsi\debug_refs.txt', 'w', encoding='utf-8') as dbg:
    for ref in formatted_refs:
        if ref['is_book']:
             dbg.write(f"BOOK: {ref['authors']}. {ref['year']}. {ref['title']}. {ref['rest']}\n")
        else:
             dbg.write(f"JOURNAL: {ref['authors']}. {ref['year']}. {ref['title']}. {ref.get('journal_name', '')} {ref.get('journal_vol', '')}\n")

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

p_title = doc.add_paragraph('DAFTAR PUSTAKA')
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.runs[0].bold = True
doc.add_paragraph('')

for ref in formatted_refs:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(24)
    
    p.add_run(f"{ref['authors']}. {ref['year']}. ")
    
    for text, is_italic in ref['title_runs']:
        p.add_run(text).italic = is_italic
    
    p.add_run('. ')
    
    if ref['is_book']:
        rest_part = ref['rest'].rstrip('.')
        if rest_part:
             p.add_run(f"{rest_part}.").italic = False
    else:
        j_name = ref.get('journal_name', '')
        if j_name:
             p.add_run(f"{j_name} ").italic = True
        j_vol = ref.get('journal_vol', '').rstrip('.')
        if j_vol:
             p.add_run(f"{j_vol}.").italic = False

try:
    doc.save(r'c:\laragon\www\ult-fkip-unila\docs\skripsi\Daftar_Pustaka_Unila_2020.docx')
    print('Data appended correctly!')
except PermissionError:
    print('PERMISSION ERROR: File is locked by MS Word.')
