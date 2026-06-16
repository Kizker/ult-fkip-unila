import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def inject(doc_path, output_path):
    doc = docx.Document(doc_path)
    
    # We will iterate through body elements to find exactly what we want.
    # But python-docx doesn't allow easy insertion after an element if it's not a paragraph.
    # Fortunately, Table has `_element` and we can insert a new paragraph XML after it.
    
    # For Implementation phase: fix the "Berdasarkan Gambar 4.27" text
    for p in doc.paragraphs:
        if 'Berdasarkan Gambar 4.27 di atas, dapat dilihat bahwa tanggapan' in p.text:
            p.text = p.text.replace('Berdasarkan Gambar 4.27 di atas', 'Berdasarkan tabel di atas')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Now let's find the exact tables and charts by looking at the XML elements sequentially
    # It's easier to find the exact paragraphs and insert around them
    
    # 1. Dev Phase: Insert AFTER table. The table has cell "Validasi Sistem"
    dev_table = None
    for t in doc.tables:
        try:
            if 'Validasi Sistem' in t.cell(0,0).text or 'Validasi Sistem' in t.cell(1,0).text or 'Ahli' in t.cell(0,0).text:
                # let's find the specific table in dev phase
                text = ""
                for r in t.rows:
                    for c in r.cells:
                        text += c.text + " "
                if 'Content Security Policy' in text:
                    dev_table = t
                    break
        except:
            pass
            
    if dev_table:
        # Create a new paragraph element
        new_p = docx.oxml.OxmlElement('w:p')
        dev_table._element.addnext(new_p)
        p = docx.text.paragraph.Paragraph(new_p, doc)
        p.text = "Berdasarkan tabel di atas, para ahli validator memberikan berbagai masukan kualitatif yang konstruktif untuk penyempurnaan sistem. Beberapa perbaikan utama yang disarankan meliputi penguatan keamanan sistem melalui penerapan Content Security Policy (CSP), perbaikan logika validasi form, serta penambahan fitur navigasi pendukung guna meningkatkan pengalaman pengguna (user experience). Seluruh masukan tersebut telah diakomodasi dan ditindaklanjuti secara langsung pada tahap pengembangan."
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.style = doc.styles['Normal']
        p.paragraph_format.first_line_indent = Pt(36)
        print("Injected Dev explanation after table.")

    # 2. Impl Phase: Insert AFTER the chart.
    # The chart is right after the text "Berdasarkan tabel di atas, dapat dilihat bahwa tanggapan..."
    impl_chart_p = None
    found_text = False
    for p in doc.paragraphs:
        if 'Berdasarkan tabel di atas, dapat dilihat bahwa tanggapan' in p.text:
            found_text = True
        elif found_text:
            if 'graphic' in p._p.xml or 'chart' in p._p.xml or 'drawing' in p._p.xml:
                impl_chart_p = p
                break
            elif len(p.text.strip()) > 0:
                # if we hit text without finding chart, break to be safe
                pass
                
    if impl_chart_p:
        new_p = impl_chart_p.insert_paragraph_before() # wait, we want AFTER the chart.
        # to insert after impl_chart_p, we can create an oxml element and add it next
        new_p_xml = docx.oxml.OxmlElement('w:p')
        impl_chart_p._p.addnext(new_p_xml)
        new_p = docx.text.paragraph.Paragraph(new_p_xml, doc)
        new_p.text = "Berdasarkan diagram persentase kepraktisan di atas, mayoritas responden (94,44% atau 17 orang) menyimpulkan bahwa sistem ini berkategori 'Sangat Praktis' dan 'Praktis' untuk digunakan secara langsung dalam kegiatan operasional sehari-hari. Hanya 1 responden (5,56%) yang memberikan penilaian 'Cukup Praktis'. Secara keseluruhan pada tahap implementasi ini, sistem yang diuji terbukti sangat layak, praktis, dan siap diimplementasikan secara penuh untuk digitalisasi pelayanan akademik."
        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        new_p.style = doc.styles['Normal']
        new_p.paragraph_format.first_line_indent = Pt(36)
        print("Injected Impl explanation after chart.")
        
    doc.save(output_path)

inject(
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx',
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
)

inject(
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx',
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
)
