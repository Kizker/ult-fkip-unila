import docx
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def replace_in_doc(doc_path, output_path):
    doc = docx.Document(doc_path)
    img_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram\22_diagram_kepraktisan.png'
    
    # 1. Delete the table
    for table in doc.tables:
        try:
            first_cell = table.cell(0,0).text.strip()
            if 'No' in first_cell or 'Responden' in first_cell:
                # Let's check if it's the right table by checking text in it
                text = ""
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                if 'Sangat Praktis' in text and '18' in text:
                    # Found it
                    table._element.getparent().remove(table._element)
                    print(f"Table removed from {doc_path}")
        except Exception as e:
            pass

    # 2. Find caption and insert image
    for i, p in enumerate(doc.paragraphs):
        if 'Rekapitulasi Distribusi Kategori Uji Kepraktisan Pengguna' in p.text and 'Tabel' in p.text:
            p.text = 'Gambar 4.27. Diagram Distribusi Kategori Uji Kepraktisan Pengguna.'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Insert image before this paragraph
            # python-docx doesn't have insert_picture on paragraph, only on run.
            # To insert before, we create a new paragraph before p
            new_p = p.insert_paragraph_before()
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = new_p.add_run()
            run.add_picture(img_path, width=Inches(5))
            
            print(f"Image inserted in {doc_path}")
            break
            
    doc.save(output_path)

replace_in_doc(
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx',
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
)

replace_in_doc(
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx',
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
)
