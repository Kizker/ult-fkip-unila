import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil

doc_clean_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
doc_hl_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'

use_case_paras = [
    "Diagram Use Case memvisualisasikan interaksi fungsional antara pengguna dengan sistem yang melibatkan empat aktor utama. Aktor pertama adalah Mahasiswa yang bertindak sebagai pemohon utama layanan persuratan akademik secara mandiri. Aktor ini memiliki otorisasi penuh untuk mengakses tiga use case inti, yaitu mengajukan permohonan dengan mengisi borang dinamis, melacak status berkas secara langsung melalui linimasa (timeline) interaktif, dan mengunduh berkas dokumen resmi yang telah selesai secara privat.",
    "Pada tataran operasional, Aktor Staf ULT bertindak sebagai pengelola utama fungsionalitas fakultas. Staf ULT berinteraksi dengan tiga use case krusial yang meliputi validasi permohonan tingkat awal, pengelolaan katalog layanan beserta templat dokumen persuratan akademik, hingga tahap eksekusi perakitan dokumen hasil berbasis server. Dokumen yang telah divalidasi oleh staf ini kemudian diteruskan kepada Aktor Pejabat. Aktor ini mewakili pimpinan tingkat fakultas maupun jurusan yang memiliki kewenangan untuk menelaah dokumen secara substantif dan membubuhkan tanda tangan elektronik (digital signature) sebagai bentuk legalisasi dokumen sebelum proses perakitan akhir dilakukan.",
    "Keseluruhan ekosistem interaksi fungsional tersebut diawasi secara penuh oleh Aktor Admin Utama. Aktor ini memegang kendali atas manajemen sistem keseluruhan, bertanggung jawab atas pengelolaan pengguna beserta konfigurasi hak akses berbasis Role-Based Access Control (RBAC), serta memantau rekaman jejak audit (audit trail) seluruh aktivitas. Kehadiran aktor ini memastikan bahwa setiap perubahan data pada sistem dapat dilacak, sehingga akuntabilitas dan stabilitas operasional platform tetap terjaga dengan baik."
]

arch_paras = [
    "Diagram arsitektur sistem pada Gambar 5 mengadopsi model berlapis (layered architecture) yang membagi komponen sistem secara vertikal menjadi lima lapisan utama. Pembagian ini bertujuan untuk memisahkan tanggung jawab masing-masing komponen, mulai dari antarmuka interaksi pengguna hingga pada lapisan mesin perakit dokumen tingkat rendah, sehingga menghasilkan sistem yang loosely coupled, mudah dipelihara, dan memiliki tingkat keamanan yang terisolasi dengan baik.",
    "Secara visual, lapisan teratas adalah Client & Interface Layer yang dirancang agar kompatibel dengan teknologi Progressive Web App (PWA). Lapisan ini menaungi empat portal utama: Public Portal sebagai sarana informasi umum, Student Portal yang menjadi dasbor utama mahasiswa untuk mengajukan permohonan dokumen, Staff/Admin Portal untuk manajemen operasional oleh staf, dan Signer Portal yang didedikasikan khusus untuk pejabat dalam memberikan persetujuan elektronik. Pengguna akan berinteraksi langsung dengan antarmuka ini sebelum permintaan diteruskan ke sistem.",
    "Lapisan kedua, Security & Routing Gatekeeper, bertindak sebagai perisai pelindung utama sistem. Lapisan ini memuat Laravel Router yang dikonfigurasi dengan Content Security Policy (CSP) untuk mencegah eksekusi skrip berbahaya. Selain itu, terdapat Spatie RBAC (Role-Based Access Control) sebagai middleware otorisasi yang secara dinamis memvalidasi hak akses pengguna berdasarkan peran mereka. Fitur Anti-IDOR juga disematkan untuk memproteksi akses berkas agar hanya dapat diunduh oleh pemilik sah, serta HtmlSanitizer yang secara ketat membersihkan setiap input berformat teks kaya (WYSIWYG) dari potensi serangan Cross-Site Scripting (XSS).",
    "Pada lapisan ketiga, Laravel 12 Application Core memegang kendali penuh atas logika bisnis aplikasi. Controllers Core, khususnya RequestAdminController, menjadi pusat pengatur lalu lintas data permohonan persuratan akademik. Komponen ini berkoordinasi langsung dengan DocumentAssembler sebagai fasilitator orkestrasi perakitan dokumen, serta HtmlToOpenXMLParser yang bertugas secara spesifik untuk membedah dan menerjemahkan tag-tag HTML dari editor teks ke dalam format XML mentah tanpa merusak gaya bawaan.",
    "Lapisan keempat merupakan Database & Storage Layer yang menangani persistensi data. Di dalamnya, MySQL Database berfungsi untuk menyimpan data operasional yang terstruktur seperti formulir pengguna, profil, dan log rekam jejak (audit trail). Secara paralel, Private Storage Disk bertindak sebagai ruang penyimpanan terenkripsi yang sepenuhnya terisolasi dari akses publik, dikhususkan untuk mengamankan fisik berkas lampiran maupun dokumen hasil perakitan berformat .docx dan PDF.",
    "Lapisan paling bawah adalah Document Assembly Engine yang ditenagai oleh pustaka PHPOffice. Lapisan ini bertanggung jawab mengeksekusi proses perakitan tingkat rendah. Word Templates berisi placeholder statis digunakan sebagai landasan surat. Selanjutnya, modul PhpWord & ZipArchive bekerja memanipulasi struktur XML mentah secara langsung di dalam arsip dokumen, menginjeksi data mahasiswa yang telah diproses. Terakhir, PDF Converter menghasilkan luaran dokumen dengan format siap cetak yang dapat diunduh langsung melalui antarmuka."
]

def format_paragraph(p, text, is_hl=False):
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if is_hl:
        from docx.enum.text import WD_COLOR_INDEX
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def fix_document(file_path, is_hl=False):
    doc = docx.Document(file_path)
    
    uc_idx = -1
    arch_idx = -1
    
    # 1. Identify indices
    for i, p in enumerate(doc.paragraphs):
        if 'Gambar 4. Diagram Use Case' in p.text:
            uc_idx = i
        elif 'Gambar 5. Diagram Arsitektur' in p.text:
            arch_idx = i

    print(f"File: {file_path}")
    print(f"Gambar 4 at: {uc_idx}, Gambar 5 at: {arch_idx}")
    
    # 2. Process Architecture first to avoid shifting indices for Use Case if Use Case is before Architecture
    # It is expected:
    # arch_idx is "Gambar 5. Diagram Arsitektur..."
    # Next paragraphs:
    # arch_idx + 1: (empty)
    # arch_idx + 2: Logika bisnis...
    # arch_idx + 3: 1) Client & Interface
    # arch_idx + 4: 3) Laravel 12
    # arch_idx + 5: 5) Document Assembly Engine
    
    # Let's delete the old texts. We will look forward from arch_idx until we hit "Perancangan Alur Kerja Persuratan"
    arch_end_idx = arch_idx + 1
    while arch_end_idx < len(doc.paragraphs) and 'Perancangan Alur Kerja Persuratan' not in doc.paragraphs[arch_end_idx].text:
        arch_end_idx += 1
        
    print(f"Arch text range: {arch_idx+1} to {arch_end_idx-1}")
    
    # We will insert new paragraphs before arch_end_idx
    for text in reversed(arch_paras):
        new_p = doc.paragraphs[arch_end_idx].insert_paragraph_before()
        format_paragraph(new_p, text, is_hl)
        
    # Delete the old ones (from arch_end_idx-1 down to arch_idx+1)
    for i in range(arch_end_idx - 1, arch_idx, -1):
        delete_paragraph(doc.paragraphs[i])
        
    # 3. Process Use Case
    # uc_idx is "Gambar 4. Diagram Use Case..."
    # uc_end_idx should be "Perancangan Arsitektur Sistem"
    uc_end_idx = uc_idx + 1
    while uc_end_idx < len(doc.paragraphs) and 'Perancangan Arsitektur Sistem' not in doc.paragraphs[uc_end_idx].text:
        uc_end_idx += 1
        
    print(f"Use Case text range: {uc_idx+1} to {uc_end_idx-1}")
    
    # Insert new paragraphs before uc_end_idx
    for text in reversed(use_case_paras):
        new_p = doc.paragraphs[uc_end_idx].insert_paragraph_before()
        format_paragraph(new_p, text, is_hl)
        
    # Delete old ones
    for i in range(uc_end_idx - 1, uc_idx, -1):
        delete_paragraph(doc.paragraphs[i])

    doc.save(file_path)
    print("Saved paragraphs.")

fix_document(doc_clean_path, is_hl=False)
fix_document(doc_hl_path, is_hl=True)
