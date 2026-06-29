import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Cm

original_text = [
    "Hasil pengembangan Website ULT FKIP Unila ini membuktikan secara empiris bahwa penerapan model R&D dengan sintaks ADDIE mampu menghasilkan produk teknologi manajemen pelayanan kampus yang terarah, efisien, aman, dan tepat guna bagi seluruh civitas akademika FKIP Unila. Melalui analisis kebutuhan awal pada tahap Analysis, ditemukan bahwa kendala utama layanan administrasi di FKIP Unila adalah lambatnya turn-around time pengurusan dokumen, hilangnya transparansi status dokumen, serta kerentanan fisik berkas yang diajukan. Sistem yang dikembangkan ini secara terarah menutup pain point tersebut dengan memindahkan seluruh proses fisik ke dalam sistem berbasis web yang dinamis dan aman.",
    "Keunggulan",
    "Website ULT FKIP Unila memiliki sejumlah keunggulan teknis dan akademis yang membedakannya dengan platform administrasi konvensional. Keunggulan-keunggulan utama tersebut dijabarkan di bawah ini:",
    "Integrasi Empat Portal Berbasis RBAC Ketat: Keberadaan portal khusus (Public, Student, Admin, Signer) memastikan setiap pengguna berinteraksi hanya dengan fungsi dan data yang relevan dengan kewenangannya. Model ini memblokir akses ilegal antarportal dan menjamin keamanan transaksi data.",
    "Otomatisasi Perakitan Dokumen Dinamis (OpenXML): Integrasi pustaka PHPWord untuk merakit dokumen Word (.docx) langsung dari basis data menggunakan template dinamis terbukti sangat efisien. Logika parser HTML-to-OpenXML dinamis terpusat pada file XML mentah (`document.xml`) menjamin bahwa run teks yang dimasukkan dari WYSIWYG editor (Tiptap Editor) pada formulir mahasiswa tetap mewarisi gaya visual visual bawaan template (font, ukuran, warna) tanpa mengalami degradasi layout visual orisinal.",
    "Keamanan Data Tingkat Tinggi dan Proteksi Otorisasi: Penggunaan private storage disk Laravel menjamin berkas berkas persyaratan sensitif mahasiswa tidak dapat diakses langsung secara publik tanpa login. Dilindungi oleh middleware otorisasi anti-IDOR yang ketat, platform berhasil menutup risiko bypass endpoint berkas mahasiswa lain.",
    "Jejak Transaksi Digital Komprehensif (Audit Trail): Kehadiran modul linimasa timeline yang mencatat setiap aksi pengguna (waktu pengajuan, nama verifikator, catatan revisi, dan tanggal tanda tangan) memberikan akuntabilitas yang tinggi bagi instansi pelayanan publik fakultas.",
    "Kompatibilitas Progressive Web App (PWA): Platform dikembangkan agar kompatibel untuk diakses via perangkat mobile secara ringan dan responsif, mempermudah mahasiswa memantau status dokumen kapan saja dan di mana saja.",
    "Kendala",
    "Selama proses riset rekayasa perangkat lunak ini, peneliti menghadapi sejumlah kendala teknis dan implementasi di lapangan. Kendala-kendala tersebut beserta solusi pemecahannya dijabarkan di bawah ini:",
    "Kendala Render HTML WYSIWYG pada Dokumen Word (.docx): Input teks formulir mahasiswa yang berasal dari WYSIWYG editor (Tiptap Editor) menyimpan tag-tag HTML (seperti `<p>`, `<b>`, `<i>`, `<u>`, `<br>`, `<ul>`, `<ol>`, `<li>`). Saat proses perakitan dokumen Word (`.docx`), tag HTML tersebut tercetak mentah sebagai teks biasa tanpa format visual. Solusi: Peneliti mengembangkan parser HTML-to-OpenXML dinamis terpusat pada file XML mentah (`document.xml`, `header*.xml`, `footer*.xml`) sebelum file DOCX dikemas ulang oleh PHPWord. Tag inline (b, strong, i, em, u) diubah menjadi elemen formatting OpenXML (`<w:b/>`, `<w:i/>`, `<w:u w:val=\"single\"/>`), sedangkan tag paragraf dan list dipetakan menggunakan pemutus baris dinamis `<w:br/>` agar gaya visual bawaan template Word tetap terjaga tanpa terjadi degradasi tata letak visual.",
    "Kendala Ancaman Keamanan Cross-Site Scripting (XSS): Kerentanan injeksi skrip jahat luar pada formulir pengajuan mahasiswa merupakan ancaman krusial pada sistem pelayanan publik online yang disoroti oleh validator sistem (Validator Ahli Sistem 2). Solusi: Peneliti memperketat pertahanan web dengan memperketat middleware HTTP Header Content Security Policy (CSP) default pada Laravel untuk memblokir celah unsafe-eval dan unsafe-inline secara total, menyaring input teks richtext menggunakan sanitasi XSS ketat (HtmlSanitizer class), dan menerapkan enkripsi nama berkas privat mahasiswa.",
    "Secara keseluruhan, hasil validasi ahli (persentase rata-rata 91,95% - Sangat Valid) dan uji kepraktisan pengguna terbatas (92,13% - Sangat Praktis) membuktikan bahwa Website ULT FKIP Unila telah berhasil dirancang dan dikembangkan dengan standar kualitas rekayasa perangkat lunak yang sangat baik. Sistem ini terbukti andal dalam menyederhanakan alur birokrasi persuratan akademik, menjamin keamanan berkas sensitif, dan memberikan tingkat akuntabilitas pelayanan publik yang tinggi bagi Fakultas Keguruan dan Ilmu Pendidikan Universitas Lampung."
]

def update_docx(filepath):
    doc = docx.Document(filepath)
    
    start_el = None
    end_el = None
    
    for p in doc.paragraphs:
        text = p.text.strip().lower()
        if "pembahasan hasil penelitian" == text:
            start_el = p._element
        elif text == "bab v" or "kesimpulan dan saran" == text:
            if start_el is not None:
                end_el = p._element
                break
                
    if start_el is None or end_el is None:
        print(f"Failed to find bounds in {filepath}")
        return
        
    curr = start_el.getnext()
    while curr is not None and curr != end_el:
        nxt = curr.getnext()
        curr.getparent().remove(curr)
        curr = nxt
        
    for txt in original_text:
        new_p = docx.text.paragraph.Paragraph(docx.oxml.OxmlElement('w:p'), doc)
        end_el.addprevious(new_p._element)
        
        run = new_p.add_run(txt)
        if txt in ["Keunggulan", "Kendala"]:
            run.bold = True
            
        new_p.style = doc.styles['Normal']
        new_p.paragraph_format.line_spacing = 1.5
        new_p.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(filepath)
    print(f"Updated {filepath}")

import glob
files = [
    r"c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx",
    r"c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx"
]

for f in files:
    update_docx(f)
