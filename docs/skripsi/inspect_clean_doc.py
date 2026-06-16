import docx
from pathlib import Path

def main():
    path = Path(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx")
    if not path.exists():
        print("Not found.")
        return
        
    doc = docx.Document(path)
    print("Total paragraphs:", len(doc.paragraphs))
    print("Total tables:", len(doc.tables))
    
    # Print the headings in the document to see the outline
    print("\n=== Document Headings ===")
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading") or p.text.isupper() and len(p.text.strip()) > 3:
            # check if it's heading-like
            print(f"P {i} ({p.style.name}): '{p.text}'")

if __name__ == "__main__":
    main()
