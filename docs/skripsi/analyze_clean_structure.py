import sys
import docx
from pathlib import Path

# Set stdout to UTF-8
sys.stdout.reconfigure(encoding='utf-8')

def main():
    path = Path(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx")
    if not path.exists():
        print("Not found.")
        return
        
    doc = docx.Document(path)
    print("Total paragraphs in Clean.docx:", len(doc.paragraphs))
    print("Total tables in Clean.docx:", len(doc.tables))
    
    # Save the paragraphs outline to a text file to read safely
    out_path = Path("docs/skripsi/clean_structure_output.txt")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(f"Total paragraphs: {len(doc.paragraphs)}\n")
        f.write(f"Total tables: {len(doc.tables)}\n\n")
        f.write("=== Paragraph Outline ===\n")
        
        in_bab3 = False
        in_bab4 = False
        in_bab5 = False
        
        for i, p in enumerate(doc.paragraphs):
            txt = p.text.strip()
            style = p.style.name
            
            is_heading = style.startswith("Heading") or (txt.isupper() and len(txt) > 3)
            
            if "METODE PENELITIAN" in txt.upper() and is_heading:
                in_bab3 = True
                f.write(f"\n[START BAB III] P {i} ({style}): '{txt}'\n")
                continue
            if "HASIL DAN PEMBAHASAN" in txt.upper() and is_heading:
                in_bab3 = False
                in_bab4 = True
                f.write(f"\n[START BAB IV] P {i} ({style}): '{txt}'\n")
                continue
            if "KESIMPULAN DAN SARAN" in txt.upper() and is_heading:
                in_bab4 = False
                in_bab5 = True
                f.write(f"\n[START BAB V] P {i} ({style}): '{txt}'\n")
                continue
            if "DAFTAR PUSTAKA" in txt.upper() and is_heading:
                in_bab5 = False
                f.write(f"\n[START DAFTAR PUSTAKA] P {i} ({style}): '{txt}'\n")
                continue
                
            if is_heading:
                f.write(f"P {i} ({style}): '{txt}'\n")
            elif any(kw in txt.lower() for kw in ["aiken", "purposive", "validator", "9 orang", "18 orang", "kepraktisan", "tabel c.1", "tabel c.2", "gambar d.1", "gambar d.2"]):
                f.write(f"  P {i} ({style}) [KEYWORD MATCH]: '{txt[:150]}...'\n")
                
    print(f"Structure saved successfully to {out_path}")

if __name__ == "__main__":
    main()
