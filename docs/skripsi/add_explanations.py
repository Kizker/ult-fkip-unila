import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def process_doc(doc_path, output_path):
    doc = docx.Document(doc_path)
    
    # 1. Update text for Implementation phase
    for p in doc.paragraphs:
        if 'Berdasarkan Tabel 4.2 di atas' in p.text:
            p.text = p.text.replace('Berdasarkan Tabel 4.2', 'Berdasarkan Gambar 4.27')
            # Fix alignment if needed (Justify)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            print(f"Updated Implementation text in {doc_path}")

    # 2. Insert new paragraph at the end of Development phase
    # Find heading 'Tahap Implementasi (Implementation)'
    for p in doc.paragraphs:
        if p.text.strip() == 'Tahap Implementasi (Implementation)' or p.text.strip() == '4.1.4 Tahap Implementasi (Implementation)':
            # Insert before this paragraph
            new_p = p.insert_paragraph_before()
            new_p.text = "Berdasarkan hasil rekapitulasi penilaian dari ketiga ahli validator, baik dari aspek materi, media, maupun sistem, perancangan antarmuka dan fungsionalitas Web ULT FKIP Unila ini memperoleh persentase rata-rata kelayakan sebesar 91,95%. Pencapaian tersebut menempatkan sistem pada kategori 'Sangat Valid'. Hal ini menunjukkan bahwa sistem yang dibangun telah secara kokoh memenuhi standar kelayakan operasional dan secara teknis sangat mumpuni untuk diujicobakan secara langsung kepada pengguna akhir guna mengevaluasi tingkat kepraktisannya."
            new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Apply first line indent to match thesis style (1.5 cm -> 0.59 inches, but let's just match style)
            new_p.style = doc.styles['Normal']
            
            # Since Normal style in this doc might have indent, if not we can add it
            new_p.paragraph_format.first_line_indent = Pt(36) # Approx 1.27 cm
            print(f"Inserted Development text in {doc_path}")
            break
            
    doc.save(output_path)

process_doc(
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx',
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
)

process_doc(
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx',
    r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
)
