import docx
import sys

def update_table(doc_path):
    print(f"Processing {doc_path}...")
    doc = docx.Document(doc_path)
    target_table = None
    for table in doc.tables:
        try:
            if 'Jenis Validasi' in table.cell(0,0).text or 'Gambar Desain Awal' in table.cell(0,0).text:
                target_table = table
                break
        except:
            pass
            
    if not target_table:
        print("Table not found")
        return
        
    # Remove all rows except the first (header)
    while len(target_table.rows) > 1:
        tr = target_table.rows[-1]._tr
        target_table._tbl.remove(tr)
        
    # Update headers
    p0 = target_table.cell(0,0).paragraphs[0]
    p0.text = ""
    p0.add_run("Gambar Desain Awal (Pra-Validasi)").bold = True
    
    p1 = target_table.cell(0,1).paragraphs[0]
    p1.text = ""
    p1.add_run("Rincian Kritik Ahli & Tindak Lanjut Eksekusi").bold = True
    
    p2 = target_table.cell(0,2).paragraphs[0]
    p2.text = ""
    p2.add_run("Gambar Desain Akhir (Pasca-Validasi)").bold = True
    
    data = [
        (
            "[MASUKKAN SCREENSHOT: UI Beranda dengan warna gradasi mencolok]",
            "Aspek Komposisi Warna (Validasi Media)\nPakar media mengkritisi penggunaan skema warna gradasi pada antarmuka utama yang dinilai terlalu mencolok dan mendistraksi fokus visual pengguna. Sebagai tindak lanjut korektif, desain antarmuka direstrukturisasi menggunakan pendekatan flat design dengan palet warna solid biru yang lebih elegan dan selaras dengan identitas visual (corporate identity) Fakultas Keguruan dan Ilmu Pendidikan.",
            "[MASUKKAN SCREENSHOT: UI Beranda flat design biru yang elegan]"
        ),
        (
            "[MASUKKAN SCREENSHOT: Tabel Dasbor Staf ULT yang padat/sempit]",
            "Aspek Tata Letak/Spasi Tabel (Validasi Media/Sistem)\nValidator menyoroti kepadatan baris pada matriks tabel operasional staf yang berpotensi memicu keletihan visual dan menyulitkan proses pemindaian data (eye-scanning). Mengakomodasi kritik tersebut, dilakukan ekspansi white space dengan memperbesar padding sel dan mengatur spasi antarbaris menjadi 1,5, sehingga keterbacaan data tabular meningkat secara signifikan.",
            "[MASUKKAN SCREENSHOT: Tabel Dasbor Staf ULT yang renggang dan rapi]"
        ),
        (
            "[MASUKKAN SCREENSHOT: Katalog layanan yang hanya berisi teks]",
            "Aspek Fitur Ikon Navigasi (Validasi Media)\nKetiadaan penanda visual pada direktori layanan dinilai memperlambat proses identifikasi fungsi oleh pengguna akhir. Solusi teknis diwujudkan melalui pengintegrasian ikon vektor spesifik pada setiap kartu layanan (service tile). Penambahan elemen semiotik ini berfungsi sebagai jangkar kognitif yang mempercepat navigasi dan pengenalan layanan secara intuitif.",
            "[MASUKKAN SCREENSHOT: Katalog layanan dilengkapi ikon]"
        ),
        (
            "[MASUKKAN SCREENSHOT: Halaman panjang tanpa tombol navigasi balik]",
            "Aspek Fitur Pendukung (Validasi Sistem)\nPakar sistem mengidentifikasi inefisiensi navigasi pada halaman dengan gulir panjang (long-scroll pages) akibat absennya pintasan untuk kembali ke area navigasi utama. Tindak lanjut eksekusi dilakukan dengan mengimplementasikan tombol melayang (floating button) 'Back to Top'. Fitur ini mengeliminasi kebutuhan gulir manual yang repetitif, sehingga mengoptimalkan mobilitas interaksi pengguna.",
            "[MASUKKAN SCREENSHOT: Halaman dengan tombol floating Back to Top]"
        )
    ]
    
    for row_data in data:
        row = target_table.add_row()
        row.cells[0].text = row_data[0]
        
        # Split text and add bold to title
        title, desc = row_data[1].split('\n')
        p = row.cells[1].paragraphs[0]
        p.text = ""
        run = p.add_run(title)
        run.bold = True
        p.add_run('\n' + desc)
        
        row.cells[2].text = row_data[2]

    doc.save(doc_path)
    print(f"Saved {doc_path}")

if __name__ == "__main__":
    file1 = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
    file2 = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
    update_table(file1)
    update_table(file2)
