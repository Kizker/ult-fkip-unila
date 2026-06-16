import sys
from docx import Document

def fix_addie_paragraph(doc_path, output_path):
    doc = Document(doc_path)
    modified = False
    
    target_text = "Berdasarkan model Four-D yang diadaptasi dari Thiagarajan (1974), tahap perancangan ini meliputi empat langkah sistematis utama, yaitu:"
    replacement_text = "Berdasarkan kerangka desain instruksional dan sistem dalam model ADDIE, tahap perancangan ini meliputi empat langkah sistematis utama, yaitu:"
    
    for p in doc.paragraphs:
        if "Thiagarajan" in p.text or "Four-D" in p.text:
            # We will clear the paragraph runs and reconstruct it to ensure we don't have fragmented runs.
            # Wait, the paragraph is:
            # "Pada tahap perancangan (Design) dalam model pengembangan ADDIE, fokus utama peneliti dialihkan dari perumusan masalah konseptual menuju perancangan solusi teknis. Berdasarkan model Four-D yang diadaptasi dari Thiagarajan (1974), tahap perancangan ini meliputi empat langkah sistematis utama, yaitu:"
            # Let's just find the exact string in the paragraph text and do a replace, then recreate runs.
            # But earlier we decided it's risky. In this specific case, it's a simple text paragraph without equations or Mendeley fields (it's just a transition sentence).
            # We can check if it has complex elements first.
            if len(p._element.xpath('.//w:fldSimple | .//w:instrText | .//m:oMath')) > 0:
                print(f"Warning: Paragraph has complex elements: {p.text}")
                continue
                
            full_text = p.text
            new_text = full_text.replace(
                "Berdasarkan model Four-D yang diadaptasi dari Thiagarajan (1974), tahap perancangan ini meliputi empat langkah sistematis utama, yaitu:",
                "Berdasarkan kerangka perancangan dalam model ADDIE, tahap perancangan ini meliputi empat langkah sistematis utama, yaitu:"
            )
            
            # Since my previous script italicized "Design" and "Four-D", if we reconstruct, we need to apply italics.
            # A simpler way: just clear runs, add text, and italicize "Design" and "ADDIE" if needed.
            p.clear()
            
            parts = new_text.split("(Design)")
            if len(parts) == 2:
                r1 = p.add_run(parts[0] + "(")
                r2 = p.add_run("Design")
                r2.italic = True
                r3 = p.add_run(")" + parts[1])
            else:
                p.add_run(new_text)
                
            modified = True
            
    if modified:
        doc.save(output_path)
        print(f"Successfully modified {output_path}")
    else:
        print(f"Target text not found in {doc_path}")

def main():
    clean_path = r"c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx"
    hl_path = r"c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx"
    
    fix_addie_paragraph(clean_path, clean_path)
    fix_addie_paragraph(hl_path, hl_path)

if __name__ == "__main__":
    main()
