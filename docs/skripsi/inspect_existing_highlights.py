import sys
import docx
from pathlib import Path

# Set stdout to UTF-8
sys.stdout.reconfigure(encoding='utf-8')

def main():
    path = Path(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx")
    if not path.exists():
        print("Not found.")
        return
        
    doc = docx.Document(path)
    print("Total paragraphs in Highlighted:", len(doc.paragraphs))
    print("Total tables in Highlighted:", len(doc.tables))
    
    print("\n=== Paragraphs with Highlights ===")
    in_bab3 = False
    in_bab4 = False
    in_bab5 = False
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if "METODE PENELITIAN" in txt.upper() and p.style.name.startswith("Heading"):
            in_bab3 = True
            continue
        if "HASIL DAN PEMBAHASAN" in txt.upper() and p.style.name.startswith("Heading"):
            in_bab3 = False
            in_bab4 = True
            continue
        if "KESIMPULAN DAN SARAN" in txt.upper() and p.style.name.startswith("Heading"):
            in_bab4 = False
            in_bab5 = True
            continue
        if "DAFTAR PUSTAKA" in txt.upper() and (p.style.name.startswith("Heading") or p.style.name == "Normal"):
            in_bab5 = False
            continue
            
        # Check if any run in paragraph has highlight
        has_highlight = False
        highlight_types = []
        for r in p.runs:
            if r.font.highlight_color is not None:
                has_highlight = True
                highlight_types.append(str(r.font.highlight_color))
                
        if has_highlight:
            loc = "OTHER"
            if in_bab3:
                loc = "BAB III"
            elif in_bab4:
                loc = "BAB IV"
            elif in_bab5:
                loc = "BAB V"
            print(f"P {i} ({p.style.name}) in {loc}: '{txt[:100]}...'")

if __name__ == "__main__":
    main()
