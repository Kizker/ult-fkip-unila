import docx
import sys

def remove_425(doc_path):
    print(f"Processing {doc_path}...")
    doc = docx.Document(doc_path)
    
    start_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "4.2.5 Kendala Pelaksanaan Lapangan dan Solusi Eksekusi" in p.text:
            start_idx = i
            break
            
    if start_idx == -1:
        print("4.2.5 not found")
        return
        
    end_idx = -1
    for i in range(start_idx + 1, len(doc.paragraphs)):
        if "KESIMPULAN" in doc.paragraphs[i].text.upper() or "BAB V" in doc.paragraphs[i].text.upper():
            end_idx = i
            break
            
    if end_idx == -1:
        end_idx = len(doc.paragraphs)
        
    print(f"Removing paragraphs from index {start_idx} to {end_idx - 1}")
    
    # We delete the xml elements of these paragraphs
    for i in range(start_idx, end_idx):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)
        
    doc.save(doc_path)
    print(f"Saved {doc_path}")

if __name__ == "__main__":
    file1 = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
    file2 = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
    remove_425(file1)
    remove_425(file2)
