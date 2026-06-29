import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Cm

new_text = [
    "Hasil pengembangan Website ULT FKIP Unila ini membuktikan secara empiris bahwa digitalisasi bukan sekadar memindahkan dokumen fisik ke dalam bentuk digital, melainkan sebuah transformasi tata kelola pelayanan yang holistik. Berangkat dari kondisi empiris yang diuraikan pada Latar Belakang (Bab I), akar permasalahan administrasi fakultas bermuara pada ketiadaan sistem pencatatan terpusat yang memicu penumpukan berkas fisik di loket, redundansi proses, hingga tingginya risiko kehilangan dokumen. Sebagai respons teknis atas urgensi tersebut, penerapan model Research and Development (R&D) bersintaks ADDIE telah berhasil melahirkan sebuah aplikasi web monolitik yang andal. Melalui analisis kebutuhan, perancangan arsitektur, hingga implementasi yang iteratif, platform digital ini diarsiteki dengan mekanisme Role-Based Access Control (RBAC) yang ketat untuk mengisolasi fungsionalitas pengguna secara proporsional. Keseluruhan proses rekayasa perangkat lunak ini secara komprehensif menggeser loket fisik konvensional ke dalam ekosistem digital yang sepenuhnya otonom.",
    "Salah satu temuan paling krusial dari implementasi sistem ini adalah keberhasilan pemecahan masalah krisis transparansi dan kebuntuan komunikasi birokrasi yang sebelumnya membelenggu pelayanan fakultas. Persoalan dokumen yang hilang atau menumpuk di tingkat program studi berhasil dieliminasi secara total melalui penerapan fitur pelacakan status surat (tracking). Fitur ini bekerja dengan menyajikan linimasa audit (auditable timeline) secara real-time yang merunut secara presisi setiap pergerakan dokumen mulai dari tahap diajukan oleh mahasiswa, direview oleh staf atau admin jurusan, ditandatangani secara elektronik oleh pimpinan, hingga kemungkinan ditolak beserta catatan perbaikannya. Melalui mekanisme ini, mahasiswa tidak lagi diharuskan melakukan konfirmasi fisik secara bolak-balik karena setiap perpindahan status layanan langsung terekam dan memicu notifikasi pembaruan. Dengan demikian, fitur pelacakan status ini membuktikan secara logis dan sistematis bahwa transparansi alur birokrasi dapat ditegakkan, sekaligus memastikan akuntabilitas setiap pejabat yang berwenang dalam memproses dokumen tersebut tanpa penundaan yang tidak beralasan.",
    "Keberhasilan pemecahan masalah transparansi tersebut terefleksi kuat pada penerimaan pengguna akhir, yang dibuktikan dengan capaian persentase kepraktisan sebesar 92,13% (kategori Sangat Praktis). Angka ini bukan sekadar kuantifikasi keberhasilan teknis, melainkan representasi empiris dari tingginya kualitas User Experience (UX) dan efisiensi sistem yang berhasil dihadirkan. Jika disandingkan dengan landasan teori sistem informasi pada Bab II, sebagaimana dikaji oleh Ferdiansyah et al. (2022) dan Maulana et al. (2023), tingginya persentase kepraktisan tersebut bermakna bahwa antarmuka sistem secara signifikan berhasil meminimalkan beban kognitif pengguna dalam menavigasi informasi dan menyelesaikan formulir dinamis. Desain tata letak yang asimetris dan modular memastikan pengguna dapat menemukan fitur yang mereka perlukan dengan cepat tanpa kebingungan. Oleh karena itu, capaian 92,13% ini mengonfirmasi landasan teori digitalisasi bahwa antarmuka yang intuitif berbanding lurus dengan kepuasan pengguna. Lebih jauh, hasil ini melahirkan paradigma baru dalam administrasi kampus Universitas Lampung bahwa perpaduan antara otomatisasi pemrosesan dokumen, privasi penyimpanan berlapis, dan kesederhanaan desain adalah kunci utama dalam mengoptimalkan penerimaan mahasiswa terhadap teknologi layanan publik baru.",
    "Di balik tingginya validitas dan kepraktisan yang diraih, proses pengembangan perangkat lunak ini tidak luput dari dinamika rintangan teknis dan non-teknis di lapangan yang harus diselesaikan secara presisi. Secara teknis, kendala paling berat terjadi pada modul perakitan dokumen akibat adanya konflik render HTML dari komponen editor teks kaya (Tiptap Editor) ke dalam format dokumen Word (.docx). Masukan teks mahasiswa yang memuat tag-tag visual awalnya tercetak mentah tanpa gaya format apa pun pada dokumen luaran, sehingga merusak estetika naskah surat resmi. Untuk menanggulangi anomali ini, peneliti merancang dan mengimplementasikan algoritma parser HTML-to-OpenXML khusus secara terpusat pada kerangka XML bawaan. Mesin parser ini bertugas menduplikasi pewarisan atribut gaya dari placeholder aslinya sebelum mengganti simpul teks, sehingga hasil akhir dokumen mampu mempertahankan integritas visual tata naskah dinas fakultas secara sempurna. Selain masalah perakitan, sistem juga menghadapi kerentanan serangan injeksi keamanan Cross-Site Scripting (XSS) pada formulir pengajuan publik, yang dengan sigap ditanggulangi melalui rekonfigurasi HTTP Header Content Security Policy (CSP) dan implementasi modul sanitasi input yang ketat pada kerangka kerja sistem.",
    "Pada dimensi non-teknis, tantangan terbesar muncul pada fase uji coba lapangan, khususnya terkait dengan manajemen sumber daya manusia. Dalam pelaksanaan evaluasi skala terbatas, peneliti berhadapan dengan fenomena resistensi dan lambatnya respons dari beberapa subjek mahasiswa serta staf operasional saat diminta kesediaannya melakukan pengujian. Kesibukan akademik mahasiswa serta beban kerja staf fakultas yang tinggi membuat jadwal pengujian kerap mengalami penundaan. Menghadapi situasi lapangan yang dinamis tersebut, peneliti mengeksekusi strategi manajerial yang berfokus pada pendekatan persuasif dan fleksibilitas jadwal. Peneliti mengadakan sesi pendampingan secara komprehensif dari meja ke meja, memberikan panduan visual ringkas tentang cara penggunaan sistem, dan membuka jalur komunikasi responsif penuh untuk memandu para responden. Solusi proaktif ini terbukti efektif dalam membangun kembali antusiasme serta kepercayaan partisipan, yang pada akhirnya memastikan seluruh data uji kepraktisan dapat terhimpun secara akurat, valid, dan merepresentasikan pengalaman pengguna yang sebenarnya. Rangkaian analisis, desain, penyelesaian kendala, hingga evaluasi ini mengukuhkan bahwa Website ULT FKIP Universitas Lampung telah mencapai maturitas teknis yang tinggi untuk mendukung digitalisasi operasional kampus secara berkelanjutan."
]

def update_docx(filepath, highlight=False):
    doc = docx.Document(filepath)
    
    start_el = None
    end_el = None
    
    for p in doc.paragraphs:
        text = p.text.strip().lower()
        if "pembahasan hasil penelitian" == text:
            start_el = p._element
        elif text == "bab v" or "kesimpulan dan saran" == text:
            if start_el is not None:
                end_el = p._element
                break
                
    if start_el is None or end_el is None:
        print(f"Failed to find bounds in {filepath}")
        return
        
    curr = start_el.getnext()
    while curr is not None and curr != end_el:
        nxt = curr.getnext()
        curr.getparent().remove(curr)
        curr = nxt
        
    for txt in new_text:
        new_p = docx.text.paragraph.Paragraph(docx.oxml.OxmlElement('w:p'), doc)
        end_el.addprevious(new_p._element)
        
        run = new_p.add_run(txt)
        if highlight:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
        new_p.style = doc.styles['Normal']
        new_p.paragraph_format.first_line_indent = Cm(1.27)
        new_p.paragraph_format.line_spacing = 1.5
        new_p.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(filepath)
    print(f"Updated {filepath}")

import glob
files = [
    r"c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx",
    r"c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx"
]

for f in files:
    is_hl = "Highlighted" in f
    update_docx(f, is_hl)
