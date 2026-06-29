import docx
import sys

def insert_bold_heading(doc, paragraph_idx, text):
    # Insert a new paragraph before the given index
    target_p = doc.paragraphs[paragraph_idx]
    new_p = target_p.insert_paragraph_before("")
    new_p.style = target_p.style # inherit style to keep spacing consistent if any
    
    # Add bold run
    run = new_p.add_run(text)
    run.bold = True
    
    # We might want to keep the alignment justified if the style defaults to something else
    # or left-aligned. Usually headings are left aligned or justified. Let's just use the style's default.

def process_document(doc_path):
    print(f"Processing {doc_path}...")
    doc = docx.Document(doc_path)
    
    # Definitions of target paragraphs and their new headings
    # We match by checking if the paragraph text starts with a specific string
    
    keunggulan_points = [
        (
            "Manajemen hak akses pada sistem konvensional",
            "a. Manajemen Hak Akses Berbasis Peran (Role-Based Access Control)"
        ),
        (
            "Siklus perakitan surat resmi pada operasional konvensional",
            "b. Otomatisasi Perakitan Dokumen Resmi (OpenXML)"
        ),
        (
            "Infrastruktur pengelolaan keamanan data pada era",
            "c. Keamanan Infrastruktur Penyimpanan Privat"
        ),
        (
            "Prosedur penelusuran rekam jejak progres",
            "d. Pelacakan Riwayat Transaksi Layanan (Audit Trail)"
        ),
        (
            "Keterbatasan purwarupa portal peladen",
            "e. Aksesibilitas Sistem Lintas Perangkat (Progressive Web App)"
        )
    ]
    
    kendala_points = [
        (
            "Evaluasi tahapan purwarupa awal sistem web",
            "a. Resolusi Degradasi Pemformatan Teks Dokumen"
        ),
        (
            "Proses pemindaian kelayakan infrastruktur keamanan",
            "b. Penguatan Validasi Keamanan Input Data (Anti-XSS)"
        ),
        (
            "Tim evaluator antarmuka interaktif memberikan",
            "c. Penyesuaian Komposisi Warna dan Tata Letak Antarmuka"
        ),
        (
            "Tahapan asesmen identifikasi observasi operasional",
            "d. Otomatisasi Penomoran Surat dan Penjadwalan"
        ),
        (
            "Tanggapan rujukan pengawasan instrumen penilaian",
            "e. Optimalisasi Tata Letak dan Ruang Kosong Tabel"
        )
    ]
    
    all_points = keunggulan_points + kendala_points
    
    # We need to find the indices of these paragraphs.
    # Since inserting a paragraph shifts the indices, we should find them all first,
    # sort by index in descending order, and then insert so we don't mess up the indices.
    
    insertions = []
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        for target_text, heading in all_points:
            if text.startswith(target_text):
                insertions.append((i, heading))
                break
                
    # Sort in reverse order
    insertions.sort(key=lambda x: x[0], reverse=True)
    
    for idx, heading in insertions:
        print(f"Inserting '{heading}' before paragraph {idx}")
        insert_bold_heading(doc, idx, heading)
        
    doc.save(doc_path)
    print(f"Saved {doc_path}\n")

if __name__ == "__main__":
    file1 = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
    file2 = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
    process_document(file1)
    process_document(file2)
