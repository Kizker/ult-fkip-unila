import re
from docx import Document
import nltk
from Sastrawi.Stemmer.StemmerFactory import StemmerFactory

def main():
    print("Downloading NLTK words...")
    try:
        nltk.download('words', quiet=True)
        english_words = set(w.lower() for w in nltk.corpus.words.words())
    except:
        print("Failed to load NLTK words. Please download manually.")
        return

    print("Loading KBBI words from Sastrawi...")
    factory = StemmerFactory()
    kbbi_words = set(factory.get_words())
    
    # Add common Indonesian prefixes/suffixes and words not in the basic list
    additional_id_words = {'dan', 'yang', 'di', 'ke', 'dari', 'pada', 'dalam', 'untuk', 'dengan', 
                          'ini', 'itu', 'adalah', 'merupakan', 'sebagai', 'tidak', 'akan', 
                          'telah', 'bisa', 'dapat', 'oleh', 'saat', 'setelah', 'sebelum', 'jika',
                          'maka', 'karena', 'sehingga', 'namun', 'tetapi', 'serta', 'juga', 'atau',
                          'kita', 'kami', 'saya', 'anda', 'mereka', 'dia', 'ia', 'beliau',
                          'nya', 'pun', 'lah', 'kah', 'tah'}
    kbbi_words.update(additional_id_words)

    # Some English words to skip because they are very common in Indonesian or false positives
    false_positives = {'di', 'ke', 'area', 'data', 'format', 'informasi', 'sistem', 'proses', 
                       'program', 'struktur', 'dokumen', 'akses', 'opsi', 'status', 'admin',
                       'portal', 'detail', 'menu', 'visi', 'misi', 'rekap', 'layout', 'margin',
                       'digital', 'fakultas', 'kampus', 'universitas', 'institusi', 'prosedur',
                       'verifikasi', 'validasi', 'evaluasi', 'implementasi', 'integrasi',
                       'visual', 'teks', 'paragraf', 'tabel', 'gambar', 'lampiran', 'bab',
                       'subbab', 'kalimat', 'kata', 'huruf', 'spasi', 'nomor', 'halaman',
                       'abstrak', 'daftar', 'isi', 'pustaka', 'referensi', 'penelitian',
                       'pengembangan', 'metode', 'teknik', 'analisis', 'desain', 'uji',
                       'coba', 'hasil', 'pembahasan', 'kesimpulan', 'saran', 'lampiran'}
    kbbi_words.update(false_positives)
    
    print("Reading docx...")
    doc = Document(r"c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx")
    
    found_foreign = set()
    foreign_with_context = []

    for para in doc.paragraphs:
        # Check if the paragraph has text
        text = para.text
        if not text.strip():
            continue
            
        words = re.findall(r'\b[a-zA-Z]+\b', text)
        for w in words:
            wl = w.lower()
            if wl in kbbi_words:
                continue
            if len(wl) <= 2:
                continue
            if wl in english_words:
                found_foreign.add(wl)
                foreign_with_context.append((wl, text.strip()))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    if not text.strip():
                        continue
                    words = re.findall(r'\b[a-zA-Z]+\b', text)
                    for w in words:
                        wl = w.lower()
                        if wl in kbbi_words:
                            continue
                        if len(wl) <= 2:
                            continue
                        if wl in english_words:
                            found_foreign.add(wl)
                            foreign_with_context.append((wl, text.strip()))

    # Also add words that are known IT terms even if not in NLTK words (like 'website', 'online')
    known_it_terms = {'website', 'online', 'offline', 'database', 'software', 'hardware', 
                      'framework', 'update', 'delete', 'create', 'read', 'admin', 'student', 
                      'staff', 'signer', 'gatekeeper', 'assembly', 'workflow', 'template', 
                      'download', 'upload', 'button', 'input', 'output', 'blackbox', 
                      'use', 'case', 'flowchart', 'sequence', 'activity', 'entity', 
                      'relationship', 'state', 'chart', 'class', 'component', 'deployment',
                      'reliability', 'validity', 'error', 'bug', 'user', 'interface', 'server',
                      'login', 'logout', 'dashboard', 'file', 'folder', 'browser', 'internet',
                      'web', 'app', 'application', 'mobile', 'desktop', 'responsive', 'ui', 'ux',
                      'frontend', 'backend', 'fullstack', 'api', 'json', 'xml', 'html', 'css',
                      'javascript', 'php', 'mysql', 'sql', 'query', 'table', 'row', 'column',
                      'primary', 'key', 'foreign', 'index', 'view', 'controller', 'model',
                      'routing', 'middleware', 'authentication', 'authorization', 'session',
                      'cookie', 'token', 'jwt', 'rest', 'soap', 'graphql', 'client', 'host',
                      'domain', 'hosting', 'cloud', 'aws', 'gcp', 'azure', 'docker', 'container',
                      'git', 'github', 'gitlab', 'bitbucket', 'repository', 'commit', 'push',
                      'pull', 'merge', 'branch', 'tag', 'release', 'deploy', 'build', 'test',
                      'debug', 'log', 'error', 'exception', 'warning', 'info', 'notice',
                      'local', 'storage', 'cache', 'memory', 'cpu', 'ram', 'disk', 'network',
                      'bandwidth', 'latency', 'ping', 'ip', 'address', 'port', 'protocol',
                      'http', 'https', 'ftp', 'ssh', 'ssl', 'tls', 'certificate', 'key',
                      'public', 'private', 'symmetric', 'asymmetric', 'encryption', 'decryption',
                      'hash', 'salt', 'pepper', 'password', 'credential', 'secret', 'key',
                      'security', 'vulnerability', 'exploit', 'patch', 'update', 'upgrade',
                      'install', 'uninstall', 'configure', 'setup', 'initialize', 'start',
                      'stop', 'restart', 'pause', 'resume', 'status', 'monitor', 'alert',
                      'notification', 'message', 'email', 'sms', 'push', 'pull', 'sync',
                      'async', 'await', 'promise', 'callback', 'event', 'listener', 'handler',
                      'trigger', 'hook', 'action', 'filter', 'plugin', 'theme', 'template',
                      'widget', 'module', 'package', 'library', 'dependency', 'framework',
                      'architecture', 'pattern', 'design', 'solid', 'dry', 'kiss', 'yagni',
                      'agile', 'scrum', 'kanban', 'waterfall', 'sprint', 'backlog', 'epic',
                      'story', 'task', 'bug', 'issue', 'ticket', 'board', 'chart', 'graph',
                      'report', 'analytics', 'statistics', 'metric', 'dimension', 'funnel',
                      'conversion', 'retention', 'churn', 'engagement', 'revenue', 'cost',
                      'profit', 'margin', 'roi', 'kpi', 'okr', 'goal', 'target', 'objective',
                      'strategy', 'tactic', 'plan', 'execution', 'evaluation', 'feedback',
                      'review', 'approval', 'rejection', 'comment', 'reply', 'like', 'share',
                      'follow', 'unfollow', 'subscribe', 'unsubscribe', 'friend', 'connection',
                      'profile', 'account', 'setting', 'preference', 'option', 'choice',
                      'select', 'radio', 'checkbox', 'text', 'password', 'number', 'date',
                      'time', 'datetime', 'month', 'week', 'color', 'range', 'file', 'image',
                      'video', 'audio', 'document', 'pdf', 'word', 'excel', 'powerpoint',
                      'zip', 'rar', 'tar', 'gz', 'bz2', '7z', 'iso', 'img', 'bin', 'exe',
                      'dll', 'so', 'dylib', 'jar', 'war', 'ear', 'class', 'java', 'c', 'cpp',
                      'cs', 'py', 'rb', 'js', 'ts', 'go', 'rs', 'swift', 'kt', 'dart', 'sh',
                      'bat', 'cmd', 'ps1', 'sql', 'md', 'txt', 'csv', 'tsv', 'ini', 'cfg',
                      'conf', 'yaml', 'yml', 'toml', 'xml', 'json', 'env', 'gitignore',
                      'dockerignore', 'npmignore', 'editorconfig', 'prettierrc', 'eslintrc'}
    
    # Also find known IT terms
    for para in doc.paragraphs:
        text = para.text
        if not text.strip(): continue
        words = re.findall(r'\b[a-zA-Z]+\b', text)
        for w in words:
            wl = w.lower()
            if wl in known_it_terms and wl not in kbbi_words:
                found_foreign.add(wl)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    if not text.strip(): continue
                    words = re.findall(r'\b[a-zA-Z]+\b', text)
                    for w in words:
                        wl = w.lower()
                        if wl in known_it_terms and wl not in kbbi_words:
                            found_foreign.add(wl)

    print(f"Found {len(found_foreign)} unique potential foreign words.")
    
    with open("potential_foreign_words.txt", "w", encoding="utf-8") as f:
        for w in sorted(list(found_foreign)):
            f.write(f"{w}\n")
            
    print("Saved potential_foreign_words.txt")

if __name__ == "__main__":
    main()
