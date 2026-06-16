import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

IMG_DIR = r"C:\laragon\www\ult-fkip-unila\docs\skripsi\rancangan_diagram\screenshots"

targets = {
    "Gambar 4.x. Tampilan Beranda Utama Sistem (Public Portal)": ("01-beranda-sistem.jpg", "Tampilan Beranda Utama Sistem (Public Portal)"),
    "Gambar 4.x. Tampilan Katalog Layanan Administrasi": ("02-daftar-layanan.jpg", "Tampilan Katalog Layanan Administrasi"),
    "Gambar 4.x. Tampilan Halaman Berita dan Pengumuman": ("004-blog-index.jpg", "Tampilan Halaman Berita dan Pengumuman"),
    "Gambar 4.x. Tampilan Halaman Autentikasi Pengguna (Login)": ("05-login.jpg", "Tampilan Halaman Autentikasi Pengguna (Login)"),
    "Gambar 4.x. Tampilan Dasbor Utama Mahasiswa (Student Portal)": ("06-dashboard-pemohon.jpg", "Tampilan Dasbor Utama Mahasiswa (Student Portal)"),
    "Gambar 4.x. Tampilan Formulir Pengajuan Dokumen Dinamis": ("08-form-pengajuan.jpg", "Tampilan Formulir Pengajuan Dokumen Dinamis"),
    "Gambar 4.x. Tampilan Pelacakan Riwayat dan Status (Audit Trail)": ("09-riwayat-permohonan.jpg", "Tampilan Pelacakan Riwayat dan Status (Audit Trail)"),
    "Gambar 4.x. Tampilan Dasbor Staf ULT dan Admin Prodi": ("admin_dashboard__desktop-1366.jpg", "Tampilan Dasbor Staf ULT dan Admin Prodi"),
    "Gambar 4.x. Tampilan Pratinjau Dokumen dan Review Berkas": ("admin_requests_show__desktop-1366.jpg", "Tampilan Pratinjau Dokumen dan Review Berkas"),
    "Gambar 4.x. Tampilan Manajemen Placeholder Template Word": ("admin_doc_formats_show__desktop-1366.jpg", "Tampilan Manajemen Placeholder Template Word"),
    "Gambar 4.x. Tampilan Persetujuan dan Verifikasi Pejabat (Signer)": ("07-detail-penandatanganan.jpg", "Tampilan Persetujuan dan Verifikasi Pejabat (Signer)"),
}

def inject_screenshots(doc_path):
    print(f"Injecting into: {doc_path}")
    doc = docx.Document(doc_path)
    
    modified = False
    for p in doc.paragraphs:
        text = p.text.strip()
        
        # Sometimes there's a leading newline or something in the text, so we check if the target string is in the paragraph text
        for target_key, (img_file, caption_text) in targets.items():
            if target_key in text:
                img_path = os.path.join(IMG_DIR, img_file)
                
                if os.path.exists(img_path):
                    # Clear paragraph text (this removes the old picture and caption)
                    p.text = ''
                    # Add new picture
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(5.5))
                    
                    run.add_break()
                    caption_run = p.add_run(f"Gambar 4.x. {caption_text}")
                    caption_run.font.name = 'Times New Roman'
                    caption_run.font.size = Pt(10)
                    
                    # Ensure the whole paragraph is centered
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(12)
                    
                    modified = True
                    print(f"Injected {img_file}")
                else:
                    print(f"Image not found: {img_path}")
                break # Move to next paragraph once matched
                
    if modified:
        doc.save(doc_path)
        print("Save successful.\n")

inject_screenshots(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx")
inject_screenshots(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx")
