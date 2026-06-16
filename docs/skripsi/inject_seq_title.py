import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

doc_clean_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
doc_hl_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'

title_text = "Perancangan Interaksi Dinamis Menggunakan Diagram Sequence"

def clone_num_pr(source_p, target_p):
    if source_p._p.pPr is not None and source_p._p.pPr.numPr is not None:
        target_p._p.get_or_add_pPr().append(deepcopy(source_p._p.pPr.numPr))

def fix_document(file_path):
    doc = docx.Document(file_path)
    
    erd_title_idx = -1
    seq_caption_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        if 'Perancangan Struktur Data Melalui' in p.text:
            erd_title_idx = i
        elif 'Gambar 8. Diagram Sequence Perakitan' in p.text:
            seq_caption_idx = i

    print(f"File: {file_path}")
    print(f"ERD Title at: {erd_title_idx}, Seq Caption at: {seq_caption_idx}")
    
    # Check if title already exists
    exists = False
    for i in range(erd_title_idx, seq_caption_idx):
        if 'Perancangan Interaksi Dinamis' in doc.paragraphs[i].text:
            exists = True
            break
            
    if exists:
        print("Sequence title already exists.")
        return
        
    # Find the image paragraph (should be right before seq_caption_idx)
    img_idx = seq_caption_idx - 1
    # We will insert the title right before the image.
    new_p = doc.paragraphs[img_idx].insert_paragraph_before()
    
    # Format text
    new_p.text = ""
    new_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    new_p.paragraph_format.line_spacing = 1.5
    # The ERD title might not have first_line_indent, it relies on numbering. Let's see:
    source_p = doc.paragraphs[erd_title_idx]
    if source_p.paragraph_format.first_line_indent:
        new_p.paragraph_format.first_line_indent = source_p.paragraph_format.first_line_indent
    if source_p.paragraph_format.left_indent:
        new_p.paragraph_format.left_indent = source_p.paragraph_format.left_indent
        
    run = new_p.add_run(title_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    # The words "Sequence" needs to be italicized
    # We can do this manually
    new_p.clear()
    
    parts = title_text.split("Sequence")
    run1 = new_p.add_run(parts[0])
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)
    
    run2 = new_p.add_run("Sequence")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    run2.italic = True
    
    # Add numbering properties by cloning from ERD title
    clone_num_pr(source_p, new_p)
    
    doc.save(file_path)
    print("Saved sequence title.")

fix_document(doc_clean_path)
fix_document(doc_hl_path)
