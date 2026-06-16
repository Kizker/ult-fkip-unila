import docx
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_clean_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
doc_hl_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'

def add_table_caption(file_path):
    doc = docx.Document(file_path)
    target_p = None
    
    for p in doc.paragraphs:
        if 'Rekapitulasi tanggapan kualitatif dari responden beserta tindakan' in p.text:
            target_p = p
            break
            
    if target_p is None:
        print(f"Target paragraph not found in {file_path}")
        return
        
    parent = target_p._element.getparent()
    idx = parent.index(target_p._element)
    
    # Check if we already inserted it
    next_el = parent[idx+1]
    if next_el.tag.endswith('p'):
        # Check text
        para = docx.text.paragraph.Paragraph(next_el, target_p._parent)
        if 'Tabel 4.2' in para.text:
            print(f"Caption already exists in {file_path}")
            return
            
    # Insert new paragraph
    new_p_elem = OxmlElement('w:p')
    target_p._element.addnext(new_p_elem)
    
    para = docx.text.paragraph.Paragraph(new_p_elem, target_p._parent)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = 1.0
    
    run = para.add_run("Tabel 4.2 Rekapitulasi Komentar dan Tindak Lanjut Uji Kepraktisan Pengguna")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.save(file_path)
    print(f"Table caption added successfully to {file_path}")

add_table_caption(doc_clean_path)
add_table_caption(doc_hl_path)
