import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_clean_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx'
doc_hl_path = r'c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'

erd_paras = [
    "Diagram Entity-Relationship (ERD) pada Gambar 7 memvisualisasikan arsitektur dan relasi antara delapan entitas utama di dalam basis data MySQL yang menopang keseluruhan sistem. Setiap entitas direpresentasikan sebagai sebuah tabel persegi panjang dengan tajuk gelap, di mana baris-baris di bawahnya menjabarkan daftar atribut (kolom) beserta spesifikasi tipe datanya, seperti BIGINT, VARCHAR, TEXT, BOOLEAN, maupun TIMESTAMP. Garis-garis solid yang menghubungkan tabel-tabel ini menunjukkan struktur relasi fisik berupa foreign key constraints.",
    "Entitas users, roles, beserta tabel pivot model_has_roles membentuk sebuah ekosistem manajemen identitas yang solid untuk mendukung fungsionalitas Role-Based Access Control (RBAC). Entitas users menyimpan kredensial pengguna, sementara roles mendefinisikan level akses, yang kemudian dipetakan secara dinamis melalui tabel pivot. Tabel services berperan sebagai katalog sentral yang menampung informasi jenis layanan persuratan beserta jalur referensi lokasi templat fisik dokumen (template_path).",
    "Sebagai pusat operasional transaksional, entitas requests bertindak untuk merekam setiap permohonan yang masuk, menyimpan status alur, kode pelacakan (tracking_code), dan menerbitkan nomor surat (issue_number). Mengingat setiap layanan memiliki kebutuhan isian yang berbeda-beda, sistem mengimplementasikan entitas request_inputs berformat pasangan kunci-nilai (field_name dan field_value) yang berelasi One-to-Many dengan requests. Mekanisme dinamis ini memastikan skema database tidak perlu direstrukturisasi setiap kali terdapat penambahan bidang isian baru. Selain itu, entitas request_documents secara fungsional dikhususkan untuk menangani metadata berkas, meliputi jalur file fisik, tipe, dan ukuran file lampiran.",
    "Alur kerja relasional ERD ini menjamin integritas data dari awal hingga akhir. Saat mahasiswa membuat permohonan, rekam induk (parent record) tercipta di tabel requests dengan status awal. Data borang dinamis terhubung masuk berantai ke request_inputs, dan seluruh lampiran diindeks di request_documents. Segala aktivitas perubahan status yang dipicu oleh intervensi staf maupun pejabat akan selalu memicu perekaman baris data baru di entitas audit_trails. Entitas audit log ini menyimpan tindakan riwayat aksi (action) beserta alamat IP aktor bersangkutan, berfungsi ganda sebagai lapisan pengamanan historis untuk memastikan akuntabilitas operasional sistem administrasi."
]

seq_paras = [
    "Diagram sequence pada Gambar 8 menyajikan visualisasi interaksi dinamis berorientasi waktu (time-oriented) yang mendeskripsikan secara spesifik bagaimana objek-objek sistem berkolaborasi dalam menyelesaikan proses perakitan dokumen tingkat rendah. Antarmuka diagram tersusun atas elemen lifeline vertikal yang merepresentasikan kelima entitas atau aktor yang berpartisipasi, direpresentasikan oleh kotak hitam di posisi teratas. Garis-garis panah horizontal bersiklus menggambarkan urutan spesifik pengiriman pesan (message passing) atau pemanggilan metode (method calls) yang mengalir kronologis dari atas ke bawah.",
    "Komponen yang terlibat secara menyeluruh meliputi aktor Staf ULT / User sebagai inisiator pemicu, RequestController sebagai antarmuka pengatur rute permintaan pada level kerangka kerja aplikasi, DocumentAssemblerService yang memegang mandat utama orkestrasi perakitan dan ekstraksi templat, kelas utilitas independen HtmlToOpenXMLParser sebagai mesin murni penerjemah tag HTML kaya ke dalam blok XML, serta Private Storage sebagai representasi infrastruktur tujuan penyimpanan fisik berkas keluaran.",
    "Skenario perakitan dokumen dimulai pada siklus langkah 1, saat staf ULT menekan tombol perakitan dari antarmuka dasbor yang mengirimkan sinyal clickAssemble(requestId) menuju RequestController. Pada langkah 2, pengendali mendelegasikan perintah komputasi transaksional ini seutuhnya ke kelas DocumentAssemblerService lewat instruksi assembleDocument. Objek service ini kemudian mempersiapkan lingkungan kerja di langkah 3 dengan memuat fail templat dokumen mentah dan mengekstraksi formulir berisikan data dari basis data (loadTemplateAndFormData). Pada titik delegasi kritis di langkah 4, service mengoper beban terjemahan teks kaya (WYSIWYG) seutuhnya kepada HtmlToOpenXMLParser dengan mengeksekusi metode writeHtmlToWordRun.",
    "HtmlToOpenXMLParser lantas mengambil alih proses eksekusi mendalam di langkah 5 secara rekursif (traverseHtmlAndInsertRuns) untuk membedah satu demi satu elemen DOM HTML, seperti tag paragraf, tebal, maupun miring. Guna mempertahankan nilai estetik dan gaya visual orisinal templat Word, parser tidak memformat teks dari nol melainkan mengeksekusi metode cloneRunPropertiesAndFormat (langkah 6) untuk menduplikasi atribut font pewarisan dari placeholder aslinya, menghindari degradasi dokumen. Setelah seluruh format OpenXML beserta atribut silsilahnya berhasil dirajut (langkah 7), himpunan instruksi simpul dikembalikan kepada DocumentAssemblerService.",
    "Pada siklus penyelesaian di langkah 8, DocumentAssemblerService merajut kembali seluruh struktur berkas XML yang telah dimodifikasi menjadi satu kesatuan arsip biner ZIP berbentuk fail .docx yang utuh (compileAndZipArchive). Berkas fisik ini selanjutnya ditransfer ke objek Private Storage pada langkah 9 melalui perintah putFileInPrivateDisk untuk dikarantina dan diamankan dari akses anonim. Siklus diakhiri dengan mekanisme sinyal balasan (return value) yang merambat mundur; status keberhasilan persisten dari penyimpanan di langkah 10, diteruskan sebagai respons HTTP sukses menuju kontroler (langkah 11), dan bermuara kembali pada layar peramban aktor utama sebagai pemanggilan fungsi redirectWithSuccessToast (langkah 12), mengkonfirmasi kepada pengguna bahwa rantai proses perakitan dokumen digital sistem telah paripurna."
]

def format_paragraph(p, text, is_hl=False):
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if is_hl:
        from docx.enum.text import WD_COLOR_INDEX
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

def fix_document(file_path, is_hl=False):
    doc = docx.Document(file_path)
    
    erd_idx = -1
    seq_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        if 'Gambar 7. Diagram Hubungan Entitas (ERD)' in p.text:
            erd_idx = i
        elif 'Gambar 8. Diagram Sequence Perakitan' in p.text:
            seq_idx = i

    print(f"File: {file_path}")
    print(f"Gambar 7 at: {erd_idx}, Gambar 8 at: {seq_idx}")
    
    # We must insert text below the titles.
    # Diagram Sequence is after ERD, so let's process Sequence first to not mess up ERD indices.
    
    # Sequence insertion point: right after `seq_idx`. Since we use insert_paragraph_before, we insert before `seq_idx + 1`
    seq_insert_idx = seq_idx + 1
    # Check if there's already text (to avoid duplicating if we run it twice)
    if seq_insert_idx < len(doc.paragraphs) and "Diagram sequence pada Gambar 8" in doc.paragraphs[seq_insert_idx].text:
        print("Sequence text already exists.")
    else:
        for text in reversed(seq_paras):
            new_p = doc.paragraphs[seq_insert_idx].insert_paragraph_before()
            format_paragraph(new_p, text, is_hl)
            
    # ERD insertion point: right after `erd_idx`
    erd_insert_idx = erd_idx + 1
    if erd_insert_idx < len(doc.paragraphs) and "Diagram Entity-Relationship (ERD)" in doc.paragraphs[erd_insert_idx].text:
        print("ERD text already exists.")
    else:
        for text in reversed(erd_paras):
            new_p = doc.paragraphs[erd_insert_idx].insert_paragraph_before()
            format_paragraph(new_p, text, is_hl)
            
    doc.save(file_path)
    print("Saved paragraphs.")

fix_document(doc_clean_path, is_hl=False)
fix_document(doc_hl_path, is_hl=True)
