import docx
import os

FILES = [
    r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx',
    r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
]

SARAN_TEXT = "Integrasi Penomoran Surat dan Penjadwalan Pimpinan: Disarankan untuk mengembangkan fitur kalender digital terintegrasi pada dasbor pimpinan fakultas yang tersinkronisasi langsung dengan mesin penomoran surat. Hal ini bertujuan untuk mencegah terjadinya tumpang tindih alokasi waktu pimpinan dan meminimalisir kesalahan cetak pada nomor dokumen kelulusan yang saat ini masih rentan terjadi."

def fix_saran():
    for filepath in FILES:
        print(f"Processing: {filepath}")
        doc = docx.Document(filepath)
        
        # We already removed the misplaced one. Now just insert the new one.
        for i, p in enumerate(doc.paragraphs):
            if 'Pengujian Keamanan Tingkat Lanjut:' in p.text:
                print(f"Found point 2 at line {i}")
                
                # Insert a new paragraph after point 2
                # We can do this by finding the next paragraph and inserting before it
                new_p = doc.paragraphs[i+1].insert_paragraph_before(SARAN_TEXT)
                
                # Copy the style (for MS Word numbering)
                new_p.style = p.style
                
                # Apply highlight if needed
                if 'Highlighted' in filepath:
                    for run in new_p.runs:
                        run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
                        
                print("Inserted point 3.")
                break
                
        doc.save(filepath)
        print("Saved.")

if __name__ == '__main__':
    fix_saran()
