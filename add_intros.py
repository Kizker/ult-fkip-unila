import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_intro(file_path, out_path, highlight=False):
    doc = docx.Document(file_path)
    
    # We will search for the headings "4.2.3 Keunggulan Sistem" and "4.2.4 Resolusi Kendala Teknis"
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        
        # 4.2.3 Keunggulan Sistem
        if text.startswith("4.2.3") and "Keunggulan Sistem" in text:
            # Check if next paragraph is already the intro
            if i + 1 < len(doc.paragraphs) and not doc.paragraphs[i+1].text.strip().startswith("a."):
                continue # Already has intro
            
            # Insert intro paragraph
            intro_p = p.insert_paragraph_before("")
            # Move it after the heading
            p._p.addnext(intro_p._p)
            
            intro_text = ("Pengembangan sistem informasi berbasis web pada Unit Layanan Terpadu (ULT) FKIP Unila "
                          "telah diupayakan sedemikian rupa untuk menghadirkan berbagai fitur yang solutif dan "
                          "efisien. Rancang bangun arsitektur yang dirumuskan pada tahap desain tidak hanya berfokus "
                          "pada fungsionalitas dasar, melainkan juga mengintegrasikan berbagai aspek mutakhir dalam "
                          "keamanan data dan fleksibilitas aksesibilitas. Berbagai kelebihan ini secara langsung "
                          "menjadi inovasi pembeda dibandingkan dengan proses pelayanan manual sebelumnya. Adapun "
                          "rincian keunggulan utama dari sistem yang telah diimplementasikan dapat dijabarkan sebagai berikut.")
            
            run = intro_p.add_run(intro_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if highlight:
                run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
                
            intro_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            intro_p.paragraph_format.first_line_indent = Pt(36) # 1.27 cm
            intro_p.paragraph_format.line_spacing = 1.5

        # 4.2.4 Resolusi Kendala Teknis
        if text.startswith("4.2.4") and "Kendala Teknis" in text:
            # Check if next paragraph is already the intro
            if i + 1 < len(doc.paragraphs) and not doc.paragraphs[i+1].text.strip().startswith("a."):
                continue # Already has intro
            
            intro_p = p.insert_paragraph_before("")
            p._p.addnext(intro_p._p)
            
            intro_text = ("Dalam proses pengembangan dan implementasi website ULT FKIP Unila, tidak terlepas dari "
                          "sejumlah hambatan teknis yang memerlukan penanganan khusus agar sistem dapat beroperasi "
                          "secara optimal. Kendala-kendala tersebut mencakup isu kompatibilitas format dokumen, "
                          "celah keamanan siber, hingga belum terakomodasinya infrastruktur otentikasi tunggal dari "
                          "pihak universitas. Namun demikian, berbagai tantangan ini telah dianalisis secara komprehensif "
                          "dan diberikan resolusi teknis yang memadai guna meminimalisasi risiko pada saat sistem digunakan "
                          "oleh sivitas akademika secara luas. Penjabaran mengenai kendala yang dihadapi beserta solusi "
                          "yang diterapkan adalah sebagai berikut.")
            
            run = intro_p.add_run(intro_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if highlight:
                run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
                
            intro_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            intro_p.paragraph_format.first_line_indent = Pt(36)
            intro_p.paragraph_format.line_spacing = 1.5

    doc.save(out_path)
    print(f"Processed {out_path}")

base_dir = r"c:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update"
clean = f"{base_dir}\\001_Skripsi_Andricha Dea Mitra_Clean.docx"
highlight = f"{base_dir}\\001_Skripsi_Andricha Dea Mitra_Highlighted.docx"

add_intro(clean, clean, highlight=False)
add_intro(highlight, highlight, highlight=True)
