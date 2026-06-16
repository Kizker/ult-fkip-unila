import docx
from docx.shared import Pt, Inches

def update_table(doc_path, is_highlighted=False):
    print(f"Updating table in {doc_path}")
    doc = docx.Document(doc_path)
    
    table_index = -1
    for i, table in enumerate(doc.tables):
        if not table.rows: continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if "Nama Validator" in header or "Komentar / Saran" in header:
            table_index = i
            break
            
    if table_index != -1:
        old_table = doc.tables[table_index]
        parent = old_table._element.getparent()
        index = parent.index(old_table._element)
        
        # Create a new table
        new_table = doc.add_table(rows=1, cols=3)
        new_table.style = old_table.style
        
        # Header
        new_table.cell(0, 0).text = "Jenis Validasi"
        new_table.cell(0, 1).text = "Saran dan Masukan"
        new_table.cell(0, 2).text = "Perbaikan"
        
        # Make header bold
        for cell in new_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
        
        # Data rows
        data = [
            ("Validasi Materi", 
             "1. Tambah penjelasan, cek kalimat panduan, dan pastikan screenshot beresolusi jelas.\n2. Beri garis pembatas antar bagian header, content, dan footer beserta gambar.", 
             "1. Penjelasan operasional dan struktur kalimat panduan telah direvisi agar lebih komunikatif, serta tangkapan layar (screenshot) diperbarui dengan resolusi tinggi.\n2. Telah ditambahkan pembatas visual (divider) yang jelas untuk memisahkan area header, konten utama, dan footer."),
             
            ("Validasi Media",
             "1. Komposisi warna sebaiknya jangan menggunakan gradasi yang terlalu mencolok.\n2. Berikan penanda ikon pada setiap tile/kartu layanan agar lebih representatif.",
             "1. Warna gradasi telah ditiadakan dan antarmuka disesuaikan menjadi warna solid (flat design) yang elegan dan konsisten.\n2. Ikon representatif yang relevan telah ditambahkan pada masing-masing kartu layanan untuk meningkatkan User Experience (UX)."),
             
            ("Validasi Sistem",
             "1. Terdapat beberapa form validasi input yang masih terlewat.\n2. Perhatikan masalah keamanan Content Security Policy (CSP), hindari penggunaan unsafe-eval/inline secara serampangan, dan atur default-src dengan ketat.\n3. Tambahkan fitur tombol 'Kembali ke atas' (Back to Top).",
             "1. Logika validasi form pada sisi klien (frontend) dan server (backend) telah dicek ulang secara komprehensif dan diperbaiki.\n2. Kebijakan CSP (Content Security Policy) pada HTTP Headers telah diperketat (strict) untuk mencegah kerentanan XSS dan sejenisnya.\n3. Tombol 'Back to Top' mengambang (floating) telah diimplementasikan untuk mempermudah navigasi pengguna.")
        ]
        
        for jenis, saran, perbaikan in data:
            row = new_table.add_row()
            row.cells[0].text = jenis
            row.cells[1].text = saran
            row.cells[2].text = perbaikan
            
            # Apply styling
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                        if is_highlighted:
                            from docx.enum.text import WD_COLOR_INDEX
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        
        # Insert new table into the document at the old table's location
        parent.insert(index, new_table._element)
        
        # Remove old table
        parent.remove(old_table._element)
        
        doc.save(doc_path)
        print("Success")

update_table(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx", False)
update_table(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx", True)
