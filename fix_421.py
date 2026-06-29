import docx
import os

FILES = [
    r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx',
    r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
]

NEW_PARAGRAPHS = [
    "Pencapaian skor kepraktisan sebesar 92,13% pada uji coba lapangan mengindikasikan keberhasilan fungsional sistem sekaligus keunggulan ergonomi antarmuka pengguna. Tingginya angka ini secara empiris mengonfirmasi teori User Experience (Maulana et al., 2023) yang menekankan urgensi perancangan antarmuka untuk meminimalkan beban kognitif (cognitive load) bagi pengguna pemula. Antarmuka portal mahasiswa (Student Portal) terbukti mampu memandu proses pengajuan dokumen secara intuitif tanpa memerlukan kurva pembelajaran yang curam.",
    "Evaluasi berkelanjutan terhadap pengalaman pengguna ini sejalan dengan temuan Ferdiansyah et al. (2022) mengenai pentingnya aspek navigasi dan kejelasan informasi dalam menentukan kualitas layanan digital. Keterlibatan sivitas akademika dalam memberikan umpan balik selama fase uji coba juga merepresentasikan penerapan prinsip User-Centered Design sebagaimana digagas oleh Sastranegara et al. (2023). Integrasi masukan langsung dari mahasiswa dan staf ULT ini secara krusial berhasil meningkatkan skor kepraktisan produk akhir setelah melalui proses desain ulang yang terarah.",
    "Transformasi digital pada sistem administrasi ini pada akhirnya bukan sekadar proses pemindahan dokumen fisik ke format elektronik, melainkan sebuah rekayasa ulang proses bisnis (business process reengineering). Perubahan fundamental ini mendukung teori Purwani et al. (2024) mengenai pemanfaatan teknologi digital untuk menciptakan model layanan pendidikan yang lebih efisien dan terbebas dari rantai birokrasi tradisional. Efektivitas adaptasi teknologi ini turut membuktikan pandangan Poernamawatie et al. (2023) bahwa otomatisasi layanan administrasi mampu mempercepat laju koordinasi lintas divisi meski dihadapkan pada hambatan literasi digital pengguna awal.",
    "Keberhasilan implementasi aplikasi pelayanan terpadu ini sangat bergantung pada dukungan struktural dan manajerial dari pihak pimpinan fakultas. Sinergi antara inovasi teknologi dan kebijakan institusi ini mencerminkan prinsip tata kelola adaptif yang dijabarkan oleh Bitchikashvili et al. (2023) dalam manajemen institusi pendidikan tinggi modern. Pengembangan infrastruktur digital yang berkelanjutan ini dipastikan mampu memfasilitasi kebutuhan administrasi berskala besar dengan tetap mempertahankan standar kualitas pelayanan prima."
]

def fix_421():
    for filepath in FILES:
        print(f"Processing: {filepath}")
        doc = docx.Document(filepath)
        
        insert_index = -1
        target_style = None
        
        # Identify where to insert and remove old paragraphs
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text.startswith("Lebih jauh, capaian ini melahirkan") or text.startswith("Pencapaian skor kepraktisan sebesar 92,13% pada uji coba lapangan mengindikasikan lebih dari sekadar"):
                if insert_index == -1:
                    insert_index = i
                    target_style = p.style
                p.text = "" # Delete the old paragraph content
        
        if insert_index != -1:
            # We insert the new paragraphs backwards at the insert_index so they appear in correct order
            # Actually, `insert_paragraph_before` is easier
            reference_p = doc.paragraphs[insert_index]
            for new_text in NEW_PARAGRAPHS:
                new_p = reference_p.insert_paragraph_before(new_text)
                new_p.style = target_style
                
                # Apply highlight for the highlighted version
                if 'Highlighted' in filepath:
                    for run in new_p.runs:
                        run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
            
            print(f"Successfully replaced text in {filepath}")
        
        doc.save(filepath)

if __name__ == '__main__':
    fix_421()
