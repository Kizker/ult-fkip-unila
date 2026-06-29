import docx
import os

FILES = [
    r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx',
    r'docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx'
]

SARAN_TEXT = "3. Integrasi Penomoran Surat dan Penjadwalan Pimpinan: Disarankan untuk mengembangkan fitur kalender digital terintegrasi pada dasbor pimpinan fakultas yang tersinkronisasi langsung dengan mesin penomoran surat. Hal ini bertujuan untuk mencegah terjadinya tumpang tindih alokasi waktu pimpinan dan meminimalisir kesalahan cetak pada nomor dokumen kelulusan yang saat ini masih rentan terjadi."
OLD_SARAN_TEXT = "Integrasi Penomoran Surat dan Penjadwalan Pimpinan: Disarankan untuk mengembangkan fitur kalender digital terintegrasi pada dasbor pimpinan fakultas yang tersinkronisasi langsung dengan mesin penomoran surat."

def fix_saran():
    for filepath in FILES:
        print(f"Processing: {filepath}")
        doc = docx.Document(filepath)
        
        # Step 1: Find and delete the incorrectly placed paragraph
        for p in doc.paragraphs:
            if OLD_SARAN_TEXT in p.text and not p.text.startswith("3."):
                p.text = "" # Delete it by clearing text (safer than p._element.getparent().remove(p._element) which sometimes breaks XML)
                print("Removed misplaced paragraph.")

        # Step 2: Find point 2 in Saran and insert point 3 after it
        for i, p in enumerate(doc.paragraphs):
            if '2. Pengujian Keamanan Tingkat Lanjut:' in p.text:
                print(f"Found point 2 at line {i}")
                # We can either use insert_paragraph_before on the next paragraph
                new_p = doc.paragraphs[i+1].insert_paragraph_before(SARAN_TEXT, style=p.style)
                
                # Apply highlight if needed
                if 'Highlighted' in filepath:
                    for run in new_p.runs:
                        run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
                        
                print("Inserted point 3.")
                break
                
        doc.save(filepath)
        print("Saved.")

if __name__ == '__main__':
    fix_saran()
