import re

def get_context():
    doc_path = r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx"
    from docx import Document
    doc = Document(doc_path)
    
    with open("context.txt", "w", encoding="utf-8") as f:
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if "Gambar 43" in text or "Gambar 44" in text or "Gambar 45" in text or "Gambar 46" in text or "Gambar 47" in text or "Gambar 48" in text or "Gambar 49" in text or "Gambar 50" in text or "Gambar 51" in text or "Gambar 52" in text:
                f.write(f"\n\n=== {text} ===\n")
                start = max(0, i - 4)
                end = min(len(doc.paragraphs), i + 5)
                for j in range(start, end):
                    f.write(f"[{j}]: {doc.paragraphs[j].text.strip()}\n")
                    
get_context()
