import docx
from docx.shared import Pt
import re

def update_respondents_to_table(doc_path, is_highlighted=False):
    print(f"Processing {doc_path}")
    doc = docx.Document(doc_path)
    
    # Find the paragraphs
    start_idx = -1
    end_idx = -1
    paragraphs_to_convert = []
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith("Responden 1 (Admin Program Studi):"):
            start_idx = i
        if start_idx != -1 and text.startswith("Responden 17 (Staf Unit"):
            end_idx = i
            break
            
    if start_idx == -1 or end_idx == -1:
        print("Could not find the target paragraphs.")
        return
        
    print(f"Found paragraphs from {start_idx} to {end_idx}")
    
    for i in range(start_idx, end_idx + 1):
        paragraphs_to_convert.append(doc.paragraphs[i].text.strip())
        
    # Get the parent element
    first_p = doc.paragraphs[start_idx]
    parent = first_p._element.getparent()
    insert_index = parent.index(first_p._element)
    
    # 1. Insert Table Caption ABOVE the table
    caption_p = docx.text.paragraph.Paragraph(docx.oxml.shared.OxmlElement('w:p'), doc._body)
    caption_p.text = "Tabel 4.2. Tanggapan Kualitatif Responden dan Tindak Lanjut" # Just a placeholder number, usually user will fix it or we use 4.x
    caption_p.alignment = 1 # Center
    
    for run in caption_p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if is_highlighted:
            from docx.enum.text import WD_COLOR_INDEX
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
    parent.insert(insert_index, caption_p._element)
    insert_index += 1
    
    # 2. Insert Table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Header
    table.cell(0, 0).text = "Responden"
    table.cell(0, 1).text = "Komentar / Saran"
    table.cell(0, 2).text = "Tindak Lanjut"
    
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            p.alignment = 1 # Center
            for run in p.runs:
                run.font.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                
    # Data extraction
    for text in paragraphs_to_convert:
        # Format: "Responden X (Role): [Comment]. [Action]."
        # We need to split intelligently.
        parts = text.split("): ", 1)
        if len(parts) == 2:
            responden = parts[0] + ")"
            rest = parts[1]
            
            # Try to split by "Peneliti" or similar
            if "Peneliti menindaklanjutinya" in rest:
                saran, tindak = rest.split("Peneliti menindaklanjutinya", 1)
                tindak = "Peneliti menindaklanjutinya" + tindak
            elif "Peneliti merespons" in rest:
                saran, tindak = rest.split("Peneliti merespons", 1)
                tindak = "Peneliti merespons" + tindak
            elif "Komentar ini menjadi bukti" in rest:
                saran, tindak = rest.split("Komentar ini menjadi bukti", 1)
                tindak = "Komentar ini menjadi bukti" + tindak
            elif "Peneliti menindaklanjuti" in rest:
                saran, tindak = rest.split("Peneliti menindaklanjuti", 1)
                tindak = "Peneliti menindaklanjuti" + tindak
            elif "Peneliti menyambut baik" in rest:
                saran, tindak = rest.split("Peneliti menyambut baik", 1)
                tindak = "Peneliti menyambut baik" + tindak
            else:
                saran = rest
                tindak = "-"
                
            row = table.add_row()
            row.cells[0].text = responden
            row.cells[1].text = saran.strip()
            row.cells[2].text = tindak.strip()
            
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                        if is_highlighted:
                            from docx.enum.text import WD_COLOR_INDEX
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                            
    parent.insert(insert_index, table._element)
    insert_index += 1
    
    # 3. Insert Explanatory Paragraph
    explanation_text = (
        "Berdasarkan Tabel 4.2 di atas, dapat dilihat bahwa tanggapan dari berbagai kelompok responden sangat positif "
        "dan membangun. Saran-saran yang diberikan, seperti penambahan fitur visualisasi jadwal seminar, otomatisasi "
        "penomoran surat, hingga pengaturan kepadatan teks pada antarmuka, telah ditindaklanjuti secara langsung oleh "
        "peneliti dalam proses pengembangan. Perbaikan-perbaikan ini memastikan bahwa sistem informasi yang dibangun "
        "tidak hanya berfungsi dengan baik secara teknis, tetapi juga sangat adaptif dan responsif terhadap kebutuhan "
        "nyata pengguna di lapangan."
    )
    exp_p = docx.text.paragraph.Paragraph(docx.oxml.shared.OxmlElement('w:p'), doc._body)
    exp_p.text = explanation_text
    exp_p.alignment = 3 # Justify
    
    # Apply indent like other paragraphs (First line indent)
    exp_p.paragraph_format.first_line_indent = Pt(36) # 1.27 cm roughly
    exp_p.paragraph_format.line_spacing = 1.5
    
    for run in exp_p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if is_highlighted:
            from docx.enum.text import WD_COLOR_INDEX
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
    parent.insert(insert_index, exp_p._element)
    
    # Remove old paragraphs
    for i in range(start_idx, end_idx + 1):
        p_elem = doc.paragraphs[i]._element
        p_elem.getparent().remove(p_elem)
        
    doc.save(doc_path)
    print("Success")

update_respondents_to_table(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Clean.docx", False)
update_respondents_to_table(r"C:\laragon\www\ult-fkip-unila\docs\skripsi\hasil_update\001_Skripsi_Andricha Dea Mitra_Highlighted.docx", True)
